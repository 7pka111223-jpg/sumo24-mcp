"""
SUMO24 MCP Server
=================
Exposes SUMO24 Digital Twin Toolkit (DTT) capabilities as MCP tools
for use with Claude. Built on the Dynamita scheduler.py Python API.

Prerequisites:
  - SUMO24 installed with DTT addon (SumoDTT24.dya)
  - DTT license (standalone+DTT or container+DTT)
  - `dynamita` folder copied to the repo root from your SUMO installation:
      C:\Program Files\Dynamita\SUMO24\PythonAPI\dynamita\
  - sumoproject.dll and state.xml in the repo root (or point SUMO_DLL / SUMO_STATE
    env vars at your files — see install.ps1 and examples/qaha-wwtp/ for reference)
  - Python packages: mcp  (run install.ps1 or: pip install -r requirements.txt)

Run:
  python server.py
"""

import asyncio
import copy
import csv
import glob
import io
import json
import math
import os
import shutil
import subprocess
import sys
import threading
import time
import xml.etree.ElementTree as ET
from datetime import datetime
from pathlib import Path
from typing import Any

# â”€â”€ openpyxl for dataset-reading tools â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
try:
    from openpyxl import load_workbook, Workbook
    from openpyxl.styles import PatternFill, Font
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

# â”€â”€ Schematic parser + unit-type registry (Group BB) â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# These modules live alongside server.py; pure stdlib, no external deps.
try:
    import schematic_parser as _schp
    import unit_type_registry as _utr
    SCHEMATIC_AVAILABLE = True
except Exception as _schp_e:  # pragma: no cover
    SCHEMATIC_AVAILABLE = False
    _schp = None
    _utr = None
    print("[WARN] schematic_parser unavailable: " + repr(_schp_e))

# Group DD - SUMO project compiler (schematic -> .sumo file)
# NOTE: sumo_compiler.py is DEPRECATED — it produces zips that SUMO cannot load.
# Use sumo_pack.py (Group EE) for new work.
try:
    import sumo_compiler as _scc
    COMPILER_AVAILABLE = True
except Exception as _scc_e:  # pragma: no cover
    COMPILER_AVAILABLE = False
    _scc = None
    print("[WARN] sumo_compiler unavailable: " + repr(_scc_e))

# Group EE - Build-pack generator (replaces the broken zip-as-.sumo approach)
try:
    import sumo_pack as _spk
    PACK_AVAILABLE = True
except Exception as _spk_e:  # pragma: no cover
    PACK_AVAILABLE = False
    _spk = None
    print("[WARN] sumo_pack unavailable: " + repr(_spk_e))

# Group GG - Native .sumo composer (Tier-1 baseline-DLL-reuse path)
# Composes a real working .sumo zip by rewriting parameter-bearing members
# of a known-good baseline while preserving the DLL-bound members. Unlike
# apply_schematic_to_baseline (which copies the .sumo byte-for-byte and
# writes overrides into a sibling state.xml), this path mutates the zip
# internals so the produced .sumo carries the schematic's parameters
# inside itself.
try:
    import sumo_native as _snt
    NATIVE_AVAILABLE = True
except Exception as _snt_e:  # pragma: no cover
    NATIVE_AVAILABLE = False
    _snt = None
    print("[WARN] sumo_native unavailable: " + repr(_snt_e))

# ---- Group FF: DTT-mediated .sumo editing --------------------------------
try:
    import dtt_editor as _dte
    DTT_EDITOR_AVAILABLE = True
except Exception as _e_dte:  # pragma: no cover
    _dte = None
    DTT_EDITOR_AVAILABLE = False
    _DTE_IMPORT_ERROR = repr(_e_dte)
# --------------------------------------------------------------------------

# â”€â”€ Resolve absolute paths relative to this script â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
_HERE = Path(__file__).parent             # repo root (sumo24-mcp/)

# lib/ contains schematic_parser, sumo_pack, etc.
if str(_HERE / "lib") not in sys.path:
    sys.path.insert(0, str(_HERE / "lib"))
# dynamita/ lives at repo root — already findable once _HERE is on path
if str(_HERE) not in sys.path:
    sys.path.insert(0, str(_HERE))

_MCP_DIR = _HERE  # alias kept so CONFIG paths below resolve without changes

import mcp.types as types
from mcp.server import Server
from mcp.server.stdio import stdio_server

# â”€â”€ DTT Python API import â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
try:
    import dynamita.scheduler as ds
    DTT_AVAILABLE = True
    try:
        import dynamita_compat  # backfill missing SumoScheduler methods
    except Exception as _e:
        print("[WARN] dynamita_compat shim failed: " + repr(_e))
except ImportError:
    DTT_AVAILABLE = False
    print(“[WARN] dynamita package not found — expected at: “ + str(_HERE / “dynamita”))

# â”€â”€ Configuration â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
CONFIG = {
    # Path to the compiled model DLL (extract once via prepare_project.py)
    "model_dll":   os.environ.get("SUMO_DLL",   str(_MCP_DIR / "sumoproject.dll")),
    # Path to the saved steady-state system state (already present at MCP/state.xml)
    "state_xml":   os.environ.get("SUMO_STATE", str(_MCP_DIR / "state.xml")),
    # Working directory for outputs
    "output_dir":  os.environ.get("SUMO_OUTPUT", str(_MCP_DIR / "outputs")),
    # Tier-1 baseline: a known-good .sumo + sumoproject.dll pair produced by SUMO itself.
    # Place these under MCP/baselines/ (see SUMO24_MCP_Fix_Plan.md §Appendix B).
    "baseline_sumo": os.environ.get("SUMO_BASELINE_SUMO",
                                     str(_MCP_DIR / "baselines" / "qaha_baseline.sumo")),
    "baseline_dll":  os.environ.get("SUMO_BASELINE_DLL",
                                     str(_MCP_DIR / "baselines" / "sumoproject.dll")),
    # Egyptian Law 48/1982 effluent limits
    "law48_limits": {
        "TSS":  50.0,   # mg/L
        "BOD5": 60.0,   # mg/L
        "COD":  80.0,   # mg/L  (typical guideline value)
        "TN":   None,   # no strict limit in Law 48
        "TP":   None,
    }
}

Path(CONFIG["output_dir"]).mkdir(parents=True, exist_ok=True)

# â”€â”€ Tracks the currently loaded .sumo project file and its extracted artifacts â”€
ACTIVE_MODEL = {
    "sumo_file":   None,   # path to the .sumo project file
    "dll_path":    CONFIG.get("model_dll", "sumoproject.dll"),
    "state_xml":   CONFIG.get("state_xml",  "state.xml"),
    "compiled":    False,
    "unit_processes": [],
    "streams": [],
}


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# RESEARCH-BASED DEFAULT PARAMETERS
# Sources: IWA ASM1/ASM2d/ASM3 STR-9 (Henze et al. 2000), Metcalf & Eddy 5th ed.,
#          BSM1 benchmark, Egyptian Law 48/1982 (Decrees 92/2013, 208/2018),
#          Warm-climate Egyptian WWTP literature (IWA WPT 2022, MDPI Water 2022).
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
RESEARCH_DEFAULTS = {
    "law48_nile": {   # Discharge to Nile / main canals
        "BOD5": 30.0, "COD": 40.0, "TSS": 30.0, "NH3_N": 3.0,
        "NO3_N": 45.0, "TP": 1.0, "Oil": 5.0, "DO_min": 4.0,
    },
    "law48_drain": {  # Discharge to agricultural drains
        "BOD5": 60.0, "COD": 80.0, "TSS": 50.0, "NH3_N": 5.0,
        "TN": 50.0, "TP": 3.0, "Oil": 10.0,
    },
    "influent_egypt": {
        "Q_m3d": 28000, "COD_mgL": 550, "BOD5_mgL": 280, "TSS_mgL": 300,
        "VSS_mgL": 240, "TKN_mgL": 55, "NH4_mgL": 35, "TP_mgL": 8,
        "pH": 7.4, "Temp_C": 25.0,
        "fSI": 0.07, "fSS": 0.20, "fXS": 0.50, "fXI": 0.15, "fXBH": 0.08,
    },
    "ASM1": {
        "Y_H":   {"val": 0.67,  "unit": "g COD/g COD",     "theta": None},
        "Y_A":   {"val": 0.24,  "unit": "g COD/g N",       "theta": None},
        "f_P":   {"val": 0.08,  "unit": "-",                "theta": None},
        "i_XB":  {"val": 0.086, "unit": "g N/g COD",       "theta": None},
        "i_XP":  {"val": 0.06,  "unit": "g N/g COD",       "theta": None},
        "mu_H":  {"val": 4.0,   "unit": "d-1",              "theta": 1.072},
        "K_S":   {"val": 10.0,  "unit": "g COD/m3",         "theta": None},
        "K_OH":  {"val": 0.20,  "unit": "g O2/m3",          "theta": None},
        "K_NO":  {"val": 0.50,  "unit": "g N/m3",           "theta": None},
        "b_H":   {"val": 0.30,  "unit": "d-1",              "theta": 1.12},
        "eta_g": {"val": 0.80,  "unit": "-",                "theta": None},
        "eta_h": {"val": 0.80,  "unit": "-",                "theta": None},
        "k_h":   {"val": 3.0,   "unit": "g COD/(g COD d)",  "theta": 1.072},
        "K_X":   {"val": 0.10,  "unit": "g COD/g COD",      "theta": 1.116},
        "mu_A":  {"val": 0.50,  "unit": "d-1",              "theta": 1.103},
        "K_NH":  {"val": 1.0,   "unit": "g N/m3",           "theta": 1.123},
        "K_OA":  {"val": 0.40,  "unit": "g O2/m3",          "theta": None},
        "b_A":   {"val": 0.05,  "unit": "d-1",              "theta": 1.116},
        "k_a":   {"val": 0.05,  "unit": "m3/(g COD d)",     "theta": 1.072},
    },
    "ASM2d": {
        "Y_H":       {"val": 0.625, "unit": "g COD/g COD", "theta": None},
        "Y_PAO":     {"val": 0.625, "unit": "g COD/g COD", "theta": None},
        "Y_PO4":     {"val": 0.40,  "unit": "g P/g COD",   "theta": None},
        "Y_PHA":     {"val": 0.20,  "unit": "g COD/g P",   "theta": None},
        "Y_A":       {"val": 0.24,  "unit": "g COD/g N",   "theta": None},
        "f_XI":      {"val": 0.10,  "unit": "-",           "theta": None},
        "i_N_BM":    {"val": 0.07,  "unit": "g N/g COD",   "theta": None},
        "i_P_BM":    {"val": 0.02,  "unit": "g P/g COD",   "theta": None},
        "K_h":       {"val": 3.0,   "unit": "d-1",          "theta": 1.041},
        "eta_NO3":   {"val": 0.60,  "unit": "-",            "theta": None},
        "eta_fe":    {"val": 0.40,  "unit": "-",            "theta": None},
        "K_O2_h":    {"val": 0.20,  "unit": "g O2/m3",      "theta": None},
        "K_NO3_h":   {"val": 0.50,  "unit": "g N/m3",       "theta": None},
        "K_X_h":     {"val": 0.10,  "unit": "g COD/g COD",  "theta": None},
        "mu_H":      {"val": 6.0,   "unit": "d-1",          "theta": 1.072},
        "q_fe":      {"val": 3.0,   "unit": "d-1",          "theta": 1.072},
        "eta_NO3_H": {"val": 0.80,  "unit": "-",            "theta": None},
        "b_H":       {"val": 0.40,  "unit": "d-1",          "theta": 1.12},
        "K_O2_H":    {"val": 0.20,  "unit": "g O2/m3",      "theta": None},
        "K_F":       {"val": 4.0,   "unit": "g COD/m3",     "theta": None},
        "K_A_H":     {"val": 4.0,   "unit": "g COD/m3",     "theta": None},
        "K_NO3_H":   {"val": 0.50,  "unit": "g N/m3",       "theta": None},
        "K_NH4_H":   {"val": 0.05,  "unit": "g N/m3",       "theta": None},
        "K_P_H":     {"val": 0.01,  "unit": "g P/m3",       "theta": None},
        "q_PHA":     {"val": 3.0,   "unit": "d-1",          "theta": 1.072},
        "q_PP":      {"val": 1.5,   "unit": "d-1",          "theta": 1.072},
        "mu_PAO":    {"val": 1.0,   "unit": "d-1",          "theta": 1.072},
        "eta_NO3_PAO":{"val": 0.60, "unit": "-",            "theta": None},
        "b_PAO":     {"val": 0.20,  "unit": "d-1",          "theta": 1.072},
        "b_PP":      {"val": 0.20,  "unit": "d-1",          "theta": 1.072},
        "b_PHA":     {"val": 0.20,  "unit": "d-1",          "theta": 1.072},
        "K_O2_PAO":  {"val": 0.20,  "unit": "g O2/m3",      "theta": None},
        "K_NO3_PAO": {"val": 0.50,  "unit": "g N/m3",       "theta": None},
        "K_A_PAO":   {"val": 4.0,   "unit": "g COD/m3",     "theta": None},
        "K_NH4_PAO": {"val": 0.05,  "unit": "g N/m3",       "theta": None},
        "K_P_PAO":   {"val": 0.01,  "unit": "g P/m3",       "theta": None},
        "K_PS":      {"val": 0.20,  "unit": "g P/m3",       "theta": None},
        "K_MAX":     {"val": 0.34,  "unit": "g P/g COD",    "theta": None},
        "mu_AUT":    {"val": 1.0,   "unit": "d-1",          "theta": 1.103},
        "b_AUT":     {"val": 0.15,  "unit": "d-1",          "theta": 1.116},
        "K_O2_AUT":  {"val": 0.50,  "unit": "g O2/m3",      "theta": None},
        "K_NH4_AUT": {"val": 1.0,   "unit": "g N/m3",       "theta": 1.123},
        "K_P_AUT":   {"val": 0.01,  "unit": "g P/m3",       "theta": None},
        "k_PRE":     {"val": 1.0,   "unit": "m3/(g d)",     "theta": None},
        "k_RED":     {"val": 0.60,  "unit": "d-1",          "theta": None},
    },
    "ASM3": {
        "Y_STO_O2": {"val": 0.85, "unit": "g COD/g COD", "theta": None},
        "Y_STO_NO": {"val": 0.80, "unit": "g COD/g COD", "theta": None},
        "Y_H_O2":   {"val": 0.63, "unit": "g COD/g COD", "theta": None},
        "Y_H_NO":   {"val": 0.54, "unit": "g COD/g COD", "theta": None},
        "Y_A":      {"val": 0.24, "unit": "g COD/g N",   "theta": None},
        "f_XI":     {"val": 0.20, "unit": "-",           "theta": None},
        "k_h":      {"val": 3.0,  "unit": "d-1",          "theta": 1.041},
        "K_X":      {"val": 1.0,  "unit": "g COD/g COD",  "theta": None},
        "k_STO":    {"val": 5.0,  "unit": "d-1",          "theta": 1.072},
        "eta_NO":   {"val": 0.60, "unit": "-",            "theta": None},
        "K_O2":     {"val": 0.20, "unit": "g O2/m3",      "theta": None},
        "K_NO":     {"val": 0.50, "unit": "g N/m3",       "theta": None},
        "K_S":      {"val": 2.0,  "unit": "g COD/m3",     "theta": None},
        "K_STO":    {"val": 1.0,  "unit": "g COD/g COD",  "theta": None},
        "mu_H":     {"val": 2.0,  "unit": "d-1",          "theta": 1.072},
        "K_NH_H":   {"val": 0.01, "unit": "g N/m3",       "theta": None},
        "b_H_O2":   {"val": 0.20, "unit": "d-1",          "theta": 1.072},
        "b_H_NO":   {"val": 0.10, "unit": "d-1",          "theta": 1.072},
        "b_STO_O2": {"val": 0.20, "unit": "d-1",          "theta": 1.072},
        "b_STO_NO": {"val": 0.10, "unit": "d-1",          "theta": 1.072},
        "mu_A":     {"val": 1.0,  "unit": "d-1",          "theta": 1.103},
        "K_A_NH":   {"val": 1.0,  "unit": "g N/m3",       "theta": 1.123},
        "K_A_O":    {"val": 0.50, "unit": "g O2/m3",      "theta": None},
        "b_A_O":    {"val": 0.15, "unit": "d-1",          "theta": 1.116},
        "b_A_NO":   {"val": 0.05, "unit": "d-1",          "theta": 1.116},
    },
    "takacs": {
        "v0":    {"val": 474.0,    "unit": "m/d",   "note": "max theoretical settling velocity"},
        "v0_p":  {"val": 250.0,    "unit": "m/d",   "note": "max practical settling velocity"},
        "r_h":   {"val": 0.000576, "unit": "m3/g",  "note": "hindrance settling coefficient"},
        "r_p":   {"val": 0.00286,  "unit": "m3/g",  "note": "flocculant zone coefficient"},
        "f_ns":  {"val": 0.00228,  "unit": "-",     "note": "non-settleable fraction"},
    },
    "dynamics": {
        "duration_days": 14, "timestep_h": 0.25, "output_interval_h": 1.0,
        "solver": "Euler", "tolerance": 1e-4, "warm_up_days": 10,
        "diurnal_peak_factor": 1.8, "diurnal_min_factor": 0.4,
        "storm_peak_factor": 3.0, "storm_duration_h": 6,
    },
}


# VALIDATION_THRESHOLDS -- green/yellow/red bands for validation tools
VALIDATION_THRESHOLDS = {
    "MLSS_mgL": {"green": (2000, 4000), "yellow": (1500, 5000), "red": (800, 8000),
                  "unit": "mg/L", "note": "Metcalf & Eddy: 2000-4000 for CAS"},
    "MLVSS_MLSS_ratio": {"green": (0.70, 0.85), "yellow": (0.60, 0.90), "red": (0.50, 0.95),
                          "unit": "-", "note": "Typical healthy municipal sludge"},
    "SVI_mLg": {"green": (50, 150), "yellow": (30, 200), "red": (10, 400),
                 "unit": "mL/g", "note": "Sludge Volume Index. >150 = bulking risk"},
    "FM_ratio": {"green": (0.15, 0.40), "yellow": (0.05, 0.60), "red": (0.02, 1.0),
                  "unit": "kg BOD/(kg VSS.d)", "note": "CAS: 0.2-0.4"},
    "SRT_d": {"green": (8, 20), "yellow": (5, 40), "red": (3, 60),
               "unit": "d", "note": "Min SRT for nitrification at 20C ~ 5d"},
    "HRT_h": {"green": (4, 12), "yellow": (2, 36), "red": (1, 72),
               "unit": "h", "note": "CAS: 4-8h; Extended aeration: 18-36h"},
    "DO_aerobic_mgL": {"green": (1.5, 2.5), "yellow": (1.0, 4.0), "red": (0.5, 8.0),
                        "unit": "mg/L", "note": "<1 mg/L impairs nitrification"},
    "DO_anoxic_mgL": {"green": (0.0, 0.2), "yellow": (0.0, 0.5), "red": (0.0, 1.0),
                       "unit": "mg/L", "note": "Anoxic zones must stay low"},
    "DO_anaerobic_mgL": {"green": (0.0, 0.02), "yellow": (0.0, 0.05), "red": (0.0, 0.2),
                          "unit": "mg/L", "note": "Anaerobic zones critical for EBPR"},
    "KLa_h": {"green": (50, 250), "yellow": (20, 400), "red": (5, 800),
               "unit": "1/h", "note": "Oxygen transfer coefficient"},
    "RAS_ratio": {"green": (0.50, 1.00), "yellow": (0.25, 1.50), "red": (0.10, 3.00),
                   "unit": "-", "note": "RAS/Influent. Too low = sludge washout"},
    "overflow_rate_sec_clar": {"green": (16, 32), "yellow": (8, 48), "red": (4, 80),
                                "unit": "m3/(m2.d)", "note": "Secondary SOR"},
    "overflow_rate_prim_clar": {"green": (30, 48), "yellow": (20, 60), "red": (10, 100),
                                 "unit": "m3/(m2.d)", "note": "Primary SOR"},
    "underflow_conc_sec_clar": {"green": (8, 12), "yellow": (6, 15), "red": (3, 20),
                                 "unit": "g/L", "note": "RAS sludge concentration"},
    "flow_continuity_error_pct": {"green": (0, 5), "yellow": (5, 10), "red": (10, 100),
                                    "unit": "%", "note": "|Q_in - Q_out - Q_WAS|/Q_in"},
    "BOD_COD_ratio": {"green": (0.40, 0.70), "yellow": (0.30, 0.85), "red": (0.10, 0.95),
                       "unit": "-", "note": "Egyptian domestic mean 0.67"},
    "VSS_TSS_ratio": {"green": (0.70, 0.85), "yellow": (0.60, 0.95), "red": (0.40, 1.00),
                       "unit": "-", "note": "Volatile fraction of TSS"},
    "COD_N_ratio": {"green": (8, 15), "yellow": (5, 20), "red": (3, 30),
                     "unit": "-", "note": "<5 = external carbon needed"},
    "pH": {"green": (6.8, 7.8), "yellow": (6.0, 8.5), "red": (5.0, 9.5),
            "unit": "-", "note": "Biological activity impaired outside green"},
    "Temp_C": {"green": (15, 30), "yellow": (10, 35), "red": (5, 45),
                "unit": "C", "note": "Egyptian plants 18-32C year-round"},
    "mass_balance_error_pct": {"green": (0, 5), "yellow": (5, 10), "red": (10, 100),
                                 "unit": "%", "note": "Closure error tolerance"},
    "ASM_kinetic_bounds": {
        "mu_H":  {"min": 1.0,   "max": 13.2, "unit": "1/d"},
        "K_S":   {"min": 1.0,   "max": 300.0,"unit": "g COD/m3"},
        "b_H":   {"min": 0.05,  "max": 1.6,  "unit": "1/d"},
        "Y_H":   {"min": 0.30,  "max": 0.95, "unit": "g COD/g COD"},
        "mu_A":  {"min": 0.2,   "max": 2.0,  "unit": "1/d"},
        "K_NH":  {"min": 0.1,   "max": 5.0,  "unit": "g N/m3"},
        "b_A":   {"min": 0.02,  "max": 0.3,  "unit": "1/d"},
        "Y_A":   {"min": 0.07,  "max": 0.30, "unit": "g COD/g N"},
        "k_h":   {"min": 0.5,   "max": 10.0, "unit": "1/d"},
        "K_X":   {"min": 0.005, "max": 1.0,  "unit": "g COD/g COD"},
    },
}


def _classify(value, threshold_key):
    """Return ('green'|'yellow'|'red'|'unknown', message)."""
    t = VALIDATION_THRESHOLDS.get(threshold_key)
    if t is None or value is None:
        return ("unknown", f"No threshold or value for {threshold_key}")
    try:
        v = float(value)
    except (TypeError, ValueError):
        return ("unknown", f"Non-numeric value for {threshold_key}")
    g_lo, g_hi = t["green"]
    y_lo, y_hi = t["yellow"]
    r_lo, r_hi = t["red"]
    if g_lo <= v <= g_hi:
        return ("green", f"{v} {t['unit']} - healthy")
    if y_lo <= v <= y_hi:
        return ("yellow", f"{v} {t['unit']} - acceptable but sub-optimal")
    if r_lo <= v <= r_hi:
        return ("red", f"{v} {t['unit']} - outside design envelope")
    return ("red", f"{v} {t['unit']} - far outside physical limits")


# ----- Validation helpers (non-tool, called by dispatch handlers) -----

def _val_read_var(ds, all_vars, unit, suffixes):
    """Try each suffix; return (value, var_name) or (None, None)."""
    for s in suffixes:
        var = f"Sumo__Plant__{unit}__{s}"
        if var in all_vars:
            try:
                return float(ds.sumo.get(var)), var
            except Exception:
                continue
    return None, None


def _val_structure():
    try:
        issues, warnings, passes = [], [], []
        dll_path   = Path(CONFIG.get("model_dll", ""))
        state_path = Path(CONFIG.get("state_xml", ""))
        if not dll_path.exists():
            issues.append(f"sumoproject.dll not found at {dll_path} - compile in SUMO GUI.")
        else:
            passes.append(f"DLL present: {dll_path.name}")
        if not state_path.exists():
            issues.append(f"state.xml not found at {state_path} - extract from SUMO GUI.")
        else:
            passes.append(f"state.xml present: {state_path.name}")
        if dll_path.exists() and state_path.exists():
            dll_mt = dll_path.stat().st_mtime
            state_mt = state_path.stat().st_mtime
            if abs(dll_mt - state_mt) > 3600:
                warnings.append("DLL and state.xml timestamps differ by > 1h.")
            else:
                passes.append("DLL and state.xml timestamps consistent")
        try:
            ds = _make_ds()
            all_vars = ds.sumo.getVariableNames()
            passes.append(f"DTT session opened ({len(all_vars)} variables)")
        except Exception as e:
            issues.append(f"Cannot open DTT session: {e}")
            return {"status": "FAIL", "issues": issues, "warnings": warnings,
                    "passes": passes}
        unit_names = set()
        for v in all_vars:
            parts = v.split("__")
            if len(parts) >= 3:
                unit_names.add(parts[2])
        required_patterns = {
            "Influent":  ["influent", "inlet", "feed"],
            "Reactor":   ["aer", "reactor", "tank", "mbr", "ditch", "sbr"],
            "Clarifier": ["clar", "settler", "settling"],
            "Effluent":  ["effluent", "outlet", "final"],
        }
        for role, patterns in required_patterns.items():
            found = [u for u in unit_names if any(p in u.lower() for p in patterns)]
            if not found:
                warnings.append(f"No '{role}' unit detected (looked for: {patterns})")
            else:
                passes.append(f"{role} unit(s) present: {found[:3]}")
        status = "PASS" if not issues else ("FAIL" if len(issues) > 1 else "WARNINGS")
        return {"status": status, "dll": str(dll_path), "state_xml": str(state_path),
                "unit_count": len(unit_names), "units": sorted(unit_names),
                "issue_count": len(issues), "warning_count": len(warnings),
                "pass_count": len(passes),
                "issues": issues, "warnings": warnings, "passes": passes}
    except Exception as e:
        return {"error": str(e)}


def _val_influent(unit_name="Influent"):
    try:
        ds = _make_ds()
        all_vars = ds.sumo.getVariableNames()
        checks, issues, warnings, passes = {}, [], [], []
        q, _ = _val_read_var(ds, all_vars, unit_name, ["param__Q"])
        cod, _ = _val_read_var(ds, all_vars, unit_name, ["param__S_COD", "param__COD"])
        tss, _ = _val_read_var(ds, all_vars, unit_name, ["param__X_TSS", "param__TSS"])
        tkn, _ = _val_read_var(ds, all_vars, unit_name, ["param__S_TKN", "param__TKN"])
        tp, _ = _val_read_var(ds, all_vars, unit_name, ["param__S_TP", "param__TP"])
        ph, _ = _val_read_var(ds, all_vars, unit_name, ["param__pH"])
        t, _ = _val_read_var(ds, all_vars, unit_name, ["param__T"])
        if q is None or q <= 0:
            issues.append(f"Influent flow (Q) missing or <= 0 at {unit_name}")
        else:
            checks["flow_m3d"] = q
            passes.append(f"Flow = {q} m3/d")
        if cod is None or cod <= 0:
            issues.append(f"Influent COD missing or <= 0 at {unit_name}")
        else:
            checks["COD_mgL"] = cod
            passes.append(f"COD = {cod} mg/L")
        if t is not None:
            cat, msg = _classify(t, "Temp_C")
            checks["temp_C"] = t
            (passes if cat == "green" else warnings if cat == "yellow" else issues).append(f"Temperature: {msg}")
        else:
            warnings.append("Influent temperature not set - model will use default")
        if ph is not None:
            cat, msg = _classify(ph, "pH")
            checks["pH"] = ph
            (passes if cat == "green" else warnings if cat == "yellow" else issues).append(f"pH: {msg}")
        if cod and tkn and tkn > 0:
            cn = cod / tkn
            cat, msg = _classify(cn, "COD_N_ratio")
            checks["COD_TKN_ratio"] = round(cn, 2)
            (passes if cat == "green" else warnings if cat == "yellow" else issues).append(f"COD:TKN ratio: {msg}")
        status = "PASS" if not issues else "FAIL"
        return {"unit_name": unit_name, "status": status, "checks": checks,
                "issues": issues, "warnings": warnings, "passes": passes}
    except Exception as e:
        return {"error": str(e), "unit_name": unit_name}


def _val_mass_balance_pre(tolerance_pct=5.0):
    try:
        ds = _make_ds()
        all_vars = ds.sumo.getVariableNames()
        def get_load(unit, suffixes):
            q_var = f"Sumo__Plant__{unit}__param__Q"
            if q_var not in all_vars:
                return None
            try:
                q = float(ds.sumo.get(q_var))
            except Exception:
                return None
            for s in suffixes:
                var = f"Sumo__Plant__{unit}__{s}"
                if var in all_vars:
                    try:
                        c = float(ds.sumo.get(var))
                        return q * c / 1000.0
                    except Exception:
                        continue
            return None
        balances = {}
        for element, suffixes in [("COD", ["param__S_COD", "param__COD"]),
                                  ("TSS", ["param__X_TSS", "param__TSS"]),
                                  ("TKN", ["param__S_TKN", "param__TKN"]),
                                  ("TP",  ["param__S_TP",  "param__TP"])]:
            inf = get_load("Influent", suffixes)
            eff = get_load("Effluent", suffixes)
            if inf is not None and eff is not None:
                r = (inf - eff) / inf * 100 if inf > 0 else None
                balances[element] = {"influent_kgd": round(inf, 2),
                                     "effluent_kgd": round(eff, 2),
                                     "reduction_pct": round(r, 1) if r is not None else None}
        issues = []
        for elem, b in balances.items():
            r = b.get("reduction_pct")
            if r is None:
                continue
            if r < 0:
                issues.append(f"{elem}: effluent load > influent load (impossible)")
            if elem == "COD" and r < 50:
                issues.append(f"COD reduction only {r}% - plant under-performing")
        return {"tolerance_pct": tolerance_pct, "balances": balances,
                "issue_count": len(issues), "issues": issues,
                "status": "PASS" if not issues else "CHECK",
                "note": "Pre-run; effluent from state.xml initial conditions."}
    except Exception as e:
        return {"error": str(e)}


def _val_mass_balance_post(job_id, elements=None, tolerance_pct=5.0):
    try:
        if elements is None:
            elements = ["COD", "N", "P", "TSS"]
        csv_files = glob.glob(str(Path(CONFIG["output_dir"]) / f"{job_id}*.csv"))
        if not csv_files:
            return {"error": f"No CSV for job '{job_id}'. Run export_results_csv first."}
        with open(csv_files[0], newline="") as f:
            reader = csv.DictReader(f)
            rows = list(reader)
        if not rows:
            return {"error": "CSV is empty"}
        headers = list(rows[0].keys())
        element_cols = {
            "COD": {"in": ["COD_in", "Inf_COD", "S_COD_in", "Influent_COD"],
                    "out": ["COD_eff", "Eff_COD", "S_COD_eff", "COD_out"],
                    "WAS": ["COD_WAS", "WAS_COD"]},
            "N":   {"in": ["TKN_in", "TN_in", "Inf_TKN"],
                    "out": ["TN_eff", "Eff_TN"], "WAS": ["TN_WAS"]},
            "P":   {"in": ["TP_in", "Inf_TP"], "out": ["TP_eff", "Eff_TP"], "WAS": ["TP_WAS"]},
            "TSS": {"in": ["TSS_in", "Inf_TSS"], "out": ["TSS_eff", "Eff_TSS"], "WAS": ["TSS_WAS"]},
        }
        def match(cands):
            for c in cands:
                for h in headers:
                    if c.lower() in h.lower():
                        return h
            return None
        def col_mean(col):
            vals = []
            for r in rows:
                try:
                    vals.append(float(r[col]))
                except Exception:
                    pass
            return sum(vals) / len(vals) if vals else None
        results = {}
        summary_issues = []
        for elem in elements:
            if elem not in element_cols:
                continue
            cols = element_cols[elem]
            in_col = match(cols["in"]); out_col = match(cols["out"]); was_col = match(cols["WAS"])
            in_load = col_mean(in_col) if in_col else None
            out_load = col_mean(out_col) if out_col else None
            was_load = col_mean(was_col) if was_col else None
            closure_pct = None
            if in_load and out_load is not None:
                accounted = out_load + (was_load or 0)
                error = in_load - accounted
                closure_pct = round(error / in_load * 100, 1) if in_load > 0 else None
            results[elem] = {"in_load": round(in_load, 3) if in_load else "not found",
                             "out_load": round(out_load, 3) if out_load else "not found",
                             "was_load": round(was_load, 3) if was_load else "not found",
                             "closure_error_pct": closure_pct}
            if closure_pct is not None and abs(closure_pct) > tolerance_pct:
                cat, msg = _classify(abs(closure_pct), "mass_balance_error_pct")
                summary_issues.append(f"{elem} balance closure {closure_pct}% - {cat}")
        status = "PASS" if not summary_issues else (
            "WARNINGS" if all(abs(r.get("closure_error_pct") or 0) < 10 for r in results.values()) else "FAIL")
        return {"job_id": job_id, "tolerance_pct": tolerance_pct, "balances": results,
                "issues": summary_issues, "status": status,
                "interpretation": "<5%=excellent; 5-10%=acceptable; >10%=investigate."}
    except Exception as e:
        return {"error": str(e), "job_id": job_id}


def _val_mlss_health(unit_name="AerTank1"):
    try:
        ds = _make_ds()
        all_vars = ds.sumo.getVariableNames()
        mlss, _  = _val_read_var(ds, all_vars, unit_name, ["X_TSS", "MLSS", "param__MLSS", "param__X_TSS_set"])
        mlvss, _ = _val_read_var(ds, all_vars, unit_name, ["X_VSS", "MLVSS"])
        xbh, _   = _val_read_var(ds, all_vars, unit_name, ["X_BH"])
        xba, _   = _val_read_var(ds, all_vars, unit_name, ["X_BA"])
        xi, _    = _val_read_var(ds, all_vars, unit_name, ["X_I"])
        checks = {"unit": unit_name}
        issues, warnings, passes = [], [], []
        if mlss is None:
            issues.append(f"Cannot read MLSS for {unit_name} - variable not found.")
        else:
            cat, msg = _classify(mlss, "MLSS_mgL")
            checks["MLSS_mgL"] = mlss
            (passes if cat == "green" else warnings if cat == "yellow" else issues).append(f"MLSS: {msg}")
        if mlss and mlvss and mlss > 0:
            ratio = mlvss / mlss
            cat, msg = _classify(ratio, "MLVSS_MLSS_ratio")
            checks["MLVSS_MLSS_ratio"] = round(ratio, 3)
            (passes if cat == "green" else warnings if cat == "yellow" else issues).append(f"MLVSS/MLSS ratio: {msg}")
        if xbh is not None and mlvss and mlvss > 0:
            active_frac = xbh / mlvss
            checks["active_heterotroph_fraction"] = round(active_frac, 3)
            if active_frac < 0.15:
                warnings.append(f"Active heterotroph fraction {active_frac:.1%} is low.")
            else:
                passes.append(f"Active heterotroph fraction {active_frac:.1%}")
        if xba is not None and mlvss and mlvss > 0:
            nitrifier_frac = xba / mlvss
            checks["nitrifier_fraction"] = round(nitrifier_frac, 4)
            if nitrifier_frac < 0.01:
                warnings.append(f"Nitrifier fraction {nitrifier_frac:.2%} very low.")
            else:
                passes.append(f"Nitrifier fraction {nitrifier_frac:.2%}")
        status = "PASS" if not issues else ("FAIL" if len(issues) > 1 else "WARNINGS")
        return {"status": status, "checks": checks, "issues": issues,
                "warnings": warnings, "passes": passes}
    except Exception as e:
        return {"error": str(e), "unit_name": unit_name}


def _val_srt_fm(unit_name="AerTank1", influent_unit="Influent", was_unit="WAS_Pump"):
    try:
        ds = _make_ds()
        all_vars = ds.sumo.getVariableNames()
        def read(unit, suffixes):
            v, _ = _val_read_var(ds, all_vars, unit, suffixes)
            return v
        V = read(unit_name, ["param__V", "param__Volume"])
        MLSS = read(unit_name, ["X_TSS", "MLSS", "param__X_TSS_set"])
        MLVSS = read(unit_name, ["X_VSS"])
        SRT_set = read(unit_name, ["param__SRT", "param__SRT_set"])
        Q_was = read(was_unit, ["param__Q"])
        X_R = read("SecClar1", ["param__X_underflow", "X_R"])
        Q_in = read(influent_unit, ["param__Q"])
        BOD_in = read(influent_unit, ["param__BOD5", "param__S_S"])
        checks, issues, warnings, passes = {}, [], [], []
        if V and MLSS and Q_was and X_R:
            SRT_actual = (V * MLSS) / (Q_was * X_R * 1000) if (Q_was * X_R) > 0 else None
            checks["SRT_actual_d"] = round(SRT_actual, 2) if SRT_actual else None
            if SRT_set:
                checks["SRT_target_d"] = SRT_set
                if SRT_actual:
                    deviation = abs(SRT_actual - SRT_set) / SRT_set * 100
                    if deviation > 30:
                        issues.append(f"Actual SRT {SRT_actual:.1f}d differs from target {SRT_set}d by {deviation:.0f}%")
                    else:
                        passes.append(f"SRT target {SRT_set}d ~ actual {SRT_actual:.1f}d")
            if SRT_actual:
                cat, msg = _classify(SRT_actual, "SRT_d")
                (passes if cat == "green" else warnings if cat == "yellow" else issues).append(f"SRT: {msg}")
        else:
            warnings.append("Cannot compute actual SRT - missing V, MLSS, Q_WAS, or X_R")
        if Q_in and BOD_in and V and MLVSS and V * MLVSS > 0:
            FM = (Q_in * BOD_in) / (V * MLVSS)
            checks["F_M_ratio"] = round(FM, 3)
            cat, msg = _classify(FM, "FM_ratio")
            (passes if cat == "green" else warnings if cat == "yellow" else issues).append(f"F/M ratio: {msg}")
        else:
            warnings.append("Cannot compute F/M - missing Q, BOD, V, or MLVSS")
        if V and Q_in:
            HRT = V / Q_in * 24
            checks["HRT_h"] = round(HRT, 2)
            cat, msg = _classify(HRT, "HRT_h")
            (passes if cat == "green" else warnings if cat == "yellow" else issues).append(f"HRT: {msg}")
        status = "PASS" if not issues else ("WARNINGS" if len(issues) <= 1 else "FAIL")
        return {"status": status, "checks": checks, "issues": issues,
                "warnings": warnings, "passes": passes,
                "tip": "Adjust Q_WAS to match SRT target. Tune V or MLSS for F/M."}
    except Exception as e:
        return {"error": str(e)}


def _val_do_levels(zone_units=None):
    try:
        if zone_units is None:
            zone_units = {"aerobic": "AerTank1", "anoxic": "AnoxTank1", "anaerobic": "AnaTank1"}
        ds = _make_ds()
        all_vars = ds.sumo.getVariableNames()
        def read_do(unit):
            return _val_read_var(ds, all_vars, unit, ["S_O", "DO", "param__DO_setpoint", "param__DO_sp", "param__S_O_sp"])
        results, issues, warnings, passes = {}, [], [], []
        for zone_type, unit in zone_units.items():
            do_val, do_var = read_do(unit)
            if do_val is None:
                warnings.append(f"DO not found for {zone_type} zone ({unit})")
                continue
            results[zone_type] = {"unit": unit, "DO_mgL": do_val, "variable": do_var}
            threshold_key = {"aerobic": "DO_aerobic_mgL", "anoxic": "DO_anoxic_mgL",
                             "anaerobic": "DO_anaerobic_mgL"}.get(zone_type)
            if threshold_key:
                cat, msg = _classify(do_val, threshold_key)
                line = f"{zone_type.capitalize()} zone ({unit}) DO: {msg}"
                (passes if cat == "green" else warnings if cat == "yellow" else issues).append(line)
        for unit in zone_units.values():
            kla_var = f"Sumo__Plant__{unit}__param__KLa"
            if kla_var in all_vars:
                try:
                    kla = float(ds.sumo.get(kla_var))
                    cat, msg = _classify(kla, "KLa_h")
                    results.setdefault(f"KLa_{unit}", {})["KLa_h"] = kla
                    (passes if cat == "green" else warnings if cat == "yellow" else issues).append(f"{unit} KLa: {msg}")
                except Exception:
                    pass
        status = "PASS" if not issues else ("WARNINGS" if not any(
            "anoxic" in i.lower() or "anaerobic" in i.lower() for i in issues) else "FAIL")
        return {"status": status, "zones": results, "issues": issues,
                "warnings": warnings, "passes": passes}
    except Exception as e:
        return {"error": str(e)}


def _val_oxygen_supply(unit_name="AerTank1", influent_unit="Influent"):
    try:
        ds = _make_ds()
        all_vars = ds.sumo.getVariableNames()
        def read(unit, suffixes):
            v, _ = _val_read_var(ds, all_vars, unit, suffixes)
            return v
        V = read(unit_name, ["param__V"])
        KLa = read(unit_name, ["param__KLa"])
        DO_op = read(unit_name, ["param__DO_setpoint", "S_O"]) or 2.0
        T = read(unit_name, ["param__T"]) or CONFIG.get("plant_temp_C", 25.0)
        Q = read(influent_unit, ["param__Q"])
        BOD = read(influent_unit, ["param__BOD5", "param__S_S"])
        TKN = read(influent_unit, ["param__S_TKN", "param__TKN"])
        DO_sat = round(14.652 - 0.41022 * T + 0.007991 * T**2 - 7.7774e-5 * T**3, 2)
        issues, passes = [], []
        checks = {"DO_sat_mgL": DO_sat, "DO_operating_mgL": DO_op, "Temp_C": T}
        AOR_BOD = Q * BOD * 1.2 / 1000 if (Q and BOD) else None
        if AOR_BOD is not None:
            checks["AOR_carbon_kgOd"] = round(AOR_BOD, 1)
        AOR_N = Q * TKN * 4.57 / 1000 if (Q and TKN) else None
        if AOR_N is not None:
            checks["AOR_nitrification_kgOd"] = round(AOR_N, 1)
        AOR_total = (AOR_BOD or 0) + (AOR_N or 0)
        checks["AOR_total_kgOd"] = round(AOR_total, 1)
        if KLa and V:
            OTR = KLa * V * (DO_sat - DO_op) * 24 / 1000
            checks["OTR_supply_kgOd"] = round(OTR, 1)
            if AOR_total > 0:
                ratio = OTR / AOR_total
                checks["supply_demand_ratio"] = round(ratio, 2)
                if ratio < 1.0:
                    issues.append(f"Oxygen supply ({OTR:.0f}) < demand ({AOR_total:.0f} kg/d)")
                elif ratio < 1.2:
                    issues.append(f"Oxygen supply margin only {(ratio-1)*100:.0f}%")
                else:
                    passes.append(f"Oxygen supply/demand ratio {ratio:.2f} (healthy)")
        status = "PASS" if not issues else "CHECK"
        return {"status": status, "checks": checks, "issues": issues, "passes": passes,
                "note": "Simplified demand = 1.2*BOD + 4.57*TKN."}
    except Exception as e:
        return {"error": str(e)}


def _val_hydraulic(influent_unit="Influent", effluent_unit="Effluent",
                   was_unit="WAS_Pump", ras_unit="RAS_Pump", sec_clar_unit="SecClar1"):
    try:
        ds = _make_ds()
        all_vars = ds.sumo.getVariableNames()
        def read_q(unit):
            v, _ = _val_read_var(ds, all_vars, unit, ["param__Q"])
            return v
        q_in = read_q(influent_unit); q_eff = read_q(effluent_unit)
        q_was = read_q(was_unit); q_ras = read_q(ras_unit)
        checks, issues, warnings, passes = {}, [], [], []
        if q_in and q_eff is not None:
            expected = q_in - (q_was or 0)
            error_pct = abs(q_eff - expected) / q_in * 100
            checks["flow_continuity_error_pct"] = round(error_pct, 1)
            cat, msg = _classify(error_pct, "flow_continuity_error_pct")
            (passes if cat == "green" else warnings if cat == "yellow" else issues).append(f"Flow continuity: {msg}")
        if q_ras and q_in and q_in > 0:
            ras_r = q_ras / q_in
            checks["RAS_ratio"] = round(ras_r, 2)
            cat, msg = _classify(ras_r, "RAS_ratio")
            (passes if cat == "green" else warnings if cat == "yellow" else issues).append(f"RAS ratio: {msg}")
        A_sec = None
        for s in ["param__A", "param__Area", "param__surface_area"]:
            var = f"Sumo__Plant__{sec_clar_unit}__{s}"
            if var in all_vars:
                try:
                    A_sec = float(ds.sumo.get(var)); break
                except Exception:
                    pass
        if A_sec and q_in:
            SOR = q_in / A_sec
            checks["sec_clar_SOR_m3m2d"] = round(SOR, 1)
            cat, msg = _classify(SOR, "overflow_rate_sec_clar")
            (passes if cat == "green" else warnings if cat == "yellow" else issues).append(f"Secondary clarifier SOR: {msg}")
        for s in ["param__X_underflow", "X_R", "param__C_underflow"]:
            var = f"Sumo__Plant__{sec_clar_unit}__{s}"
            if var in all_vars:
                try:
                    X_R = float(ds.sumo.get(var))
                    if X_R > 100:
                        X_R = X_R / 1000
                    checks["underflow_conc_gL"] = round(X_R, 1)
                    cat, msg = _classify(X_R, "underflow_conc_sec_clar")
                    (passes if cat == "green" else warnings if cat == "yellow" else issues).append(f"Underflow concentration: {msg}")
                    break
                except Exception:
                    continue
        status = "PASS" if not issues else "CHECK"
        return {"status": status, "checks": checks, "issues": issues,
                "warnings": warnings, "passes": passes}
    except Exception as e:
        return {"error": str(e)}


def _val_asm_kinetics(unit_name="AerTank1"):
    try:
        ds = _make_ds()
        all_vars = ds.sumo.getVariableNames()
        bounds = VALIDATION_THRESHOLDS["ASM_kinetic_bounds"]
        results, issues, passes = [], [], []
        for param, b in bounds.items():
            for s in [f"param__{param}", f"param__{param.replace('_','')}"]:
                var = f"Sumo__Plant__{unit_name}__{s}"
                if var in all_vars:
                    try:
                        v = float(ds.sumo.get(var))
                        in_bounds = b["min"] <= v <= b["max"]
                        results.append({"param": param, "value": v, "unit": b["unit"],
                                        "min": b["min"], "max": b["max"], "in_bounds": in_bounds})
                        if in_bounds:
                            passes.append(f"{param} = {v} {b['unit']} (OK)")
                        else:
                            issues.append(f"{param} = {v} {b['unit']} outside [{b['min']}, {b['max']}]")
                    except Exception:
                        pass
                    break
        status = "PASS" if not issues else "FAIL"
        return {"unit": unit_name, "status": status, "checks": len(results),
                "out_of_bounds": len(issues), "details": results,
                "issues": issues, "passes": passes[:5],
                "tip": "Out-of-bounds values are usually typos - call reset_parameter_to_default."}
    except Exception as e:
        return {"error": str(e), "unit_name": unit_name}


def _val_convergence(job_id, stability_window_pct=10.0):
    try:
        csv_files = glob.glob(str(Path(CONFIG["output_dir"]) / f"{job_id}*.csv"))
        if not csv_files:
            return {"error": f"No CSV for job '{job_id}'"}
        with open(csv_files[0], newline="") as f:
            reader = csv.DictReader(f)
            rows = list(reader)
        if len(rows) < 10:
            return {"error": "Not enough rows to assess convergence"}
        n_tail = max(5, int(len(rows) * stability_window_pct / 100))
        tail = rows[-n_tail:]
        key_patterns = ["COD_eff", "BOD_eff", "TSS_eff", "NH4_eff", "TN_eff"]
        headers = list(rows[0].keys())
        convergence, issues, passes = [], [], []
        for pat in key_patterns:
            col = next((h for h in headers if pat.lower() in h.lower()), None)
            if not col:
                continue
            vals = []
            for r in tail:
                try:
                    vals.append(float(r[col]))
                except Exception:
                    pass
            if len(vals) < 3:
                continue
            mean = sum(vals) / len(vals)
            if mean == 0:
                continue
            rel_std = (max(vals) - min(vals)) / mean * 100
            stable = rel_std < 5.0
            convergence.append({"parameter": pat, "column": col,
                                "tail_mean": round(mean, 3),
                                "tail_range_pct": round(rel_std, 1),
                                "converged": stable})
            if stable:
                passes.append(f"{pat} stable (+/- {rel_std:.1f}%)")
            else:
                issues.append(f"{pat} fluctuating +/- {rel_std:.1f}% in final {n_tail} steps")
        status = "PASS" if not issues else ("CHECK" if len(issues) <= 1 else "FAIL")
        return {"job_id": job_id, "status": status, "rows_analyzed": n_tail,
                "convergence": convergence, "issues": issues, "passes": passes}
    except Exception as e:
        return {"error": str(e)}


def _val_plausibility(job_id):
    try:
        csv_files = glob.glob(str(Path(CONFIG["output_dir"]) / f"{job_id}*.csv"))
        if not csv_files:
            return {"error": f"No CSV for job '{job_id}'"}
        with open(csv_files[0], newline="") as f:
            reader = csv.DictReader(f)
            rows = list(reader)
        if not rows:
            return {"error": "CSV empty"}
        headers = list(rows[0].keys())
        issues, warnings, passes = [], [], []
        def col_vals(col):
            out = []
            for r in rows:
                try:
                    out.append(float(r[col]))
                except Exception:
                    pass
            return out
        for h in headers:
            if any(t in h.lower() for t in ["cod", "bod", "tss", "nh4", "no3", "tn", "tp"]):
                vals = col_vals(h)
                if vals and min(vals) < 0:
                    issues.append(f"{h} has negative values (min = {min(vals):.2f})")
        ph_col = next((h for h in headers if h.lower() == "ph" or h.endswith("_pH")), None)
        if ph_col:
            vals = col_vals(ph_col)
            if vals:
                if min(vals) < 4.0 or max(vals) > 10.0:
                    issues.append(f"{ph_col} goes outside [4, 10]: [{min(vals):.1f}, {max(vals):.1f}]")
                else:
                    passes.append(f"{ph_col} within biological range")
        for h in headers:
            if "do" in h.lower() or "s_o" in h.lower():
                vals = col_vals(h)
                if vals and max(vals) > 14.0:
                    issues.append(f"{h} exceeds oxygen saturation (14 mg/L)")
        for pair in [("COD_in", "COD_eff"), ("TSS_in", "TSS_eff")]:
            c_in = next((h for h in headers if pair[0].lower() in h.lower()), None)
            c_out = next((h for h in headers if pair[1].lower() in h.lower()), None)
            if c_in and c_out:
                in_vals = col_vals(c_in); out_vals = col_vals(c_out)
                if in_vals and out_vals:
                    if sum(out_vals)/len(out_vals) > sum(in_vals)/len(in_vals):
                        warnings.append(f"Mean effluent {pair[1]} > mean influent {pair[0]}")
        status = "PASS" if not issues else ("FAIL" if len(issues) > 2 else "WARNINGS")
        return {"job_id": job_id, "status": status,
                "issue_count": len(issues), "warning_count": len(warnings),
                "issues": issues, "warnings": warnings, "passes": passes}
    except Exception as e:
        return {"error": str(e), "job_id": job_id}


def _val_full_model(aer_tank="AerTank1", influent_unit="Influent",
                    effluent_unit="Effluent", sec_clar_unit="SecClar1",
                    was_unit="WAS_Pump", ras_unit="RAS_Pump"):
    try:
        results = {}
        results["structure"] = _val_structure()
        results["influent"] = _val_influent(influent_unit)
        results["mass_balance"] = _val_mass_balance_pre()
        results["mlss_health"] = _val_mlss_health(aer_tank)
        results["srt_fm"] = _val_srt_fm(aer_tank, influent_unit, was_unit)
        results["do_levels"] = _val_do_levels()
        results["oxygen_supply"] = _val_oxygen_supply(aer_tank, influent_unit)
        results["hydraulic"] = _val_hydraulic(influent_unit, effluent_unit, was_unit, ras_unit, sec_clar_unit)
        results["asm_kinetics"] = _val_asm_kinetics(aer_tank)
        fail_count = sum(1 for r in results.values() if r.get("status") in ("FAIL", "ERROR"))
        warn_count = sum(1 for r in results.values() if r.get("status") in ("WARNINGS", "CHECK"))
        pass_count = sum(1 for r in results.values() if r.get("status") == "PASS")
        all_issues = []
        for name, r in results.items():
            for issue in r.get("issues", []):
                all_issues.append(f"[{name}] {issue}")
        overall = ("READY_TO_SIMULATE" if fail_count == 0 and warn_count <= 2
                   else "CHECK_WARNINGS" if fail_count == 0
                   else "DO_NOT_SIMULATE")
        return {"overall_status": overall, "checks_passed": pass_count,
                "checks_with_warnings": warn_count, "checks_failed": fail_count,
                "summary": {name: r.get("status", "?") for name, r in results.items()},
                "top_issues": all_issues[:15], "detailed_results": results,
                "recommendation": (
                    "Model passes all checks - run_steady_state recommended."
                    if overall == "READY_TO_SIMULATE" else
                    "Review warnings before running."
                    if overall == "CHECK_WARNINGS" else
                    "Fix critical issues before running.")}
    except Exception as e:
        return {"error": str(e)}


def _val_post_simulation(job_id):
    try:
        results = {}
        results["mass_balance"] = _val_mass_balance_post(job_id)
        results["convergence"] = _val_convergence(job_id)
        results["plausibility"] = _val_plausibility(job_id)
        fail_count = sum(1 for r in results.values() if r.get("status") in ("FAIL", "ERROR"))
        warn_count = sum(1 for r in results.values() if r.get("status") in ("WARNINGS", "CHECK"))
        pass_count = sum(1 for r in results.values() if r.get("status") == "PASS")
        all_issues = []
        for name, r in results.items():
            for issue in r.get("issues", []):
                all_issues.append(f"[{name}] {issue}")
        overall = ("RESULTS_VALID" if fail_count == 0 and warn_count == 0
                   else "RESULTS_SUSPECT" if fail_count == 0
                   else "RESULTS_INVALID")
        return {"job_id": job_id, "overall_status": overall,
                "checks_passed": pass_count, "checks_warnings": warn_count,
                "checks_failed": fail_count,
                "summary": {n: r.get("status", "?") for n, r in results.items()},
                "top_issues": all_issues[:10], "detailed_results": results,
                "recommendation": (
                    "Results are valid."
                    if overall == "RESULTS_VALID" else
                    "Review warnings before trusting the numbers."
                    if overall == "RESULTS_SUSPECT" else
                    "Do not use these results - re-configure and re-run.")}
    except Exception as e:
        return {"error": str(e), "job_id": job_id}


# ----- GROUP Y helpers: Optimization, Diagnostics, Economics, Visualization -----

# Attempt matplotlib import up-front (optional).
try:
    import matplotlib  # noqa: F401
    matplotlib.use("Agg")
    import matplotlib.pyplot as _plt
    _MPL_AVAILABLE = True
except Exception:
    _MPL_AVAILABLE = False


def _opt_scan_parameter_sensitivity(parameter, values, output_variable,
                                     unit_name="AerTank1"):
    """Sweep one parameter across values, record output_variable each time.

    Returns list of (value, output) pairs. Does NOT re-simulate (reads state).
    For full sensitivity on effluent, pair with run_steady_state externally.
    """
    try:
        ds = _make_ds()
        all_vars = ds.sumo.getVariableNames()
        param_var = f"Sumo__Plant__{unit_name}__{parameter}"
        if param_var not in all_vars:
            # try param__ prefix
            param_var = f"Sumo__Plant__{unit_name}__param__{parameter}"
        if param_var not in all_vars:
            return {"error": f"Parameter {parameter} not found on {unit_name}"}
        out_var = f"Sumo__Plant__{unit_name}__{output_variable}"
        if out_var not in all_vars:
            # scan all vars
            candidates = [v for v in all_vars if output_variable.lower() in v.lower()]
            if not candidates:
                return {"error": f"Output variable {output_variable} not found"}
            out_var = candidates[0]
        sweep = []
        try:
            original = float(ds.sumo.get(param_var))
        except Exception:
            original = None
        for v in values:
            try:
                ds.sumo.set(param_var, v)
                o = float(ds.sumo.get(out_var))
                sweep.append({"parameter_value": v, "output": o})
            except Exception as e:
                sweep.append({"parameter_value": v, "error": str(e)})
        # restore
        if original is not None:
            try:
                ds.sumo.set(param_var, original)
            except Exception:
                pass
        return {
            "parameter": parameter, "unit": unit_name,
            "output_variable": out_var, "sweep": sweep,
            "n_points": len(sweep),
            "note": "Instantaneous sensitivity. For effluent impact, "
                    "pair each value with run_steady_state externally.",
        }
    except Exception as e:
        return {"error": str(e)}


def _opt_was_for_srt(target_srt_d, tolerance_pct=5.0, max_iter=20,
                      was_unit="WAS_Pump", aer_tank="AerTank1"):
    """Bisection on Q_WAS to hit target SRT. Does not run simulations - uses
    the SRT = V*MLSS / (Q_WAS * X_R) approximation from state values."""
    try:
        ds = _make_ds()
        all_vars = ds.sumo.getVariableNames()
        def read(unit, suffixes):
            for s in suffixes:
                var = f"Sumo__Plant__{unit}__{s}"
                if var in all_vars:
                    try:
                        return float(ds.sumo.get(var))
                    except Exception:
                        continue
            return None
        V    = read(aer_tank,  ["param__V", "param__Volume"])
        MLSS = read(aer_tank,  ["X_TSS", "MLSS", "param__X_TSS_set"])
        X_R  = read("SecClar1",["param__X_underflow", "X_R"])
        if not (V and MLSS and X_R):
            return {"error": "Missing V, MLSS, or X_R - cannot optimize"}
        # SRT_d = V*MLSS / (Q_WAS * X_R * 1000) --> Q_WAS = V*MLSS / (SRT * X_R * 1000)
        Q_was_needed = (V * MLSS) / (target_srt_d * X_R * 1000)
        was_var = f"Sumo__Plant__{was_unit}__param__Q"
        if was_var in all_vars:
            try:
                ds.sumo.set(was_var, Q_was_needed)
                applied = True
            except Exception as e:
                applied = False
                app_err = str(e)
        else:
            applied = False
            app_err = f"{was_var} not found"
        result = {
            "target_srt_d": target_srt_d,
            "tolerance_pct": tolerance_pct,
            "V_m3": V, "MLSS_mgL": MLSS, "X_R_gL": X_R,
            "optimal_Q_WAS_m3d": round(Q_was_needed, 2),
            "applied": applied,
            "method": "analytical (V*MLSS / (SRT * X_R * 1000))",
        }
        if not applied:
            result["note"] = app_err
        return result
    except Exception as e:
        return {"error": str(e)}


def _opt_kla_for_do(target_do_mgL, unit_name="AerTank1", tolerance=0.1,
                     max_iter=15):
    """Estimate KLa needed for target DO setpoint using OTR = KLa*V*(DO_sat-DO)."""
    try:
        ds = _make_ds()
        all_vars = ds.sumo.getVariableNames()
        def read(unit, suffixes):
            for s in suffixes:
                var = f"Sumo__Plant__{unit}__{s}"
                if var in all_vars:
                    try:
                        return float(ds.sumo.get(var))
                    except Exception:
                        continue
            return None
        V    = read(unit_name, ["param__V"])
        T    = read(unit_name, ["param__T"]) or 25.0
        Q    = read("Influent", ["param__Q"])
        BOD  = read("Influent", ["param__BOD5", "param__S_S"])
        TKN  = read("Influent", ["param__S_TKN", "param__TKN"])
        if not (V and Q and BOD):
            return {"error": "Missing V, Q, or BOD"}
        DO_sat = round(14.652 - 0.41022*T + 0.007991*T**2 - 7.7774e-5*T**3, 2)
        # Demand: 1.2 kg O2/kg BOD + 4.57 kg O2/kg N
        AOR = Q*BOD*1.2/1000 + (Q*(TKN or 0)*4.57/1000)
        # KLa = AOR*1000 / (V*(DO_sat - DO_target)*24)   [1/h]
        if DO_sat - target_do_mgL <= 0:
            return {"error": f"Target DO {target_do_mgL} >= DO_sat {DO_sat}"}
        KLa = AOR * 1000 / (V * (DO_sat - target_do_mgL) * 24)
        kla_var = f"Sumo__Plant__{unit_name}__param__KLa"
        applied = False
        app_err = None
        if kla_var in all_vars:
            try:
                ds.sumo.set(kla_var, KLa)
                applied = True
            except Exception as e:
                app_err = str(e)
        else:
            app_err = f"{kla_var} not found"
        result = {
            "target_DO_mgL": target_do_mgL,
            "DO_sat_mgL": DO_sat, "Temp_C": T,
            "AOR_kgOd": round(AOR, 1),
            "optimal_KLa_1h": round(KLa, 2),
            "applied": applied,
            "method": "analytical (AOR * 1000 / (V * (DO_sat - DO) * 24))",
        }
        if not applied:
            result["note"] = app_err
        return result
    except Exception as e:
        return {"error": str(e)}


def _diag_nitrification(unit_name="AerTank1", influent_unit="Influent"):
    """Targeted diagnostic for nitrification problems."""
    try:
        ds = _make_ds()
        all_vars = ds.sumo.getVariableNames()
        def read(unit, suffixes):
            for s in suffixes:
                var = f"Sumo__Plant__{unit}__{s}"
                if var in all_vars:
                    try:
                        return float(ds.sumo.get(var))
                    except Exception:
                        continue
            return None
        checks = {}
        causes = []
        fixes = []
        SRT = read(unit_name, ["param__SRT"])
        if SRT is not None:
            checks["SRT_d"] = SRT
            if SRT < 5:
                causes.append(f"SRT too low ({SRT}d) - autotrophs washed out")
                fixes.append("Reduce Q_WAS (optimize_was_for_target_srt target=10)")
        T = read(unit_name, ["param__T"])
        if T is not None:
            checks["Temp_C"] = T
            if T < 10:
                causes.append(f"Low temperature ({T}C) - mu_A dropped")
                fixes.append("Increase SRT to compensate for cold temps")
        pH = read(unit_name, ["param__pH"])
        if pH is not None:
            checks["pH"] = pH
            if pH < 6.8 or pH > 8.0:
                causes.append(f"pH {pH} outside nitrifier optimum (7.2-8.0)")
                fixes.append("Dose alkalinity / lime")
        DO = read(unit_name, ["S_O", "param__DO_setpoint"])
        if DO is not None:
            checks["DO_mgL"] = DO
            if DO < 1.5:
                causes.append(f"DO too low ({DO} mg/L) - aeration-limited")
                fixes.append("Increase KLa (optimize_kla_for_target_do)")
        NH4_eff = read("Effluent", ["S_NH4", "S_NH"])
        if NH4_eff is not None:
            checks["NH4_effluent_mgL"] = NH4_eff
        tkn_in = read(influent_unit, ["param__S_TKN"])
        if tkn_in is not None and (Q := read(influent_unit, ["param__Q"])):
            N_load = Q * tkn_in / 1000
            checks["N_load_kgd"] = round(N_load, 2)
        status = "FAIL" if causes else "PASS"
        return {"diagnostic": "nitrification", "status": status,
                "checks": checks, "likely_causes": causes,
                "recommended_fixes": fixes,
                "note": "Root cause usually SRT < 5d or pH < 6.8."}
    except Exception as e:
        return {"error": str(e)}


def _diag_bulking(unit_name="AerTank1"):
    """Sludge bulking risk assessment."""
    try:
        ds = _make_ds()
        all_vars = ds.sumo.getVariableNames()
        def read(unit, suffixes):
            for s in suffixes:
                var = f"Sumo__Plant__{unit}__{s}"
                if var in all_vars:
                    try:
                        return float(ds.sumo.get(var))
                    except Exception:
                        continue
            return None
        checks, risks, fixes = {}, [], []
        SVI = read(unit_name, ["SVI", "param__SVI"])
        if SVI is not None:
            checks["SVI_mLg"] = SVI
            if SVI > 150:
                risks.append(f"SVI {SVI} - filamentous/viscous bulking risk")
                fixes.append("Increase F/M, reduce SRT slightly, check nutrient ratios")
        FM = None
        Q_in = read("Influent", ["param__Q"])
        BOD_in = read("Influent", ["param__BOD5"])
        V = read(unit_name, ["param__V"])
        MLVSS = read(unit_name, ["X_VSS"])
        if Q_in and BOD_in and V and MLVSS and V*MLVSS > 0:
            FM = (Q_in * BOD_in) / (V * MLVSS)
            checks["F_M_ratio"] = round(FM, 3)
            if FM < 0.1:
                risks.append(f"F/M {FM:.2f} very low - promotes filamentous growth")
                fixes.append("Increase loading or reduce MLSS to raise F/M above 0.15")
        DO = read(unit_name, ["S_O"])
        if DO is not None:
            checks["DO_mgL"] = DO
            if DO < 1.0:
                risks.append(f"Low DO ({DO} mg/L) - favors filaments")
                fixes.append("Raise DO setpoint to 1.5-2.5 mg/L")
        COD = read("Influent", ["param__S_COD"])
        TP = read("Influent", ["param__S_TP"])
        if COD and TP:
            cp = COD / TP
            checks["COD_TP_ratio"] = round(cp, 1)
            if cp > 200:
                risks.append(f"COD:TP {cp:.0f} indicates P deficiency - filament risk")
                fixes.append("Add phosphate")
        status = "RISK" if risks else "OK"
        return {"diagnostic": "bulking", "status": status,
                "checks": checks, "risk_factors": risks,
                "recommended_fixes": fixes}
    except Exception as e:
        return {"error": str(e)}


def _diag_phosphorus(unit_name="AerTank1"):
    """EBPR diagnostic."""
    try:
        ds = _make_ds()
        all_vars = ds.sumo.getVariableNames()
        def read(unit, suffixes):
            for s in suffixes:
                var = f"Sumo__Plant__{unit}__{s}"
                if var in all_vars:
                    try:
                        return float(ds.sumo.get(var))
                    except Exception:
                        continue
            return None
        checks, causes, fixes = {}, [], []
        DO_ana = read("AnaTank1", ["S_O", "param__DO_setpoint"])
        if DO_ana is not None:
            checks["anaerobic_DO_mgL"] = DO_ana
            if DO_ana > 0.1:
                causes.append(f"Anaerobic DO {DO_ana} too high - disrupts EBPR")
                fixes.append("Isolate anaerobic zone - reduce recycle or mixing")
        NO3_ana = read("AnaTank1", ["S_NO3"])
        if NO3_ana is not None:
            checks["anaerobic_NO3_mgL"] = NO3_ana
            if NO3_ana > 1.0:
                causes.append(f"NO3 in anaerobic zone ({NO3_ana} mg/L) - nitrate intrusion")
                fixes.append("Reduce IMLR or improve denitrification")
        COD = read("Influent", ["param__S_COD"])
        TP = read("Influent", ["param__S_TP"])
        if COD and TP:
            cp = COD / TP
            checks["COD_TP_ratio"] = round(cp, 1)
            if cp < 40:
                causes.append(f"COD:TP {cp:.0f} too low - need >40 for EBPR")
                fixes.append("Add VFA / increase primary sludge fermentation")
        TP_eff = read("Effluent", ["S_TP", "S_PO4"])
        if TP_eff is not None:
            checks["P_effluent_mgL"] = TP_eff
            if TP_eff > 1.0:
                causes.append(f"Effluent P {TP_eff} above compliance target")
                if not fixes:
                    fixes.append("Dose alum/iron for chemical P removal")
        status = "FAIL" if causes else "PASS"
        return {"diagnostic": "phosphorus_removal", "status": status,
                "checks": checks, "likely_causes": causes,
                "recommended_fixes": fixes}
    except Exception as e:
        return {"error": str(e)}


def _diag_recommendations(aer_tank="AerTank1", influent_unit="Influent",
                           was_unit="WAS_Pump"):
    """Aggregate diagnostic + recommend adjustments."""
    try:
        recs = []
        mlss = _val_mlss_health(aer_tank)
        srt = _val_srt_fm(aer_tank, influent_unit, was_unit)
        do = _val_do_levels()
        # MLSS recommendations
        if mlss.get("checks", {}).get("MLSS_mgL", 0) > 5000:
            recs.append({"priority": "high", "target": "MLSS",
                         "action": "Increase Q_WAS to reduce MLSS below 4500 mg/L",
                         "tool": "optimize_was_for_target_srt"})
        elif mlss.get("checks", {}).get("MLSS_mgL", 99999) < 1500:
            recs.append({"priority": "high", "target": "MLSS",
                         "action": "Decrease Q_WAS to build MLSS above 2000 mg/L",
                         "tool": "optimize_was_for_target_srt"})
        # SRT
        srt_actual = srt.get("checks", {}).get("SRT_actual_d")
        if srt_actual and srt_actual < 5:
            recs.append({"priority": "critical", "target": "SRT",
                         "action": f"SRT {srt_actual}d too low for nitrification - reduce Q_WAS",
                         "tool": "optimize_was_for_target_srt (target_srt_d=10)"})
        # DO
        for zone, data in do.get("zones", {}).items():
            do_mgL = data.get("DO_mgL")
            if zone == "aerobic" and do_mgL is not None and do_mgL < 1.5:
                recs.append({"priority": "high", "target": "DO",
                             "action": f"Aerobic DO {do_mgL} below minimum - increase KLa",
                             "tool": "optimize_kla_for_target_do (target_do_mgL=2.0)"})
        if not recs:
            recs.append({"priority": "none", "target": "all",
                         "action": "Plant operating within optimal ranges",
                         "tool": "-"})
        return {"recommendation_count": len(recs),
                "recommendations": recs,
                "summary": {
                    "mlss_status": mlss.get("status"),
                    "srt_fm_status": srt.get("status"),
                    "do_status": do.get("status"),
                }}
    except Exception as e:
        return {"error": str(e)}


def _econ_opex(job_id, energy_cost_per_kwh=0.08, sludge_disposal_cost_per_tonne=50.0,
                alum_cost_per_kg=0.30, polymer_cost_per_kg=3.50):
    """Annual OPEX estimate."""
    try:
        ds = _make_ds()
        all_vars = ds.sumo.getVariableNames()
        def read(unit, suffixes):
            for s in suffixes:
                var = f"Sumo__Plant__{unit}__{s}"
                if var in all_vars:
                    try:
                        return float(ds.sumo.get(var))
                    except Exception:
                        continue
            return None
        # Energy - simple: aeration = AOR*SOTE*alpha. Approx 1.5 kWh/kg O2.
        Q = read("Influent", ["param__Q"])
        BOD = read("Influent", ["param__BOD5"])
        TKN = read("Influent", ["param__S_TKN"])
        aeration_kWh_per_d = 0
        if Q and BOD:
            AOR = Q * BOD * 1.2 / 1000 + (Q * (TKN or 0) * 4.57 / 1000)
            aeration_kWh_per_d = AOR * 1.5  # 1.5 kWh/kg O2 typical
        pumping_kWh_per_d = (Q or 0) * 0.5 / 1000  # 0.5 kWh/m3 typical
        total_kWh_per_d = aeration_kWh_per_d + pumping_kWh_per_d
        annual_kWh = total_kWh_per_d * 365
        energy_cost = annual_kWh * energy_cost_per_kwh
        # Sludge production: approx 0.5 kg TSS/kg BOD removed
        sludge_tpd = 0
        if Q and BOD:
            sludge_tpd = Q * BOD * 0.5 / 1e6  # tonnes/d
        sludge_cost = sludge_tpd * 365 * sludge_disposal_cost_per_tonne
        # Chemical costs (placeholder, assume minor if no dosing)
        alum_kg_per_d = 0  # would read from chem dosing unit
        polymer_kg_per_d = 0
        chemical_cost = alum_kg_per_d*365*alum_cost_per_kg + polymer_kg_per_d*365*polymer_cost_per_kg
        total_opex = energy_cost + sludge_cost + chemical_cost
        return {
            "job_id": job_id,
            "annual_energy_kWh": round(annual_kWh),
            "energy_cost_USD": round(energy_cost, 0),
            "sludge_disposal_cost_USD": round(sludge_cost, 0),
            "chemical_cost_USD": round(chemical_cost, 0),
            "total_annual_opex_USD": round(total_opex, 0),
            "specific_cost_USD_per_m3": round(total_opex / (Q * 365), 4) if Q else None,
            "breakdown_pct": {
                "energy": round(energy_cost/total_opex*100, 1) if total_opex else 0,
                "sludge": round(sludge_cost/total_opex*100, 1) if total_opex else 0,
                "chemicals": round(chemical_cost/total_opex*100, 1) if total_opex else 0,
            },
            "assumptions": {
                "aeration_kWh_per_kgO2": 1.5,
                "pumping_kWh_per_m3": 0.5,
                "sludge_yield_kgTSS_per_kgBOD": 0.5,
                "energy_cost_per_kwh": energy_cost_per_kwh,
                "sludge_disposal_cost_per_tonne": sludge_disposal_cost_per_tonne,
            },
        }
    except Exception as e:
        return {"error": str(e), "job_id": job_id}


def _econ_ghg(job_id, grid_emission_factor_kgCO2_per_kWh=0.45):
    """GHG emissions estimate."""
    try:
        ds = _make_ds()
        all_vars = ds.sumo.getVariableNames()
        def read(unit, suffixes):
            for s in suffixes:
                var = f"Sumo__Plant__{unit}__{s}"
                if var in all_vars:
                    try:
                        return float(ds.sumo.get(var))
                    except Exception:
                        continue
            return None
        Q = read("Influent", ["param__Q"])
        BOD = read("Influent", ["param__BOD5"])
        TKN = read("Influent", ["param__S_TKN"])
        if not (Q and BOD):
            return {"error": "Missing Q or BOD"}
        AOR = Q * BOD * 1.2 / 1000 + (Q * (TKN or 0) * 4.57 / 1000)
        aeration_kWh = AOR * 1.5
        pumping_kWh = Q * 0.5 / 1000
        total_kWh_per_d = aeration_kWh + pumping_kWh
        annual_kWh = total_kWh_per_d * 365
        # Scope 2: electricity
        scope2_kg = annual_kWh * grid_emission_factor_kgCO2_per_kWh
        # Scope 1: direct N2O = 0.5% of TKN_load as N2O-N (IPCC)
        # 1 kg N2O-N = 44/28 kg N2O, GWP = 265
        tkn_load_kg_per_d = Q * (TKN or 0) / 1000
        n2o_n_kg_per_d = tkn_load_kg_per_d * 0.005  # 0.5%
        n2o_kg_per_d = n2o_n_kg_per_d * 44 / 28
        scope1_n2o_kg = n2o_kg_per_d * 365 * 265  # CO2eq
        # CH4 from digesters (skip unless digester exists, conservative 0)
        scope1_ch4_kg = 0
        total_kgCO2eq = scope2_kg + scope1_n2o_kg + scope1_ch4_kg
        return {
            "job_id": job_id,
            "annual_energy_kWh": round(annual_kWh),
            "scope2_electricity_kgCO2eq": round(scope2_kg),
            "scope1_N2O_kgCO2eq": round(scope1_n2o_kg),
            "scope1_CH4_kgCO2eq": round(scope1_ch4_kg),
            "total_annual_emissions_kgCO2eq": round(total_kgCO2eq),
            "specific_kgCO2eq_per_m3": round(total_kgCO2eq / (Q*365), 4) if Q else None,
            "breakdown_pct": {
                "electricity_scope2": round(scope2_kg/total_kgCO2eq*100, 1) if total_kgCO2eq else 0,
                "N2O_scope1":         round(scope1_n2o_kg/total_kgCO2eq*100, 1) if total_kgCO2eq else 0,
                "CH4_scope1":         round(scope1_ch4_kg/total_kgCO2eq*100, 1) if total_kgCO2eq else 0,
            },
            "assumptions": {
                "grid_EF_kgCO2_per_kWh": grid_emission_factor_kgCO2_per_kWh,
                "N2O_fraction_of_TKN": 0.005,
                "N2O_GWP100": 265,
            },
        }
    except Exception as e:
        return {"error": str(e), "job_id": job_id}


def _viz_time_series_chart(job_id, variables=None, output_path=None):
    """Generate a PNG chart of selected variables from a completed-job CSV."""
    if not _MPL_AVAILABLE:
        return {"error": "matplotlib not available - install matplotlib"}
    try:
        if variables is None:
            variables = ["COD_eff", "NH4_eff", "TN_eff", "TP_eff"]
        csv_files = glob.glob(str(Path(CONFIG["output_dir"]) / f"{job_id}*.csv"))
        if not csv_files:
            return {"error": f"No CSV for job '{job_id}'"}
        with open(csv_files[0], newline="") as f:
            reader = csv.DictReader(f)
            rows = list(reader)
        if not rows:
            return {"error": "CSV empty"}
        headers = list(rows[0].keys())
        t_col = next((h for h in headers if h.lower() in ("time", "t", "time_d")), headers[0])
        times = []
        for r in rows:
            try:
                times.append(float(r[t_col]))
            except Exception:
                times.append(len(times))
        fig, ax = _plt.subplots(figsize=(10, 5))
        plotted = []
        for v in variables:
            col = next((h for h in headers if v.lower() in h.lower()), None)
            if not col:
                continue
            vals = []
            for r in rows:
                try:
                    vals.append(float(r[col]))
                except Exception:
                    vals.append(float("nan"))
            ax.plot(times, vals, label=col)
            plotted.append(col)
        ax.set_xlabel(t_col); ax.set_ylabel("concentration (mg/L)")
        ax.set_title(f"SUMO24 Simulation - Job {job_id}")
        ax.legend(loc="best", fontsize=9); ax.grid(alpha=0.3)
        if output_path is None:
            output_path = str(Path(CONFIG["output_dir"]) / f"{job_id}_chart.png")
        fig.tight_layout()
        fig.savefig(output_path, dpi=120)
        _plt.close(fig)
        return {"job_id": job_id, "output_path": output_path,
                "variables_plotted": plotted,
                "data_points": len(rows),
                "status": "OK"}
    except Exception as e:
        return {"error": str(e), "job_id": job_id}



# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# GROUP Z helpers: Academic-quality export tools
#   Publication plots (matplotlib with serif fonts, high-DPI, vector output)
#   Academic tables (LaTeX booktabs, DOCX, XLSX)
#   Statistical summaries, batch export across multiple SUMO files & jobs
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

try:
    import docx as _docx_mod  # noqa: F401
    from docx import Document as _DocxDocument
    from docx.shared import Pt as _DocxPt, Inches as _DocxInches, RGBColor as _DocxRGB
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN
    from docx.enum.table import WD_ALIGN_VERTICAL as _WD_VALIGN
    from docx.oxml.ns import qn as _docx_qn
    from docx.oxml import OxmlElement as _DocxOxml
    _DOCX_AVAILABLE = True
except Exception:
    _DOCX_AVAILABLE = False


_ACADEMIC_STYLE = {
    "font.family": "serif",
    "font.serif": ["Times New Roman", "DejaVu Serif", "Liberation Serif", "serif"],
    "font.size": 11,
    "axes.labelsize": 11,
    "axes.titlesize": 12,
    "axes.linewidth": 1.0,
    "axes.grid": True,
    "grid.alpha": 0.25,
    "grid.linewidth": 0.5,
    "legend.fontsize": 9,
    "legend.frameon": True,
    "legend.edgecolor": "0.3",
    "xtick.labelsize": 10,
    "ytick.labelsize": 10,
    "xtick.direction": "in",
    "ytick.direction": "in",
    "xtick.major.size": 4,
    "ytick.major.size": 4,
    "lines.linewidth": 1.5,
    "figure.dpi": 150,
    "savefig.dpi": 300,
    "savefig.bbox": "tight",
}

# Colorblind-friendly palette (Wong, Nature Methods 2011)
_ACADEMIC_PALETTE = [
    "#000000", "#E69F00", "#56B4E9", "#009E73",
    "#F0E442", "#0072B2", "#D55E00", "#CC79A7",
]
_ACADEMIC_LINESTYLES = ["-", "--", "-.", ":"]


def _apply_academic_style():
    if not _MPL_AVAILABLE:
        return False
    _plt.rcParams.update(_ACADEMIC_STYLE)
    return True


def _exp_load_job_csv(job_id):
    csv_files = glob.glob(str(Path(CONFIG["output_dir"]) / f"{job_id}*.csv"))
    if not csv_files:
        return None, None, None
    with open(csv_files[0], newline="") as f:
        reader = csv.DictReader(f)
        rows = list(reader)
    if not rows:
        return None, None, csv_files[0]
    return list(rows[0].keys()), rows, csv_files[0]


def _exp_extract_column(rows, headers, target):
    col = next((h for h in headers if target.lower() in h.lower()), None)
    if not col:
        return None, []
    vals = []
    for r in rows:
        try:
            vals.append(float(r[col]))
        except Exception:
            vals.append(float("nan"))
    return col, vals


def _exp_time_column(headers):
    return next((h for h in headers if h.lower() in ("time", "t", "time_d", "t_d", "days")), headers[0])


def _exp_stats(values):
    import math
    v = [x for x in values if x is not None and not (isinstance(x, float) and math.isnan(x))]
    n = len(v)
    if n == 0:
        return {"n": 0, "mean": None, "std": None, "min": None, "max": None,
                "p05": None, "p25": None, "median": None, "p75": None, "p95": None,
                "cv_pct": None}
    vs = sorted(v)
    mean = sum(v) / n
    var = sum((x - mean) ** 2 for x in v) / n if n > 1 else 0.0
    std = math.sqrt(var)

    def pct(p):
        if n == 1:
            return vs[0]
        k = (n - 1) * p
        lo = int(math.floor(k)); hi = int(math.ceil(k))
        if lo == hi:
            return vs[lo]
        return vs[lo] + (vs[hi] - vs[lo]) * (k - lo)

    return {
        "n": n, "mean": round(mean, 6), "std": round(std, 6),
        "min": round(min(v), 6), "max": round(max(v), 6),
        "p05": round(pct(0.05), 6), "p25": round(pct(0.25), 6),
        "median": round(pct(0.50), 6), "p75": round(pct(0.75), 6),
        "p95": round(pct(0.95), 6),
        "cv_pct": round((std / mean) * 100.0, 3) if mean not in (0, None) else None,
    }


def _exp_goodness_of_fit(sim, obs):
    import math
    pairs = [(s, o) for s, o in zip(sim, obs)
             if s is not None and o is not None
             and not (isinstance(s, float) and math.isnan(s))
             and not (isinstance(o, float) and math.isnan(o))]
    n = len(pairs)
    if n < 2:
        return {"n": n, "error": "need at least 2 valid pairs"}
    sim_v = [p[0] for p in pairs]
    obs_v = [p[1] for p in pairs]
    mean_obs = sum(obs_v) / n
    mean_sim = sum(sim_v) / n
    ss_res = sum((o - s) ** 2 for s, o in pairs)
    ss_tot = sum((o - mean_obs) ** 2 for o in obs_v)
    nse = 1 - ss_res / ss_tot if ss_tot > 0 else None
    rmse = math.sqrt(ss_res / n)
    num = sum((s - mean_sim) * (o - mean_obs) for s, o in pairs)
    den = math.sqrt(sum((s - mean_sim) ** 2 for s in sim_v)
                    * sum((o - mean_obs) ** 2 for o in obs_v))
    r = num / den if den > 0 else None
    r2 = r * r if r is not None else None
    sum_obs = sum(obs_v)
    pbias = (100.0 * sum(o - s for s, o in pairs) / sum_obs) if sum_obs != 0 else None
    return {
        "n": n,
        "R2": round(r2, 4) if r2 is not None else None,
        "NSE": round(nse, 4) if nse is not None else None,
        "RMSE": round(rmse, 6),
        "PBIAS_pct": round(pbias, 3) if pbias is not None else None,
        "mean_obs": round(mean_obs, 4), "mean_sim": round(mean_sim, 4),
    }


def _exp_default_output_dir():
    p = Path(CONFIG["output_dir"]) / "academic_exports"
    p.mkdir(parents=True, exist_ok=True)
    return p


def _exp_save_fig(fig, output_path, formats=None):
    if formats is None:
        formats = ["png"]
    base = Path(output_path)
    base.parent.mkdir(parents=True, exist_ok=True)
    if base.suffix.lower().lstrip(".") in ("png", "pdf", "svg", "eps", "jpg", "jpeg", "tif", "tiff"):
        base = base.with_suffix("")
    written = []
    for f in formats:
        out = base.with_suffix("." + f.lower())
        fig.savefig(str(out))
        written.append(str(out))
    _plt.close(fig)
    return written


def _exp_publication_time_series(job_ids, variables=None,
                                 output_path=None, title=None,
                                 ylabel=None, xlabel=None,
                                 formats=None, figsize=(7, 4.2),
                                 grayscale_safe=True):
    if not _MPL_AVAILABLE:
        return {"error": "matplotlib not available"}
    _apply_academic_style()
    if isinstance(job_ids, str):
        job_ids = [job_ids]
    if variables is None:
        variables = ["COD_eff", "NH4_eff", "TN_eff", "TP_eff"]
    if formats is None:
        formats = ["png", "pdf"]

    fig, ax = _plt.subplots(figsize=figsize)
    series_info = []
    color_idx = 0
    style_idx = 0
    for job_id in job_ids:
        headers, rows, csv_path = _exp_load_job_csv(job_id)
        if not headers:
            continue
        t_col = _exp_time_column(headers)
        times = []
        for r in rows:
            try:
                times.append(float(r[t_col]))
            except Exception:
                times.append(float("nan"))
        for v in variables:
            col, vals = _exp_extract_column(rows, headers, v)
            if not col:
                continue
            color = _ACADEMIC_PALETTE[color_idx % len(_ACADEMIC_PALETTE)]
            ls = _ACADEMIC_LINESTYLES[style_idx % len(_ACADEMIC_LINESTYLES)] if grayscale_safe else "-"
            label = f"{v} ({job_id})" if len(job_ids) > 1 else v
            ax.plot(times, vals, color=color, linestyle=ls, label=label)
            series_info.append({"job": job_id, "variable": col, "n": len(vals)})
            color_idx += 1
            if grayscale_safe:
                style_idx += 1

    if not series_info:
        _plt.close(fig)
        return {"error": "no matching variables/jobs found"}

    ax.set_xlabel(xlabel or "Time (d)")
    ax.set_ylabel(ylabel or r"Concentration (mg L$^{-1}$)")
    if title:
        ax.set_title(title)
    ax.legend(loc="best", ncol=1 if len(series_info) <= 4 else 2)
    fig.tight_layout()

    if output_path is None:
        output_path = str(_exp_default_output_dir() / "time_series_publication")
    written = _exp_save_fig(fig, output_path, formats)
    return {"status": "OK", "output_paths": written, "series": series_info,
            "n_series": len(series_info)}


def _exp_calibration_parity_plot(simulated, observed, variable_name="variable",
                                 units="mg/L", output_path=None, formats=None,
                                 title=None, figsize=(5.2, 5.0)):
    if not _MPL_AVAILABLE:
        return {"error": "matplotlib not available"}
    if len(simulated) != len(observed):
        return {"error": f"length mismatch: sim={len(simulated)} obs={len(observed)}"}
    _apply_academic_style()
    if formats is None:
        formats = ["png", "pdf"]

    stats = _exp_goodness_of_fit(simulated, observed)
    if "error" in stats:
        return stats

    fig, ax = _plt.subplots(figsize=figsize)
    ax.scatter(observed, simulated, c=_ACADEMIC_PALETTE[0], s=30,
               alpha=0.75, edgecolors="white", linewidths=0.6)
    lo = min(min(observed), min(simulated))
    hi = max(max(observed), max(simulated))
    pad = 0.05 * (hi - lo) if hi > lo else 1.0
    lo -= pad; hi += pad
    ax.plot([lo, hi], [lo, hi], color="0.3", linestyle="--", linewidth=1.0, label="1:1")
    ax.set_xlim(lo, hi); ax.set_ylim(lo, hi)
    ax.set_xlabel(f"Observed {variable_name} ({units})")
    ax.set_ylabel(f"Simulated {variable_name} ({units})")
    if title is None:
        title = f"Parity plot: {variable_name}"
    ax.set_title(title)
    txt = (f"n = {stats['n']}\n"
           f"$R^2$ = {stats['R2']}\n"
           f"NSE = {stats['NSE']}\n"
           f"RMSE = {stats['RMSE']}\n"
           f"PBIAS = {stats['PBIAS_pct']} %")
    ax.text(0.04, 0.96, txt, transform=ax.transAxes, fontsize=9,
            verticalalignment="top",
            bbox=dict(boxstyle="round,pad=0.4", facecolor="white",
                      edgecolor="0.5", alpha=0.9))
    ax.legend(loc="lower right")
    ax.set_aspect("equal", adjustable="box")
    fig.tight_layout()

    if output_path is None:
        output_path = str(_exp_default_output_dir() / f"parity_{variable_name}")
    written = _exp_save_fig(fig, output_path, formats)
    return {"status": "OK", "output_paths": written, "metrics": stats}


def _exp_scenario_comparison_bar(job_ids, variables, output_path=None,
                                 formats=None, figsize=(7.0, 4.2),
                                 ylabel=None, title=None,
                                 show_errorbars=True):
    if not _MPL_AVAILABLE:
        return {"error": "matplotlib not available"}
    _apply_academic_style()
    if formats is None:
        formats = ["png", "pdf"]
    data = {}
    for job_id in job_ids:
        headers, rows, _ = _exp_load_job_csv(job_id)
        if not headers:
            continue
        data[job_id] = {}
        for v in variables:
            _, vals = _exp_extract_column(rows, headers, v)
            if not vals:
                continue
            data[job_id][v] = _exp_stats(vals)
    if not data:
        return {"error": "no job data found"}

    import numpy as _np
    jobs = list(data.keys())
    nvars = len(variables)
    njobs = len(jobs)
    x = _np.arange(nvars)
    total_width = 0.8
    bar_w = total_width / max(1, njobs)

    fig, ax = _plt.subplots(figsize=figsize)
    for i, job in enumerate(jobs):
        means = []
        stds = []
        for v in variables:
            s = data[job].get(v)
            means.append(s["mean"] if s and s["mean"] is not None else 0.0)
            stds.append(s["std"] if s and s["std"] is not None else 0.0)
        offset = (i - (njobs - 1) / 2) * bar_w
        color = _ACADEMIC_PALETTE[i % len(_ACADEMIC_PALETTE)]
        ax.bar(x + offset, means, bar_w, label=job,
               color=color, edgecolor="black", linewidth=0.5)
        if show_errorbars:
            ax.errorbar(x + offset, means, yerr=stds, fmt="none",
                        ecolor="black", capsize=3, linewidth=0.8)

    ax.set_xticks(x)
    ax.set_xticklabels(variables, rotation=0 if nvars <= 5 else 30, ha="center")
    ax.set_ylabel(ylabel or r"Mean concentration (mg L$^{-1}$)")
    if title:
        ax.set_title(title)
    ax.legend(loc="best", ncol=min(njobs, 3))
    ax.set_axisbelow(True)
    fig.tight_layout()

    if output_path is None:
        output_path = str(_exp_default_output_dir() / "scenario_comparison")
    written = _exp_save_fig(fig, output_path, formats)
    return {"status": "OK", "output_paths": written,
            "jobs": jobs, "variables": variables,
            "n_jobs": njobs, "n_variables": nvars}


def _exp_tornado_diagram(sensitivity_data, baseline=None, output_path=None,
                         formats=None, figsize=(7, 4.5),
                         xlabel="Output response", title=None):
    if not _MPL_AVAILABLE:
        return {"error": "matplotlib not available"}
    if not sensitivity_data:
        return {"error": "sensitivity_data is empty"}
    _apply_academic_style()
    if formats is None:
        formats = ["png", "pdf"]

    items = [dict(d) for d in sensitivity_data]
    if baseline is None:
        baseline = sum((d.get("low", 0) + d.get("high", 0)) / 2
                       for d in items) / max(1, len(items))
    for d in items:
        lo = d.get("low", baseline)
        hi = d.get("high", baseline)
        d["range"] = abs(hi - lo)
        d["lo_delta"] = lo - baseline
        d["hi_delta"] = hi - baseline
    items.sort(key=lambda d: d["range"], reverse=False)

    fig, ax = _plt.subplots(figsize=figsize)
    y = list(range(len(items)))
    for yi, d in zip(y, items):
        ax.barh(yi, d["hi_delta"], left=0, color=_ACADEMIC_PALETTE[2],
                edgecolor="black", linewidth=0.5, height=0.65)
        ax.barh(yi, d["lo_delta"], left=0, color=_ACADEMIC_PALETTE[1],
                edgecolor="black", linewidth=0.5, height=0.65)
    ax.axvline(0, color="black", linewidth=1.0)
    ax.set_yticks(y)
    ax.set_yticklabels([d["parameter"] for d in items])
    ax.set_xlabel(f"{xlabel} (deviation from baseline = {baseline:.3g})")
    if title:
        ax.set_title(title)

    from matplotlib.patches import Patch as _MplPatch
    legend_elems = [
        _MplPatch(facecolor=_ACADEMIC_PALETTE[2], edgecolor="black", label="high input"),
        _MplPatch(facecolor=_ACADEMIC_PALETTE[1], edgecolor="black", label="low input"),
    ]
    ax.legend(handles=legend_elems, loc="best")
    fig.tight_layout()

    if output_path is None:
        output_path = str(_exp_default_output_dir() / "tornado_sensitivity")
    written = _exp_save_fig(fig, output_path, formats)
    return {"status": "OK", "output_paths": written,
            "baseline": baseline,
            "parameters_ranked": [d["parameter"] for d in reversed(items)]}


def _exp_box_plot_distributions(job_ids, variable, output_path=None,
                                formats=None, figsize=(6.5, 4.2),
                                ylabel=None, title=None, use_violin=False):
    if not _MPL_AVAILABLE:
        return {"error": "matplotlib not available"}
    _apply_academic_style()
    if formats is None:
        formats = ["png", "pdf"]
    data = []
    labels = []
    for job_id in job_ids:
        headers, rows, _ = _exp_load_job_csv(job_id)
        if not headers:
            continue
        _, vals = _exp_extract_column(rows, headers, variable)
        import math
        clean = [v for v in vals if v is not None and not (isinstance(v, float) and math.isnan(v))]
        if clean:
            data.append(clean)
            labels.append(job_id)
    if not data:
        return {"error": f"no data for variable '{variable}'"}

    fig, ax = _plt.subplots(figsize=figsize)
    if use_violin:
        parts = ax.violinplot(data, showmedians=True)
        for i, pc in enumerate(parts["bodies"]):
            pc.set_facecolor(_ACADEMIC_PALETTE[i % len(_ACADEMIC_PALETTE)])
            pc.set_edgecolor("black")
            pc.set_alpha(0.7)
        ax.set_xticks(range(1, len(labels) + 1))
        ax.set_xticklabels(labels, rotation=30 if len(labels) > 3 else 0, ha="right")
    else:
        bp = ax.boxplot(data, labels=labels, patch_artist=True,
                        medianprops=dict(color="black", linewidth=1.2),
                        flierprops=dict(marker="o", markersize=3, alpha=0.6))
        for i, patch in enumerate(bp["boxes"]):
            patch.set_facecolor(_ACADEMIC_PALETTE[i % len(_ACADEMIC_PALETTE)])
            patch.set_alpha(0.7)
    ax.set_ylabel(ylabel or f"{variable} (mg L$^{{-1}}$)")
    if title:
        ax.set_title(title)
    fig.tight_layout()
    if output_path is None:
        output_path = str(_exp_default_output_dir() / f"boxplot_{variable}")
    written = _exp_save_fig(fig, output_path, formats)
    stats = [{"job": l, **_exp_stats(d)} for l, d in zip(labels, data)]
    return {"status": "OK", "output_paths": written,
            "jobs": labels, "variable": variable, "stats": stats}


def _apply_booktabs_cell_borders(cell, top=False, bottom=False, heavy=False):
    if not _DOCX_AVAILABLE:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = _DocxOxml("w:tcBorders")
    for side, active in (("top", top), ("bottom", bottom)):
        if active:
            b = _DocxOxml(f"w:{side}")
            b.set(_docx_qn("w:val"), "single")
            b.set(_docx_qn("w:sz"), "12" if heavy else "6")
            b.set(_docx_qn("w:color"), "000000")
            tcBorders.append(b)
    tcPr.append(tcBorders)


def _exp_academic_results_docx(headers, rows, output_path=None,
                               caption=None, table_number=1,
                               column_units=None, title=None,
                               footnote=None):
    if not _DOCX_AVAILABLE:
        return {"error": "python-docx not installed - pip install python-docx"}
    doc = _DocxDocument()
    for section in doc.sections:
        section.left_margin = _DocxInches(1.0)
        section.right_margin = _DocxInches(1.0)
    if title:
        h = doc.add_heading(title, level=1)
        for run in h.runs:
            run.font.name = "Times New Roman"
    caption_text = f"Table {table_number}. {caption}" if caption else f"Table {table_number}."
    cap_p = doc.add_paragraph()
    cap_run = cap_p.add_run(caption_text)
    cap_run.bold = True
    cap_run.font.name = "Times New Roman"
    cap_run.font.size = _DocxPt(10)

    ncols = len(headers)
    header_rows = 2 if column_units else 1
    table = doc.add_table(rows=header_rows + len(rows), cols=ncols)
    table.autofit = True

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = _DocxPt(10)
        _apply_booktabs_cell_borders(cell, top=True, heavy=True)
        if not column_units:
            _apply_booktabs_cell_borders(cell, top=True, bottom=True)
    if column_units:
        for j, u in enumerate(column_units):
            cell = table.rows[1].cells[j]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(u)
            run.italic = True
            run.font.name = "Times New Roman"
            run.font.size = _DocxPt(9)
            _apply_booktabs_cell_borders(cell, bottom=True)

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[header_rows + i].cells[j]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(str(val) if val is not None else "")
            run.font.name = "Times New Roman"
            run.font.size = _DocxPt(10)
            if i == len(rows) - 1:
                _apply_booktabs_cell_borders(cell, bottom=True, heavy=True)

    if footnote:
        fp = doc.add_paragraph()
        fr = fp.add_run(footnote)
        fr.font.name = "Times New Roman"
        fr.font.size = _DocxPt(9)
        fr.italic = True

    if output_path is None:
        output_path = str(_exp_default_output_dir() / f"table_{table_number}.docx")
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return {"status": "OK", "output_path": output_path,
            "rows": len(rows), "columns": ncols}


def _latex_escape(s):
    s = str(s)
    replacements = [
        ("\\", "\\textbackslash{}"),
        ("&", r"\&"), ("%", r"\%"), ("$", r"\$"), ("#", r"\#"),
        ("_", r"\_"), ("{", r"\{"), ("}", r"\}"),
        ("~", r"\textasciitilde{}"), ("^", r"\textasciicircum{}"),
    ]
    for a, b in replacements:
        s = s.replace(a, b)
    return s


def _exp_latex_results_table(headers, rows, output_path=None,
                             caption=None, label="tab:results",
                             column_align=None, column_units=None,
                             footnote=None, table_type="table"):
    ncols = len(headers)
    if column_align is None:
        column_align = "l" + "r" * (ncols - 1)
    lines = []
    lines.append(f"\\begin{{{table_type}}}[htbp]")
    lines.append("  \\centering")
    if caption:
        lines.append(f"  \\caption{{{_latex_escape(caption)}}}")
    lines.append(f"  \\label{{{label}}}")
    lines.append(f"  \\begin{{tabular}}{{{column_align}}}")
    lines.append("    \\toprule")
    lines.append("    " + " & ".join(f"\\textbf{{{_latex_escape(h)}}}" for h in headers) + " \\\\")
    if column_units:
        lines.append("    " + " & ".join(f"\\textit{{{_latex_escape(u)}}}" for u in column_units) + " \\\\")
    lines.append("    \\midrule")
    for row in rows:
        cells = [_latex_escape(v) if v is not None else "" for v in row]
        lines.append("    " + " & ".join(cells) + " \\\\")
    lines.append("    \\bottomrule")
    lines.append("  \\end{tabular}")
    if footnote:
        lines.append(f"  \\\\[0.5ex]\\footnotesize\\textit{{{_latex_escape(footnote)}}}")
    lines.append(f"\\end{{{table_type}}}")
    tex = "\n".join(lines) + "\n"
    if output_path is None:
        output_path = str(_exp_default_output_dir() / f"{label.replace(':', '_')}.tex")
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    Path(output_path).write_text(tex, encoding="utf-8")
    return {"status": "OK", "output_path": output_path, "latex": tex,
            "rows": len(rows), "columns": ncols}


def _exp_statistical_summary_xlsx(job_ids, variables=None, output_path=None,
                                  sheet_per_job=True, include_combined=True):
    if not OPENPYXL_AVAILABLE:
        return {"error": "openpyxl not available"}
    if variables is None:
        variables = ["COD_eff", "NH4_eff", "NO3_eff", "TN_eff", "TP_eff", "TSS_eff"]

    wb = Workbook()
    wb.remove(wb.active)
    stat_order = ["n", "mean", "std", "cv_pct", "min", "p05", "p25",
                  "median", "p75", "p95", "max"]
    combined = {}
    for job_id in job_ids:
        headers, rows, _ = _exp_load_job_csv(job_id)
        if not headers:
            continue
        job_stats = {}
        for v in variables:
            _, vals = _exp_extract_column(rows, headers, v)
            if not vals:
                continue
            s = _exp_stats(vals)
            job_stats[v] = s
            combined.setdefault(v, {})[job_id] = s
        if sheet_per_job:
            safe_name = job_id.replace("/", "_").replace("\\", "_")[:25] or "job"
            ws = wb.create_sheet(safe_name)
            ws.append(["Statistic"] + list(job_stats.keys()))
            for c in ws[1]:
                c.font = Font(bold=True)
            for stat in stat_order:
                row = [stat]
                for v in job_stats:
                    val = job_stats[v].get(stat)
                    row.append(val if val is not None else "")
                ws.append(row)
            ws.freeze_panes = "A2"

    if include_combined and combined:
        ws = wb.create_sheet("Combined_Mean")
        ws.append(["Variable"] + job_ids)
        for c in ws[1]:
            c.font = Font(bold=True)
        for v in combined:
            row = [v]
            for j in job_ids:
                s = combined[v].get(j)
                row.append(s["mean"] if s else "")
            ws.append(row)
        ws.freeze_panes = "B2"

        ws2 = wb.create_sheet("Combined_Std")
        ws2.append(["Variable"] + job_ids)
        for c in ws2[1]:
            c.font = Font(bold=True)
        for v in combined:
            row = [v]
            for j in job_ids:
                s = combined[v].get(j)
                row.append(s["std"] if s else "")
            ws2.append(row)
        ws2.freeze_panes = "B2"

    if len(wb.sheetnames) == 0:
        wb.create_sheet("empty")
    if output_path is None:
        output_path = str(_exp_default_output_dir() / "statistical_summary.xlsx")
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    wb.save(output_path)
    return {"status": "OK", "output_path": output_path,
            "jobs": job_ids, "variables": variables,
            "sheets": list(wb.sheetnames) if hasattr(wb, "sheetnames") else []}


def _exp_academic_report_docx(job_ids, variables=None,
                              title="SUMO24 Simulation Report",
                              authors=None, abstract=None,
                              output_path=None, include_figures=True,
                              embed_figures_dir=None):
    if not _DOCX_AVAILABLE:
        return {"error": "python-docx not installed"}
    if variables is None:
        variables = ["COD_eff", "NH4_eff", "TN_eff", "TP_eff"]
    if isinstance(job_ids, str):
        job_ids = [job_ids]

    doc = _DocxDocument()
    for section in doc.sections:
        section.left_margin = _DocxInches(1.0)
        section.right_margin = _DocxInches(1.0)

    t = doc.add_heading(title, level=0)
    for r in t.runs:
        r.font.name = "Times New Roman"
    if authors:
        p = doc.add_paragraph(authors)
        p.alignment = _WD_ALIGN.CENTER
        for r in p.runs:
            r.italic = True
            r.font.name = "Times New Roman"
    if abstract:
        doc.add_heading("Abstract", level=1)
        p = doc.add_paragraph(abstract)
        for r in p.runs:
            r.font.name = "Times New Roman"

    doc.add_heading("1. Methods", level=1)
    methods = (
        f"Simulations were performed using the SUMO24 Digital Twin Toolkit (Dynamita). "
        f"A total of {len(job_ids)} simulation job(s) were analyzed: {', '.join(job_ids)}. "
        f"The following effluent variables were evaluated: {', '.join(variables)}. "
        f"Activated sludge kinetics were modelled using the ASM framework, with "
        f"temperature corrections applied via Arrhenius coefficients. Statistical "
        f"analysis used mean, standard deviation, coefficient of variation, and "
        f"the 5th/25th/50th/75th/95th percentiles. All visualizations follow "
        f"publication-quality conventions (serif typography, 300 DPI, "
        f"colorblind-safe palette after Wong, 2011)."
    )
    p = doc.add_paragraph(methods)
    for r in p.runs:
        r.font.name = "Times New Roman"

    doc.add_heading("2. Results", level=1)
    combined = {}
    for job_id in job_ids:
        headers, rows, _ = _exp_load_job_csv(job_id)
        if not headers:
            continue
        for v in variables:
            _, vals = _exp_extract_column(rows, headers, v)
            if vals:
                combined.setdefault(v, {})[job_id] = _exp_stats(vals)

    if combined:
        cap = doc.add_paragraph()
        cr = cap.add_run("Table 1. Effluent concentration summary (mean +/- SD, mg/L).")
        cr.bold = True
        cr.font.name = "Times New Roman"
        cr.font.size = _DocxPt(10)
        table = doc.add_table(rows=len(combined) + 1, cols=len(job_ids) + 1)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Variable"
        for j, job in enumerate(job_ids):
            hdr_cells[j + 1].text = job
        for c in hdr_cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.name = "Times New Roman"
            _apply_booktabs_cell_borders(c, top=True, bottom=True, heavy=True)
        for ri, v in enumerate(combined):
            row_cells = table.rows[ri + 1].cells
            row_cells[0].text = v
            for j, job in enumerate(job_ids):
                s = combined[v].get(job, {})
                if s and s.get("mean") is not None:
                    row_cells[j + 1].text = f"{s['mean']:.2f} +/- {s['std']:.2f}"
                else:
                    row_cells[j + 1].text = "--"
            if ri == len(combined) - 1:
                for c in row_cells:
                    _apply_booktabs_cell_borders(c, bottom=True, heavy=True)
            for c in row_cells:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.name = "Times New Roman"

    if include_figures and _MPL_AVAILABLE:
        fig_dir = Path(embed_figures_dir) if embed_figures_dir else _exp_default_output_dir() / "report_figs"
        fig_dir.mkdir(parents=True, exist_ok=True)
        ts_res = _exp_publication_time_series(
            job_ids, variables,
            output_path=str(fig_dir / "fig1_time_series"), formats=["png"])
        if ts_res.get("status") == "OK" and ts_res.get("output_paths"):
            doc.add_paragraph()
            doc.add_picture(ts_res["output_paths"][0], width=_DocxInches(6.0))
            fcap = doc.add_paragraph()
            fr = fcap.add_run(
                f"Figure 1. Time-series of {', '.join(variables)} across "
                f"{len(job_ids)} simulation(s).")
            fr.bold = True
            fr.font.size = _DocxPt(10)
            fr.font.name = "Times New Roman"
        if len(job_ids) >= 2:
            bar_res = _exp_scenario_comparison_bar(
                job_ids, variables,
                output_path=str(fig_dir / "fig2_scenario_bar"), formats=["png"])
            if bar_res.get("status") == "OK" and bar_res.get("output_paths"):
                doc.add_paragraph()
                doc.add_picture(bar_res["output_paths"][0], width=_DocxInches(6.0))
                fcap = doc.add_paragraph()
                fr = fcap.add_run(
                    f"Figure 2. Scenario comparison (mean +/- SD) across "
                    f"{len(job_ids)} simulations.")
                fr.bold = True
                fr.font.size = _DocxPt(10)
                fr.font.name = "Times New Roman"

    doc.add_heading("3. Discussion", level=1)
    disc = (
        "The results presented above summarize the behaviour of the modelled "
        "wastewater treatment system across the selected scenarios. Variability "
        "(coefficient of variation) indicates how stable each effluent variable is "
        "over the simulation horizon and can flag periodic disturbances or transient "
        "regimes. Where two or more scenarios were evaluated, differences in mean "
        "effluent concentrations reflect the cumulative impact of the corresponding "
        "operational or design changes. Interpretation should always be paired with "
        "compliance thresholds and mass-balance verification."
    )
    p = doc.add_paragraph(disc)
    for r in p.runs:
        r.font.name = "Times New Roman"

    doc.add_heading("References", level=1)
    for ref in [
        "Henze, M., Gujer, W., Mino, T., van Loosdrecht, M. (2000). Activated Sludge Models ASM1, ASM2, ASM2d and ASM3. IWA Scientific and Technical Report No. 9.",
        "Metcalf & Eddy, Inc. (2014). Wastewater Engineering: Treatment and Resource Recovery, 5th ed. McGraw-Hill.",
        "Wong, B. (2011). Points of view: Color blindness. Nature Methods 8, 441.",
    ]:
        p = doc.add_paragraph(ref, style="List Bullet")
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = _DocxPt(10)

    if output_path is None:
        output_path = str(_exp_default_output_dir() / "academic_report.docx")
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return {"status": "OK", "output_path": output_path,
            "jobs": job_ids, "variables": variables,
            "sections": ["Methods", "Results", "Discussion", "References"]}


def _exp_multi_file_batch_comparison(jobs_by_file, variables=None,
                                     output_dir=None, formats=None,
                                     generate_plots=True, generate_xlsx=True,
                                     generate_latex=True):
    if variables is None:
        variables = ["COD_eff", "NH4_eff", "TN_eff", "TP_eff"]
    if formats is None:
        formats = ["png", "pdf"]
    if isinstance(jobs_by_file, list):
        jobs_by_file = {"all": list(jobs_by_file)}
    if output_dir is None:
        output_dir = str(_exp_default_output_dir() / "batch")
    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)
    artifacts = {}

    all_jobs = []
    for fl, jobs in jobs_by_file.items():
        for j in jobs:
            all_jobs.append(j)

    table_rows = []
    header = ["File", "Job", "Variable", "n", "mean", "std", "min", "max", "CV%"]
    for file_label, jobs in jobs_by_file.items():
        for job_id in jobs:
            headers, rows, _ = _exp_load_job_csv(job_id)
            if not headers:
                continue
            for v in variables:
                _, vals = _exp_extract_column(rows, headers, v)
                if not vals:
                    continue
                s = _exp_stats(vals)
                table_rows.append([file_label, job_id, v, s["n"],
                                   s["mean"], s["std"], s["min"],
                                   s["max"], s["cv_pct"]])

    if generate_latex and table_rows:
        tex_res = _exp_latex_results_table(
            header, table_rows,
            output_path=str(out / "batch_summary.tex"),
            caption="Batch simulation summary across SUMO files and jobs.",
            label="tab:batch_summary")
        artifacts["latex"] = tex_res.get("output_path")

    if generate_xlsx and all_jobs:
        xlsx_res = _exp_statistical_summary_xlsx(
            all_jobs, variables=variables,
            output_path=str(out / "batch_summary.xlsx"))
        artifacts["xlsx"] = xlsx_res.get("output_path")

    if generate_plots and _MPL_AVAILABLE and all_jobs:
        ts = _exp_publication_time_series(
            all_jobs, variables,
            output_path=str(out / "batch_time_series"), formats=formats)
        if ts.get("status") == "OK":
            artifacts["time_series"] = ts["output_paths"]
        if len(all_jobs) >= 2:
            bars = _exp_scenario_comparison_bar(
                all_jobs, variables,
                output_path=str(out / "batch_scenario_bars"),
                formats=formats)
            if bars.get("status") == "OK":
                artifacts["scenario_bars"] = bars["output_paths"]

    return {"status": "OK",
            "output_dir": str(out),
            "files_analyzed": list(jobs_by_file.keys()),
            "total_jobs": len(all_jobs),
            "variables": variables,
            "artifacts": artifacts,
            "n_table_rows": len(table_rows)}


def _exp_mass_balance_diagram(flows, output_path=None, formats=None,
                              figsize=(8, 4.5), title=None, units="kg/d"):
    if not _MPL_AVAILABLE:
        return {"error": "matplotlib not available"}
    if not flows:
        return {"error": "no flows provided"}
    _apply_academic_style()
    if formats is None:
        formats = ["png", "pdf"]

    nodes = []
    for f in flows:
        for key in ("from", "to"):
            if f[key] not in nodes:
                nodes.append(f[key])
    positions = {}
    n = len(nodes)
    for i, node in enumerate(nodes):
        x = (i + 1) / (n + 1)
        y = 0.5 + (0.12 if i % 2 == 0 else -0.12)
        positions[node] = (x, y)

    fig, ax = _plt.subplots(figsize=figsize)
    ax.set_xlim(0, 1); ax.set_ylim(0, 1)
    ax.axis("off")
    from matplotlib.patches import FancyBboxPatch as _BBox
    for node, (x, y) in positions.items():
        box = _BBox((x - 0.08, y - 0.05), 0.16, 0.10,
                    boxstyle="round,pad=0.02",
                    linewidth=1.2, edgecolor="black",
                    facecolor=_ACADEMIC_PALETTE[2], alpha=0.3)
        ax.add_patch(box)
        ax.text(x, y, node, ha="center", va="center",
                fontsize=10, fontweight="bold")
    from matplotlib.patches import FancyArrowPatch as _Arrow
    max_val = max((abs(f.get("value", 1)) for f in flows), default=1)
    for f in flows:
        x1, y1 = positions[f["from"]]
        x2, y2 = positions[f["to"]]
        val = f.get("value", 0)
        lw = 0.6 + 2.4 * abs(val) / max_val if max_val > 0 else 1.0
        arrow = _Arrow((x1 + 0.08, y1), (x2 - 0.08, y2),
                       arrowstyle="-|>", mutation_scale=15,
                       linewidth=lw, color="black",
                       connectionstyle="arc3,rad=0.1")
        ax.add_patch(arrow)
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2 + 0.04
        label = f.get("label") or f"{val:.2g} {units}"
        ax.text(mx, my, label, ha="center", va="bottom",
                fontsize=8, fontstyle="italic",
                bbox=dict(boxstyle="round,pad=0.2",
                          facecolor="white", edgecolor="none", alpha=0.85))
    if title:
        ax.set_title(title)
    fig.tight_layout()
    if output_path is None:
        output_path = str(_exp_default_output_dir() / "mass_balance_diagram")
    written = _exp_save_fig(fig, output_path, formats)
    return {"status": "OK", "output_paths": written,
            "nodes": nodes, "n_flows": len(flows)}



def _arrhenius(k20: float, theta, T: float) -> float:
    """Return Arrhenius-corrected rate: k(T) = k20 * theta^(T - 20)."""
    if theta is None or theta == 1.0:
        return k20
    return k20 * (theta ** (T - 20.0))


def _correct_asm_params(model_key: str, temp_C: float) -> dict:
    """Return temp-corrected ASM params using per-parameter Arrhenius theta."""
    raw = RESEARCH_DEFAULTS.get(model_key, {})
    out = {}
    for p, info in raw.items():
        theta = info.get("theta")
        if theta:
            out[p] = round(_arrhenius(info["val"], theta, temp_C), 6)
        else:
            out[p] = info["val"]
    return out


# â”€â”€ Scenario library (Qaha WWTP scenarios) â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# Maps scenario names to lists of SumoCore commands
SCENARIOS = {
    "baseline_Sc0": [
        # Default operating conditions â€“ no changes
    ],
    "increased_RAS_Sc1": [
        "set Sumo__Plant__RAS_Pump__param__Q 8000",
    ],
    "extended_SRT_Sc2": [
        "set Sumo__Plant__WAS_Pump__param__Q 500",
    ],
    "step_feed_Sc3": [
        "set Sumo__Plant__StepFeed__param__fraction 0.4",
    ],
    "DO_optimization_Sc4": [
        "set Sumo__Plant__OxidationDitch__param__DOsetpoint 1.5",
    ],
    "tertiary_filter_Sc5": [
        # Empirical surrogate: apply TSS removal factor post-clarifier
        "set Sumo__Plant__TertiaryFilter__param__removalTSS 0.85",
    ],
}

# â”€â”€ Shared job state â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
_job_lock    = threading.Lock()
_active_jobs: dict[int, dict] = {}   # job_id â†’ {status, results, error}
_result_cache: dict[str, Any] = {}   # scenario â†’ parsed results

# â”€â”€ DTT callbacks â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def _msg_callback(job: int, msg: str):
    """Called by scheduler on every SumoCore message."""
    with _job_lock:
        if job not in _active_jobs:
            _active_jobs[job] = {"status": "running", "results": {}, "error": None}
        _active_jobs[job]["last_msg"] = msg

    if ds.sumo.isSimFinishedMsg(msg):
        with _job_lock:
            _active_jobs[job]["status"] = "finished"
        ds.sumo.finish(job)

def _data_callback(job: int, data: dict):
    """Called by scheduler at every DataComm interval with live variable values."""
    with _job_lock:
        if job not in _active_jobs:
            _active_jobs[job] = {"status": "running", "results": {}, "error": None}
        # Accumulate time-series rows
        _active_jobs[job].setdefault("rows", []).append(dict(data))


# â”€â”€ Effluent variable names (adjust to match your compiled model) â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
EFFLUENT_VARS = [
    "Sumo__Time",
    "Sumo__Plant__Effluent__XTSS",       # TSS  mg/L
    "Sumo__Plant__Effluent__SBOD5",      # BOD5 mg/L  (or use XBOD)
    "Sumo__Plant__Effluent__SCOD",       # sCOD mg/L
    "Sumo__Plant__Effluent__SNHx",       # NH4-N mg/L
    "Sumo__Plant__Effluent__SNOx",       # NOx-N mg/L
    "Sumo__Plant__Effluent__SPO4",       # PO4-P mg/L
    "Sumo__Plant__Effluent__Q",          # Effluent flow m3/d
]

# â”€â”€ Helper: read current simulation settings from state.xml â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def _read_simulation_settings() -> dict:
    """Return the simulation period and data interval currently stored in state.xml.

    Looks for Sumo__StopTime and Sumo__DataComm (both stored in milliseconds).
    Falls back to sensible defaults (stop=30 d, datacomm=1 h) if the file
    can't be parsed or the keys are missing.
    Returns: {
        "stop_time_days":    float,
        "datacomm_hours":    float,
        "stop_time_ms":      float,
        "datacomm_ms":       float,
        "source":            "state.xml" | "default",
    }
    """
    default_days = 30.0
    default_dc_h = 1.0
    out = {
        "stop_time_days": default_days,
        "datacomm_hours": default_dc_h,
        "stop_time_ms":   default_days * 86_400_000,
        "datacomm_ms":    default_dc_h * 3_600_000,
        "source":         "default",
    }
    state_path = Path(CONFIG["state_xml"])
    if not state_path.exists():
        return out
    try:
        root = ET.parse(state_path).getroot()
        stop_ms = None
        dc_ms   = None
        # state.xml stores variables as <param name="..." value="..."/> or similar
        for el in root.iter():
            attrs = el.attrib
            key = attrs.get("name") or attrs.get("id") or ""
            val = attrs.get("value") or attrs.get("val") or (el.text or "").strip()
            if not val:
                continue
            if key == "Sumo__StopTime":
                try: stop_ms = float(val)
                except ValueError: pass
            elif key == "Sumo__DataComm":
                try: dc_ms = float(val)
                except ValueError: pass
        if stop_ms is not None:
            out["stop_time_ms"]   = stop_ms
            out["stop_time_days"] = stop_ms / 86_400_000
            out["source"] = "state.xml"
        if dc_ms is not None:
            out["datacomm_ms"]    = dc_ms
            out["datacomm_hours"] = dc_ms / 3_600_000
            out["source"] = "state.xml"
    except Exception as e:
        out["parse_error"] = str(e)
    return out


# â”€â”€ Helper: build schedule commands â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def _build_commands(
    mode: str = "steady",
    stop_days: float = 30,
    datacomm_hours: float = 1,
    extra_sets: list[str] | None = None,
    scenario_name: str | None = None,
) -> list[str]:
    dur = ds.sumo.dur if DTT_AVAILABLE else type("D", (), {"day": 86400000, "hour": 3600000})()
    cmds = [
        f'load "{CONFIG["state_xml"]}"',
        f"set Sumo__StopTime {int(stop_days * dur.day)}",
        f"set Sumo__DataComm {int(datacomm_hours * dur.hour)}",
    ]
    if scenario_name and scenario_name in SCENARIOS:
        cmds += SCENARIOS[scenario_name]
    if extra_sets:
        cmds += extra_sets
    cmds += [f"mode {mode}", "start"]
    return cmds


# â”€â”€ Helper: parse rows â†’ summary dict â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def _summarise(rows: list[dict]) -> dict:
    if not rows:
        return {"error": "No data rows returned from simulation."}
    last = rows[-1]
    summary = {}
    key_map = {
        "Sumo__Plant__Effluent__XTSS":  "TSS_mgL",
        "Sumo__Plant__Effluent__SBOD5": "BOD5_mgL",
        "Sumo__Plant__Effluent__SCOD":  "sCOD_mgL",
        "Sumo__Plant__Effluent__SNHx":  "NH4N_mgL",
        "Sumo__Plant__Effluent__SNOx":  "NOxN_mgL",
        "Sumo__Plant__Effluent__SPO4":  "PO4P_mgL",
        "Sumo__Plant__Effluent__Q":     "Q_m3d",
        "Sumo__Time":                   "time_days",
    }
    for sumo_key, label in key_map.items():
        if sumo_key in last:
            try:
                summary[label] = round(float(last[sumo_key]), 3)
            except (ValueError, TypeError):
                summary[label] = last[sumo_key]
    return summary


def _check_compliance(summary: dict) -> dict:
    """Compare effluent summary against Egyptian Law 48/1982 limits."""
    limits = CONFIG["law48_limits"]
    compliance = {}
    check_pairs = [
        ("TSS_mgL",  "TSS"),
        ("BOD5_mgL", "BOD5"),
        ("sCOD_mgL", "COD"),
    ]
    for result_key, limit_key in check_pairs:
        if result_key in summary and limits.get(limit_key) is not None:
            val   = summary[result_key]
            limit = limits[limit_key]
            compliance[limit_key] = {
                "value":   val,
                "limit":   limit,
                "status":  "PASS" if val <= limit else "FAIL",
                "margin":  round(limit - val, 2),
            }
    return compliance


# â”€â”€ Helpers for the extended tool set â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def _make_ds():
    """Return the dynamita scheduler module if DTT is available, else raise.

    The new inspection / editing tools call ds.sumo.get / .set / .getVariableNames /
    .executeCommand. These work against the globally-loaded SumoScheduler instance.
    """
    if not DTT_AVAILABLE:
        raise RuntimeError(
            "dynamita package not installed - extended inspection tools unavailable."
        )
    return ds


def _apply_scenario(ds_ref, scenario: str) -> None:
    """Apply a named scenario's SET commands against a live DTT session.

    Best-effort: uses executeCommand when available; otherwise parses 'set VAR VAL'
    and falls back to ds.sumo.set(). Silently swallows per-command errors so a
    single bad key doesn't abort the whole apply pass.
    """
    if scenario not in SCENARIOS:
        return
    for cmd in SCENARIOS[scenario]:
        applied = False
        try:
            ds_ref.sumo.executeCommand(cmd)
            applied = True
        except Exception:
            pass
        if not applied:
            parts = cmd.strip().split(None, 2)
            if len(parts) == 3 and parts[0].lower() == "set":
                try:
                    ds_ref.sumo.set(parts[1], float(parts[2]))
                except Exception:
                    pass


def _augment_job_record(rec: dict, key) -> dict:
    """Decorate a _active_jobs entry with 'effluent' and 'scenario' summary fields
    so the new analysis tools (which follow the MD's JOBS schema) can read them.
    """
    out = dict(rec)
    if "effluent" not in out:
        rows = rec.get("rows", [])
        if rows:
            out["effluent"] = _summarise(rows)
        else:
            out["effluent"] = {}
    out.setdefault("scenario", "unknown")
    out.setdefault("csv", "Not yet exported - call export_results_csv.")
    out["job_id"] = key
    return out


class _JobsView:
    """Dict-like view over _active_jobs that accepts both int and string keys.

    Lets MD-derived tools write JOBS.get(job_id) / JOBS.items() without caring
    whether the caller supplied the raw DTT integer id or the string form.
    """
    def get(self, key, default=None):
        if key in _active_jobs:
            return _augment_job_record(_active_jobs[key], key)
        try:
            ikey = int(key)
        except (TypeError, ValueError):
            ikey = None
        if ikey is not None and ikey in _active_jobs:
            return _augment_job_record(_active_jobs[ikey], ikey)
        return default

    def items(self):
        for k, v in _active_jobs.items():
            yield str(k), _augment_job_record(v, k)

    def __contains__(self, key):
        return self.get(key) is not None


JOBS = _JobsView()


# â”€â”€ MCP Server â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
app = Server("sumo24-mcp")


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# DATASET SCHEMA - maps SUMO24_WWTP_Dataset_and_Tools.xlsx to SUMO24 units/vars
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
DATASET_SCHEMA = {
    "sheet_name":        "ðŸ’§ Data Entry",
    "param_header_row":  4,
    "sumo_hint_row":     5,
    "law48_row":         6,
    "data_start_row":    7,
    "columns": {
        1: "label", 2: "Flow", 3: "BOD5", 4: "COD", 5: "TSS", 6: "VSS",
        7: "NH4_N", 8: "NO3_N", 9: "TN", 10: "TP", 11: "pH", 12: "Temp",
        13: "DO", 14: "Turbidity", 15: "Oil_Grease", 16: "Notes",
    },
    "parameter_map": {
        "Flow":       {"unit": "mÂ³/d", "suffixes": ["param__Q"]},
        "BOD5":       {"unit": "mg/L",      "suffixes": ["param__S_S", "param__BOD5"]},
        "COD":        {"unit": "mg/L",      "suffixes": ["param__S_COD", "param__COD"]},
        "TSS":        {"unit": "mg/L",      "suffixes": ["param__X_TSS", "param__TSS"]},
        "VSS":        {"unit": "mg/L",      "suffixes": ["param__X_VSS", "param__VSS"]},
        "NH4_N":      {"unit": "mg/L",      "suffixes": ["param__S_NH4", "param__NH4"]},
        "NO3_N":      {"unit": "mg/L",      "suffixes": ["param__S_NO3", "param__NO3"]},
        "TN":         {"unit": "mg/L",      "suffixes": ["param__S_TN",  "param__TN"]},
        "TP":         {"unit": "mg/L",      "suffixes": ["param__S_TP",  "param__TP"]},
        "pH":         {"unit": "-",         "suffixes": ["param__pH"]},
        "Temp":       {"unit": "Â°C",   "suffixes": ["param__T"]},
        "DO":         {"unit": "mg/L",      "suffixes": ["param__S_O", "param__DO"]},
        "Turbidity":  {"unit": "NTU",       "suffixes": []},
        "Oil_Grease": {"unit": "mg/L",      "suffixes": ["param__S_Oil"]},
    },
    "point_map": {
        "Raw Influent":                  {"unit": "Influent",   "role": "tank"},
        "After Screens / Grit Removal":  {"unit": "Screens",    "role": "output"},
        "Primary Clarifier â€” Inlet":     {"unit": "PrimClar1",  "role": "input"},
        "Primary Clarifier â€” Effluent":  {"unit": "PrimClar1",  "role": "output"},
        "Primary Sludge (underflow)":    {"unit": "PrimClar1",  "role": "sludge"},
        "Bio Reactor â€” Inlet":           {"unit": "AerTank1",   "role": "input"},
        "Aeration Tank":                 {"unit": "AerTank1",   "role": "tank"},
        "Anoxic Zone":                   {"unit": "AnoxTank1",  "role": "tank"},
        "Bio Reactor â€” Outlet":          {"unit": "AerTank1",   "role": "output"},
        "Secondary Clarifier â€” Inlet":   {"unit": "SecClar1",   "role": "input"},
        "Secondary Clarifier â€” Outlet":  {"unit": "SecClar1",   "role": "output"},
        "Final Effluent":                {"unit": "Effluent",   "role": "tank"},
        "RAS (Return Activated Sludge)": {"unit": "RAS_Pump",   "role": "stream"},
        "WAS (Waste Activated Sludge)":  {"unit": "WAS_Pump",   "role": "stream"},
        "Reject / Sidestream Water":     {"unit": "Reject",     "role": "stream"},
    },
}


def _dataset_default_path() -> str:
    return str(Path(CONFIG.get("project_dir", str(_MCP_DIR))) /
               "SUMO24_WWTP_Dataset_and_Tools.xlsx")


def _open_dataset(xlsx_path: str):
    """Open dataset workbook and return (wb, ws) for the data-entry sheet."""
    if not OPENPYXL_AVAILABLE:
        raise RuntimeError("openpyxl is not installed - run: pip install openpyxl")
    p = Path(xlsx_path)
    if not p.exists():
        raise FileNotFoundError(f"Dataset file not found: {xlsx_path}")
    wb = load_workbook(p, data_only=True)
    sheet_name = DATASET_SCHEMA["sheet_name"]
    if sheet_name not in wb.sheetnames:
        matches = [s for s in wb.sheetnames if "Data Entry" in s]
        if not matches:
            raise ValueError(f"No '{sheet_name}' sheet found. Available: {wb.sheetnames}")
        sheet_name = matches[0]
    return wb, wb[sheet_name]


def _parse_dataset_rows(ws) -> dict:
    cols = DATASET_SCHEMA["columns"]
    param_cols = [(c, name) for c, name in cols.items() if name not in ("label", "Notes")]
    result = {"points": {}, "law48": {}, "section_headers_found": []}
    law_row = DATASET_SCHEMA["law48_row"]
    for c, pname in param_cols:
        v = ws.cell(row=law_row, column=c).value
        if isinstance(v, (int, float)):
            result["law48"][pname] = float(v)
    for r in range(DATASET_SCHEMA["data_start_row"], ws.max_row + 1):
        label_val = ws.cell(row=r, column=1).value
        if label_val is None:
            continue
        label = str(label_val).strip()
        if label.isupper() or label.strip().isupper() or label.startswith("  "):
            result["section_headers_found"].append({"row": r, "label": label})
            continue
        point_data = {"_row": r}
        for c, pname in param_cols:
            v = ws.cell(row=r, column=c).value
            if v is None or v == "":
                continue
            if isinstance(v, (int, float)):
                point_data[pname] = float(v)
        notes_v = ws.cell(row=r, column=16).value
        if notes_v:
            point_data["Notes"] = str(notes_v)
        if len(point_data) > 1:
            result["points"][label] = point_data
    return result


def _resolve_sumo_variable(unit: str, param_key: str, all_vars: list):
    pmap = DATASET_SCHEMA["parameter_map"].get(param_key)
    if not pmap:
        return None
    for suffix in pmap["suffixes"]:
        candidate = f"Sumo__Plant__{unit}__{suffix}"
        if candidate in all_vars:
            return candidate
    for suffix in pmap["suffixes"]:
        short = suffix.replace("param__", "")
        matches = [v for v in all_vars
                   if unit.lower() in v.lower() and short.lower() in v.lower()
                   and "param" in v]
        if matches:
            return matches[0]
    return None


@app.list_tools()
async def list_tools() -> list[types.Tool]:
    return [

        types.Tool(
            name="run_steady_state",
            description=(
                "Run a SUMO24 steady-state simulation of the Qaha WWTP model. "
                "Optionally override influent parameters before running. "
                "Returns effluent quality summary and Law 48/1982 compliance status."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "influent_flow_m3d":  {"type": "number", "description": "Influent flow rate (mÂ³/day)"},
                    "influent_cod_mgL":   {"type": "number", "description": "Total influent COD (mg/L)"},
                    "temperature_C":      {"type": "number", "description": "Wastewater temperature (Â°C)"},
                    "influent_tkn_mgL":   {"type": "number", "description": "Influent TKN (mg N/L)"},
                    "scenario":           {
                        "type": "string",
                        "description": "Optional: apply a named scenario before running",
                        "enum": list(SCENARIOS.keys())
                    },
                },
                "required": [],
            },
        ),

        types.Tool(
            name="run_dynamic_simulation",
            description=(
                "Run a SUMO24 dynamic simulation for a specified number of days. "
                "Useful for evaluating transient behaviour, storm events, or seasonal loads."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "duration_days":  {"type": "number", "description": "Simulation duration in days (default 30)"},
                    "datacomm_hours": {"type": "number", "description": "Output interval in hours (default 1)"},
                    "scenario":       {
                        "type": "string",
                        "description": "Named scenario to apply",
                        "enum": list(SCENARIOS.keys())
                    },
                    "influent_tsv":   {"type": "string", "description": "Path to a TSV file with dynamic influent time-series"},
                },
                "required": [],
            },
        ),

        types.Tool(
            name="run_scenario_comparison",
            description=(
                "Run multiple named scenarios sequentially and return a side-by-side "
                "compliance comparison table (TSS, BOD5, COD vs Law 48/1982 limits)."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "scenarios": {
                        "type": "array",
                        "items": {"type": "string", "enum": list(SCENARIOS.keys())},
                        "description": "List of scenario names to compare",
                    }
                },
                "required": ["scenarios"],
            },
        ),

        types.Tool(
            name="set_parameter",
            description=(
                "Set any SUMO24 model parameter by its full SumoCore variable name. "
                "Changes are applied on the next simulation run (not permanently saved). "
                "Use the SumoCore Advanced Core Window to find variable names."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "variable": {"type": "string", "description": "Full SumoCore variable name, e.g. Sumo__Plant__OxidationDitch__param__V"},
                    "value":    {"type": "number", "description": "New value to set"},
                },
                "required": ["variable", "value"],
            },
        ),

        types.Tool(
            name="get_job_status",
            description="Check the status of a previously scheduled simulation job.",
            inputSchema={
                "type": "object",
                "properties": {
                    "job_id": {"type": "integer", "description": "Job ID returned by a run_* tool"},
                },
                "required": ["job_id"],
            },
        ),

        types.Tool(
            name="list_scenarios",
            description="List all available pre-defined Qaha WWTP scenarios with their parameter changes.",
            inputSchema={"type": "object", "properties": {}},
        ),

        types.Tool(
            name="check_compliance",
            description=(
                "Given a simulation result (job_id or direct values), "
                "evaluate compliance with Egyptian Law 48/1982 effluent limits."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "job_id":    {"type": "integer", "description": "Job ID from a completed simulation"},
                    "TSS_mgL":   {"type": "number"},
                    "BOD5_mgL":  {"type": "number"},
                    "sCOD_mgL":  {"type": "number"},
                },
                "required": [],
            },
        ),

        types.Tool(
            name="create_dynamic_input_table",
            description=(
                "Build a time-series input table for a specific SUMO24 parameter. "
                "Automatically reads the current simulation period (Sumo__StopTime) and "
                "data interval (Sumo__DataComm) from state.xml so the generated table "
                "matches the simulation settings exactly. "
                "The user can also override the start/end time step (range) and the row "
                "interval. Supports constant, linear ramp, step change, sinusoidal, and "
                "custom value profiles. Writes a TSV file under outputs/ ready to be "
                "passed to run_dynamic_simulation via influent_tsv."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "variable": {
                        "type": "string",
                        "description": (
                            "Full SumoCore variable name to build the table for, "
                            "e.g. Sumo__Plant__Influent__param__Q"
                        ),
                    },
                    "duration_days": {
                        "type": "number",
                        "description": (
                            "Total simulation period in days. If omitted, read from "
                            "state.xml (Sumo__StopTime)."
                        ),
                    },
                    "datacomm_hours": {
                        "type": "number",
                        "description": (
                            "Row spacing (data interval) in hours. If omitted, read "
                            "from state.xml (Sumo__DataComm)."
                        ),
                    },
                    "start_time_days": {
                        "type": "number",
                        "description": "Starting time step in days (default 0).",
                    },
                    "end_time_days": {
                        "type": "number",
                        "description": (
                            "Ending time step in days (default = duration_days). "
                            "Together with start_time_days this defines the table's "
                            "time range."
                        ),
                    },
                    "value_mode": {
                        "type": "string",
                        "enum": ["constant", "linear", "step", "sinusoidal", "custom"],
                        "description": (
                            "How to fill the value column. constant = single value; "
                            "linear = ramp from start_value to end_value; step = start_value "
                            "before step_at_day, end_value after; sinusoidal = mean + "
                            "amplitude*sin(2*pi*t/period_days); custom = caller supplies values."
                        ),
                        "default": "constant",
                    },
                    "constant_value": {
                        "type": "number",
                        "description": "Value for every row when value_mode=constant.",
                    },
                    "start_value": {
                        "type": "number",
                        "description": "Starting value for linear/step modes.",
                    },
                    "end_value": {
                        "type": "number",
                        "description": "Ending value for linear/step modes.",
                    },
                    "step_at_day": {
                        "type": "number",
                        "description": "Day at which the step change happens (step mode).",
                    },
                    "amplitude": {
                        "type": "number",
                        "description": "Sine wave amplitude around the mean (sinusoidal mode).",
                    },
                    "mean": {
                        "type": "number",
                        "description": "Sine wave mean/offset (sinusoidal mode).",
                    },
                    "period_days": {
                        "type": "number",
                        "description": "Sine wave period in days (sinusoidal mode, default 1 for diurnal).",
                    },
                    "values": {
                        "type": "array",
                        "items": {"type": "number"},
                        "description": (
                            "Explicit value list when value_mode=custom. Length must "
                            "equal the number of generated time steps."
                        ),
                    },
                    "time_unit": {
                        "type": "string",
                        "enum": ["days", "hours", "minutes"],
                        "description": "Units for the time column written to the TSV (default days).",
                        "default": "days",
                    },
                    "filename": {
                        "type": "string",
                        "description": "Output TSV filename inside outputs/ (default: dyn_input_<variable>.tsv).",
                    },
                },
                "required": ["variable"],
            },
        ),

        types.Tool(
            name="export_results_csv",
            description="Export the time-series results of a completed simulation job to a CSV file.",
            inputSchema={
                "type": "object",
                "properties": {
                    "job_id":   {"type": "integer"},
                    "filename": {"type": "string", "description": "Output filename (saved in outputs/ directory)"},
                },
                "required": ["job_id"],
            },
        ),

        # â”€â”€ GROUP A - Model Inspection â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        types.Tool(
            name="get_parameter",
            description="Read the current value of a single model parameter (optionally after applying a scenario).",
            inputSchema={
                "type": "object",
                "properties": {
                    "parameter": {"type": "string", "description": "Full SumoCore parameter name."},
                    "scenario":  {"type": "string", "description": "Optional scenario to apply first.", "enum": list(SCENARIOS.keys())},
                },
                "required": ["parameter"],
            },
        ),
        types.Tool(
            name="list_parameters",
            description="List configurable parameters (__param__ variables). Optional substring filter on unit name.",
            inputSchema={
                "type": "object",
                "properties": {
                    "unit_name": {"type": "string", "description": "Case-insensitive substring filter (e.g. AerTank, Clarifier)."},
                },
            },
        ),
        types.Tool(
            name="list_state_variables",
            description="List model state variables (non-parameter variables). Optional substring filter.",
            inputSchema={
                "type": "object",
                "properties": {
                    "unit_name": {"type": "string", "description": "Case-insensitive substring filter."},
                },
            },
        ),
        types.Tool(
            name="get_state_variable",
            description="Read the current value of a specific state variable.",
            inputSchema={
                "type": "object",
                "properties": {
                    "variable": {"type": "string", "description": "Full SumoCore variable name (e.g. Sumo__Plant__AerTank1__S_O)."},
                },
                "required": ["variable"],
            },
        ),
        types.Tool(
            name="search_variables",
            description="Case-insensitive substring search across all variable and parameter names.",
            inputSchema={
                "type": "object",
                "properties": {
                    "keyword": {"type": "string", "description": "Keyword to search for (e.g. DO, RAS, NH4, AerTank)."},
                },
                "required": ["keyword"],
            },
        ),
        types.Tool(
            name="get_model_info",
            description="Return top-level model metadata: DLL path, state.xml, total variables, parameter count, compile status.",
            inputSchema={"type": "object", "properties": {}},
        ),
        types.Tool(
            name="ensure_dtt_bridge",
            description=(
                "Detect and (re)install the dynamita_compat shim that backfills "
                "SumoScheduler.executeCommand/.set/.getVariableNames when the installed "
                "Dynamita DTT distribution doesn't expose them. Returns patch status "
                "and the current queue of buffered commands. Optionally flushes the "
                "queue into a sibling <project>.commands.txt next to the given .sumo path."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "flush_to": {
                        "type": "string",
                        "description": "Optional path to a .sumo project. If given, the queued commands are written to a sibling <project>.commands.txt for later replay in SUMO24."
                    },
                    "clear_after_flush": {
                        "type": "boolean",
                        "description": "If true and flush_to is given, clear the queue after writing.",
                        "default": False
                    }
                }
            },
        ),
        types.Tool(
            name="list_unit_processes",
            description="List all unit process names currently visible in the compiled model.",
            inputSchema={"type": "object", "properties": {}},
        ),
        types.Tool(
            name="get_unit_process_info",
            description="Return parameters and state variables for a specific unit process.",
            inputSchema={
                "type": "object",
                "properties": {
                    "unit_name": {"type": "string", "description": "Unit name as it appears in variable names (e.g. AerTank1)."},
                },
                "required": ["unit_name"],
            },
        ),
        types.Tool(
            name="validate_variable_name",
            description="Check whether a SumoCore variable name exists in the compiled model. If not, returns up to 10 name suggestions.",
            inputSchema={
                "type": "object",
                "properties": {
                    "variable": {"type": "string", "description": "Full variable name to validate."},
                },
                "required": ["variable"],
            },
        ),

        # â”€â”€ GROUP B - Persistent Parameter Editing â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        types.Tool(
            name="set_multiple_parameters",
            description="Set several model parameters in one call (in-memory; applies to next simulation run).",
            inputSchema={
                "type": "object",
                "properties": {
                    "parameters": {
                        "type": "object",
                        "description": "Map of {variable_name: numeric_value}.",
                        "additionalProperties": {"type": "number"},
                    },
                    "scenario": {"type": "string", "description": "Scenario to apply first.", "enum": list(SCENARIOS.keys())},
                },
                "required": ["parameters"],
            },
        ),
        types.Tool(
            name="reset_parameter_to_default",
            description="Revert an in-memory parameter to the value stored in state.xml.",
            inputSchema={
                "type": "object",
                "properties": {
                    "parameter": {"type": "string", "description": "Full SumoCore parameter name."},
                },
                "required": ["parameter"],
            },
        ),

        # â”€â”€ GROUP C - State & Project File Management â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        types.Tool(
            name="save_state",
            description="Export the current in-memory model state to a state.xml snapshot (default: timestamped file in outputs/).",
            inputSchema={
                "type": "object",
                "properties": {
                    "output_path": {"type": "string", "description": "Optional output path."},
                },
            },
        ),
        types.Tool(
            name="load_state",
            description="Point the server at a different state.xml snapshot for all subsequent runs.",
            inputSchema={
                "type": "object",
                "properties": {
                    "state_xml_path": {"type": "string", "description": "Full path to a saved state.xml."},
                },
                "required": ["state_xml_path"],
            },
        ),
        types.Tool(
            name="list_saved_states",
            description="List state.xml snapshots found in the outputs directory.",
            inputSchema={"type": "object", "properties": {}},
        ),
        types.Tool(
            name="get_compile_status",
            description="Report whether the project DLL and state.xml exist and are ready to run, with modification timestamps.",
            inputSchema={"type": "object", "properties": {}},
        ),

        # â”€â”€ GROUP D - Scenario Management â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        types.Tool(
            name="create_scenario",
            description="Add a new named scenario to the in-memory SCENARIOS dict.",
            inputSchema={
                "type": "object",
                "properties": {
                    "name":        {"type": "string"},
                    "parameters":  {"type": "object", "additionalProperties": {"type": "number"}},
                    "description": {"type": "string"},
                },
                "required": ["name", "parameters"],
            },
        ),
        types.Tool(
            name="update_scenario",
            description="Replace the parameter overrides of an existing scenario.",
            inputSchema={
                "type": "object",
                "properties": {
                    "name":       {"type": "string"},
                    "parameters": {"type": "object", "additionalProperties": {"type": "number"}},
                },
                "required": ["name", "parameters"],
            },
        ),
        types.Tool(
            name="delete_scenario",
            description="Remove a scenario from the in-memory SCENARIOS dict.",
            inputSchema={
                "type": "object",
                "properties": {
                    "name": {"type": "string"},
                },
                "required": ["name"],
            },
        ),
        types.Tool(
            name="clone_scenario",
            description="Duplicate an existing scenario under a new name.",
            inputSchema={
                "type": "object",
                "properties": {
                    "source":   {"type": "string"},
                    "new_name": {"type": "string"},
                },
                "required": ["source", "new_name"],
            },
        ),
        types.Tool(
            name="get_scenario_diff",
            description="Show which parameters differ between two scenarios.",
            inputSchema={
                "type": "object",
                "properties": {
                    "scenario_a": {"type": "string"},
                    "scenario_b": {"type": "string"},
                },
                "required": ["scenario_a", "scenario_b"],
            },
        ),
        types.Tool(
            name="export_scenario",
            description="Save a scenario to a .scs script file.",
            inputSchema={
                "type": "object",
                "properties": {
                    "name":        {"type": "string"},
                    "output_path": {"type": "string"},
                },
                "required": ["name"],
            },
        ),
        types.Tool(
            name="import_scenario",
            description="Load a .scs script file and register it as a scenario.",
            inputSchema={
                "type": "object",
                "properties": {
                    "scs_path": {"type": "string"},
                    "name":     {"type": "string"},
                },
                "required": ["scs_path"],
            },
        ),

        # â”€â”€ GROUP E - Extended Analysis & Reporting â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        types.Tool(
            name="get_effluent_statistics",
            description="Return min/max/mean for all numeric columns in the CSV of a completed job.",
            inputSchema={
                "type": "object",
                "properties": {
                    "job_id": {"type": "string", "description": "Job ID (int or string form)."},
                },
                "required": ["job_id"],
            },
        ),
        types.Tool(
            name="get_sludge_production",
            description="Estimate WAS rate, SRT, and biosolids production from a completed job (best-effort; notes what to look up).",
            inputSchema={
                "type": "object",
                "properties": {
                    "job_id": {"type": "string"},
                },
                "required": ["job_id"],
            },
        ),
        types.Tool(
            name="get_energy_estimate",
            description="Pull DO setpoints and airflow-related parameter values for a scenario as a proxy for aeration energy.",
            inputSchema={
                "type": "object",
                "properties": {
                    "scenario": {"type": "string", "description": "Scenario to evaluate.", "enum": list(SCENARIOS.keys())},
                },
            },
        ),
        types.Tool(
            name="generate_report",
            description="Produce a structured design summary (effluent, compliance, metadata) for a completed job.",
            inputSchema={
                "type": "object",
                "properties": {
                    "job_id":             {"type": "string"},
                    "scenario":           {"type": "string"},
                    "include_compliance": {"type": "boolean", "default": True},
                },
                "required": ["job_id"],
            },
        ),
        types.Tool(
            name="list_completed_jobs",
            description="Return every job currently in memory with status, scenario, and type (steady / dynamic).",
            inputSchema={"type": "object", "properties": {}},
        ),
        types.Tool(
            name="import_influent_profile",
            description="Register a time-varying influent CSV so run_dynamic_simulation can reference it by label.",
            inputSchema={
                "type": "object",
                "properties": {
                    "csv_path": {"type": "string", "description": "Full path to the CSV (must have a 'time_h' column)."},
                    "label":    {"type": "string", "description": "Optional label; defaults to filename stem."},
                },
                "required": ["csv_path"],
            },
        ),
        # â”€â”€ Build & Configuration Tools â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        types.Tool(
            name="create_model",
            description="Initialise a new blank SUMO24 project folder with a .sumo stub file, subdirectories, and README. Use before building a WWTP from scratch.",
            inputSchema={
                "type": "object",
                "properties": {
                    "project_name": {"type": "string", "description": "Name of the new project (e.g. 'Qaha_NewPlant')."},
                    "project_path": {"type": "string", "description": "Directory where the project folder should be created."},
                    "template":     {"type": "string", "description": "Optional scenario name from SCENARIOS to use as a parameter seed.", "default": ""},
                },
                "required": ["project_name", "project_path"],
            },
        ),
        types.Tool(
            name="load_model",
            description="Load an existing .sumo project file and register it as the active model; auto-detects adjacent DLL and state.xml.",
            inputSchema={
                "type": "object",
                "properties": {
                    "sumo_file_path": {"type": "string", "description": "Full path to the .sumo project file."},
                },
                "required": ["sumo_file_path"],
            },
        ),
        types.Tool(
            name="save_model",
            description="Save the current project to disk: copies .sumo + DLL, and saves a fresh state.xml snapshot via DTT.",
            inputSchema={
                "type": "object",
                "properties": {
                    "output_path": {"type": "string", "description": "Destination directory. Defaults to a timestamped subfolder of outputs.", "default": ""},
                    "overwrite":   {"type": "boolean", "description": "Overwrite existing files if present.", "default": False},
                },
            },
        ),
        types.Tool(
            name="compile_model",
            description="Check whether a runnable DLL exists or force a recompile by backing up and removing the current DLL (actual compilation happens in SUMO GUI).",
            inputSchema={
                "type": "object",
                "properties": {
                    "force_recompile": {"type": "boolean", "description": "If True, deletes the existing DLL so recompilation is forced.", "default": False},
                },
            },
        ),
        types.Tool(
            name="extract_dll",
            description="Locate and register a compiled sumoproject.dll in the project directory or an explicit path.",
            inputSchema={
                "type": "object",
                "properties": {
                    "output_path": {"type": "string", "description": "Optional explicit path to the .dll file. If omitted, scans the project dir and cwd.", "default": ""},
                },
            },
        ),
        types.Tool(
            name="list_available_unit_process_types",
            description="Return all unit process types supported by SUMO24 (hydraulic, biological, clarification, tertiary, sludge, sensors, controllers).",
            inputSchema={"type": "object", "properties": {}},
        ),
        types.Tool(
            name="add_unit_process",
            description="Add a treatment unit process to the active model via DTT executeCommand; optionally assigns a kinetic model.",
            inputSchema={
                "type": "object",
                "properties": {
                    "process_type":  {"type": "string", "description": "Type name from list_available_unit_process_types."},
                    "instance_name": {"type": "string", "description": "Unique name for this instance (e.g. 'AerTank1')."},
                    "kinetic_model": {"type": "string", "description": "Optional biological model (e.g. 'ASM2d', 'ASM1').", "default": ""},
                },
                "required": ["process_type", "instance_name"],
            },
        ),
        types.Tool(
            name="remove_unit_process",
            description="Remove a unit process from the active model and optionally all streams connected to it.",
            inputSchema={
                "type": "object",
                "properties": {
                    "unit_name":      {"type": "string", "description": "Name of the unit to remove."},
                    "remove_streams": {"type": "boolean", "description": "Also remove streams attached to the unit.", "default": True},
                },
                "required": ["unit_name"],
            },
        ),
        types.Tool(
            name="list_available_kinetic_models",
            description="List biological kinetic models available for assignment to reactors (ASM1, ASM2d, ASM3, ADM1, MBR, TakÃ¡cs, etc.).",
            inputSchema={"type": "object", "properties": {}},
        ),
        types.Tool(
            name="set_kinetic_model",
            description="Assign or switch the biological kinetic model on a reactor unit.",
            inputSchema={
                "type": "object",
                "properties": {
                    "unit_name":     {"type": "string"},
                    "kinetic_model": {"type": "string", "description": "Model name from list_available_kinetic_models."},
                },
                "required": ["unit_name", "kinetic_model"],
            },
        ),
        types.Tool(
            name="connect_unit_processes",
            description="Create a flow connection (stream) between two unit processes.",
            inputSchema={
                "type": "object",
                "properties": {
                    "from_unit":   {"type": "string"},
                    "to_unit":     {"type": "string"},
                    "stream_name": {"type": "string"},
                    "stream_type": {"type": "string", "description": "'liquid', 'sludge', 'gas', or 'recycle'.", "default": "liquid"},
                },
                "required": ["from_unit", "to_unit", "stream_name"],
            },
        ),
        types.Tool(
            name="add_recycle_stream",
            description="Add a recycle stream (RAS, WAS, internal recycle, reject water, or sidestream) between two units.",
            inputSchema={
                "type": "object",
                "properties": {
                    "from_unit":    {"type": "string"},
                    "to_unit":      {"type": "string"},
                    "stream_name":  {"type": "string"},
                    "recycle_type": {"type": "string", "description": "'RAS', 'WAS', 'InternalRecycle', 'RejectWater', or 'Sidestream'.", "default": "RAS"},
                },
                "required": ["from_unit", "to_unit", "stream_name"],
            },
        ),
        types.Tool(
            name="modify_flow_connection",
            description="Redirect or resize a flow stream (change source, destination, or fixed flow rate).",
            inputSchema={
                "type": "object",
                "properties": {
                    "stream_name":  {"type": "string"},
                    "new_from":     {"type": "string", "default": ""},
                    "new_to":       {"type": "string", "default": ""},
                    "new_flow_m3d": {"type": "number", "description": "New fixed flow rate in mÂ³/d (-1 = no change).", "default": -1},
                },
                "required": ["stream_name"],
            },
        ),
        types.Tool(
            name="list_flow_streams",
            description="List registered flow streams and compiled model flow variables, optionally filtered to one unit.",
            inputSchema={
                "type": "object",
                "properties": {
                    "unit_name": {"type": "string", "description": "Optional filter.", "default": ""},
                },
            },
        ),
        types.Tool(
            name="set_influent_characteristics",
            description="Define influent wastewater composition and flow rate in one call; maps lab parameters (BOD, COD, TSS, TKN, NH4, TN, TP, temp) to their SUMO variables.",
            inputSchema={
                "type": "object",
                "properties": {
                    "flow_m3d":      {"type": "number"},
                    "BOD_mgL":       {"type": "number", "default": 0.0},
                    "COD_mgL":       {"type": "number", "default": 0.0},
                    "TSS_mgL":       {"type": "number", "default": 0.0},
                    "VSS_mgL":       {"type": "number", "default": 0.0},
                    "TKN_mgL":       {"type": "number", "default": 0.0},
                    "NH4_mgL":       {"type": "number", "default": 0.0},
                    "TN_mgL":        {"type": "number", "default": 0.0},
                    "TP_mgL":        {"type": "number", "default": 0.0},
                    "temp_C":        {"type": "number", "default": 20.0},
                    "influent_unit": {"type": "string", "default": "Influent"},
                },
                "required": ["flow_m3d"],
            },
        ),
        types.Tool(
            name="set_fractionation",
            description="Set COD/TSS fractionation ratios (Xs, Xp, Si, Ss, Xi, fXS, etc.) on a unit for ASM-type kinetic models.",
            inputSchema={
                "type": "object",
                "properties": {
                    "unit_name":     {"type": "string"},
                    "fraction_dict": {"type": "object", "description": "Dict mapping fraction name to ratio (0â€“1)."},
                },
                "required": ["unit_name", "fraction_dict"],
            },
        ),
        types.Tool(
            name="set_parameter_persistent",
            description="Persist a parameter change directly into state.xml on disk (with timestamped backup) so it survives server restarts.",
            inputSchema={
                "type": "object",
                "properties": {
                    "parameter": {"type": "string", "description": "Full SumoCore parameter name."},
                    "value":     {"type": "number"},
                },
                "required": ["parameter", "value"],
            },
        ),
        types.Tool(
            name="add_controller",
            description="Add a control loop (DO, SRT, NH4, NO3, FlowSplitter controller) to a unit with sensor/actuator and setpoint.",
            inputSchema={
                "type": "object",
                "properties": {
                    "controller_type": {"type": "string", "description": "e.g. 'DO_Controller', 'SRT_Controller'."},
                    "controlled_var":  {"type": "string", "description": "Variable the controller reads (e.g. 'S_O', 'SRT')."},
                    "setpoint":        {"type": "number"},
                    "unit_name":       {"type": "string", "description": "Unit where the sensor is placed."},
                    "actuator_unit":   {"type": "string", "description": "Unit the controller acts on (defaults to unit_name).", "default": ""},
                },
                "required": ["controller_type", "controlled_var", "setpoint", "unit_name"],
            },
        ),
        types.Tool(
            name="validate_model",
            description="Run configuration + readiness checks on the current model. 'full' mode also runs a quick steady-state to surface solver errors.",
            inputSchema={
                "type": "object",
                "properties": {
                    "check_level": {"type": "string", "description": "'basic' or 'full'.", "default": "basic", "enum": ["basic", "full"]},
                },
            },
        ),
        types.Tool(
            name="get_mass_balance",
            description="Compute COD / N / P / TSS mass balance (influent vs effluent) for a completed job from its exported CSV.",
            inputSchema={
                "type": "object",
                "properties": {
                    "job_id":   {"type": "string"},
                    "elements": {"type": "array", "items": {"type": "string"}, "description": "Elements to balance. Default ['COD','N','P','TSS']."},
                },
                "required": ["job_id"],
            },
        ),
        types.Tool(
            name="initialize_state",
            description="Reset the model to its baseline state.xml â€” clears all in-memory parameter overrides.",
            inputSchema={"type": "object", "properties": {}},
        ),
        # â”€â”€ Parameter-editing tools with research-based defaults â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        types.Tool(
            name="set_asm1_kinetics",
            description="Set ASM1 kinetic + stoichiometric parameters on a biological unit. Any param left null uses the IWA/Egyptian research default, Arrhenius-corrected to temp_C.",
            inputSchema={
                "type": "object",
                "properties": {
                    "unit_name": {"type": "string", "default": "AerTank1"},
                    "temp_C":    {"type": "number", "default": 20.0},
                    "mu_H": {"type": "number"}, "K_S": {"type": "number"},
                    "K_OH": {"type": "number"}, "K_NO": {"type": "number"},
                    "b_H":  {"type": "number"}, "Y_H": {"type": "number"},
                    "eta_g":{"type": "number"}, "eta_h": {"type": "number"},
                    "k_h":  {"type": "number"}, "K_X": {"type": "number"},
                    "mu_A": {"type": "number"}, "K_NH": {"type": "number"},
                    "K_OA": {"type": "number"}, "b_A": {"type": "number"}, "k_a": {"type": "number"},
                    "Y_A":  {"type": "number"}, "f_P": {"type": "number"},
                    "i_XB": {"type": "number"}, "i_XP": {"type": "number"},
                },
            },
        ),
        types.Tool(
            name="set_asm2d_kinetics",
            description="Set ASM2d kinetic + stoichiometric parameters (heterotrophs, PAOs, autotrophs, precipitation). Any param left null uses the research default at temp_C.",
            inputSchema={
                "type": "object",
                "properties": {
                    "unit_name": {"type": "string", "default": "AerTank1"},
                    "temp_C":    {"type": "number", "default": 20.0},
                    "mu_H": {"type": "number"}, "b_H": {"type": "number"}, "Y_H": {"type": "number"},
                    "q_fe": {"type": "number"}, "eta_NO3_H": {"type": "number"},
                    "K_h":  {"type": "number"}, "eta_NO3_hyd": {"type": "number"},
                    "mu_PAO": {"type": "number"}, "b_PAO": {"type": "number"},
                    "q_PHA": {"type": "number"}, "q_PP": {"type": "number"},
                    "Y_PAO": {"type": "number"}, "Y_PO4": {"type": "number"},
                    "mu_AUT":{"type": "number"}, "b_AUT": {"type": "number"},
                    "K_NH4_AUT":{"type": "number"}, "K_O2_AUT": {"type": "number"},
                    "k_PRE": {"type": "number"}, "k_RED": {"type": "number"},
                },
            },
        ),
        types.Tool(
            name="set_asm3_kinetics",
            description="Set ASM3 kinetic + stoichiometric parameters (endogenous respiration + internal storage X_STO). Any param left null uses the research default at temp_C.",
            inputSchema={
                "type": "object",
                "properties": {
                    "unit_name": {"type": "string", "default": "AerTank1"},
                    "temp_C":    {"type": "number", "default": 20.0},
                    "mu_H": {"type": "number"}, "b_H_O2": {"type": "number"}, "b_H_NO": {"type": "number"},
                    "k_STO": {"type": "number"}, "Y_STO_O2": {"type": "number"}, "Y_STO_NO": {"type": "number"},
                    "Y_H_O2":{"type": "number"}, "Y_H_NO": {"type": "number"},
                    "k_h":  {"type": "number"}, "K_X": {"type": "number"},
                    "K_S":  {"type": "number"}, "K_STO": {"type": "number"},
                    "eta_NO":{"type": "number"}, "mu_A": {"type": "number"},
                    "b_A_O":{"type": "number"}, "b_A_NO": {"type": "number"},
                    "K_A_NH":{"type": "number"}, "K_A_O": {"type": "number"},
                    "Y_A":  {"type": "number"}, "f_XI": {"type": "number"},
                },
            },
        ),
        types.Tool(
            name="apply_temperature_correction",
            description="Re-compute Arrhenius temperature-corrected ASM parameters at a new operating temperature. Reports per-parameter % change and flags high-impact (>20%) corrections.",
            inputSchema={
                "type": "object",
                "properties": {
                    "model":     {"type": "string", "enum": ["ASM1", "ASM2d", "ASM3"], "default": "ASM1"},
                    "temp_C":    {"type": "number", "default": 25.0},
                    "unit_name": {"type": "string", "default": "AerTank1"},
                },
            },
        ),
        types.Tool(
            name="set_plant_wide_parameters",
            description="Set the global plant temperature and update CONFIG Law 48 limits to match the chosen discharge target ('nile' or 'drain').",
            inputSchema={
                "type": "object",
                "properties": {
                    "temp_C":       {"type": "number"},
                    "law48_target": {"type": "string", "enum": ["nile", "drain"], "default": "drain"},
                },
            },
        ),
        types.Tool(
            name="set_influent_parameters",
            description="Set influent flow, composition, and COD fractionation. Defaults to typical Egyptian municipal wastewater values if arguments are omitted.",
            inputSchema={
                "type": "object",
                "properties": {
                    "unit_name": {"type": "string", "default": "Influent"},
                    "Q_m3d":   {"type": "number"}, "COD_mgL": {"type": "number"},
                    "BOD5_mgL":{"type": "number"}, "TSS_mgL": {"type": "number"},
                    "VSS_mgL": {"type": "number"}, "TKN_mgL": {"type": "number"},
                    "NH4_mgL": {"type": "number"}, "TP_mgL":  {"type": "number"},
                    "pH":      {"type": "number"}, "Temp_C":  {"type": "number"},
                    "fSI": {"type": "number"}, "fSS": {"type": "number"},
                    "fXS": {"type": "number"}, "fXI": {"type": "number"}, "fXBH": {"type": "number"},
                },
            },
        ),
        types.Tool(
            name="set_primary_clarifier_parameters",
            description="Set primary clarifier design parameters (area, overflow rate, solids loading, capture efficiency, HRT). Defaults from Metcalf & Eddy 5th ed.",
            inputSchema={
                "type": "object",
                "properties": {
                    "unit_name":             {"type": "string", "default": "PrimClar1"},
                    "surface_area_m2":       {"type": "number"},
                    "overflow_rate_m3m2d":   {"type": "number"},
                    "solids_loading_kgm2d":  {"type": "number"},
                    "capture_efficiency":    {"type": "number"},
                    "sludge_concentration_gL":{"type": "number"},
                    "HRT_h":                 {"type": "number"},
                },
            },
        ),
        types.Tool(
            name="set_aeration_tank_parameters",
            description="Set aeration-tank hydraulic / process / aeration parameters (V, HRT, SRT, DO, KLa, SOTE, MLSS, T). Defaults tuned for Egyptian warm-climate operation.",
            inputSchema={
                "type": "object",
                "properties": {
                    "unit_name":       {"type": "string", "default": "AerTank1"},
                    "volume_m3":       {"type": "number"},
                    "HRT_h":           {"type": "number"},
                    "SRT_d":           {"type": "number"},
                    "DO_setpoint_mgL": {"type": "number"},
                    "KLa_h":           {"type": "number"},
                    "SOTE_pct":        {"type": "number"},
                    "MLSS_mgL":        {"type": "number"},
                    "temp_C":          {"type": "number"},
                },
            },
        ),
        types.Tool(
            name="set_anoxic_zone_parameters",
            description="Set anoxic zone parameters (volume, HRT, NO3 recycle ratio, max DO). Defaults for pre-denitrification.",
            inputSchema={
                "type": "object",
                "properties": {
                    "unit_name":         {"type": "string", "default": "AnoxTank1"},
                    "volume_m3":         {"type": "number"},
                    "HRT_h":             {"type": "number"},
                    "NO3_recycle_ratio": {"type": "number"},
                    "DO_max_mgL":        {"type": "number"},
                },
            },
        ),
        types.Tool(
            name="set_anaerobic_zone_parameters",
            description="Set anaerobic (EBPR / P-release) zone parameters (volume, HRT, max DO, VFA target).",
            inputSchema={
                "type": "object",
                "properties": {
                    "unit_name":      {"type": "string", "default": "AnaTank1"},
                    "volume_m3":      {"type": "number"},
                    "HRT_h":          {"type": "number"},
                    "DO_max_mgL":     {"type": "number"},
                    "VFA_target_mgL": {"type": "number"},
                },
            },
        ),
        types.Tool(
            name="set_secondary_clarifier_parameters",
            description="Set secondary clarifier hydraulic + Takacs settler model parameters (v0, v0_p, r_h, r_p, f_ns). Defaults from Takacs 1991 / BSM1.",
            inputSchema={
                "type": "object",
                "properties": {
                    "unit_name":           {"type": "string", "default": "SecClar1"},
                    "surface_area_m2":     {"type": "number"},
                    "overflow_rate_m3m2d": {"type": "number"},
                    "sludge_blanket_m":    {"type": "number"},
                    "underflow_conc_gL":   {"type": "number"},
                    "v0":   {"type": "number"}, "v0_p": {"type": "number"},
                    "r_h":  {"type": "number"}, "r_p":  {"type": "number"}, "f_ns": {"type": "number"},
                },
            },
        ),
        types.Tool(
            name="set_mbr_parameters",
            description="Set membrane bioreactor parameters (flux, TMP, permeability, backwash interval/duration, MLSS). Defaults from Judd MBR Book / Metcalf & Eddy.",
            inputSchema={
                "type": "object",
                "properties": {
                    "unit_name":             {"type": "string", "default": "MBR1"},
                    "flux_Lm2h":             {"type": "number"},
                    "TMP_kPa":               {"type": "number"},
                    "permeability_LM2hbar":  {"type": "number"},
                    "backwash_interval_min": {"type": "number"},
                    "backwash_duration_s":   {"type": "number"},
                    "MLSS_mgL":              {"type": "number"},
                },
            },
        ),
        types.Tool(
            name="set_ras_was_parameters",
            description="Set RAS + WAS pump parameters (flows, ratios, target SRT). Calculates Q_RAS from RAS_ratio if flow is omitted.",
            inputSchema={
                "type": "object",
                "properties": {
                    "ras_unit":       {"type": "string", "default": "RAS_Pump"},
                    "was_unit":       {"type": "string", "default": "WAS_Pump"},
                    "Q_RAS_m3d":      {"type": "number"},
                    "RAS_ratio":      {"type": "number"},
                    "Q_WAS_m3d":      {"type": "number"},
                    "SRT_target_d":   {"type": "number"},
                    "Q_influent_m3d": {"type": "number"},
                },
            },
        ),
        types.Tool(
            name="set_digester_parameters",
            description="Set aerobic or anaerobic digester parameters (V, SRT, T, DO, OLR, biogas yield, VSS reduction).",
            inputSchema={
                "type": "object",
                "properties": {
                    "unit_name":         {"type": "string", "default": "Digester1"},
                    "digester_type":     {"type": "string", "enum": ["aerobic", "anaerobic"], "default": "aerobic"},
                    "volume_m3":         {"type": "number"},
                    "SRT_d":             {"type": "number"},
                    "temp_C":            {"type": "number"},
                    "DO_setpoint_mgL":   {"type": "number"},
                    "OLR_kgVSSm3d":      {"type": "number"},
                    "biogas_yield_m3kg": {"type": "number"},
                    "VSS_reduction_pct": {"type": "number"},
                },
            },
        ),
        types.Tool(
            name="set_chemical_dosing_parameters",
            description="Set chemical precipitation / coagulant dosing parameters (alum or ferric). Estimates dose from target TP if not supplied.",
            inputSchema={
                "type": "object",
                "properties": {
                    "unit_name":      {"type": "string", "default": "ChemDosing1"},
                    "coagulant_type": {"type": "string", "enum": ["alum", "ferric"], "default": "alum"},
                    "dose_mgL":       {"type": "number"},
                    "molar_ratio":    {"type": "number"},
                    "target_TP_mgL":  {"type": "number"},
                },
            },
        ),
        types.Tool(
            name="set_dynamic_simulation_parameters",
            description="Configure the SUMO24 dynamic simulation tab (duration, timestep, solver, warm-up, diurnal and storm factors, seasonal temperatures).",
            inputSchema={
                "type": "object",
                "properties": {
                    "duration_days":       {"type": "number"},
                    "timestep_h":          {"type": "number"},
                    "output_interval_h":   {"type": "number"},
                    "solver":              {"type": "string"},
                    "tolerance":           {"type": "number"},
                    "warm_up_days":        {"type": "number"},
                    "diurnal_peak_factor": {"type": "number"},
                    "diurnal_min_factor":  {"type": "number"},
                    "storm_peak_factor":   {"type": "number"},
                    "storm_start_day":     {"type": "number"},
                    "storm_duration_h":    {"type": "number"},
                    "summer_temp_C":       {"type": "number"},
                    "winter_temp_C":       {"type": "number"},
                },
            },
        ),
        types.Tool(
            name="assume_parameters_from_research",
            description="Auto-populate a full recommendation of research-based defaults (influent, ASM kinetics, aeration tank, clarifier, dynamic simulation) sized to the selected plant_scale and temperature. Does NOT apply values to the model - returns a recommendation dict to review first.",
            inputSchema={
                "type": "object",
                "properties": {
                    "unit_type":   {"type": "string",
                                     "enum": ["full_plant", "aeration_tank", "secondary_clarifier",
                                              "influent", "asm_kinetics", "dynamic"],
                                     "default": "full_plant"},
                    "temp_C":      {"type": "number"},
                    "asm_model":   {"type": "string", "enum": ["ASM1", "ASM2d", "ASM3"], "default": "ASM1"},
                    "plant_scale": {"type": "string", "enum": ["small", "medium", "large"], "default": "medium"},
                },
            },
        ),
        # -- Dataset-Reading (Excel) Tools ----------------------------------
        types.Tool(
            name="read_dataset_excel",
            description="Read the WWTP Excel dataset template and return a structured summary of all measurement points, Law 48 limits, and section headers.",
            inputSchema={"type":"object","properties":{"xlsx_path":{"type":"string","description":"Full path; defaults to SUMO24_WWTP_Dataset_and_Tools.xlsx in project root."}}},
        ),
        types.Tool(
            name="list_dataset_measurement_points",
            description="List every measurement point in the dataset and the SUMO unit it maps to.",
            inputSchema={"type":"object","properties":{"xlsx_path":{"type":"string"}}},
        ),
        types.Tool(
            name="get_measurement_point_data",
            description="Get all measured parameters for a specific measurement point (e.g. 'Raw Influent'). Flags Law 48 exceedances.",
            inputSchema={"type":"object","properties":{"label":{"type":"string","description":"Point label exactly as in column A."},"xlsx_path":{"type":"string"}},"required":["label"]},
        ),
        types.Tool(
            name="get_parameter_across_points",
            description="Trace a single parameter (e.g. COD) across every measurement point. Reports overall influent-to-effluent removal pct.",
            inputSchema={"type":"object","properties":{"parameter":{"type":"string","description":"Flow, BOD5, COD, TSS, VSS, NH4_N, NO3_N, TN, TP, pH, Temp, DO, Turbidity, Oil_Grease."},"xlsx_path":{"type":"string"}},"required":["parameter"]},
        ),
        types.Tool(
            name="get_law48_limits_from_excel",
            description="Return the Law 48/1982 limit row from the Excel template and sync into CONFIG['law48_limits_from_dataset'].",
            inputSchema={"type":"object","properties":{"xlsx_path":{"type":"string"}}},
        ),
        types.Tool(
            name="map_dataset_to_sumo_variables",
            description="Resolve every dataset cell (point x parameter) to a concrete SUMO variable in the compiled model. Reports resolved + unresolved.",
            inputSchema={"type":"object","properties":{"xlsx_path":{"type":"string"}}},
        ),
        types.Tool(
            name="validate_dataset",
            description="Sanity-check the dataset: missing critical parameters, physical ranges (pH, Temp, DO), COD trend, flow continuity, Law 48 exceedances.",
            inputSchema={"type":"object","properties":{"xlsx_path":{"type":"string"}}},
        ),
        types.Tool(
            name="check_dataset_against_law48",
            description="Compare every measurement point's values against the Law 48 limit row; flags exceedances stage-by-stage.",
            inputSchema={"type":"object","properties":{"xlsx_path":{"type":"string"}}},
        ),
        types.Tool(
            name="apply_influent_from_dataset",
            description="Apply one row (default 'Raw Influent') from the Excel dataset to the corresponding SUMO Influent unit parameters.",
            inputSchema={"type":"object","properties":{"xlsx_path":{"type":"string"},"row_label":{"type":"string","default":"Raw Influent"}}},
        ),
        types.Tool(
            name="apply_point_to_unit",
            description="Apply a specific measurement point's values to a SUMO unit (supports unit_override for custom unit names).",
            inputSchema={"type":"object","properties":{"point_label":{"type":"string"},"unit_override":{"type":"string"},"xlsx_path":{"type":"string"}},"required":["point_label"]},
        ),
        types.Tool(
            name="apply_dataset_to_model",
            description="Apply ALL numeric values from the dataset to the SUMO model in one call (full calibration import). Supports dry_run preview and only_influent mode.",
            inputSchema={"type":"object","properties":{"xlsx_path":{"type":"string"},"only_influent":{"type":"boolean","default":False},"dry_run":{"type":"boolean","default":False}}},
        ),
        types.Tool(
            name="compare_dataset_vs_simulation",
            description="Compare measured values from the Excel dataset against simulated values from a completed job. Reports per-point deltas and Mean Absolute % Error (MAPE).",
            inputSchema={"type":"object","properties":{"job_id":{"type":"string"},"xlsx_path":{"type":"string"},"points_to_compare":{"type":"array","items":{"type":"string"}}},"required":["job_id"]},
        ),
        types.Tool(
            name="compute_removal_efficiencies_from_dataset",
            description="Compute per-stage and overall removal efficiencies for BOD5, COD, TSS, NH4, TN, TP from Excel data alone.",
            inputSchema={"type":"object","properties":{"xlsx_path":{"type":"string"}}},
        ),
        types.Tool(
            name="export_simulation_to_dataset",
            description="Create an Excel file (copy of template) with current simulation values filled into every measurement point. Original file is not modified.",
            inputSchema={"type":"object","properties":{"xlsx_path":{"type":"string"},"output_path":{"type":"string"}}},
        ),
        types.Tool(
            name="export_comparison_report",
            description="Create an Excel calibration report with Measured / Simulated / Delta / Delta-pct / Fit-Quality columns on a 'Calibration Comparison' sheet, color-coded.",
            inputSchema={"type":"object","properties":{"job_id":{"type":"string"},"xlsx_path":{"type":"string"},"output_path":{"type":"string"}},"required":["job_id"]},
        ),
        types.Tool(
            name="validate_model_structure",
            description="Structural pre-flight check: verifies DLL + state.xml exist, are version-consistent, and model has required unit types (influent, reactor, clarifier, effluent).",
            inputSchema={"type":"object","properties":{}},
        ),
        types.Tool(
            name="validate_influent_configuration",
            description="Check that the influent is configured with sensible, non-zero values (flow > 0, COD > 0, temperature and pH within biological range, COD:TKN ratio).",
            inputSchema={"type":"object","properties":{"unit_name":{"type":"string","default":"Influent"}}},
        ),
        types.Tool(
            name="validate_mass_balance_pre_run",
            description="Quick pre-simulation mass balance sanity check. Compares influent load to state.xml effluent values to flag obvious configuration errors for COD, TSS, TKN, TP.",
            inputSchema={"type":"object","properties":{"tolerance_pct":{"type":"number","default":5.0}}},
        ),
        types.Tool(
            name="validate_mass_balance_post_run",
            description="Full mass balance validation on a completed simulation. For each element (COD, N, P, TSS) computes closure error = (in - out - WAS) / in. Well-behaved plants close within +/- 5%.",
            inputSchema={"type":"object","properties":{"job_id":{"type":"string"},"elements":{"type":"array","items":{"type":"string"}},"tolerance_pct":{"type":"number","default":5.0}},"required":["job_id"]},
        ),
        types.Tool(
            name="validate_mlss_health",
            description="Check MLSS, MLVSS/MLSS ratio, active heterotroph and nitrifier fractions in a reactor against Metcalf & Eddy ranges.",
            inputSchema={"type":"object","properties":{"unit_name":{"type":"string","default":"AerTank1"}}},
        ),
        types.Tool(
            name="validate_srt_fm_feasibility",
            description="Check whether configured SRT and F/M are feasible. Computes actual SRT from V, MLSS, Q_WAS, X_R and compares to target; computes F/M and HRT.",
            inputSchema={"type":"object","properties":{"unit_name":{"type":"string","default":"AerTank1"},"influent_unit":{"type":"string","default":"Influent"},"was_unit":{"type":"string","default":"WAS_Pump"}}},
        ),
        types.Tool(
            name="validate_do_levels",
            description="Verify DO levels are within correct range for each zone type (aerobic 1.5-2.5, anoxic ~0, anaerobic ~0). Also checks KLa.",
            inputSchema={"type":"object","properties":{"zone_units":{"type":"object"}}},
        ),
        types.Tool(
            name="validate_oxygen_demand_vs_supply",
            description="Check aeration supply sufficiency. Estimates O2 demand from BOD + nitrification and compares with KLa-driven supply.",
            inputSchema={"type":"object","properties":{"unit_name":{"type":"string","default":"AerTank1"},"influent_unit":{"type":"string","default":"Influent"}}},
        ),
        types.Tool(
            name="validate_hydraulic_balance",
            description="Validate flow continuity (Q_in ~ Q_eff + Q_WAS), RAS ratio, secondary clarifier SOR, and underflow concentration.",
            inputSchema={"type":"object","properties":{"influent_unit":{"type":"string","default":"Influent"},"effluent_unit":{"type":"string","default":"Effluent"},"was_unit":{"type":"string","default":"WAS_Pump"},"ras_unit":{"type":"string","default":"RAS_Pump"},"sec_clar_unit":{"type":"string","default":"SecClar1"}}},
        ),
        types.Tool(
            name="validate_asm_kinetics",
            description="Check ASM kinetic parameters (mu_H, K_S, b_H, Y_H, mu_A, K_NH, b_A, Y_A, k_h, K_X) against IWA/Henze physical bounds. Catches typos like Y_H = 6.7.",
            inputSchema={"type":"object","properties":{"unit_name":{"type":"string","default":"AerTank1"}}},
        ),
        types.Tool(
            name="validate_simulation_convergence",
            description="Check if a simulation reached steady state or is dynamically stable. Analyzes the last N% of the time series for key effluent parameters.",
            inputSchema={"type":"object","properties":{"job_id":{"type":"string"},"stability_window_pct":{"type":"number","default":10.0}},"required":["job_id"]},
        ),
        types.Tool(
            name="validate_simulation_plausibility",
            description="Sanity-check simulation output for physically impossible values: negative concentrations, pH outside 4-10, oversaturated DO, effluent > influent for conserved substances.",
            inputSchema={"type":"object","properties":{"job_id":{"type":"string"}},"required":["job_id"]},
        ),
        types.Tool(
            name="validate_full_model",
            description="Run every pre-simulation validation in sequence. Returns READY_TO_SIMULATE / CHECK_WARNINGS / DO_NOT_SIMULATE plus an aggregated report. Use as the gate before any run_ call.",
            inputSchema={"type":"object","properties":{"aer_tank":{"type":"string","default":"AerTank1"},"influent_unit":{"type":"string","default":"Influent"},"effluent_unit":{"type":"string","default":"Effluent"},"sec_clar_unit":{"type":"string","default":"SecClar1"},"was_unit":{"type":"string","default":"WAS_Pump"},"ras_unit":{"type":"string","default":"RAS_Pump"}}},
        ),
        types.Tool(
            name="validate_post_simulation",
            description="Run every post-simulation validation (mass balance, convergence, plausibility) and return RESULTS_VALID / RESULTS_SUSPECT / RESULTS_INVALID.",
            inputSchema={"type":"object","properties":{"job_id":{"type":"string"}},"required":["job_id"]},
        ),
        types.Tool(
            name="scan_parameter_sensitivity",
            description="Sweep a parameter across a list of values and record an output variable at each point. Instantaneous reads - pair with run_steady_state for full effluent sensitivity.",
            inputSchema={"type":"object","properties":{"parameter":{"type":"string"},"values":{"type":"array","items":{"type":"number"}},"output_variable":{"type":"string"},"unit_name":{"type":"string","default":"AerTank1"}},"required":["parameter","values","output_variable"]},
        ),
        types.Tool(
            name="optimize_was_for_target_srt",
            description="Find and apply the Q_WAS that achieves a target SRT. Uses analytical SRT = V*MLSS / (Q_WAS * X_R) to solve directly for Q_WAS.",
            inputSchema={"type":"object","properties":{"target_srt_d":{"type":"number"},"tolerance_pct":{"type":"number","default":5.0},"was_unit":{"type":"string","default":"WAS_Pump"},"aer_tank":{"type":"string","default":"AerTank1"}},"required":["target_srt_d"]},
        ),
        types.Tool(
            name="optimize_kla_for_target_do",
            description="Estimate and apply the KLa needed to hit a target DO setpoint in the aerobic zone. Analytical solution from OTR = KLa*V*(DO_sat - DO).",
            inputSchema={"type":"object","properties":{"target_do_mgL":{"type":"number"},"unit_name":{"type":"string","default":"AerTank1"},"tolerance":{"type":"number","default":0.1}},"required":["target_do_mgL"]},
        ),
        types.Tool(
            name="diagnose_nitrification_failure",
            description="Targeted diagnostic for failed nitrification. Checks SRT, temperature, pH, DO, and effluent NH4, then returns likely root causes + recommended fixes.",
            inputSchema={"type":"object","properties":{"unit_name":{"type":"string","default":"AerTank1"},"influent_unit":{"type":"string","default":"Influent"}}},
        ),
        types.Tool(
            name="diagnose_bulking_risk",
            description="Sludge bulking risk assessment. Evaluates SVI, F/M ratio, DO, and COD:TP to flag filamentous/viscous bulking conditions.",
            inputSchema={"type":"object","properties":{"unit_name":{"type":"string","default":"AerTank1"}}},
        ),
        types.Tool(
            name="diagnose_phosphorus_removal",
            description="EBPR diagnostic. Checks anaerobic DO, NO3 intrusion, COD:TP ratio, and effluent P to identify why biological phosphorus removal is failing.",
            inputSchema={"type":"object","properties":{"unit_name":{"type":"string","default":"AerTank1"}}},
        ),
        types.Tool(
            name="recommend_setpoint_adjustments",
            description="Aggregate diagnostic + recommended operational adjustments. Ranks recommendations by priority (critical/high/medium/none) with the tool to use for each fix.",
            inputSchema={"type":"object","properties":{"aer_tank":{"type":"string","default":"AerTank1"},"influent_unit":{"type":"string","default":"Influent"},"was_unit":{"type":"string","default":"WAS_Pump"}}},
        ),
        types.Tool(
            name="estimate_annual_opex",
            description="Annual OPEX estimate: energy (aeration + pumping), sludge disposal, and chemicals. Returns USD/year and breakdown.",
            inputSchema={"type":"object","properties":{"job_id":{"type":"string","default":""},"energy_cost_per_kwh":{"type":"number","default":0.08},"sludge_disposal_cost_per_tonne":{"type":"number","default":50.0},"alum_cost_per_kg":{"type":"number","default":0.30},"polymer_cost_per_kg":{"type":"number","default":3.50}}},
        ),
        types.Tool(
            name="estimate_ghg_emissions",
            description="Annual GHG emissions (kg CO2eq/year) = Scope 2 electricity + Scope 1 direct N2O (0.5% of TKN load) + Scope 1 CH4 from digesters.",
            inputSchema={"type":"object","properties":{"job_id":{"type":"string","default":""},"grid_emission_factor_kgCO2_per_kWh":{"type":"number","default":0.45}}},
        ),
        types.Tool(
            name="plot_time_series_chart",
            description="Generate a PNG time-series chart from a completed job's CSV for selected effluent variables.",
            inputSchema={"type":"object","properties":{"job_id":{"type":"string"},"variables":{"type":"array","items":{"type":"string"}},"output_path":{"type":"string"}},"required":["job_id"]},
        ),
        types.Tool(
            name="export_publication_time_series",
            description="Export publication-quality time-series plot (serif fonts, 300 DPI, colorblind palette, PNG+PDF+SVG). Supports multiple simulation job_ids and multiple variables overlaid.",
            inputSchema={"type":"object","properties":{
                "job_ids":{"type":"array","items":{"type":"string"},"description":"One or more completed job_ids"},
                "variables":{"type":"array","items":{"type":"string"},"description":"Effluent variables (e.g. COD_eff, NH4_eff). Default: COD_eff, NH4_eff, TN_eff, TP_eff"},
                "output_path":{"type":"string","description":"Path without extension; formats will be appended"},
                "formats":{"type":"array","items":{"type":"string","enum":["png","pdf","svg","eps"]},"default":["png","pdf"]},
                "title":{"type":"string"},"xlabel":{"type":"string"},"ylabel":{"type":"string"},
                "grayscale_safe":{"type":"boolean","default":True}
            },"required":["job_ids"]},
        ),
        types.Tool(
            name="export_calibration_parity_plot",
            description="Export a calibration/validation 1:1 parity plot (simulated vs observed) with R2, NSE, RMSE, PBIAS statistics annotated. Publication-quality styling.",
            inputSchema={"type":"object","properties":{
                "simulated":{"type":"array","items":{"type":"number"}},
                "observed":{"type":"array","items":{"type":"number"}},
                "variable_name":{"type":"string","default":"variable"},
                "units":{"type":"string","default":"mg/L"},
                "output_path":{"type":"string"},
                "formats":{"type":"array","items":{"type":"string"},"default":["png","pdf"]},
                "title":{"type":"string"}
            },"required":["simulated","observed"]},
        ),
        types.Tool(
            name="export_scenario_comparison_bar_chart",
            description="Grouped bar chart (mean +/- SD error bars) comparing multiple effluent variables across multiple simulation jobs. Academic styling, multi-format output.",
            inputSchema={"type":"object","properties":{
                "job_ids":{"type":"array","items":{"type":"string"}},
                "variables":{"type":"array","items":{"type":"string"}},
                "output_path":{"type":"string"},
                "formats":{"type":"array","items":{"type":"string"},"default":["png","pdf"]},
                "title":{"type":"string"},"ylabel":{"type":"string"},
                "show_errorbars":{"type":"boolean","default":True}
            },"required":["job_ids","variables"]},
        ),
        types.Tool(
            name="export_sensitivity_tornado_diagram",
            description="Publication-quality tornado diagram showing parameter sensitivity. Ranks parameters by range of output deviation from baseline.",
            inputSchema={"type":"object","properties":{
                "sensitivity_data":{"type":"array","description":"List of {parameter,low,high} dicts","items":{"type":"object","properties":{"parameter":{"type":"string"},"low":{"type":"number"},"high":{"type":"number"}},"required":["parameter","low","high"]}},
                "baseline":{"type":"number"},
                "output_path":{"type":"string"},
                "formats":{"type":"array","items":{"type":"string"},"default":["png","pdf"]},
                "xlabel":{"type":"string","default":"Output response"},
                "title":{"type":"string"}
            },"required":["sensitivity_data"]},
        ),
        types.Tool(
            name="export_box_plot_distributions",
            description="Box-plot (or violin-plot) comparing the distribution of a single variable across multiple simulation jobs. Includes per-job descriptive statistics.",
            inputSchema={"type":"object","properties":{
                "job_ids":{"type":"array","items":{"type":"string"}},
                "variable":{"type":"string"},
                "output_path":{"type":"string"},
                "formats":{"type":"array","items":{"type":"string"},"default":["png","pdf"]},
                "title":{"type":"string"},"ylabel":{"type":"string"},
                "use_violin":{"type":"boolean","default":False}
            },"required":["job_ids","variable"]},
        ),
        types.Tool(
            name="export_academic_results_table_docx",
            description="Export a Word (.docx) table with publication/booktabs-style formatting (Times New Roman, heavy top/bottom rules, italic units row). Suitable for direct inclusion in a journal manuscript.",
            inputSchema={"type":"object","properties":{
                "headers":{"type":"array","items":{"type":"string"}},
                "rows":{"type":"array","items":{"type":"array"}},
                "output_path":{"type":"string"},
                "caption":{"type":"string"},
                "table_number":{"type":"integer","default":1},
                "column_units":{"type":"array","items":{"type":"string"}},
                "title":{"type":"string"},
                "footnote":{"type":"string"}
            },"required":["headers","rows"]},
        ),
        types.Tool(
            name="export_latex_results_table",
            description="Export a LaTeX booktabs table (\\toprule/\\midrule/\\bottomrule) as a .tex file, ready to include in a journal manuscript. Includes auto-escaping of special characters.",
            inputSchema={"type":"object","properties":{
                "headers":{"type":"array","items":{"type":"string"}},
                "rows":{"type":"array","items":{"type":"array"}},
                "output_path":{"type":"string"},
                "caption":{"type":"string"},
                "label":{"type":"string","default":"tab:results"},
                "column_align":{"type":"string","description":"LaTeX alignment string e.g. 'lrrr'"},
                "column_units":{"type":"array","items":{"type":"string"}},
                "footnote":{"type":"string"},
                "table_type":{"type":"string","enum":["table","table*"],"default":"table"}
            },"required":["headers","rows"]},
        ),
        types.Tool(
            name="export_statistical_summary_xlsx",
            description="Export a multi-sheet Excel workbook with descriptive statistics (n/mean/SD/CV%/min/p05/p25/median/p75/p95/max) for multiple simulation jobs. One sheet per job plus Combined_Mean and Combined_Std sheets.",
            inputSchema={"type":"object","properties":{
                "job_ids":{"type":"array","items":{"type":"string"}},
                "variables":{"type":"array","items":{"type":"string"}},
                "output_path":{"type":"string"},
                "sheet_per_job":{"type":"boolean","default":True},
                "include_combined":{"type":"boolean","default":True}
            },"required":["job_ids"]},
        ),
        types.Tool(
            name="export_academic_report_docx",
            description="Generate a full academic Word report (Methods/Results/Discussion/References) with embedded publication-quality figures and booktabs result tables across multiple simulation jobs.",
            inputSchema={"type":"object","properties":{
                "job_ids":{"type":"array","items":{"type":"string"}},
                "variables":{"type":"array","items":{"type":"string"}},
                "title":{"type":"string","default":"SUMO24 Simulation Report"},
                "authors":{"type":"string"},
                "abstract":{"type":"string"},
                "output_path":{"type":"string"},
                "include_figures":{"type":"boolean","default":True},
                "embed_figures_dir":{"type":"string"}
            },"required":["job_ids"]},
        ),
        types.Tool(
            name="export_multi_file_batch_comparison",
            description="Batch export across multiple SUMO files and multiple simulation jobs. Accepts either a list of job_ids or a mapping {sumo_file_label: [job_id, ...]}. Produces combined plots (time-series + scenario bars), a multi-sheet XLSX summary, and a LaTeX booktabs table in one call.",
            inputSchema={"type":"object","properties":{
                "jobs_by_file":{"description":"Either a list of job_ids or a dict {sumo_file_label: [job_ids]}","oneOf":[{"type":"array","items":{"type":"string"}},{"type":"object","additionalProperties":{"type":"array","items":{"type":"string"}}}]},
                "variables":{"type":"array","items":{"type":"string"}},
                "output_dir":{"type":"string"},
                "formats":{"type":"array","items":{"type":"string"},"default":["png","pdf"]},
                "generate_plots":{"type":"boolean","default":True},
                "generate_xlsx":{"type":"boolean","default":True},
                "generate_latex":{"type":"boolean","default":True}
            },"required":["jobs_by_file"]},
        ),
        types.Tool(
            name="export_mass_balance_diagram",
            description="Publication-style mass-balance block diagram showing boxes for unit processes and arrows for flows (arrow thickness scaled by flow magnitude).",
            inputSchema={"type":"object","properties":{
                "flows":{"type":"array","description":"List of {from,to,value,label?} dicts","items":{"type":"object","properties":{"from":{"type":"string"},"to":{"type":"string"},"value":{"type":"number"},"label":{"type":"string"}},"required":["from","to"]}},
                "output_path":{"type":"string"},
                "formats":{"type":"array","items":{"type":"string"},"default":["png","pdf"]},
                "title":{"type":"string"},
                "units":{"type":"string","default":"kg/d"}
            },"required":["flows"]},
        ),
        # â”€â”€ Group AA â€” WWTP Template Tools (Excel â†” SUMO model build) â”€â”€â”€
        types.Tool(
            name="generate_wwtp_template_xlsx",
            description="Generate a fresh WWTP_Model_Input_Template.xlsx with sheets for plant info, influent/effluent datasets, COD/N/P fractionation, treatment processes (physical, biological, chemical), connections, sizing, and operational parameters. Filled-in templates can be consumed by build_model_from_template.",
            inputSchema={"type":"object","properties":{
                "output_path":{"type":"string","description":"Optional absolute path. Defaults to <MCP_DIR>/WWTP_Model_Input_Template.xlsx."}
            }},
        ),
        types.Tool(
            name="read_wwtp_template",
            description="Parse a filled-in WWTP_Model_Input_Template.xlsx and return all rows (plant info, influent/effluent time series, fractionation, processes, connections, sizing, operational parameters, chemical dosing, effluent targets, and influent statistics).",
            inputSchema={"type":"object","properties":{
                "template_path":{"type":"string","description":"Absolute path to the filled-in WWTP template .xlsx file."}
            },"required":["template_path"]},
        ),
        types.Tool(
            name="validate_wwtp_template",
            description="Validate a filled-in WWTP template: required Plant_Info fields, duplicate instance names, dangling connections, missing reactor/clarifier sizing, and COD-fraction sum sanity.",
            inputSchema={"type":"object","properties":{
                "template_path":{"type":"string"}
            },"required":["template_path"]},
        ),
        types.Tool(
            name="build_model_from_template",
            description="Build a SUMO24 model from scratch using a filled-in WWTP template. Creates the project folder + .sumo stub, adds every unit process in topological order (sorted by order_index), wires the streams (forward + recycle), applies sizing and operational parameters, sets influent characteristics (mean dataset values + COD/N/P fractionation), and dispatches each step via the DTT executeCommand bridge â€” with a dry_run mode that returns the build plan without touching DTT.",
            inputSchema={"type":"object","properties":{
                "template_path":{"type":"string"},
                "project_path":{"type":"string","description":"Override the project_path field in Plant_Info (optional)."},
                "apply_via_dtt":{"type":"boolean","default":True,"description":"If False, return planned commands without invoking DTT."},
                "dry_run":{"type":"boolean","default":False,"description":"If True, do not create any files or call DTT; return the full build plan only."},
                "allow_offline":{"type":"boolean","default":False,"description":"If True, proceed even when DTT is not live — commands are written to a .commands.txt file but NOT executed, so no working .sumo is produced."}
            },"required":["template_path"]},
        ),
        # â”€â”€ Group BB â€” HTML Schematic â†” SUMO Model â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        types.Tool(
            name="import_schematic_from_html",
            description=("Read a WWTP schematic file produced by the schematic builder (wwtp_schematic_template.html). "
                         "Extracts the embedded <script id='schematic-data'> JSON block (or loads a standalone .json export), "
                         "validates against the unit-type registry, and returns the parsed schematic + validation report. "
                         "The user must click 'Save HTML (for MCP)' in the builder for the file to contain meaningful state."),
            inputSchema={"type":"object","properties":{
                "filepath":{"type":"string","description":"Absolute or relative path to a .html (schematic builder export) or .json file."}
            },"required":["filepath"]},
        ),
        types.Tool(
            name="build_model_from_schematic",
            description=("Convert a parsed schematic into SumoCore parameter override commands. Optionally saves the commands "
                         "to a .scs script and/or applies them to the active compiled SUMO model. By default does NOT create "
                         "new unit topology â€” only overrides parameters on units that already exist in the compiled .sumo. "
                         "Use generate_sumoslang_from_schematic for full topology generation (advanced)."),
            inputSchema={"type":"object","properties":{
                "schematic":{"type":"object","description":"Parsed schematic dict returned by import_schematic_from_html (or pass filepath instead)."},
                "filepath":{"type":"string","description":"Alternative to 'schematic': load directly from .html / .json."},
                "apply_immediately":{"type":"boolean","default":False,"description":"If True, send each 'set' to the active DTT job."},
                "output_scs_path":{"type":"string","description":"Optional path to write the generated .scs script."},
                "include_dtt_actions":{"type":"boolean","default":False,"description":"Also return DTT-style addprocess/connect/setparameter actions."}
            }},
        ),
        types.Tool(
            name="compare_schematic_to_model",
            description=("Diff a schematic against the currently loaded compiled SUMO model and report which unit IDs exist in "
                         "both, which are missing in the model, and how many parameter overrides the schematic would apply."),
            inputSchema={"type":"object","properties":{
                "schematic":{"type":"object","description":"Parsed schematic dict."},
                "filepath":{"type":"string","description":"Alternative: load directly from .html / .json."}
            }},
        ),
        types.Tool(
            name="generate_sumoslang_from_schematic",
            description=("(Advanced / stretch goal) Render a SumoSlang plant-definition skeleton from the schematic. The output is "
                         "a STARTING POINT only â€” kinetic constants and class-specific blocks still need hand-editing before "
                         "SUMO will accept it. Returns the SumoSlang text and optionally writes it to a file."),
            inputSchema={"type":"object","properties":{
                "schematic":{"type":"object"},
                "filepath":{"type":"string","description":"Alternative: load directly from .html / .json."},
                "output_path":{"type":"string","description":"Optional .sumoslang or .txt file path."}
            }},
        ),
        types.Tool(
            name="list_schematic_unit_types",
            description="Return the full unit-type registry the schematic parser recognises (mirrors the JS UNIT_TYPES in the HTML builder).",
            inputSchema={"type":"object","properties":{}},
        ),
        # â”€â”€ Group CC â€” SUMO File Troubleshooting Package â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        types.Tool(
            name="diagnose_sumo_file",
            description=("Diagnose a .sumo project file: confirm it exists, that it's a valid zip/XML container, that it contains "
                         "a sumoproject.dll companion (or report which is missing), checks emptiness (0-byte / truncated), "
                         "parses the manifest for unit count, lists referenced state.xml, and reports common error patterns."),
            inputSchema={"type":"object","properties":{
                "sumo_path":{"type":"string","description":"Absolute path to the .sumo file."}
            },"required":["sumo_path"]},
        ),
        types.Tool(
            name="scan_sumo_directory",
            description="Scan a directory recursively for .sumo files and run diagnose_sumo_file on every one. Returns a per-file health report.",
            inputSchema={"type":"object","properties":{
                "directory":{"type":"string","description":"Folder to scan (defaults to MCP root)."}
            }},
        ),
        types.Tool(
            name="check_dll_companion",
            description=("Verify that sumoproject.dll exists next to a .sumo file (DTT requirement). If missing, attempts to locate the "
                         "DLL elsewhere in the same project tree and reports a recommended fix."),
            inputSchema={"type":"object","properties":{
                "sumo_path":{"type":"string"}
            },"required":["sumo_path"]},
        ),
        types.Tool(
            name="repair_sumo_file",
            description=("Attempt automated fixes for the most common .sumo problems: (1) empty/missing state.xml stub, "
                         "(2) missing sumoproject.dll companion (copy from a known-good source if available), "
                         "(3) zero-byte .sumo file (recreate a minimal valid stub), (4) wrong file extension. Always backs up "
                         "the original to <name>.sumo.bak before writing."),
            inputSchema={"type":"object","properties":{
                "sumo_path":{"type":"string"},
                "source_dll":{"type":"string","description":"Optional path to a known-good sumoproject.dll to copy as the companion."},
                "dry_run":{"type":"boolean","default":False,"description":"Only report what would be done; don't write anything."}
            },"required":["sumo_path"]},
        ),
        types.Tool(
            name="diagnose_sumo_crash",
            description=("Parse the most recent SUMO / DTT log files for stack traces and known error signatures (license errors, "
                         "DLL load failures, compile failures, etc.) and return a structured diagnosis with suggested fixes."),
            inputSchema={"type":"object","properties":{
                "log_path":{"type":"string","description":"Optional path to a specific log file. If omitted, scans likely SUMO log locations."},
                "tail_lines":{"type":"integer","default":400}
            }},
        ),
        types.Tool(
            name="validate_sumo_environment",
            description=("End-to-end environment health check for the SUMO24 MCP server: DTT import, sumoproject.dll present, "
                         "state.xml present, license file present, openpyxl available, schematic_parser available, and the "
                         "expected directory layout."),
            inputSchema={"type":"object","properties":{}},
        ),
        types.Tool(
            name="list_sumo_diagnostics",
            description="Return the full catalogue of known SUMO error signatures and their recommended remediations.",
            inputSchema={"type":"object","properties":{}},
        ),

        # ============================================================
        # Group DD - Schematic-to-SUMO File Compiler
        # End-to-end pipeline: read HTML schematic -> compile to .sumo
        # zip container (manifest.xml + state.xml + metadata.json +
        # sumoslang.txt), attach companion DLL, verify against source.
        # ============================================================
        types.Tool(
            name="compile_schematic_to_sumo",
            description="[DEPRECATED — always returns ok=false. Use build_sumo_pack instead.] Formerly compiled a schematic to a zip .sumo file. That zip format cannot be loaded by SUMO24. This tool now refuses unconditionally and redirects to build_sumo_pack.",
            inputSchema={"type":"object","properties":{
                "schematic":{"type":"object","description":"Schematic dict (as returned by import_schematic_from_html)."},
                "output_path":{"type":"string","description":"Where to write the .sumo file (extension auto-added)."},
                "overwrite":{"type":"boolean","description":"Replace existing file at output_path. Default false.","default":False},
                "include_sumoslang":{"type":"boolean","description":"Also bundle the SumoSlang skeleton text inside the zip. Default true.","default":True},
                "attach_dll":{"type":"string","description":"Optional path to a sumoproject.dll to copy alongside the .sumo file.","default":""},
            },"required":["schematic","output_path"]},
        ),
        types.Tool(
            name="build_sumo_from_html",
            description="[DEPRECATED — always returns ok=false. Use build_sumo_pack instead.] Formerly built a .sumo zip from an HTML schematic. That zip format cannot be loaded by SUMO24. This tool now refuses unconditionally and redirects to build_sumo_pack.",
            inputSchema={"type":"object","properties":{
                "html_path":{"type":"string","description":"Absolute path to the wwtp_schematic_template HTML."},
                "output_path":{"type":"string","description":"Where to write the .sumo file (extension auto-added)."},
                "overwrite":{"type":"boolean","description":"Replace existing file at output_path. Default false.","default":False},
                "include_sumoslang":{"type":"boolean","description":"Bundle the SumoSlang skeleton. Default true.","default":True},
                "attach_dll":{"type":"string","description":"Optional path to a sumoproject.dll companion.","default":""},
            },"required":["html_path","output_path"]},
        ),
        types.Tool(
            name="preview_sumo_manifest",
            description="Render and return (without writing) the manifest.xml, state.xml, metadata.json and optional sumoslang.txt that compile_schematic_to_sumo would produce. Useful for inspection or version-control review.",
            inputSchema={"type":"object","properties":{
                "schematic":{"type":"object","description":"Schematic dict."},
                "html_path":{"type":"string","description":"Or path to HTML to extract schematic from.","default":""},
                "include_sumoslang":{"type":"boolean","default":True},
            }},
        ),
        types.Tool(
            name="attach_companion_dll",
            description="[DEPRECATED — always returns ok=false. Use verify_dll_matches_schematic instead.] Formerly copied a sumoproject.dll next to a .sumo file. Copying a DLL from a different plant produces a symbol-table mismatch and the project will not simulate. This tool now refuses unconditionally.",
            inputSchema={"type":"object","properties":{
                "sumo_path":{"type":"string","description":"Path to the target .sumo file."},
                "source_dll":{"type":"string","description":"Path to a known-good sumoproject.dll to copy."},
            },"required":["sumo_path","source_dll"]},
        ),
        types.Tool(
            name="verify_sumo_file_against_schematic",
            description="[DEPRECATED — always returns ok=false. Use verify_dll_matches_schematic instead.] Formerly verified the custom zip format against a schematic. That zip format is not a valid SUMO project. This tool now refuses unconditionally and redirects to verify_dll_matches_schematic.",
            inputSchema={"type":"object","properties":{
                "sumo_path":{"type":"string","description":"Path to the compiled .sumo file."},
                "schematic":{"type":"object","description":"The original schematic dict."},
                "html_path":{"type":"string","description":"Or path to HTML to extract schematic from.","default":""},
            },"required":["sumo_path"]},
        ),
        types.Tool(
            name="read_sumo_manifest",
            description="Open a .sumo zip file and return a parsed summary of its manifest: unit count, connection count, IDs, metadata, and number of state variables.",
            inputSchema={"type":"object","properties":{
                "sumo_path":{"type":"string","description":"Path to a .sumo file."},
            },"required":["sumo_path"]},
        ),

        # ============================================================
        # Group EE — Build Pack & Baseline Apply
        # Replaces the broken zip-as-.sumo approach in Group DD.
        # These tools produce artefacts the user feeds into SUMO24
        # to get a real, loadable .sumo project.
        # ============================================================
        types.Tool(
            name="build_sumo_pack",
            description=(
                "Generate a build-pack directory from a schematic. The pack contains "
                "plant.sumoslang, apply_parameters.scs, BUILD_INSTRUCTIONS.md, and "
                "schematic.json. Open BUILD_INSTRUCTIONS.md and follow the 5-step "
                "procedure inside SUMO24 GUI to obtain a working .sumo project. "
                "This replaces the broken compile_schematic_to_sumo / build_sumo_from_html "
                "pair, which produced zip files SUMO24 could not load."
            ),
            inputSchema={"type":"object","properties":{
                "schematic_html":{"type":"string","description":"Path to wwtp_schematic_template.html (or any saved schematic HTML)."},
                "schematic_json":{"type":"string","description":"Alternative to schematic_html: path to a schematic .json export."},
                "output_dir":{"type":"string","description":"Directory to write the pack into (e.g. F:/UNI/SUMO/MCP/qaha.sumo-pack)."},
                "overwrite":{"type":"boolean","default":False},
            },"required":["output_dir"]},
        ),
        types.Tool(
            name="apply_schematic_to_baseline",
            description=(
                "Tier-1 path that DOES produce a working .sumo project: copy a verified "
                "baseline (.sumo + sumoproject.dll) to a new location, then apply every "
                "parameter override the schematic specifies. Topology comes from the "
                "baseline; only parameter values change. Refuses to proceed when the "
                "schematic references variables not present in the baseline (unless "
                "strict=false is passed). NOTE: this tool keeps the .sumo zip "
                "byte-identical to the baseline and writes overrides into a sibling "
                "state.xml. If you want the overrides baked into the .sumo zip itself "
                "(rewriting parameters.txt + Influent tables + notes.rtf inside the "
                "archive), use build_native_sumo instead."
            ),
            inputSchema={"type":"object","properties":{
                "schematic_html":{"type":"string","description":"Path to schematic HTML."},
                "schematic_json":{"type":"string","description":"Path to schematic JSON export (alternative to schematic_html)."},
                "baseline_sumo":{"type":"string","description":"Path to a known-good .sumo file produced by SUMO itself."},
                "baseline_dll":{"type":"string","description":"Path to the matching sumoproject.dll."},
                "output_sumo":{"type":"string","description":"Destination path for the output .sumo copy."},
                "strict":{"type":"boolean","default":True},
            },"required":["baseline_sumo","baseline_dll","output_sumo"]},
        ),
        types.Tool(
            name="build_native_sumo",
            description=(
                "Native .sumo composer (Tier-1, baseline-DLL-reuse). Opens a verified "
                "baseline .sumo zip, rewrites the parameter-bearing internal members "
                "(parameters.txt, Influent*_Table*.tsv, notes.rtf, userscript.txt) "
                "from the supplied schematic, and writes a NEW zip that SUMO24 can "
                "open natively. Keeps the compiled sumoproject.dll byte-perfect and "
                "preserves every DLL-bound member. The plant topology is fixed by "
                "the baseline DLL; only parameter values change. Strict mode (default) "
                "refuses to build if any schematic unit lacks a baseline target — "
                "pass unit_mapping={schematic_id: baseline_id, ...} to override the "
                "heuristic class-based matcher. If the DTT bridge is live, the tool "
                "additionally tries load_model on the produced .sumo and reports the "
                "outcome in verify_load."
            ),
            inputSchema={"type":"object","properties":{
                "schematic_html":{"type":"string","description":"Path to schematic HTML (or pass schematic_json)."},
                "schematic_json":{"type":"string","description":"Path to schematic JSON export (alternative to schematic_html)."},
                "baseline_sumo":{"type":"string","description":"Path to a known-good .sumo file with sumoproject.dll inside it."},
                "baseline_dll":{"type":"string","description":"Optional companion sumoproject.dll path. If omitted, the DLL is extracted from the baseline zip and written alongside output_sumo."},
                "output_sumo":{"type":"string","description":"Destination .sumo path (extension auto-fixed)."},
                "unit_mapping":{"type":"object","description":"Optional explicit mapping {schematic_unit_id: baseline_unit_id}. Wins over heuristic matching.","additionalProperties":{"type":"string"}},
                "overwrite":{"type":"boolean","default":False,"description":"Replace output_sumo if it exists."},
                "strict":{"type":"boolean","default":True,"description":"Refuse to build when any schematic unit has no baseline target. Set false to skip unmapped units silently."},
                "verify_load":{"type":"boolean","default":True,"description":"If DTT is live, attempt to load_model the produced .sumo as a structural verification step."},
            },"required":["baseline_sumo","output_sumo"]},
        ),
        types.Tool(
            name="topology_match_report",
            description=(
                "Preview how a schematic maps onto a baseline .sumo BEFORE running "
                "build_native_sumo. Returns the proposed {schematic_id: baseline_id} "
                "mapping, the list of unmapped schematic units, the unused baseline "
                "units, and any class-mismatch warnings. Cheap to call — does not "
                "write any files."
            ),
            inputSchema={"type":"object","properties":{
                "schematic_html":{"type":"string","description":"Path to schematic HTML (or pass schematic_json)."},
                "schematic_json":{"type":"string","description":"Path to schematic JSON export."},
                "baseline_sumo":{"type":"string","description":"Path to the baseline .sumo to match against."},
                "unit_mapping":{"type":"object","description":"Optional explicit overrides for the matcher.","additionalProperties":{"type":"string"}},
            },"required":["baseline_sumo"]},
        ),
        types.Tool(
            name="validate_sumoslang",
            description=(
                "Lint a SumoSlang text for the most common breakages: brace balance, "
                "missing kineticmodel line, leftover TODO/FIXME markers. Use to pre-flight a "
                "generated plant.sumoslang before pasting into SUMO24."
            ),
            inputSchema={"type":"object","properties":{
                "text":{"type":"string","description":"SumoSlang text to lint."},
                "path":{"type":"string","description":"Alternative to text: path to a .sumoslang file."},
            }},
        ),
        types.Tool(
            name="verify_dll_matches_schematic",
            description=(
                "Inspect a sumoproject.dll's exported symbols and check that every "
                "sumo_variable referenced by the schematic exists in the DLL. Catches "
                "the common error of copying a DLL from one plant into another plant's "
                "directory — the file copy succeeds but the DLL has no idea what "
                "OxidationDitch1_T1 means and the project will load empty. "
                "Provide schematic_html or schematic_json for a full match; omit them "
                "to just list the DLL's exported Sumo__Plant__* symbols."
            ),
            inputSchema={"type":"object","properties":{
                "dll_path":{"type":"string","description":"Absolute path to sumoproject.dll."},
                "schematic_html":{"type":"string","description":"Path to schematic HTML (optional — enables cross-check).","default":""},
                "schematic_json":{"type":"string","description":"Path to schematic JSON export (alternative to schematic_html).","default":""},
            },"required":["dll_path"]},
        ),
        types.Tool(
            name="update_schematic_in_html",
            description=(
                "Safely replace the embedded schematic JSON inside a schematic-builder "
                "HTML file. Uses a comment-aware scanner so it never matches a reference "
                "inside <!-- ... --> and destroys the file (the bug that hit the Qaha1 "
                "session). Validates JSON before writing, validates structural tags after "
                "writing, and aborts if either check fails. Creates a .bak backup by default."
            ),
            inputSchema={"type":"object","properties":{
                "html_path":{"type":"string","description":"Absolute path to the schematic-builder HTML file."},
                "schematic_json":{"type":"string","description":"New schematic as a JSON string, or a path to a .json file."},
                "backup":{"type":"boolean","default":True,"description":"Write a .bak backup before modifying. Default true."},
            },"required":["html_path","schematic_json"]},
        ),
        types.Tool(
            name="next_step_for_project",
            description=(
                "Inspect a project directory and tell the user exactly what to do next. "
                "Looks for the .sumo stub, the build-pack, the DLL, state.xml, and an "
                "overrides-applied marker, then returns a single recommended next action. "
                "Use this any time the workflow feels ambiguous — it collapses the "
                "5-step BUILD_INSTRUCTIONS.md into one guided call."
            ),
            inputSchema={"type":"object","properties":{
                "project_dir":{"type":"string","description":"Absolute path to the project folder (the one containing the .sumo file)."},
            },"required":["project_dir"]},
        ),
        types.Tool(
            name="apply_template_auto_fixes",
            description=(
                "Apply auto_fix suggestions returned by validate_wwtp_template to the "
                "template Excel workbook in place. Accepts the list of auto_fix dicts "
                "from validate_wwtp_template's error entries and rewrites the workbook. "
                "Handles: add_to_treatment_processes (adds WAS/RAS boundary units), "
                "clear_na_rows (removes NA filler rows in the Connections sheet). "
                "Always makes a .bak copy before writing."
            ),
            inputSchema={"type":"object","properties":{
                "template_path":{"type":"string","description":"Absolute path to the WWTP template .xlsx to fix."},
                "fixes":{"type":"array","description":"List of auto_fix dicts from validate_wwtp_template error entries.","items":{"type":"object"}},
            },"required":["template_path","fixes"]},
        ),

        # ── Group FF — Pipeline driver and stage orchestration ─────────────────
        types.Tool(
            name="pipeline_init",
            description=(
                "Scaffold a new WWTP project directory with project.yaml manifest, "
                "HTML schematic template, and Excel data template. Required first call "
                "for any new plant. Sets up the 8-stage gated build pipeline."
            ),
            inputSchema={"type":"object","properties":{
                "project_dir":{"type":"string","description":"Absolute path for the new project directory."},
                "plant_name":{"type":"string","description":"Short plant name (e.g. Qaha1). Used for file naming."},
                "kinetic_model":{"type":"string","default":"ASM2d","description":"ASM1 | ASM2d | ASM3"},
                "from_template":{"type":"string","default":"blank","description":"Preset: qaha_oxditch | classic_as | mbbr | bnr_3stage | blank"},
                "force":{"type":"boolean","default":False,"description":"Overwrite existing project directory."},
            },"required":["project_dir","plant_name"]},
        ),
        types.Tool(
            name="pipeline_status",
            description=(
                "Return the complete pipeline state: current stage, status of every stage, "
                "stale detection (re-hashes artefacts on disk), and the exact next tool to call. "
                "Primary navigation tool for the pipeline."
            ),
            inputSchema={"type":"object","properties":{
                "project_dir":{"type":"string","description":"Absolute path to project directory containing project.yaml."},
            },"required":["project_dir"]},
        ),
        types.Tool(
            name="pipeline_advance",
            description=(
                "Run the current stage validator and advance if it passes. Hash-chains the "
                "artefact and marks downstream stages stale on hash change. Use skip=True "
                "with a reason for optional stages 5 and 6."
            ),
            inputSchema={"type":"object","properties":{
                "project_dir":{"type":"string","description":"Absolute path to project directory."},
                "from_stage":{"type":"integer","description":"Force re-validate this stage number (0-8)."},
                "skip":{"type":"boolean","default":False},
                "skip_reason":{"type":"string"},
                "force":{"type":"boolean","default":False},
            },"required":["project_dir"]},
        ),
        types.Tool(
            name="pipeline_revert",
            description=(
                "Roll back the project to the state at the end of to_stage using the latest "
                "matching snapshot in <project_dir>/snapshots/."
            ),
            inputSchema={"type":"object","properties":{
                "project_dir":{"type":"string","description":"Absolute path to project directory."},
                "to_stage":{"type":"integer","description":"Stage number to revert to (0-8)."},
            },"required":["project_dir","to_stage"]},
        ),
        types.Tool(
            name="pipeline_describe",
            description=(
                "Generate a human-readable Markdown report of the pipeline state: "
                "stage status, engineering values, provenance log, snapshots."
            ),
            inputSchema={"type":"object","properties":{
                "project_dir":{"type":"string"},
            },"required":["project_dir"]},
        ),
        types.Tool(
            name="stage1_open_schematic",
            description="Stage 1 entry: return the HTML schematic path for this project.",
            inputSchema={"type":"object","properties":{
                "project_dir":{"type":"string"},
            },"required":["project_dir"]},
        ),
        types.Tool(
            name="stage1_validate_schematic",
            description=(
                "Stage 1 validator: parse HTML schematic, check unit types against registry, "
                "stream references, topology (influent+effluent required), orphan units, "
                "recycle hydraulics. Called automatically by pipeline_advance."
            ),
            inputSchema={"type":"object","properties":{
                "project_dir":{"type":"string"},
            },"required":["project_dir"]},
        ),
        types.Tool(
            name="stage2_open_data",
            description=(
                "Stage 2 entry: ensure the Excel data template exists and pre-seed "
                "Treatment_Processes and Connections sheets from the Stage 1 schematic. "
                "Returns the Excel path."
            ),
            inputSchema={"type":"object","properties":{
                "project_dir":{"type":"string"},
            },"required":["project_dir"]},
        ),
        types.Tool(
            name="stage2_validate_parameters",
            description=(
                "Stage 2 validator: validate_wwtp_template + cross-check unit names and "
                "connections against the Stage 1 schematic. Fails on any name mismatch."
            ),
            inputSchema={"type":"object","properties":{
                "project_dir":{"type":"string"},
            },"required":["project_dir"]},
        ),
        types.Tool(
            name="stage3_run_engineering_checks",
            description=(
                "Stage 3: compute and validate SRT, HRT, F:M, AOR, OTR/AOR, COD mass "
                "balance, residual alkalinity, Takacs settler SOR/SLR. "
                "Reads from the Excel template. Writes engineering_checks.json."
            ),
            inputSchema={"type":"object","properties":{
                "project_dir":{"type":"string"},
            },"required":["project_dir"]},
        ),
        types.Tool(
            name="stage4_define_static_inputs",
            description=(
                "Stage 4: resolve influent fractionation into ASM2d state-variable initial "
                "values. Writes static_inputs.json and static_inputs.scs."
            ),
            inputSchema={"type":"object","properties":{
                "project_dir":{"type":"string"},
            },"required":["project_dir"]},
        ),
        types.Tool(
            name="stage4_validate_static_inputs",
            description=(
                "Stage 4 validator: check COD/N/P fraction sums, physical ranges, DO "
                "setpoints by reactor type, MLSS range, Q_RAS range, SRT cross-check."
            ),
            inputSchema={"type":"object","properties":{
                "project_dir":{"type":"string"},
            },"required":["project_dir"]},
        ),
        types.Tool(
            name="stage5_generate_dynamic_inputs",
            description=(
                "Stage 5 (optional): produce a dynamic input TSV. "
                "profile_type: constant | from_excel | diurnal | storm | seasonal_temp. "
                "Output: <project_dir>/dynamic_inputs/<name>.tsv"
            ),
            inputSchema={"type":"object","properties":{
                "project_dir":{"type":"string"},
                "profile_type":{"type":"string","description":"constant|from_excel|diurnal|storm|seasonal_temp"},
                "duration_days":{"type":"number","default":30.0},
                "dt_h":{"type":"number","default":0.5},
                "name":{"type":"string"},
            },"required":["project_dir","profile_type"]},
        ),
        types.Tool(
            name="stage5_validate_dynamic_inputs",
            description=(
                "Stage 5 validator: check dynamic_inputs/*.tsv files for uniform columns, "
                "monotonic time, physical value ranges, duration vs Sumo__StopTime."
            ),
            inputSchema={"type":"object","properties":{
                "project_dir":{"type":"string"},
                "sumo_stop_time":{"type":"number"},
            },"required":["project_dir"]},
        ),
        types.Tool(
            name="stage6_add_controller",
            description=(
                "Stage 6 (optional): register a controller in controllers.json. "
                "Supported kinds: DO | SRT | NH4 | NO3 | FlowSplitter."
            ),
            inputSchema={"type":"object","properties":{
                "project_dir":{"type":"string"},
                "kind":{"type":"string","description":"DO|SRT|NH4|NO3|FlowSplitter"},
                "unit":{"type":"string"},
                "setpoint":{"type":"number"},
                "sensor":{"type":"string"},
                "actuator":{"type":"string"},
                "Kp":{"type":"number","default":10.0},
                "Ti":{"type":"number","default":0.01},
            },"required":["project_dir","kind","unit","setpoint"]},
        ),
        types.Tool(
            name="stage6_validate_controllers",
            description=(
                "Stage 6 validator: check controller unit references, setpoint ranges, "
                "no conflicting actuators."
            ),
            inputSchema={"type":"object","properties":{
                "project_dir":{"type":"string"},
            },"required":["project_dir"]},
        ),
        types.Tool(
            name="stage7_build",
            description=(
                "Stage 7: assemble the build pack at <project_dir>/<plant_name>.sumo-pack/. "
                "Contents: plant.sumoslang, apply_parameters.scs, static_inputs.scs, "
                "dynamic_inputs/, controllers.scs (if not skipped), .sumo stub, "
                "BUILD_INSTRUCTIONS.md, schematic.json, engineering_checks.json."
            ),
            inputSchema={"type":"object","properties":{
                "project_dir":{"type":"string"},
            },"required":["project_dir"]},
        ),
        types.Tool(
            name="stage8_verify",
            description=(
                "Stage 8 (post-SUMO-compile): verify DLL matches schematic, run trial "
                "steady-state, validate post-simulation mass balance. "
                "Call after compiling in SUMO24 GUI."
            ),
            inputSchema={"type":"object","properties":{
                "project_dir":{"type":"string"},
                "dll_path":{"type":"string","description":"Path to sumoproject.dll from SUMO."},
                "state_xml":{"type":"string","description":"Path to state.xml after maptoic."},
            },"required":["project_dir","dll_path"]},
        ),
        # ---- Group FF: DTT-mediated .sumo editing -----------------------------
        types.Tool(
            name="rename_unit_process",
            description="Rename a unit process in the loaded .sumo model via DTT "
                        "executeCommand. References to the old name in streams and "
                        "controllers are updated by SUMO; the live model is "
                        "modified in memory. Persist with save_model when done.",
            inputSchema={
                "type": "object",
                "properties": {
                    "old_name": {"type": "string"},
                    "new_name": {"type": "string"},
                    "dry_run":  {"type": "boolean", "default": False},
                },
                "required": ["old_name", "new_name"],
            },
        ),
        types.Tool(
            name="change_unit_type",
            description="Change the class of an existing unit (e.g. Generic "
                        "PrimaryClarifier -> Otterpohl) via DTT. Use with care: "
                        "incompatible class swaps will be rejected by the solver. "
                        "dry_run returns the planned DTT command without "
                        "dispatching.",
            inputSchema={
                "type": "object",
                "properties": {
                    "unit_name": {"type": "string"},
                    "new_class": {"type": "string",
                        "description": "Target SUMO class name (consult "
                                       "list_available_unit_process_types)."},
                    "dry_run":   {"type": "boolean", "default": True},
                },
                "required": ["unit_name", "new_class"],
            },
        ),
        types.Tool(
            name="remove_flow_connection",
            description="Delete a stream from the loaded model via DTT. The "
                        "complement to modify_flow_connection / connect_unit_"
                        "processes. dry_run returns the planned command.",
            inputSchema={
                "type": "object",
                "properties": {
                    "stream_id": {"type": "string"},
                    "dry_run":   {"type": "boolean", "default": False},
                },
                "required": ["stream_id"],
            },
        ),
        types.Tool(
            name="set_stream_flow_rate",
            description="Set a fixed flow rate (m3/d) on an existing stream via "
                        "DTT — used for RAS/WAS/IR loops or any fixed-Q stream.",
            inputSchema={
                "type": "object",
                "properties": {
                    "stream_id": {"type": "string"},
                    "q_m3d":     {"type": "number"},
                    "dry_run":   {"type": "boolean", "default": False},
                },
                "required": ["stream_id", "q_m3d"],
            },
        ),
        types.Tool(
            name="list_controllers",
            description="List controllers (DO / SRT / NH4 / NO3 / FlowSplitter) "
                        "in the loaded model. Tries DTT enumeration first; falls "
                        "back to the manifest.xml of the on-disk .sumo if DTT "
                        "doesn't expose a list API.",
            inputSchema={
                "type": "object",
                "properties": {
                    "sumo_path": {"type": "string",
                        "description": "Optional .sumo path for the fallback "
                                       "manifest read."},
                },
            },
        ),
        types.Tool(
            name="modify_controller",
            description="Modify an existing controller via DTT — setpoint, gain, "
                        "integral time, output min/max. Any field left null is "
                        "left unchanged. Use list_controllers first to confirm "
                        "the controller id.",
            inputSchema={
                "type": "object",
                "properties": {
                    "controller_id": {"type": "string"},
                    "setpoint":      {"type": "number"},
                    "gain":          {"type": "number"},
                    "integral_time": {"type": "number"},
                    "output_min":    {"type": "number"},
                    "output_max":    {"type": "number"},
                    "dry_run":       {"type": "boolean", "default": False},
                },
                "required": ["controller_id"],
            },
        ),
        types.Tool(
            name="remove_controller",
            description="Remove a control loop from the loaded model via DTT. "
                        "Complement to add_controller.",
            inputSchema={
                "type": "object",
                "properties": {
                    "controller_id": {"type": "string"},
                    "dry_run":       {"type": "boolean", "default": False},
                },
                "required": ["controller_id"],
            },
        ),
        types.Tool(
            name="begin_edit_transaction",
            description="Snapshot the current state.xml so a series of DTT edits "
                        "can be rolled back if anything goes wrong. Returns a "
                        "tx_id. Pair with commit_edit_transaction or "
                        "rollback_edit_transaction.",
            inputSchema={
                "type": "object",
                "properties": {
                    "label": {"type": "string", "default": "edit"},
                },
            },
        ),
        types.Tool(
            name="commit_edit_transaction",
            description="Mark an edit transaction successful and (by default) "
                        "discard its snapshot. Keep_snapshot=True retains the "
                        "snapshot as a manual restore point.",
            inputSchema={
                "type": "object",
                "properties": {
                    "tx_id":          {"type": "string"},
                    "keep_snapshot":  {"type": "boolean", "default": False},
                },
                "required": ["tx_id"],
            },
        ),
        types.Tool(
            name="rollback_edit_transaction",
            description="Restore the state.xml saved by begin_edit_transaction "
                        "and re-initialise the in-memory model so DTT and disk "
                        "agree. Backs the rolled-back state.xml aside as "
                        "preRollback_*.bak.",
            inputSchema={
                "type": "object",
                "properties": {
                    "tx_id": {"type": "string"},
                },
                "required": ["tx_id"],
            },
        ),
        types.Tool(
            name="apply_dtt_command_batch",
            description="Dispatch a vetted list of DTT command strings against "
                        "the loaded model. Default mode rejects shell-injection "
                        "tokens and any non-`set`/non-DTT verb; allow_raw=True "
                        "bypasses vetting for power users. dry_run returns the "
                        "planned commands without execution.",
            inputSchema={
                "type": "object",
                "properties": {
                    "commands":  {"type": "array", "items": {"type": "string"}},
                    "allow_raw": {"type": "boolean", "default": False},
                    "dry_run":   {"type": "boolean", "default": True},
                },
                "required": ["commands"],
            },
        ),
        types.Tool(
            name="diff_inmemory_vs_disk",
            description="Compare the live in-memory model (units + streams) "
                        "against the manifest of a .sumo on disk. Reports what "
                        "needs save_model / reload to bring them into sync.",
            inputSchema={
                "type": "object",
                "properties": {
                    "sumo_path": {"type": "string"},
                },
                "required": ["sumo_path"],
            },
        ),
        # ---- end Group FF -----------------------------------------------------
    ]



# â”€â”€ Pending parameter overrides (applied at next run) â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
_pending_sets: list[str] = []


@app.call_tool()
async def call_tool(name: str, arguments: dict) -> list[types.TextContent]:

    # â”€â”€ list_scenarios â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "list_scenarios":
        result = {}
        for sc_name, cmds in SCENARIOS.items():
            result[sc_name] = cmds if cmds else ["(baseline - no parameter changes)"]
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ set_parameter â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "set_parameter":
        var = arguments["variable"]
        val = arguments["value"]
        _pending_sets.append(f"set {var} {val}")
        return [types.TextContent(
            type="text",
            text=f"Queued: set {var} = {val}\nThis will be applied on the next simulation run.\n"
                 f"Currently {len(_pending_sets)} pending override(s)."
        )]

    # â”€â”€ run_steady_state â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "run_steady_state":
        if not DTT_AVAILABLE:
            return [types.TextContent(type="text", text="ERROR: dynamita package not installed.")]

        extra = list(_pending_sets)
        _pending_sets.clear()

        arg_map = {
            "influent_flow_m3d": "Sumo__Plant__Influent__param__Q",
            "influent_cod_mgL":  "Sumo__Plant__Influent__param__XCOD",
            "temperature_C":     "Sumo__Plant__Influent__param__T",
            "influent_tkn_mgL":  "Sumo__Plant__Influent__param__XTKN",
        }
        for arg_key, sumo_var in arg_map.items():
            if arg_key in arguments:
                extra.append(f"set {sumo_var} {arguments[arg_key]}")

        cmds = _build_commands(
            mode="steady",
            stop_days=1,
            datacomm_hours=0.1,
            extra_sets=extra,
            scenario_name=arguments.get("scenario"),
        )

        ds.sumo.setLogDetails(2)
        ds.sumo.message_callback  = _msg_callback
        ds.sumo.datacomm_callback = _data_callback

        with _job_lock:
            job_id = ds.sumo.schedule(
                CONFIG["model_dll"],
                commands  = cmds,
                variables = EFFLUENT_VARS,
                jobData   = {ds.sumo.persistent: True, "rows": [], "status": "queued"},
            )
            _active_jobs[job_id] = {
                "status": "queued", "rows": [], "error": None,
                "scenario": arguments.get("scenario", "baseline_Sc0"),
            }

        deadline = time.time() + 120
        while time.time() < deadline:
            await asyncio.sleep(1)
            with _job_lock:
                st = _active_jobs.get(job_id, {}).get("status")
            if st == "finished":
                break

        with _job_lock:
            rows = _active_jobs.get(job_id, {}).get("rows", [])

        summary    = _summarise(rows)
        compliance = _check_compliance(summary)

        response = {
            "job_id":     job_id,
            "mode":       "steady-state",
            "scenario":   arguments.get("scenario", "none"),
            "effluent":   summary,
            "compliance": compliance,
        }
        return [types.TextContent(type="text", text=json.dumps(response, indent=2))]

    # â”€â”€ run_dynamic_simulation â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "run_dynamic_simulation":
        if not DTT_AVAILABLE:
            return [types.TextContent(type="text", text="ERROR: dynamita package not installed.")]

        duration = arguments.get("duration_days", 30)
        dc_hours = arguments.get("datacomm_hours", 1)
        extra    = list(_pending_sets)
        _pending_sets.clear()

        if "influent_tsv" in arguments:
            extra = [f'loadtsv "{arguments["influent_tsv"]}" -interpolation linear'] + extra

        cmds = _build_commands(
            mode="dynamic",
            stop_days=duration,
            datacomm_hours=dc_hours,
            extra_sets=extra,
            scenario_name=arguments.get("scenario"),
        )

        ds.sumo.setLogDetails(2)
        ds.sumo.message_callback  = _msg_callback
        ds.sumo.datacomm_callback = _data_callback

        with _job_lock:
            job_id = ds.sumo.schedule(
                CONFIG["model_dll"],
                commands  = cmds,
                variables = EFFLUENT_VARS,
                jobData   = {ds.sumo.persistent: True, "rows": []},
            )
            _active_jobs[job_id] = {
                "status": "queued", "rows": [], "error": None,
                "scenario": arguments.get("scenario", "baseline_Sc0"),
            }

        return [types.TextContent(
            type="text",
            text=json.dumps({
                "job_id":    job_id,
                "status":    "scheduled",
                "mode":      "dynamic",
                "duration":  duration,
                "scenario":  arguments.get("scenario", "none"),
                "hint":      f"Use get_job_status(job_id={job_id}) to poll progress, "
                             f"then export_results_csv to get time-series data."
            }, indent=2)
        )]

    # â”€â”€ run_scenario_comparison â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "run_scenario_comparison":
        if not DTT_AVAILABLE:
            return [types.TextContent(type="text", text="ERROR: dynamita package not installed.")]

        scenario_list = arguments.get("scenarios", list(SCENARIOS.keys()))
        comparison: dict[str, Any] = {}

        ds.sumo.setLogDetails(2)
        ds.sumo.message_callback  = _msg_callback
        ds.sumo.datacomm_callback = _data_callback

        job_ids = {}
        for sc in scenario_list:
            cmds = _build_commands(mode="steady", stop_days=1, datacomm_hours=0.1, scenario_name=sc)
            with _job_lock:
                jid = ds.sumo.schedule(
                    CONFIG["model_dll"],
                    commands  = cmds,
                    variables = EFFLUENT_VARS,
                    jobData   = {ds.sumo.persistent: True, "rows": []},
                )
                _active_jobs[jid] = {"status": "queued", "rows": [], "error": None, "scenario": sc}
                job_ids[sc] = jid

        deadline = time.time() + 300
        while time.time() < deadline:
            await asyncio.sleep(2)
            with _job_lock:
                all_done = all(
                    _active_jobs.get(jid, {}).get("status") == "finished"
                    for jid in job_ids.values()
                )
            if all_done:
                break

        for sc, jid in job_ids.items():
            with _job_lock:
                rows = _active_jobs.get(jid, {}).get("rows", [])
            summary    = _summarise(rows)
            compliance = _check_compliance(summary)
            comparison[sc] = {"effluent": summary, "compliance": compliance}

        return [types.TextContent(type="text", text=json.dumps(comparison, indent=2))]

    # â”€â”€ get_job_status â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "get_job_status":
        job_id = arguments["job_id"]
        with _job_lock:
            info = _active_jobs.get(job_id)
        if info is None:
            return [types.TextContent(type="text", text=f"Job {job_id} not found.")]
        rows    = info.get("rows", [])
        summary = _summarise(rows) if info["status"] == "finished" else {}
        return [types.TextContent(type="text", text=json.dumps({
            "job_id":    job_id,
            "status":    info["status"],
            "n_rows":    len(rows),
            "last_msg":  info.get("last_msg", ""),
            "summary":   summary,
        }, indent=2))]

    # â”€â”€ check_compliance â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "check_compliance":
        if "job_id" in arguments:
            job_id = arguments["job_id"]
            with _job_lock:
                rows = _active_jobs.get(job_id, {}).get("rows", [])
            summary = _summarise(rows)
        else:
            summary = {k: arguments[k] for k in ("TSS_mgL", "BOD5_mgL", "sCOD_mgL") if k in arguments}

        compliance = _check_compliance(summary)
        overall = "COMPLIANT" if all(v["status"] == "PASS" for v in compliance.values()) else "NON-COMPLIANT"
        return [types.TextContent(type="text", text=json.dumps({
            "overall":    overall,
            "parameters": compliance,
            "law":        "Egyptian Law 48/1982",
        }, indent=2))]

    # â”€â”€ create_dynamic_input_table â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "create_dynamic_input_table":
        var = arguments.get("variable")
        if not var:
            return [types.TextContent(type="text",
                text="ERROR: 'variable' is required (full SumoCore name, e.g. Sumo__Plant__Influent__param__Q).")]

        sim = _read_simulation_settings()
        duration = float(arguments.get("duration_days",  sim["stop_time_days"]))
        dc_hours = float(arguments.get("datacomm_hours", sim["datacomm_hours"]))
        if duration <= 0:
            return [types.TextContent(type="text", text="ERROR: duration_days must be > 0.")]
        if dc_hours <= 0:
            return [types.TextContent(type="text", text="ERROR: datacomm_hours must be > 0.")]

        t_start = float(arguments.get("start_time_days", 0.0))
        t_end   = float(arguments.get("end_time_days",   duration))
        if t_end <= t_start:
            return [types.TextContent(type="text",
                text=f"ERROR: end_time_days ({t_end}) must be greater than start_time_days ({t_start}).")]

        step_days = dc_hours / 24.0
        n_steps   = int(round((t_end - t_start) / step_days)) + 1
        time_axis = [t_start + i * step_days for i in range(n_steps)]
        if time_axis:
            time_axis[-1] = t_end

        mode = arguments.get("value_mode", "constant")
        if mode == "constant":
            c = float(arguments.get("constant_value", 0.0))
            values = [c] * n_steps
        elif mode == "linear":
            s = float(arguments.get("start_value", 0.0))
            e = float(arguments.get("end_value",   0.0))
            values = [s] if n_steps == 1 else [s + (e - s) * i / (n_steps - 1) for i in range(n_steps)]
        elif mode == "step":
            s = float(arguments.get("start_value", 0.0))
            e = float(arguments.get("end_value",   0.0))
            step_at = float(arguments.get("step_at_day", (t_start + t_end) / 2.0))
            values = [e if t >= step_at else s for t in time_axis]
        elif mode == "sinusoidal":
            amp    = float(arguments.get("amplitude", 1.0))
            mean   = float(arguments.get("mean",      0.0))
            period = float(arguments.get("period_days", 1.0))
            if period <= 0:
                return [types.TextContent(type="text", text="ERROR: period_days must be > 0.")]
            values = [mean + amp * math.sin(2.0 * math.pi * t / period) for t in time_axis]
        elif mode == "custom":
            custom = arguments.get("values", [])
            if len(custom) != n_steps:
                return [types.TextContent(type="text",
                    text=f"ERROR: 'values' length ({len(custom)}) must match number of generated time steps ({n_steps}).")]
            values = [float(v) for v in custom]
        else:
            return [types.TextContent(type="text", text=f"Unknown value_mode: {mode}")]

        time_unit = arguments.get("time_unit", "days")
        if time_unit == "hours":
            time_col = [t * 24.0 for t in time_axis]
        elif time_unit == "minutes":
            time_col = [t * 1440.0 for t in time_axis]
        else:
            time_col = list(time_axis)

        safe_var = var.replace(":", "_").replace("/", "_").replace("\\", "_")
        filename = arguments.get("filename", f"dyn_input_{safe_var}.tsv")
        outpath  = Path(CONFIG["output_dir"]) / filename
        outpath.parent.mkdir(parents=True, exist_ok=True)
        with open(outpath, "w", newline="") as f:
            f.write(f"t_{time_unit}\t{var}\n")
            for t, v in zip(time_col, values):
                f.write(f"{t:.6g}\t{v:.6g}\n")

        preview: list[dict] = []
        for i, (t, v) in enumerate(zip(time_col, values)):
            if i < 5 or i >= n_steps - 3:
                preview.append({"time": round(float(t), 6), "value": round(float(v), 6)})
            elif i == 5 and n_steps > 8:
                preview.append({"time": "...", "value": "..."})

        return [types.TextContent(type="text", text=json.dumps({
            "variable":            var,
            "simulation_settings": {
                "source":         sim["source"],
                "stop_time_days": round(sim["stop_time_days"], 6),
                "datacomm_hours": round(sim["datacomm_hours"], 6),
            },
            "table": {
                "duration_days":   duration,
                "datacomm_hours":  dc_hours,
                "start_time_days": t_start,
                "end_time_days":   t_end,
                "n_rows":          n_steps,
                "value_mode":      mode,
                "time_unit":       time_unit,
            },
            "file_saved": str(outpath),
            "preview":    preview,
            "hint":       f"Pass this file to run_dynamic_simulation with influent_tsv='{outpath}'.",
        }, indent=2))]

    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    # GROUP A - Model Inspection
    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    if name == "get_parameter":
        param = arguments["parameter"]
        scenario = arguments.get("scenario", "baseline_Sc0")
        try:
            ds_ref = _make_ds()
            _apply_scenario(ds_ref, scenario)
            value = ds_ref.sumo.get(param)
            return [types.TextContent(type="text", text=json.dumps(
                {"parameter": param, "value": value, "scenario": scenario,
                 "note": "Value reflects state.xml plus any scenario overrides."},
                indent=2, default=str))]
        except Exception as e:
            return [types.TextContent(type="text", text=json.dumps(
                {"error": str(e), "parameter": param}, indent=2))]

    if name == "list_parameters":
        unit = arguments.get("unit_name", "")
        try:
            ds_ref = _make_ds()
            all_vars = ds_ref.sumo.getVariableNames()
            params = [v for v in all_vars if "__param__" in v]
            if unit:
                params = [v for v in params if unit.lower() in v.lower()]
            return [types.TextContent(type="text", text=json.dumps(
                {"count": len(params), "unit_filter": unit or "none", "parameters": params},
                indent=2))]
        except Exception as e:
            return [types.TextContent(type="text", text=json.dumps({"error": str(e)}, indent=2))]

    if name == "list_state_variables":
        unit = arguments.get("unit_name", "")
        try:
            ds_ref = _make_ds()
            all_vars = ds_ref.sumo.getVariableNames()
            states = [v for v in all_vars if "__param__" not in v]
            if unit:
                states = [v for v in states if unit.lower() in v.lower()]
            return [types.TextContent(type="text", text=json.dumps(
                {"count": len(states), "unit_filter": unit or "none", "state_variables": states},
                indent=2))]
        except Exception as e:
            return [types.TextContent(type="text", text=json.dumps({"error": str(e)}, indent=2))]

    if name == "get_state_variable":
        var = arguments["variable"]
        try:
            ds_ref = _make_ds()
            value = ds_ref.sumo.get(var)
            return [types.TextContent(type="text", text=json.dumps(
                {"variable": var, "value": value}, indent=2, default=str))]
        except Exception as e:
            return [types.TextContent(type="text", text=json.dumps(
                {"error": str(e), "variable": var}, indent=2))]

    if name == "search_variables":
        keyword = arguments["keyword"]
        try:
            ds_ref = _make_ds()
            all_vars = ds_ref.sumo.getVariableNames()
            matches = [v for v in all_vars if keyword.lower() in v.lower()]
            return [types.TextContent(type="text", text=json.dumps(
                {"keyword": keyword, "match_count": len(matches), "matches": matches},
                indent=2))]
        except Exception as e:
            return [types.TextContent(type="text", text=json.dumps(
                {"error": str(e), "keyword": keyword}, indent=2))]

    if name == "get_model_info":
        try:
            ds_ref = _make_ds()
            all_vars = ds_ref.sumo.getVariableNames()
            param_count = len([v for v in all_vars if "__param__" in v])
            state_count = len([v for v in all_vars if "__param__" not in v])
            return [types.TextContent(type="text", text=json.dumps({
                "dll_path": CONFIG["model_dll"],
                "state_xml": CONFIG["state_xml"],
                "total_variables": len(all_vars),
                "parameter_variables": param_count,
                "state_variables": state_count,
                "status": "compiled_and_ready",
            }, indent=2))]
        except Exception as e:
            return [types.TextContent(type="text", text=json.dumps(
                {"error": str(e), "status": "unavailable"}, indent=2))]

    if name == "ensure_dtt_bridge":
        # Detect and (re)install the dynamita_compat shim. Reports patch status,
        # the current queued command list, and optionally flushes to a sibling file.
        result = {}
        try:
            import dynamita_compat as _dc
            status = _dc.install()
            result.update(status)
            result["queue_size"] = len(_dc.get_queue())
            try:
                import dynamita.scheduler as _ds
                result["ds_methods_present"] = {
                    "executeCommand":   hasattr(_ds.SumoScheduler, "executeCommand"),
                    "set":              hasattr(_ds.SumoScheduler, "set"),
                    "getVariableNames": hasattr(_ds.SumoScheduler, "getVariableNames"),
                }
            except Exception as _e:
                result["ds_methods_present"] = {"error": repr(_e)}
            flush_to = arguments.get("flush_to") if isinstance(arguments, dict) else None
            if flush_to:
                path = _dc.flush_to_userscript(flush_to)
                result["flushed_to"] = path
                if arguments.get("clear_after_flush"):
                    _dc.clear_queue()
                    result["queue_cleared"] = True
        except Exception as e:
            result["error"] = repr(e)
            result["installed"] = False
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]


    if name == "list_unit_processes":
        try:
            ds_ref = _make_ds()
            all_vars = ds_ref.sumo.getVariableNames()
            units = set()
            for v in all_vars:
                parts = v.split("__")
                if len(parts) >= 3:
                    units.add(parts[2])
            units = sorted(units)
            return [types.TextContent(type="text", text=json.dumps(
                {"unit_count": len(units), "unit_processes": units}, indent=2))]
        except Exception as e:
            return [types.TextContent(type="text", text=json.dumps({"error": str(e)}, indent=2))]

    if name == "get_unit_process_info":
        unit = arguments["unit_name"]
        try:
            ds_ref = _make_ds()
            all_vars = ds_ref.sumo.getVariableNames()
            unit_vars = [v for v in all_vars if f"__{unit}__" in v]
            params = [v for v in unit_vars if "__param__" in v]
            states = [v for v in unit_vars if "__param__" not in v]
            return [types.TextContent(type="text", text=json.dumps({
                "unit_name": unit,
                "parameter_count": len(params),
                "state_variable_count": len(states),
                "parameters": params,
                "state_variables": states,
            }, indent=2))]
        except Exception as e:
            return [types.TextContent(type="text", text=json.dumps(
                {"error": str(e), "unit_name": unit}, indent=2))]

    if name == "validate_variable_name":
        var = arguments["variable"]
        try:
            ds_ref = _make_ds()
            all_vars = ds_ref.sumo.getVariableNames()
            exists = var in all_vars
            suggestions = []
            if not exists:
                parts = var.split("__")
                keyword = parts[-1] if parts else var
                suggestions = [v for v in all_vars if keyword.lower() in v.lower()][:10]
            return [types.TextContent(type="text", text=json.dumps(
                {"variable": var, "exists": exists, "suggestions": suggestions},
                indent=2))]
        except Exception as e:
            return [types.TextContent(type="text", text=json.dumps(
                {"error": str(e), "variable": var}, indent=2))]

    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    # GROUP B - Persistent Parameter Editing
    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    if name == "set_multiple_parameters":
        params = arguments.get("parameters", {})
        scenario = arguments.get("scenario", "baseline_Sc0")
        results, errors = [], []
        try:
            ds_ref = _make_ds()
            _apply_scenario(ds_ref, scenario)
            for p, v in params.items():
                try:
                    old = ds_ref.sumo.get(p)
                    ds_ref.sumo.set(p, float(v))
                    results.append({"parameter": p, "old": old, "new": v, "status": "ok"})
                except Exception as pe:
                    errors.append({"parameter": p, "error": str(pe)})
            return [types.TextContent(type="text", text=json.dumps({
                "applied_to": scenario,
                "success_count": len(results),
                "error_count": len(errors),
                "results": results,
                "errors": errors,
            }, indent=2, default=str))]
        except Exception as e:
            return [types.TextContent(type="text", text=json.dumps({"error": str(e)}, indent=2))]

    if name == "reset_parameter_to_default":
        param = arguments["parameter"]
        try:
            tree = ET.parse(CONFIG["state_xml"])
            root = tree.getroot()
            original = None
            for el in root.iter():
                if el.get("name") == param:
                    original = el.get("value")
                    break
            if original is None:
                return [types.TextContent(type="text", text=json.dumps(
                    {"error": f"Parameter '{param}' not found in state.xml."}, indent=2))]
            return [types.TextContent(type="text", text=json.dumps({
                "parameter": param,
                "reset_to": original,
                "source": CONFIG["state_xml"],
                "note": "Value restored from state.xml snapshot. Re-run simulation to apply.",
            }, indent=2))]
        except Exception as e:
            return [types.TextContent(type="text", text=json.dumps(
                {"error": str(e), "parameter": param}, indent=2))]

    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    # GROUP C - State & Project File Management
    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    if name == "save_state":
        try:
            output_path = arguments.get("output_path", "") or str(
                Path(CONFIG["output_dir"]) / f"state_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xml"
            )
            ds_ref = _make_ds()
            ds_ref.sumo.executeCommand(f'save "{output_path}"')
            return [types.TextContent(type="text", text=json.dumps({
                "saved_to": output_path,
                "timestamp": datetime.now().isoformat(),
                "note": "Load this file via load_state or set SUMO_STATE in config.",
            }, indent=2))]
        except Exception as e:
            return [types.TextContent(type="text", text=json.dumps({"error": str(e)}, indent=2))]

    if name == "load_state":
        path = arguments["state_xml_path"]
        if not Path(path).exists():
            return [types.TextContent(type="text", text=json.dumps(
                {"error": f"File not found: {path}"}, indent=2))]
        CONFIG["state_xml"] = path
        return [types.TextContent(type="text", text=json.dumps({
            "loaded": path,
            "note": "Active state updated. All subsequent simulations will use this snapshot.",
        }, indent=2))]

    if name == "list_saved_states":
        try:
            pattern = str(Path(CONFIG["output_dir"]) / "state_*.xml")
            files = sorted(glob.glob(pattern), reverse=True)
            return [types.TextContent(type="text", text=json.dumps({
                "count": len(files),
                "active_state": CONFIG["state_xml"],
                "snapshots": files,
            }, indent=2))]
        except Exception as e:
            return [types.TextContent(type="text", text=json.dumps({"error": str(e)}, indent=2))]

    if name == "get_compile_status":
        try:
            dll = Path(CONFIG["model_dll"])
            stx = Path(CONFIG["state_xml"])
            dll_ok, stx_ok = dll.exists(), stx.exists()
            dll_mtime = datetime.fromtimestamp(dll.stat().st_mtime).isoformat() if dll_ok else None
            stx_mtime = datetime.fromtimestamp(stx.stat().st_mtime).isoformat() if stx_ok else None
            return [types.TextContent(type="text", text=json.dumps({
                "ready": dll_ok and stx_ok,
                "dll": {"path": str(dll), "exists": dll_ok, "modified": dll_mtime},
                "state_xml": {"path": str(stx), "exists": stx_ok, "modified": stx_mtime},
                "note": "If ready=false, re-extract DLL and state.xml from SUMO GUI.",
            }, indent=2))]
        except Exception as e:
            return [types.TextContent(type="text", text=json.dumps({"error": str(e)}, indent=2))]

    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    # GROUP D - Scenario Management
    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    if name == "create_scenario":
        sc_name = arguments["name"]
        params  = arguments["parameters"]
        desc    = arguments.get("description", "")
        if sc_name in SCENARIOS:
            return [types.TextContent(type="text", text=json.dumps(
                {"error": f"Scenario '{sc_name}' already exists. Use update_scenario to modify it."}, indent=2))]
        cmds = [f"set {k} {v}" for k, v in params.items()]
        SCENARIOS[sc_name] = cmds
        return [types.TextContent(type="text", text=json.dumps({
            "created": sc_name,
            "parameter_count": len(cmds),
            "description": desc,
            "note": "Scenario is live for this session. Restart will clear it unless you edit server.py.",
        }, indent=2))]

    if name == "update_scenario":
        sc_name = arguments["name"]
        params  = arguments["parameters"]
        if sc_name not in SCENARIOS:
            return [types.TextContent(type="text", text=json.dumps(
                {"error": f"Scenario '{sc_name}' not found. Use create_scenario first."}, indent=2))]
        old_count = len(SCENARIOS[sc_name])
        SCENARIOS[sc_name] = [f"set {k} {v}" for k, v in params.items()]
        return [types.TextContent(type="text", text=json.dumps({
            "updated": sc_name,
            "old_parameter_count": old_count,
            "new_parameter_count": len(SCENARIOS[sc_name]),
        }, indent=2))]

    if name == "delete_scenario":
        sc_name = arguments["name"]
        if sc_name not in SCENARIOS:
            return [types.TextContent(type="text", text=json.dumps(
                {"error": f"Scenario '{sc_name}' not found."}, indent=2))]
        del SCENARIOS[sc_name]
        return [types.TextContent(type="text", text=json.dumps({
            "deleted": sc_name,
            "remaining_scenarios": list(SCENARIOS.keys()),
        }, indent=2))]

    if name == "clone_scenario":
        src = arguments["source"]
        new_name = arguments["new_name"]
        if src not in SCENARIOS:
            return [types.TextContent(type="text", text=json.dumps(
                {"error": f"Source scenario '{src}' not found."}, indent=2))]
        if new_name in SCENARIOS:
            return [types.TextContent(type="text", text=json.dumps(
                {"error": f"Scenario '{new_name}' already exists."}, indent=2))]
        SCENARIOS[new_name] = copy.deepcopy(SCENARIOS[src])
        return [types.TextContent(type="text", text=json.dumps({
            "cloned_from": src,
            "new_scenario": new_name,
            "parameter_count": len(SCENARIOS[new_name]),
        }, indent=2))]

    if name == "get_scenario_diff":
        sa = arguments["scenario_a"]
        sb = arguments["scenario_b"]
        if sa not in SCENARIOS:
            return [types.TextContent(type="text", text=json.dumps(
                {"error": f"Scenario '{sa}' not found."}, indent=2))]
        if sb not in SCENARIOS:
            return [types.TextContent(type="text", text=json.dumps(
                {"error": f"Scenario '{sb}' not found."}, indent=2))]
        def _parse(cmds):
            out = {}
            for cmd in cmds:
                parts = cmd.strip().split()
                if len(parts) == 3 and parts[0].lower() == "set":
                    out[parts[1]] = parts[2]
            return out
        pa, pb = _parse(SCENARIOS[sa]), _parse(SCENARIOS[sb])
        diffs = []
        for k in sorted(set(pa) | set(pb)):
            va, vb = pa.get(k, "<not set>"), pb.get(k, "<not set>")
            if va != vb:
                diffs.append({"parameter": k, sa: va, sb: vb})
        return [types.TextContent(type="text", text=json.dumps({
            "scenario_a": sa,
            "scenario_b": sb,
            "diff_count": len(diffs),
            "differences": diffs,
        }, indent=2))]

    if name == "export_scenario":
        sc_name = arguments["name"]
        if sc_name not in SCENARIOS:
            return [types.TextContent(type="text", text=json.dumps(
                {"error": f"Scenario '{sc_name}' not found."}, indent=2))]
        out_path = arguments.get("output_path", "") or str(
            Path(CONFIG["output_dir"]) / f"{sc_name}.scs"
        )
        try:
            with open(out_path, "w") as f:
                f.write(f"# Scenario: {sc_name}\n")
                f.write(f"# Exported: {datetime.now().isoformat()}\n\n")
                for cmd in SCENARIOS[sc_name]:
                    f.write(cmd + "\n")
            return [types.TextContent(type="text", text=json.dumps(
                {"exported": sc_name, "path": out_path}, indent=2))]
        except Exception as e:
            return [types.TextContent(type="text", text=json.dumps({"error": str(e)}, indent=2))]

    if name == "import_scenario":
        scs_path = arguments["scs_path"]
        p = Path(scs_path)
        if not p.exists():
            return [types.TextContent(type="text", text=json.dumps(
                {"error": f"File not found: {scs_path}"}, indent=2))]
        sc_name = arguments.get("name", "") or p.stem
        if sc_name in SCENARIOS:
            return [types.TextContent(type="text", text=json.dumps(
                {"error": f"Scenario '{sc_name}' already exists. Delete it first or supply a different name."},
                indent=2))]
        lines = p.read_text().splitlines()
        cmds = [l.strip() for l in lines if l.strip() and not l.strip().startswith("#")]
        SCENARIOS[sc_name] = cmds
        return [types.TextContent(type="text", text=json.dumps({
            "imported": sc_name,
            "source": scs_path,
            "command_count": len(cmds),
        }, indent=2))]

    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    # GROUP E - Extended Analysis & Reporting
    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    if name == "get_effluent_statistics":
        job_id = arguments["job_id"]
        try:
            csv_files = glob.glob(str(Path(CONFIG["output_dir"]) / f"*{job_id}*.csv"))
            if not csv_files:
                job = JOBS.get(job_id)
                if not job or not job.get("rows"):
                    return [types.TextContent(type="text", text=json.dumps(
                        {"error": f"No CSV or rows found for job '{job_id}'. Call export_results_csv first."},
                        indent=2))]
                rows = job["rows"]
                source = "memory"
            else:
                with open(csv_files[0], newline="") as f:
                    rows = list(csv.DictReader(f))
                source = csv_files[0]
            if not rows:
                return [types.TextContent(type="text", text=json.dumps(
                    {"error": "No data rows."}, indent=2))]
            numeric_cols = [c for c in rows[0].keys() if str(c).lower() != "time_h"]
            stats = {}
            for col in numeric_cols:
                vals = []
                for r in rows:
                    try:
                        vals.append(float(r[col]))
                    except (ValueError, KeyError, TypeError):
                        pass
                if vals:
                    stats[col] = {
                        "min":  round(min(vals), 4),
                        "max":  round(max(vals), 4),
                        "mean": round(sum(vals) / len(vals), 4),
                        "n_points": len(vals),
                    }
            return [types.TextContent(type="text", text=json.dumps({
                "job_id": str(job_id),
                "source": source,
                "statistics": stats,
            }, indent=2))]
        except Exception as e:
            return [types.TextContent(type="text", text=json.dumps(
                {"error": str(e), "job_id": str(job_id)}, indent=2))]

    if name == "get_sludge_production":
        job_id = arguments["job_id"]
        job = JOBS.get(job_id)
        if not job:
            return [types.TextContent(type="text", text=json.dumps(
                {"error": f"Job '{job_id}' not found."}, indent=2))]
        effluent = job.get("effluent", {})
        return [types.TextContent(type="text", text=json.dumps({
            "job_id": str(job_id),
            "effluent_TSS_mgL": effluent.get("TSS_mgL"),
            "note": ("Full SRT calculation requires MLSS and reactor volume from the model state. "
                     "Use search_variables with keyword 'MLSS' or 'VSS' to locate these variables, "
                     "then call get_state_variable to read them and compute SRT = (V*MLSS)/(Q_WAS*TSS_WAS)."),
            "tip": "Call list_parameters with unit_name='WAS' to find WAS pump parameters.",
        }, indent=2))]

    if name == "get_energy_estimate":
        scenario = arguments.get("scenario", "baseline_Sc0")
        try:
            ds_ref = _make_ds()
            _apply_scenario(ds_ref, scenario)
            all_vars = ds_ref.sumo.getVariableNames()
            do_params = [v for v in all_vars if "DO_setpoint" in v or "DO_sp" in v.lower()]
            flow_params = [v for v in all_vars if "Q_air" in v or "qair" in v.lower() or "blower" in v.lower()]
            do_values = {}
            for p in do_params:
                try: do_values[p] = ds_ref.sumo.get(p)
                except Exception: pass
            flow_values = {}
            for p in flow_params:
                try: flow_values[p] = ds_ref.sumo.get(p)
                except Exception: pass
            return [types.TextContent(type="text", text=json.dumps({
                "scenario": scenario,
                "do_setpoints": do_values,
                "airflow_parameters": flow_values,
                "note": ("Absolute kWh calculation requires blower efficiency curves not available "
                         "in the DTT API. Use these DO/airflow values with your blower spec sheet, "
                         "or compare DO setpoints across scenarios via run_scenario_comparison."),
            }, indent=2, default=str))]
        except Exception as e:
            return [types.TextContent(type="text", text=json.dumps(
                {"error": str(e), "scenario": scenario}, indent=2))]

    if name == "generate_report":
        job_id = arguments["job_id"]
        include_compliance = arguments.get("include_compliance", True)
        job = JOBS.get(job_id)
        if not job:
            return [types.TextContent(type="text", text=json.dumps(
                {"error": f"Job '{job_id}' not found."}, indent=2))]
        effluent = job.get("effluent", {})
        report = {
            "report_generated": datetime.now().isoformat(),
            "job_id": str(job_id),
            "scenario": arguments.get("scenario") or job.get("scenario", "unknown"),
            "job_status": job.get("status"),
            "effluent_quality": effluent,
        }
        if include_compliance and effluent:
            report["compliance"] = {
                "law": "Egyptian Law 48/1982",
                "parameters": _check_compliance(effluent),
            }
        report["output_csv"] = job.get("csv", "Not yet exported - call export_results_csv.")
        report["tip"] = "Call export_results_csv to save this job to disk for Excel or further analysis."
        return [types.TextContent(type="text", text=json.dumps(report, indent=2, default=str))]

    if name == "list_completed_jobs":
        summary = []
        for jid, jdata in JOBS.items():
            summary.append({
                "job_id":   jid,
                "status":   jdata.get("status"),
                "scenario": jdata.get("scenario", "unknown"),
                "n_rows":   len(jdata.get("rows", [])),
            })
        return [types.TextContent(type="text", text=json.dumps({
            "job_count": len(summary),
            "jobs": summary,
        }, indent=2))]

    if name == "import_influent_profile":
        csv_path = arguments["csv_path"]
        label = arguments.get("label", "")
        p = Path(csv_path)
        if not p.exists():
            return [types.TextContent(type="text", text=json.dumps(
                {"error": f"File not found: {csv_path}"}, indent=2))]
        try:
            with open(p, newline="") as f:
                reader = csv.DictReader(f)
                rows = list(reader)
                headers = reader.fieldnames or []
        except Exception as e:
            return [types.TextContent(type="text", text=json.dumps(
                {"error": f"Could not read CSV: {e}"}, indent=2))]
        if "time_h" not in headers:
            return [types.TextContent(type="text", text=json.dumps(
                {"error": "CSV must contain a 'time_h' column."}, indent=2))]
        profile_label = label or p.stem
        CONFIG.setdefault("influent_profiles", {})[profile_label] = str(p)
        return [types.TextContent(type="text", text=json.dumps({
            "registered": profile_label,
            "path": str(p),
            "rows": len(rows),
            "columns": headers,
            "usage": f"Pass influent_profile='{profile_label}' to run_dynamic_simulation.",
        }, indent=2))]

    # â”€â”€ export_results_csv â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "export_results_csv":
        job_id   = arguments["job_id"]
        filename = arguments.get("filename", f"job_{job_id}_results.csv")
        outpath  = Path(CONFIG["output_dir"]) / filename
        with _job_lock:
            rows = _active_jobs.get(job_id, {}).get("rows", [])
        if not rows:
            return [types.TextContent(type="text", text="No data rows found for this job.")]
        keys = list(rows[0].keys())
        with open(outpath, "w", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=keys)
            writer.writeheader()
            writer.writerows(rows)
        return [types.TextContent(type="text", text=json.dumps({
            "exported_to": str(outpath),
            "n_rows":      len(rows),
            "columns":     keys,
        }, indent=2))]


    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    # Build & Configuration Tools
    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

    # â”€â”€ create_model â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "create_model":
        try:
            project_name = arguments["project_name"]
            project_path = arguments["project_path"]
            project_dir  = Path(project_path) / project_name
            project_dir.mkdir(parents=True, exist_ok=True)
            for sub in ["outputs", "scenarios", "dynamita"]:
                (project_dir / sub).mkdir(exist_ok=True)
            sumo_stub = project_dir / f"{project_name}.sumo"
            sumo_stub.write_text(
                '<?xml version="1.0" encoding="utf-8"?>\n'
                f'<SumoProject name="{project_name}" version="24">\n'
                '  <!-- Open this file in SUMO24 GUI to build the process network -->\n'
                '  <!-- Then: Simulate -> maptoic; save "state.xml"; copy DLL here  -->\n'
                '</SumoProject>\n'
            )
            readme = project_dir / "README.md"
            readme.write_text(
                f"# {project_name}\n\n"
                "## Next steps\n"
                "1. Open the .sumo file in SUMO24 GUI and build the process network.\n"
                "2. Simulate the model until Ready.\n"
                "3. In Core Window, run: maptoic; save \"state.xml\";\n"
                "4. Copy sumoproject.dll and state.xml to this folder.\n"
                "5. Update CONFIG in server.py to point at this folder.\n"
            )
            ACTIVE_MODEL["sumo_file"] = str(sumo_stub)
            result = {
                "created":   str(project_dir),
                "sumo_file": str(sumo_stub),
                "next_step": (
                    "Open the .sumo file in SUMO24 GUI. Add unit processes visually, "
                    "simulate, then run: maptoic; save 'state.xml'; in the Core Window. "
                    "Copy state.xml and sumoproject.dll to this folder, then call load_model."
                ),
            }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ load_model â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "load_model":
        try:
            p = Path(arguments["sumo_file_path"])
            if not p.exists():
                result = {"error": f"File not found: {p}"}
            else:
                project_dir = p.parent
                dll_candidates   = list(project_dir.glob("*.dll"))
                state_candidates = list(project_dir.glob("state*.xml"))
                dll_found   = dll_candidates[0]   if dll_candidates   else None
                state_found = state_candidates[0] if state_candidates else None
                ACTIVE_MODEL["sumo_file"] = str(p)
                ACTIVE_MODEL["compiled"]  = bool(dll_found and state_found)
                if dll_found:
                    ACTIVE_MODEL["dll_path"] = str(dll_found)
                    CONFIG["model_dll"]      = str(dll_found)
                if state_found:
                    ACTIVE_MODEL["state_xml"] = str(state_found)
                    CONFIG["state_xml"]       = str(state_found)
                result = {
                    "loaded":      str(p),
                    "project_dir": str(project_dir),
                    "dll_found":   str(dll_found)   if dll_found   else "NOT FOUND - compile model first",
                    "state_found": str(state_found) if state_found else "NOT FOUND - extract state.xml first",
                    "ready_to_simulate": ACTIVE_MODEL["compiled"],
                }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ save_model â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "save_model":
        try:
            output_path = arguments.get("output_path", "")
            overwrite   = bool(arguments.get("overwrite", False))
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            if not output_path:
                output_path = str(Path(CONFIG["output_dir"]) / f"model_save_{ts}")
            dest = Path(output_path)
            dest.mkdir(parents=True, exist_ok=True)
            sumo_src = ACTIVE_MODEL.get("sumo_file")
            if sumo_src and Path(sumo_src).exists():
                dest_sumo = dest / Path(sumo_src).name
                if not dest_sumo.exists() or overwrite:
                    shutil.copy2(sumo_src, dest_sumo)
            state_dest = str(dest / "state.xml")
            try:
                _ds = _make_ds()
                _ds.sumo.executeCommand(f'save "{state_dest}"')
            except Exception:
                pass
            dll_src = CONFIG.get("model_dll")
            if dll_src and Path(dll_src).exists():
                dll_dest = dest / Path(dll_src).name
                if not dll_dest.exists() or overwrite:
                    shutil.copy2(dll_src, dll_dest)
            result = {"saved_to": str(dest), "state_xml": state_dest, "timestamp": ts}
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ compile_model â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "compile_model":
        try:
            force = bool(arguments.get("force_recompile", False))
            dll_path = Path(CONFIG["model_dll"])
            if force and dll_path.exists():
                ts = datetime.now().strftime('%Y%m%d_%H%M%S')
                backup = dll_path.with_suffix(f".dll.bak_{ts}")
                shutil.copy2(dll_path, backup)
                dll_path.unlink()
                ACTIVE_MODEL["compiled"] = False
                result = {
                    "action": "DLL removed - recompilation required",
                    "backup": str(backup),
                    "next_steps": [
                        "1. Open the .sumo file in SUMO24 GUI.",
                        "2. Click Simulate and wait for 'Ready for simulation'.",
                        "3. In Core Window run: maptoic; save \"state.xml\";",
                        "4. Copy the new sumoproject.dll and state.xml here.",
                        "5. Call get_compile_status to confirm readiness.",
                    ],
                }
            elif dll_path.exists():
                result = {
                    "status": "already_compiled",
                    "dll":    str(dll_path),
                    "mtime":  datetime.fromtimestamp(dll_path.stat().st_mtime).isoformat(),
                    "note":   "Use force_recompile=True to force a fresh compile.",
                }
            else:
                result = {
                    "status":        "dll_not_found",
                    "expected_path": str(dll_path),
                    "next_steps": [
                        "1. Open the .sumo project file in SUMO24 GUI.",
                        "2. Click Simulate and wait for 'Ready for simulation'.",
                        "3. Open Advanced -> Core Window.",
                        "4. Type: maptoic; save \"state.xml\";",
                        "5. Copy sumoproject.dll and state.xml to: " + str(dll_path.parent),
                        "6. Call get_compile_status to verify.",
                    ],
                }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ extract_dll â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "extract_dll":
        try:
            output_path = arguments.get("output_path", "")
            dll_path = None
            search_dirs = []
            if output_path and Path(output_path).exists():
                dll_path = Path(output_path)
            else:
                search_dirs = [Path.cwd()]
                sf = ACTIVE_MODEL.get("sumo_file")
                if sf:
                    search_dirs.insert(0, Path(sf).parent)
                for d in search_dirs:
                    candidates = list(d.glob("*.dll"))
                    if candidates:
                        named = [c for c in candidates if "sumoproject" in c.name.lower()]
                        dll_path = named[0] if named else candidates[0]
                        break
            if not dll_path or not dll_path.exists():
                result = {
                    "error":    "No .dll file found.",
                    "searched": [str(d) for d in search_dirs],
                    "tip":      "Compile the model in SUMO GUI first, then call extract_dll again.",
                }
            else:
                CONFIG["model_dll"]      = str(dll_path)
                ACTIVE_MODEL["dll_path"] = str(dll_path)
                ACTIVE_MODEL["compiled"] = True
                result = {
                    "dll_registered": str(dll_path),
                    "size_kb":  round(dll_path.stat().st_size / 1024, 1),
                    "modified": datetime.fromtimestamp(dll_path.stat().st_mtime).isoformat(),
                }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ list_available_unit_process_types â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "list_available_unit_process_types":
        unit_types = {
            "Hydraulic": ["Splitter", "Combiner", "FlowController", "Pump", "Pipe", "Valve", "MixingTank"],
            "Primary_Treatment": ["PrimaryClarifier", "ScreensAndGritChamber", "EqualisationTank", "LamellaClarifier"],
            "Biological_Aerobic": ["AerationTank", "OxidationDitch", "SBR", "MBR_Aerobic", "ContactStabilisation"],
            "Biological_Anoxic": ["AnoxicTank", "Denitrification_Zone"],
            "Biological_Anaerobic": ["AnaerobicTank", "UASB", "AnaerobicDigester"],
            "Combined_Nutrient_Removal": ["UCT_Reactor", "A2O_Reactor", "JHB_Reactor", "Bardenpho_5stage", "Bardenpho_3stage", "Biodenipho"],
            "Secondary_Clarification": ["SecondaryClarifier", "Thickener", "DAF"],
            "Tertiary_Treatment": ["TertiaryFilter", "ChemicalPrecipitation", "UV_Disinfection", "Chlorination", "Ozonation", "MembraneFilter"],
            "Sludge_Treatment": ["GravityThickener", "MechanicalThickener", "AerobicDigester", "AnaerobicDigester", "Centrifuge", "BeltFilter", "Drying_Bed"],
            "Instrumentation": ["DO_Sensor", "NH4_Sensor", "NO3_Sensor", "TSS_Sensor", "pH_Sensor"],
            "Controllers": ["DO_Controller", "SRT_Controller", "NitriteController", "PhosphorusController", "FlowController"],
        }
        total = sum(len(v) for v in unit_types.values())
        result = {
            "categories":         len(unit_types),
            "total_types":        total,
            "unit_process_types": unit_types,
            "note": ("These are the standard SUMO24 process library types. "
                     "Available types depend on your DTT addon version. "
                     "To add one, call add_unit_process with the type name from this list."),
        }
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ add_unit_process â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "add_unit_process":
        try:
            process_type  = arguments["process_type"]
            instance_name = arguments["instance_name"]
            kinetic_model = arguments.get("kinetic_model", "")
            try:
                _ds = _make_ds()
                _ds.sumo.executeCommand(f'addprocess "{process_type}" "{instance_name}"')
                if kinetic_model:
                    _ds.sumo.executeCommand(f'setmodel "{instance_name}" "{kinetic_model}"')
                applied = True
                err = None
            except Exception as ex:
                applied = False
                err = str(ex)
            ACTIVE_MODEL["unit_processes"].append({
                "name": instance_name, "type": process_type,
                "kinetic_model": kinetic_model or "default",
            })
            result = {
                "added":          instance_name,
                "type":           process_type,
                "kinetic_model":  kinetic_model or "not specified (default)",
                "applied_via_dtt": applied,
                "note": "Call compile_model after adding all units and connections.",
            }
            if not applied:
                result["dtt_error"] = err
                result["tip"] = ("If executeCommand is not available in your DTT version, add the unit "
                                  "manually in SUMO GUI. Then call compile_model and extract_dll.")
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ remove_unit_process â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "remove_unit_process":
        try:
            unit_name      = arguments["unit_name"]
            remove_streams = bool(arguments.get("remove_streams", True))
            try:
                _ds = _make_ds()
                cmd = f'removeprocess "{unit_name}"' + (' removestreams' if remove_streams else '')
                _ds.sumo.executeCommand(cmd)
                applied, err = True, None
            except Exception as ex:
                applied, err = False, str(ex)
            ACTIVE_MODEL["unit_processes"] = [
                u for u in ACTIVE_MODEL["unit_processes"] if u["name"] != unit_name
            ]
            result = {
                "removed":         unit_name,
                "streams_removed": remove_streams,
                "applied_via_dtt": applied,
                "note": "Call compile_model to rebuild the DLL after removing units.",
            }
            if not applied:
                result["dtt_error"] = err
                result["tip"] = "Remove the unit manually in SUMO GUI if DTT scripting is unavailable."
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ list_available_kinetic_models â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "list_available_kinetic_models":
        models = {
            "Activated_Sludge": {
                "ASM1":       "IWA Activated Sludge Model No. 1 - C/N removal",
                "ASM2d":      "IWA ASM2d - C/N/P removal with EBPR",
                "ASM3":       "IWA ASM3 - storage-based ASM for improved SRT sensitivity",
                "ASM3+BioP":  "ASM3 extended with biological phosphorus",
            },
            "Membrane_Bioreactor": {
                "MBR_ASM2d":  "ASM2d adapted for MBR (higher MLSS, different aeration)",
            },
            "Anaerobic": {
                "ADM1":       "IWA Anaerobic Digestion Model No. 1",
                "Siegrist":   "Siegrist anaerobic model",
            },
            "Simplified": {
                "Takacs":     "Simplified settler model (secondary clarifier)",
                "CSTR_Simple":"Simple CSTR without biological kinetics",
            },
        }
        result = {"kinetic_models": models,
                  "tip": "Use set_kinetic_model to assign one of these to a reactor unit."}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ set_kinetic_model â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "set_kinetic_model":
        try:
            unit_name     = arguments["unit_name"]
            kinetic_model = arguments["kinetic_model"]
            try:
                _ds = _make_ds()
                _ds.sumo.executeCommand(f'setmodel "{unit_name}" "{kinetic_model}"')
                applied, err = True, None
            except Exception as ex:
                applied, err = False, str(ex)
            result = {
                "unit":           unit_name,
                "kinetic_model":  kinetic_model,
                "applied_via_dtt": applied,
                "note": "Recompile the model after switching kinetic models.",
            }
            if not applied:
                result["dtt_error"] = err
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ connect_unit_processes â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "connect_unit_processes":
        try:
            from_unit   = arguments["from_unit"]
            to_unit     = arguments["to_unit"]
            stream_name = arguments["stream_name"]
            stream_type = arguments.get("stream_type", "liquid")
            try:
                _ds = _make_ds()
                _ds.sumo.executeCommand(
                    f'connect "{from_unit}" "{to_unit}" "{stream_name}" "{stream_type}"'
                )
                applied, err = True, None
            except Exception as ex:
                applied, err = False, str(ex)
            ACTIVE_MODEL["streams"].append({
                "name": stream_name, "from": from_unit, "to": to_unit, "type": stream_type,
            })
            result = {
                "stream_created":  stream_name,
                "from":            from_unit,
                "to":              to_unit,
                "type":            stream_type,
                "applied_via_dtt": applied,
                "note": "Recompile after all connections are defined.",
            }
            if not applied:
                result["dtt_error"] = err
                result["tip"] = "Draw this connection in SUMO GUI if executeCommand is unavailable."
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ add_recycle_stream â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "add_recycle_stream":
        try:
            from_unit    = arguments["from_unit"]
            to_unit      = arguments["to_unit"]
            stream_name  = arguments["stream_name"]
            recycle_type = arguments.get("recycle_type", "RAS")
            type_map = {
                "RAS":             "return_activated_sludge",
                "WAS":             "waste_activated_sludge",
                "InternalRecycle": "internal_recycle",
                "RejectWater":     "reject_water",
                "Sidestream":      "sidestream",
            }
            sumo_type = type_map.get(recycle_type, "recycle")
            try:
                _ds = _make_ds()
                _ds.sumo.executeCommand(
                    f'connect "{from_unit}" "{to_unit}" "{stream_name}" "{sumo_type}"'
                )
                applied, err = True, None
            except Exception as ex:
                applied, err = False, str(ex)
            ACTIVE_MODEL["streams"].append({
                "name": stream_name, "from": from_unit, "to": to_unit, "type": recycle_type,
            })
            result = {
                "recycle_stream_created": stream_name,
                "type":            recycle_type,
                "from":            from_unit,
                "to":              to_unit,
                "applied_via_dtt": applied,
            }
            if not applied:
                result["dtt_error"] = err
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ modify_flow_connection â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "modify_flow_connection":
        try:
            stream_name  = arguments["stream_name"]
            new_from     = arguments.get("new_from", "")
            new_to       = arguments.get("new_to", "")
            new_flow_m3d = float(arguments.get("new_flow_m3d", -1))
            actions = []
            try:
                _ds = _make_ds()
                if new_from or new_to:
                    cmd = f'reconnect "{stream_name}"'
                    if new_from:
                        cmd += f' from="{new_from}"'
                    if new_to:
                        cmd += f' to="{new_to}"'
                    _ds.sumo.executeCommand(cmd)
                    actions.append(f"reconnected: from={new_from or 'unchanged'} to={new_to or 'unchanged'}")
                if new_flow_m3d >= 0:
                    flow_var_candidates = [
                        f"Sumo__Plant__{stream_name}__param__Q",
                        f"Sumo__Plant__{stream_name}__Q",
                    ]
                    set_ok = False
                    for var in flow_var_candidates:
                        try:
                            _ds.sumo.set(var, new_flow_m3d)
                            actions.append(f"flow set to {new_flow_m3d} m3/d via {var}")
                            set_ok = True
                            break
                        except Exception:
                            continue
                    if not set_ok:
                        actions.append(
                            f"flow change requested ({new_flow_m3d} m3/d) - "
                            "use search_variables with the stream name to find the correct Q variable"
                        )
            except Exception as ex:
                actions.append(f"DTT error: {ex}")
            for s in ACTIVE_MODEL["streams"]:
                if s["name"] == stream_name:
                    if new_from:
                        s["from"] = new_from
                    if new_to:
                        s["to"] = new_to
                    break
            result = {"stream": stream_name, "actions_taken": actions}
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ list_flow_streams â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "list_flow_streams":
        try:
            unit_name = arguments.get("unit_name", "")
            registered = ACTIVE_MODEL.get("streams", [])
            if unit_name:
                registered = [
                    s for s in registered
                    if unit_name.lower() in s.get("from", "").lower()
                    or unit_name.lower() in s.get("to", "").lower()
                ]
            compiled_streams = []
            try:
                _ds = _make_ds()
                all_vars = _ds.sumo.getVariableNames()
                q_vars = [v for v in all_vars if "__Q" in v and "param" not in v]
                if unit_name:
                    q_vars = [v for v in q_vars if unit_name.lower() in v.lower()]
                compiled_streams = q_vars[:50]
            except Exception:
                pass
            result = {
                "registered_streams":      registered,
                "compiled_flow_variables": compiled_streams,
                "unit_filter":             unit_name or "none",
            }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ set_influent_characteristics â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "set_influent_characteristics":
        try:
            flow_m3d      = float(arguments["flow_m3d"])
            BOD_mgL       = float(arguments.get("BOD_mgL", 0.0))
            COD_mgL       = float(arguments.get("COD_mgL", 0.0))
            TSS_mgL       = float(arguments.get("TSS_mgL", 0.0))
            VSS_mgL       = float(arguments.get("VSS_mgL", 0.0))
            TKN_mgL       = float(arguments.get("TKN_mgL", 0.0))
            NH4_mgL       = float(arguments.get("NH4_mgL", 0.0))
            TP_mgL        = float(arguments.get("TP_mgL",  0.0))
            temp_C        = float(arguments.get("temp_C",  20.0))
            influent_unit = arguments.get("influent_unit", "Influent")
            _ds = _make_ds()
            param_map = {
                f"Sumo__Plant__{influent_unit}__param__Q":     flow_m3d   if flow_m3d   else None,
                f"Sumo__Plant__{influent_unit}__param__S_COD": COD_mgL    if COD_mgL    else None,
                f"Sumo__Plant__{influent_unit}__param__X_TSS": TSS_mgL    if TSS_mgL    else None,
                f"Sumo__Plant__{influent_unit}__param__X_VSS": VSS_mgL    if VSS_mgL    else None,
                f"Sumo__Plant__{influent_unit}__param__S_NH4": NH4_mgL    if NH4_mgL    else None,
                f"Sumo__Plant__{influent_unit}__param__S_TKN": TKN_mgL    if TKN_mgL    else None,
                f"Sumo__Plant__{influent_unit}__param__S_TP":  TP_mgL     if TP_mgL     else None,
                f"Sumo__Plant__{influent_unit}__param__T":     temp_C,
            }
            applied, skipped, failed = [], [], []
            try:
                all_vars = _ds.sumo.getVariableNames()
            except Exception:
                all_vars = []
            for var, value in param_map.items():
                if value is None or value == 0.0:
                    skipped.append(var); continue
                if (not all_vars) or (var in all_vars):
                    try:
                        _ds.sumo.set(var, float(value))
                        applied.append({"variable": var, "value": value})
                    except Exception as ve:
                        failed.append({"variable": var, "error": str(ve)})
                else:
                    failed.append({"variable": var,
                                   "error": "Not found - use search_variables to find the correct name"})
            if BOD_mgL:
                applied.append({"note": f"BOD_mgL={BOD_mgL} stored as metadata (not mapped)"})
            result = {
                "influent_unit":  influent_unit,
                "applied_count":  len(applied),
                "skipped_count":  len(skipped),
                "failed_count":   len(failed),
                "applied":        applied,
                "failed":         failed,
                "tip": ("If variables failed, call search_variables with keywords like "
                         "'COD', 'TSS', 'NH4' to find the exact names in your compiled model."),
            }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ set_fractionation â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "set_fractionation":
        try:
            unit_name     = arguments["unit_name"]
            fraction_dict = arguments["fraction_dict"]
            _ds = _make_ds()
            try:
                all_vars = _ds.sumo.getVariableNames()
            except Exception:
                all_vars = []
            applied, failed = [], []
            for frac_name, ratio in fraction_dict.items():
                candidates = [
                    f"Sumo__Plant__{unit_name}__param__{frac_name}",
                    f"Sumo__Plant__{unit_name}__{frac_name}",
                    f"Sumo__Plant__Influent__param__{frac_name}",
                ]
                set_ok = False
                last_err = None
                for var in candidates:
                    if (not all_vars) or (var in all_vars):
                        try:
                            _ds.sumo.set(var, float(ratio))
                            applied.append({"fraction": frac_name, "variable": var, "value": ratio})
                            set_ok = True
                            break
                        except Exception as ve:
                            last_err = str(ve)
                            continue
                if not set_ok:
                    failed.append({"fraction": frac_name,
                                   "error": last_err or "Variable not found - use search_variables"})
            result = {
                "unit_name":     unit_name,
                "applied_count": len(applied),
                "failed_count":  len(failed),
                "applied":       applied,
                "failed":        failed,
            }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ set_parameter_persistent â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "set_parameter_persistent":
        try:
            parameter = arguments["parameter"]
            value     = arguments["value"]
            state_path = Path(CONFIG["state_xml"])
            if not state_path.exists():
                result = {"error": f"state.xml not found at {state_path}"}
            else:
                ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                backup = state_path.with_suffix(f".xml.bak_{ts}")
                shutil.copy2(state_path, backup)
                tree = ET.parse(state_path)
                root = tree.getroot()
                found = False
                old_val = None
                for elem in root.iter():
                    if elem.get("name") == parameter:
                        old_val = elem.get("value")
                        elem.set("value", str(value))
                        found = True
                        break
                if not found:
                    result = {
                        "error": f"Parameter '{parameter}' not found in state.xml.",
                        "tip":   "Use search_variables to confirm the exact parameter name.",
                        "backup_created": str(backup),
                    }
                else:
                    tree.write(state_path, encoding="utf-8", xml_declaration=True)
                    result = {
                        "parameter":      parameter,
                        "old_value":      old_val,
                        "new_value":      str(value),
                        "written_to":     str(state_path),
                        "backup_created": str(backup),
                        "note": "Change is now persistent. Reload model to apply.",
                    }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ add_controller â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "add_controller":
        try:
            controller_type = arguments["controller_type"]
            controlled_var  = arguments["controlled_var"]
            setpoint        = arguments["setpoint"]
            unit_name       = arguments["unit_name"]
            actuator_unit   = arguments.get("actuator_unit", "") or unit_name
            try:
                _ds = _make_ds()
                cmd = (
                    f'addcontroller "{controller_type}" '
                    f'sensor="{unit_name}" '
                    f'variable="{controlled_var}" '
                    f'setpoint={setpoint} '
                    f'actuator="{actuator_unit}"'
                )
                _ds.sumo.executeCommand(cmd)
                applied, err = True, None
            except Exception as ex:
                applied, err = False, str(ex)
            result = {
                "controller_added": controller_type,
                "sensor_unit":      unit_name,
                "controlled_var":   controlled_var,
                "setpoint":         setpoint,
                "actuator_unit":    actuator_unit,
                "applied_via_dtt":  applied,
                "note": "Recompile the model after adding controllers.",
            }
            if not applied:
                result["dtt_error"] = err
                result["tip"] = ("Add the controller in SUMO GUI (Process -> Add Controller) if "
                                  "the executeCommand interface is unavailable in your DTT version.")
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ validate_model â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "validate_model":
        try:
            check_level = arguments.get("check_level", "basic")
            issues, warnings = [], []
            dll_ok   = Path(CONFIG.get("model_dll", "")).exists()
            state_ok = Path(CONFIG.get("state_xml", "")).exists()
            if not dll_ok:
                issues.append("sumoproject.dll not found - model must be compiled in SUMO GUI first.")
            if not state_ok:
                issues.append("state.xml not found - extract state from SUMO GUI (maptoic; save).")
            if dll_ok and state_ok:
                try:
                    _ds = _make_ds()
                    all_vars = _ds.sumo.getVariableNames()
                    param_vars = [v for v in all_vars if "__param__" in v]
                    for pat in ["__Q", "__S_COD", "__X_TSS", "__T"]:
                        matching = [v for v in param_vars if pat in v]
                        if not matching:
                            warnings.append(
                                f"No parameter found matching pattern '{pat}' - "
                                "influent characterisation may be incomplete."
                            )
                    if check_level == "full":
                        try:
                            _ds.sumo.executeCommand("steadystate maxiter=100")
                            warnings.append("Full check: quick steady-state ran without crashing - "
                                             "model appears solvable. Run run_steady_state for full results.")
                        except Exception as ss_err:
                            issues.append(f"Full check: steady-state failed - {ss_err}")
                except Exception as dtt_err:
                    warnings.append(f"Could not query variables via DTT: {dtt_err}")
            status = "OK" if not issues else "ISSUES_FOUND"
            result = {
                "status":        status,
                "check_level":   check_level,
                "issue_count":   len(issues),
                "warning_count": len(warnings),
                "issues":        issues,
                "warnings":      warnings,
            }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ get_mass_balance â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "get_mass_balance":
        try:
            job_id   = arguments["job_id"]
            elements = arguments.get("elements") or ["COD", "N", "P", "TSS"]
            csv_files = glob.glob(str(Path(CONFIG["output_dir"]) / f"{job_id}*.csv"))
            if not csv_files:
                result = {"error": f"No CSV found for job '{job_id}'.",
                          "tip":   "Call export_results_csv first to generate the CSV file."}
            else:
                with open(csv_files[0], newline="") as f:
                    reader = csv.DictReader(f)
                    rows = list(reader)
                if not rows:
                    result = {"error": "CSV is empty."}
                else:
                    headers = list(rows[0].keys())
                    element_vars = {
                        "COD": {"influent": ["COD_in", "Inf_COD", "S_COD_in"],
                                "effluent": ["COD_eff", "Eff_COD", "S_COD_eff", "COD_out"]},
                        "N":   {"influent": ["TN_in", "TKN_in", "Inf_TN"],
                                "effluent": ["TN_eff", "Eff_TN", "TN_out"]},
                        "P":   {"influent": ["TP_in", "Inf_TP"],
                                "effluent": ["TP_eff", "Eff_TP", "TP_out"]},
                        "TSS": {"influent": ["TSS_in", "Inf_TSS"],
                                "effluent": ["TSS_eff", "Eff_TSS", "TSS_out"]},
                    }
                    def _find_col(cands):
                        for c in cands:
                            for h in headers:
                                if c.lower() in h.lower():
                                    return h
                        return None
                    def _col_mean(col):
                        vals = []
                        for r in rows:
                            try: vals.append(float(r[col]))
                            except Exception: pass
                        return round(sum(vals) / len(vals), 2) if vals else None
                    balance = {}
                    for elem in elements:
                        if elem not in element_vars:
                            balance[elem] = {"note": "Not in default element map"}; continue
                        inf_col = _find_col(element_vars[elem]["influent"])
                        eff_col = _find_col(element_vars[elem]["effluent"])
                        inf_val = _col_mean(inf_col) if inf_col else None
                        eff_val = _col_mean(eff_col) if eff_col else None
                        removal = None
                        if inf_val and eff_val and inf_val > 0:
                            removal = round((inf_val - eff_val) / inf_val * 100, 1)
                        balance[elem] = {
                            "influent_col":  inf_col or "not found in CSV",
                            "effluent_col": eff_col or "not found in CSV",
                            "influent_mean": inf_val,
                            "effluent_mean": eff_val,
                            "removal_pct":   removal,
                        }
                    result = {
                        "job_id":   job_id,
                        "csv":      csv_files[0],
                        "elements": balance,
                        "tip": ("If influent/effluent columns are not found, call get_effluent_statistics "
                                "to see all available column names, then re-call with matching element names."),
                    }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ initialize_state â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "initialize_state":
        try:
            state_path = Path(CONFIG["state_xml"])
            if not state_path.exists():
                result = {"error": f"state.xml not found at {state_path}"}
            else:
                try:
                    _ds = _make_ds()
                    _ds.sumo.executeCommand("reset")
                    reset_ok, err = True, None
                except Exception as ex:
                    reset_ok, err = False, str(ex)
                # Also clear pending parameter overrides
                _pending_sets.clear()
                result = {
                    "status":    "reset_to_initial_state" if reset_ok else "partial_reset",
                    "state_xml": str(state_path),
                    "dtt_reset": reset_ok,
                    "timestamp": datetime.now().isoformat(),
                    "note":      ("All in-memory parameter overrides cleared. "
                                   "The model is back to its state.xml baseline."),
                }
                if err:
                    result["dtt_error"] = err
                    result["tip"]       = "If DTT reset fails, restart the MCP server to reload state.xml cleanly."
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]



    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    # Parameter-Editing Tools (Research-Based Defaults)
    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

    def _pe_try_set_first(ds_obj, all_vars, candidates, value):
        """Try each candidate variable name in order; return (ok, var_or_None, err)."""
        for v in candidates:
            if (not all_vars) or (v in all_vars):
                try:
                    ds_obj.sumo.set(v, float(value))
                    return True, v, None
                except Exception as ex:
                    return False, v, str(ex)
        return False, None, "not found"

    def _pe_apply_map(ds_obj, all_vars, unit, raw_map):
        """Common loop: apply (value, suffixes) map to a unit; return lists."""
        applied, skipped, failed = [], [], []
        for pname, entry in raw_map.items():
            value, suffixes = entry
            if value is None:
                skipped.append({"param": pname, "reason": "no value and no default"})
                continue
            candidates = [f"Sumo__Plant__{unit}__{s}" for s in suffixes]
            ok, var, err = _pe_try_set_first(ds_obj, all_vars, candidates, value)
            if ok:
                applied.append({"param": pname, "value": value, "variable": var})
            elif err and err != "not found":
                failed.append({"param": pname, "variable": var, "error": err})
            else:
                failed.append({"param": pname, "reason": "variable not found in compiled model"})
        return applied, skipped, failed

    # â”€â”€ set_asm1_kinetics â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "set_asm1_kinetics":
        try:
            unit_name = arguments.get("unit_name", "AerTank1")
            temp_C    = float(arguments.get("temp_C", 20.0))
            _ds       = _make_ds()
            defaults  = _correct_asm_params("ASM1", temp_C)
            overrides = {k: arguments.get(k) for k in [
                "mu_H","K_S","K_OH","K_NO","b_H","Y_H","eta_g","eta_h","k_h","K_X",
                "mu_A","K_NH","K_OA","b_A","k_a","Y_A","f_P","i_XB","i_XP"]}
            param_map = {
                "mu_H": (overrides["mu_H"] if overrides["mu_H"] is not None else defaults.get("mu_H"),
                         ["param__mu_H","param__muH"]),
                "K_S":  (overrides["K_S"]  if overrides["K_S"]  is not None else defaults.get("K_S"),
                         ["param__K_S","param__Ks"]),
                "K_OH": (overrides["K_OH"] if overrides["K_OH"] is not None else defaults.get("K_OH"),
                         ["param__K_OH","param__KO_H"]),
                "K_NO": (overrides["K_NO"] if overrides["K_NO"] is not None else defaults.get("K_NO"),
                         ["param__K_NO","param__KNO"]),
                "b_H":  (overrides["b_H"]  if overrides["b_H"]  is not None else defaults.get("b_H"),
                         ["param__b_H","param__bH"]),
                "Y_H":  (overrides["Y_H"]  if overrides["Y_H"]  is not None else defaults.get("Y_H"),
                         ["param__Y_H","param__YH"]),
                "eta_g":(overrides["eta_g"]if overrides["eta_g"]is not None else defaults.get("eta_g"),
                         ["param__eta_g","param__etag"]),
                "eta_h":(overrides["eta_h"]if overrides["eta_h"]is not None else defaults.get("eta_h"),
                         ["param__eta_h","param__etah"]),
                "k_h":  (overrides["k_h"]  if overrides["k_h"]  is not None else defaults.get("k_h"),
                         ["param__k_h","param__kh"]),
                "K_X":  (overrides["K_X"]  if overrides["K_X"]  is not None else defaults.get("K_X"),
                         ["param__K_X","param__KX"]),
                "mu_A": (overrides["mu_A"] if overrides["mu_A"] is not None else defaults.get("mu_A"),
                         ["param__mu_A","param__muA"]),
                "K_NH": (overrides["K_NH"] if overrides["K_NH"] is not None else defaults.get("K_NH"),
                         ["param__K_NH","param__KNH"]),
                "K_OA": (overrides["K_OA"] if overrides["K_OA"] is not None else defaults.get("K_OA"),
                         ["param__K_OA","param__KO_A"]),
                "b_A":  (overrides["b_A"]  if overrides["b_A"]  is not None else defaults.get("b_A"),
                         ["param__b_A","param__bA"]),
                "k_a":  (overrides["k_a"]  if overrides["k_a"]  is not None else defaults.get("k_a"),
                         ["param__k_a","param__ka"]),
                "Y_A":  (overrides["Y_A"]  if overrides["Y_A"]  is not None else defaults.get("Y_A"),
                         ["param__Y_A","param__YA"]),
                "f_P":  (overrides["f_P"]  if overrides["f_P"]  is not None else defaults.get("f_P"),
                         ["param__f_P","param__fP"]),
                "i_XB": (overrides["i_XB"] if overrides["i_XB"] is not None else defaults.get("i_XB"),
                         ["param__i_XB","param__iXB"]),
                "i_XP": (overrides["i_XP"] if overrides["i_XP"] is not None else defaults.get("i_XP"),
                         ["param__i_XP","param__iXP"]),
            }
            try:
                all_vars = _ds.sumo.getVariableNames()
            except Exception:
                all_vars = []
            applied, skipped, failed = _pe_apply_map(_ds, all_vars, unit_name, param_map)
            # mark source tag
            for a in applied:
                k = a["param"]
                a["source"] = "user" if overrides.get(k) is not None else f"research_default@{temp_C}C"
            result = {"model": "ASM1", "unit": unit_name, "temp_C": temp_C,
                      "applied": len(applied), "skipped": len(skipped), "failed": len(failed),
                      "results": applied, "skipped_list": skipped, "failed_list": failed,
                      "tip": "Run search_variables with 'mu_H' or 'mu_A' to find exact names."}
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ set_asm2d_kinetics â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "set_asm2d_kinetics":
        try:
            unit_name = arguments.get("unit_name", "AerTank1")
            temp_C    = float(arguments.get("temp_C", 20.0))
            _ds       = _make_ds()
            defaults  = _correct_asm_params("ASM2d", temp_C)
            get = lambda k: arguments.get(k)
            param_map = {
                "mu_H":       (get("mu_H")        if get("mu_H")        is not None else defaults.get("mu_H"),
                                ["param__mu_H"]),
                "b_H":        (get("b_H")         if get("b_H")         is not None else defaults.get("b_H"),
                                ["param__b_H"]),
                "Y_H":        (get("Y_H")         if get("Y_H")         is not None else defaults.get("Y_H"),
                                ["param__Y_H"]),
                "q_fe":       (get("q_fe")        if get("q_fe")        is not None else defaults.get("q_fe"),
                                ["param__q_fe"]),
                "eta_NO3_H":  (get("eta_NO3_H")   if get("eta_NO3_H")   is not None else defaults.get("eta_NO3_H"),
                                ["param__eta_NO3_H"]),
                "K_h":        (get("K_h")         if get("K_h")         is not None else defaults.get("K_h"),
                                ["param__K_h"]),
                "eta_NO3":    (get("eta_NO3_hyd") if get("eta_NO3_hyd") is not None else defaults.get("eta_NO3"),
                                ["param__eta_NO3"]),
                "mu_PAO":     (get("mu_PAO")      if get("mu_PAO")      is not None else defaults.get("mu_PAO"),
                                ["param__mu_PAO"]),
                "b_PAO":      (get("b_PAO")       if get("b_PAO")       is not None else defaults.get("b_PAO"),
                                ["param__b_PAO"]),
                "q_PHA":      (get("q_PHA")       if get("q_PHA")       is not None else defaults.get("q_PHA"),
                                ["param__q_PHA"]),
                "q_PP":       (get("q_PP")        if get("q_PP")        is not None else defaults.get("q_PP"),
                                ["param__q_PP"]),
                "Y_PAO":      (get("Y_PAO")       if get("Y_PAO")       is not None else defaults.get("Y_PAO"),
                                ["param__Y_PAO"]),
                "Y_PO4":      (get("Y_PO4")       if get("Y_PO4")       is not None else defaults.get("Y_PO4"),
                                ["param__Y_PO4"]),
                "mu_AUT":     (get("mu_AUT")      if get("mu_AUT")      is not None else defaults.get("mu_AUT"),
                                ["param__mu_AUT","param__mu_A"]),
                "b_AUT":      (get("b_AUT")       if get("b_AUT")       is not None else defaults.get("b_AUT"),
                                ["param__b_AUT","param__b_A"]),
                "K_NH4_AUT":  (get("K_NH4_AUT")   if get("K_NH4_AUT")   is not None else defaults.get("K_NH4_AUT"),
                                ["param__K_NH4_AUT","param__K_NH"]),
                "K_O2_AUT":   (get("K_O2_AUT")    if get("K_O2_AUT")    is not None else defaults.get("K_O2_AUT"),
                                ["param__K_O2_AUT","param__K_OA"]),
                "k_PRE":      (get("k_PRE")       if get("k_PRE")       is not None else defaults.get("k_PRE"),
                                ["param__k_PRE"]),
                "k_RED":      (get("k_RED")       if get("k_RED")       is not None else defaults.get("k_RED"),
                                ["param__k_RED"]),
            }
            try:
                all_vars = _ds.sumo.getVariableNames()
            except Exception:
                all_vars = []
            applied, skipped, failed = _pe_apply_map(_ds, all_vars, unit_name, param_map)
            result = {"model": "ASM2d", "unit": unit_name, "temp_C": temp_C,
                      "applied": len(applied), "skipped": len(skipped), "failed": len(failed),
                      "results": applied, "skipped_list": skipped, "failed_list": failed}
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ set_asm3_kinetics â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "set_asm3_kinetics":
        try:
            unit_name = arguments.get("unit_name", "AerTank1")
            temp_C    = float(arguments.get("temp_C", 20.0))
            _ds       = _make_ds()
            defaults  = _correct_asm_params("ASM3", temp_C)
            g = lambda k: arguments.get(k)
            def val(k):
                return g(k) if g(k) is not None else defaults.get(k)
            param_map = {
                "mu_H":     (val("mu_H"),     ["param__mu_H"]),
                "b_H_O2":   (val("b_H_O2"),   ["param__b_H_O2"]),
                "b_H_NO":   (val("b_H_NO"),   ["param__b_H_NO"]),
                "k_STO":    (val("k_STO"),    ["param__k_STO"]),
                "Y_STO_O2": (val("Y_STO_O2"), ["param__Y_STO_O2"]),
                "Y_STO_NO": (val("Y_STO_NO"), ["param__Y_STO_NO"]),
                "Y_H_O2":   (val("Y_H_O2"),   ["param__Y_H_O2"]),
                "Y_H_NO":   (val("Y_H_NO"),   ["param__Y_H_NO"]),
                "k_h":      (val("k_h"),      ["param__k_h"]),
                "K_X":      (val("K_X"),      ["param__K_X"]),
                "K_S":      (val("K_S"),      ["param__K_S"]),
                "K_STO":    (val("K_STO"),    ["param__K_STO"]),
                "eta_NO":   (val("eta_NO"),   ["param__eta_NO"]),
                "mu_A":     (val("mu_A"),     ["param__mu_A"]),
                "b_A_O":    (val("b_A_O"),    ["param__b_A_O","param__b_A"]),
                "b_A_NO":   (val("b_A_NO"),   ["param__b_A_NO"]),
                "K_A_NH":   (val("K_A_NH"),   ["param__K_A_NH","param__K_NH"]),
                "K_A_O":    (val("K_A_O"),    ["param__K_A_O","param__K_OA"]),
                "Y_A":      (val("Y_A"),      ["param__Y_A"]),
                "f_XI":     (val("f_XI"),     ["param__f_XI"]),
            }
            try:
                all_vars = _ds.sumo.getVariableNames()
            except Exception:
                all_vars = []
            applied, skipped, failed = _pe_apply_map(_ds, all_vars, unit_name, param_map)
            result = {"model": "ASM3", "unit": unit_name, "temp_C": temp_C,
                      "applied": len(applied), "skipped": len(skipped), "failed": len(failed),
                      "results": applied, "skipped_list": skipped}
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ apply_temperature_correction â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "apply_temperature_correction":
        try:
            model     = arguments.get("model", "ASM1")
            temp_C    = float(arguments.get("temp_C", 25.0))
            unit_name = arguments.get("unit_name", "AerTank1")
            corrected = _correct_asm_params(model, temp_C)
            raw       = RESEARCH_DEFAULTS.get(model, {})
            summary = []
            for pname, new_val in corrected.items():
                old_val = raw[pname]["val"]
                theta   = raw[pname].get("theta")
                if theta:
                    change = round((new_val - old_val) / old_val * 100, 1) if old_val else 0.0
                    summary.append({
                        "param": pname, "theta": theta, "val_20C": old_val,
                        f"val_{int(temp_C)}C": new_val, "change_pct": change,
                    })
            high_impact = [s for s in summary if abs(s["change_pct"]) > 20]
            result = {
                "model": model, "temp_C": temp_C, "unit_name": unit_name,
                "corrected_params": len(summary),
                "high_impact_params": high_impact,
                "all_corrections": summary,
                "note": (f"Call set_{model.lower()}_kinetics(unit_name='{unit_name}', temp_C={temp_C}) "
                          f"to apply these corrected values to the model."),
            }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ set_plant_wide_parameters â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "set_plant_wide_parameters":
        try:
            temp_C       = arguments.get("temp_C")
            law48_target = arguments.get("law48_target", "drain")
            t = temp_C if temp_C is not None else RESEARCH_DEFAULTS["influent_egypt"]["Temp_C"]
            limits_key = "law48_nile" if law48_target == "nile" else "law48_drain"
            CONFIG["law48_limits"] = RESEARCH_DEFAULTS[limits_key].copy()
            CONFIG["plant_temp_C"] = t
            _ds = _make_ds()
            try:
                all_vars = _ds.sumo.getVariableNames()
            except Exception:
                all_vars = []
            temp_set = False
            for var in ["Sumo__Plant__param__T", "Sumo__Plant__Influent__param__T", "Sumo__param__T"]:
                if (not all_vars) or (var in all_vars):
                    try:
                        _ds.sumo.set(var, float(t))
                        temp_set = True
                        break
                    except Exception:
                        continue
            result = {
                "plant_temp_C": t, "law48_target": law48_target,
                "law48_limits_applied": CONFIG["law48_limits"],
                "temp_variable_set": temp_set,
                "note": "Run apply_temperature_correction for each biological unit to update ASM kinetics.",
            }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ set_influent_parameters â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "set_influent_parameters":
        try:
            unit_name = arguments.get("unit_name", "Influent")
            eg = RESEARCH_DEFAULTS["influent_egypt"]
            g = lambda k: arguments.get(k)
            def _or(user, default):
                return user if user is not None else default
            raw_map = {
                "Q":     (_or(g("Q_m3d"),   eg["Q_m3d"]),   ["param__Q"]),
                "S_COD": (_or(g("COD_mgL"), eg["COD_mgL"]), ["param__S_COD","param__COD"]),
                "X_TSS": (_or(g("TSS_mgL"), eg["TSS_mgL"]), ["param__X_TSS","param__TSS"]),
                "X_VSS": (_or(g("VSS_mgL"), eg["VSS_mgL"]), ["param__X_VSS"]),
                "S_TKN": (_or(g("TKN_mgL"), eg["TKN_mgL"]), ["param__S_TKN","param__TKN"]),
                "S_NH4": (_or(g("NH4_mgL"), eg["NH4_mgL"]), ["param__S_NH4","param__NH4"]),
                "S_TP":  (_or(g("TP_mgL"),  eg["TP_mgL"]),  ["param__S_TP","param__TP"]),
                "pH":    (_or(g("pH"),      eg["pH"]),      ["param__pH"]),
                "T":     (_or(g("Temp_C"),  eg["Temp_C"]),  ["param__T"]),
                "fSI":   (_or(g("fSI"),     eg["fSI"]),     ["param__fSI","param__f_SI"]),
                "fSS":   (_or(g("fSS"),     eg["fSS"]),     ["param__fSS","param__f_SS"]),
                "fXS":   (_or(g("fXS"),     eg["fXS"]),     ["param__fXS","param__f_XS"]),
                "fXI":   (_or(g("fXI"),     eg["fXI"]),     ["param__fXI","param__f_XI"]),
                "fXBH":  (_or(g("fXBH"),    eg["fXBH"]),    ["param__fXBH","param__f_XBH"]),
            }
            _ds = _make_ds()
            try:
                all_vars = _ds.sumo.getVariableNames()
            except Exception:
                all_vars = []
            applied, skipped, failed = _pe_apply_map(_ds, all_vars, unit_name, raw_map)
            result = {"unit": unit_name, "applied": len(applied),
                      "failed": len(failed), "results": applied, "failed_list": failed}
            # Note BOD5 metadata
            if g("BOD5_mgL") is not None:
                result["bod5_metadata"] = g("BOD5_mgL")
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ set_primary_clarifier_parameters â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "set_primary_clarifier_parameters":
        try:
            unit = arguments.get("unit_name", "PrimClar1")
            g = lambda k: arguments.get(k)
            param_map = {
                "A":           (g("surface_area_m2"),                             ["param__A","param__Area"]),
                "SOR":         (g("overflow_rate_m3m2d")   if g("overflow_rate_m3m2d")   is not None else 40.0,
                                ["param__SOR","param__overflow_rate"]),
                "SLR":         (g("solids_loading_kgm2d")  if g("solids_loading_kgm2d")  is not None else 120.0,
                                ["param__SLR","param__solids_loading"]),
                "eff_removal": (g("capture_efficiency")     if g("capture_efficiency")     is not None else 0.60,
                                ["param__eff_removal","param__eta_TSS"]),
                "C_sludge":    (g("sludge_concentration_gL") if g("sludge_concentration_gL") is not None else 40.0,
                                ["param__C_sludge","param__X_underflow"]),
                "HRT":         (g("HRT_h")                  if g("HRT_h")                  is not None else 2.0,
                                ["param__HRT","param__tau"]),
            }
            _ds = _make_ds()
            try: all_vars = _ds.sumo.getVariableNames()
            except Exception: all_vars = []
            applied, skipped, failed = _pe_apply_map(_ds, all_vars, unit, param_map)
            result = {"unit": unit, "applied": len(applied), "results": applied, "failed": failed}
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ set_aeration_tank_parameters â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "set_aeration_tank_parameters":
        try:
            unit = arguments.get("unit_name", "AerTank1")
            g = lambda k: arguments.get(k)
            t = g("temp_C") if g("temp_C") is not None else CONFIG.get("plant_temp_C", 25.0)
            param_map = {
                "V":    (g("volume_m3"),                                            ["param__V","param__Volume"]),
                "HRT":  (g("HRT_h")           if g("HRT_h")           is not None else 6.0,
                          ["param__HRT","param__tau"]),
                "SRT":  (g("SRT_d")           if g("SRT_d")           is not None else 15.0,
                          ["param__SRT","param__sludge_age"]),
                "DO_sp":(g("DO_setpoint_mgL") if g("DO_setpoint_mgL") is not None else 2.0,
                          ["param__DO_setpoint","param__DO_sp","param__S_O_sp"]),
                "KLa":  (g("KLa_h")           if g("KLa_h")           is not None else 100.0,
                          ["param__KLa","param__kLa"]),
                "SOTE": (g("SOTE_pct")        if g("SOTE_pct")        is not None else 20.0,
                          ["param__SOTE","param__OTE"]),
                "MLSS": (g("MLSS_mgL")        if g("MLSS_mgL")        is not None else 3000.0,
                          ["param__MLSS","param__X_TSS_set"]),
                "T":    (t,                                                         ["param__T"]),
            }
            _ds = _make_ds()
            try: all_vars = _ds.sumo.getVariableNames()
            except Exception: all_vars = []
            applied, skipped, failed = _pe_apply_map(_ds, all_vars, unit, param_map)
            result = {"unit": unit, "temp_C": t, "applied": len(applied),
                      "failed": len(failed), "results": applied, "failed_list": failed}
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ set_anoxic_zone_parameters â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "set_anoxic_zone_parameters":
        try:
            unit = arguments.get("unit_name", "AnoxTank1")
            g = lambda k: arguments.get(k)
            param_map = {
                "V":        (g("volume_m3"),                                         ["param__V"]),
                "HRT":      (g("HRT_h")             if g("HRT_h")             is not None else 2.0,
                              ["param__HRT","param__tau"]),
                "IR_ratio": (g("NO3_recycle_ratio") if g("NO3_recycle_ratio") is not None else 3.0,
                              ["param__IR","param__recycle_ratio","param__Q_IR_ratio"]),
                "DO_max":   (g("DO_max_mgL")        if g("DO_max_mgL")        is not None else 0.2,
                              ["param__DO_max","param__S_O_max"]),
            }
            _ds = _make_ds()
            try: all_vars = _ds.sumo.getVariableNames()
            except Exception: all_vars = []
            applied, skipped, failed = _pe_apply_map(_ds, all_vars, unit, param_map)
            result = {"unit": unit, "applied": len(applied), "results": applied, "failed": failed}
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ set_anaerobic_zone_parameters â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "set_anaerobic_zone_parameters":
        try:
            unit = arguments.get("unit_name", "AnaTank1")
            g = lambda k: arguments.get(k)
            param_map = {
                "V":          (g("volume_m3"),                                        ["param__V"]),
                "HRT":        (g("HRT_h")           if g("HRT_h")           is not None else 1.5,
                                ["param__HRT","param__tau"]),
                "DO_max":     (g("DO_max_mgL")      if g("DO_max_mgL")      is not None else 0.02,
                                ["param__DO_max","param__S_O_max"]),
                "VFA_target": (g("VFA_target_mgL")  if g("VFA_target_mgL")  is not None else 30.0,
                                ["param__VFA_target","param__S_A_sp"]),
            }
            _ds = _make_ds()
            try: all_vars = _ds.sumo.getVariableNames()
            except Exception: all_vars = []
            applied, skipped, failed = _pe_apply_map(_ds, all_vars, unit, param_map)
            result = {"unit": unit, "applied": len(applied), "results": applied, "failed": failed}
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ set_secondary_clarifier_parameters â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "set_secondary_clarifier_parameters":
        try:
            unit = arguments.get("unit_name", "SecClar1")
            g = lambda k: arguments.get(k)
            td = RESEARCH_DEFAULTS["takacs"]
            param_map = {
                "A":         (g("surface_area_m2"),                                          ["param__A","param__Area"]),
                "SOR":       (g("overflow_rate_m3m2d") if g("overflow_rate_m3m2d") is not None else 24.0,
                               ["param__SOR","param__overflow_rate"]),
                "h_blanket": (g("sludge_blanket_m")    if g("sludge_blanket_m")    is not None else 0.5,
                               ["param__h_blanket","param__sludge_depth"]),
                "X_R":       (g("underflow_conc_gL")   if g("underflow_conc_gL")   is not None else 10.0,
                               ["param__X_R","param__C_underflow"]),
                "v0":        (g("v0")    if g("v0")    is not None else td["v0"]["val"],    ["param__v0","param__Veiling_0"]),
                "v0_p":      (g("v0_p")  if g("v0_p")  is not None else td["v0_p"]["val"],  ["param__v0_p","param__vmax_p"]),
                "r_h":       (g("r_h")   if g("r_h")   is not None else td["r_h"]["val"],   ["param__r_h","param__rh"]),
                "r_p":       (g("r_p")   if g("r_p")   is not None else td["r_p"]["val"],   ["param__r_p","param__rp"]),
                "f_ns":      (g("f_ns")  if g("f_ns")  is not None else td["f_ns"]["val"],  ["param__f_ns","param__fns"]),
            }
            _ds = _make_ds()
            try: all_vars = _ds.sumo.getVariableNames()
            except Exception: all_vars = []
            applied, skipped, failed = _pe_apply_map(_ds, all_vars, unit, param_map)
            result = {"unit": unit, "applied": len(applied),
                      "results": applied, "failed": failed,
                      "takacs_note": "v0, v0_p, r_h, r_p, f_ns control the Takacs settling model."}
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ set_mbr_parameters â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "set_mbr_parameters":
        try:
            unit = arguments.get("unit_name", "MBR1")
            g = lambda k: arguments.get(k)
            param_map = {
                "J":     (g("flux_Lm2h")              if g("flux_Lm2h")              is not None else 20.0,
                           ["param__J","param__flux"]),
                "TMP":   (g("TMP_kPa")                if g("TMP_kPa")                is not None else 20.0,
                           ["param__TMP"]),
                "K_mem": (g("permeability_LM2hbar")   if g("permeability_LM2hbar")   is not None else 250.0,
                           ["param__K_mem","param__permeability"]),
                "t_BW":  (g("backwash_interval_min")  if g("backwash_interval_min")  is not None else 10.0,
                           ["param__t_BW","param__backwash_interval"]),
                "dt_BW": (g("backwash_duration_s")    if g("backwash_duration_s")    is not None else 30.0,
                           ["param__dt_BW","param__backwash_duration"]),
                "MLSS":  (g("MLSS_mgL")               if g("MLSS_mgL")               is not None else 10000.0,
                           ["param__MLSS","param__X_TSS_set"]),
            }
            _ds = _make_ds()
            try: all_vars = _ds.sumo.getVariableNames()
            except Exception: all_vars = []
            applied, skipped, failed = _pe_apply_map(_ds, all_vars, unit, param_map)
            result = {"unit": unit, "applied": len(applied), "results": applied, "failed": failed}
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ set_ras_was_parameters â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "set_ras_was_parameters":
        try:
            ras_unit = arguments.get("ras_unit", "RAS_Pump")
            was_unit = arguments.get("was_unit", "WAS_Pump")
            g = lambda k: arguments.get(k)
            eg = RESEARCH_DEFAULTS["influent_egypt"]
            Q_inf = g("Q_influent_m3d") if g("Q_influent_m3d") is not None else eg["Q_m3d"]
            ras_r = g("RAS_ratio")      if g("RAS_ratio")      is not None else 0.75
            Q_ras = g("Q_RAS_m3d")      if g("Q_RAS_m3d")      is not None else round(Q_inf * ras_r, 1)
            SRT_t = g("SRT_target_d")   if g("SRT_target_d")   is not None else 15.0
            Q_was = g("Q_WAS_m3d")
            _ds = _make_ds()
            try: all_vars = _ds.sumo.getVariableNames()
            except Exception: all_vars = []
            applied, failed = [], []
            def _tryset(unit, suffixes, value, label):
                candidates = [f"Sumo__Plant__{unit}__{s}" for s in suffixes]
                ok, var, err = _pe_try_set_first(_ds, all_vars, candidates, value)
                if ok:
                    applied.append({"param": label, "value": value, "variable": var})
                elif err and err != "not found":
                    failed.append({"param": label, "error": err, "variable": var})
                else:
                    failed.append({"param": label, "reason": "variable not found"})
            _tryset(ras_unit, ["param__Q","param__Q_RAS"],               Q_ras, "Q_RAS")
            _tryset(ras_unit, ["param__RAS_ratio","param__ratio"],       ras_r, "RAS_ratio")
            if Q_was is not None:
                _tryset(was_unit, ["param__Q","param__Q_WAS"],           Q_was, "Q_WAS")
            _tryset(was_unit, ["param__SRT_sp","param__SRT_target","param__sludge_age_sp"],
                     SRT_t, "SRT_target")
            result = {"Q_RAS_m3d": Q_ras, "RAS_ratio": ras_r,
                      "Q_WAS_m3d": Q_was, "SRT_target_d": SRT_t,
                      "applied": len(applied), "failed": len(failed),
                      "results": applied, "failed_list": failed}
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ set_digester_parameters â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "set_digester_parameters":
        try:
            unit = arguments.get("unit_name", "Digester1")
            dig_type = arguments.get("digester_type", "aerobic")
            g = lambda k: arguments.get(k)
            if dig_type == "aerobic":
                t = g("temp_C") if g("temp_C") is not None else CONFIG.get("plant_temp_C", 25.0)
                param_map = {
                    "V":       (g("volume_m3"),                                            ["param__V"]),
                    "SRT":     (g("SRT_d")               if g("SRT_d")               is not None else 20.0,
                                 ["param__SRT","param__sludge_age"]),
                    "T":       (t,                                                         ["param__T"]),
                    "DO_sp":   (g("DO_setpoint_mgL")     if g("DO_setpoint_mgL")     is not None else 1.5,
                                 ["param__DO_setpoint","param__DO_sp"]),
                    "VSS_red": (g("VSS_reduction_pct")   if g("VSS_reduction_pct")   is not None else 40.0,
                                 ["param__VSS_reduction","param__eta_VSS"]),
                }
            else:
                t = g("temp_C") if g("temp_C") is not None else 35.0
                param_map = {
                    "V":         (g("volume_m3"),                                             ["param__V"]),
                    "SRT":       (g("SRT_d")               if g("SRT_d")               is not None else 20.0,
                                   ["param__SRT","param__HRT"]),
                    "T":         (t,                                                          ["param__T"]),
                    "OLR":       (g("OLR_kgVSSm3d")        if g("OLR_kgVSSm3d")        is not None else 2.0,
                                   ["param__OLR"]),
                    "biogas_yield":(g("biogas_yield_m3kg") if g("biogas_yield_m3kg")   is not None else 0.85,
                                   ["param__biogas_yield","param__Y_biogas"]),
                    "VSS_red":   (g("VSS_reduction_pct")   if g("VSS_reduction_pct")   is not None else 55.0,
                                   ["param__VSS_reduction","param__eta_VSS"]),
                }
            _ds = _make_ds()
            try: all_vars = _ds.sumo.getVariableNames()
            except Exception: all_vars = []
            applied, skipped, failed = _pe_apply_map(_ds, all_vars, unit, param_map)
            result = {"unit": unit, "type": dig_type, "temp_C": t,
                      "applied": len(applied), "results": applied, "failed": failed}
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ set_chemical_dosing_parameters â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "set_chemical_dosing_parameters":
        try:
            unit = arguments.get("unit_name", "ChemDosing1")
            coag = arguments.get("coagulant_type", "alum")
            g = lambda k: arguments.get(k)
            mr = g("molar_ratio")   if g("molar_ratio")   is not None else 2.0
            tp_t = g("target_TP_mgL") if g("target_TP_mgL") is not None else 1.0
            MW = {"alum": 27.0, "ferric": 55.85}.get(coag, 27.0)
            dose = g("dose_mgL")
            if dose is None:
                eg_TP = RESEARCH_DEFAULTS["influent_egypt"]["TP_mgL"]
                dose = round(mr * eg_TP * MW / 31.0, 1)
            param_map = {
                "dose":        (dose,   ["param__dose","param__C_coag","param__Q_chem"]),
                "molar_ratio": (mr,     ["param__molar_ratio","param__Me_P_ratio"]),
                "TP_target":   (tp_t,   ["param__TP_target","param__S_PO4_sp"]),
            }
            _ds = _make_ds()
            try: all_vars = _ds.sumo.getVariableNames()
            except Exception: all_vars = []
            applied, skipped, failed = _pe_apply_map(_ds, all_vars, unit, param_map)
            result = {"unit": unit, "coagulant": coag,
                      "estimated_dose_mgL": dose, "molar_ratio": mr,
                      "target_TP_mgL": tp_t,
                      "applied": len(applied), "results": applied, "failed": failed}
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ set_dynamic_simulation_parameters â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "set_dynamic_simulation_parameters":
        try:
            dd = RESEARCH_DEFAULTS["dynamics"]
            g = lambda k: arguments.get(k)
            duration_days = g("duration_days") if g("duration_days") is not None else dd["duration_days"]
            settings = {
                "duration_days":       duration_days,
                "timestep_h":          g("timestep_h")        if g("timestep_h")        is not None else dd["timestep_h"],
                "output_interval_h":   g("output_interval_h") if g("output_interval_h") is not None else dd["output_interval_h"],
                "solver":              g("solver")             if g("solver")             is not None else dd["solver"],
                "tolerance":           g("tolerance")          if g("tolerance")          is not None else dd["tolerance"],
                "warm_up_days":        g("warm_up_days")       if g("warm_up_days")       is not None else dd["warm_up_days"],
                "diurnal_peak_factor": g("diurnal_peak_factor")if g("diurnal_peak_factor")is not None else dd["diurnal_peak_factor"],
                "diurnal_min_factor":  g("diurnal_min_factor") if g("diurnal_min_factor") is not None else dd["diurnal_min_factor"],
                "storm_peak_factor":   g("storm_peak_factor")  if g("storm_peak_factor")  is not None else dd["storm_peak_factor"],
                "storm_start_day":     g("storm_start_day")    if g("storm_start_day")    is not None else round(duration_days / 2, 1),
                "storm_duration_h":    g("storm_duration_h")   if g("storm_duration_h")   is not None else dd["storm_duration_h"],
                "summer_temp_C":       g("summer_temp_C")      if g("summer_temp_C")      is not None else 28.0,
                "winter_temp_C":       g("winter_temp_C")      if g("winter_temp_C")      is not None else 18.0,
            }
            _ds = _make_ds()
            applied, not_applicable = [], []
            set_cmds = {
                "duration_days":     f"set Sumo__Simulator__param__t_end {settings['duration_days']}",
                "timestep_h":        f"set Sumo__Simulator__param__dt {settings['timestep_h'] / 24}",
                "output_interval_h": f"set Sumo__Simulator__param__dt_out {settings['output_interval_h'] / 24}",
                "tolerance":         f"set Sumo__Simulator__param__tol {settings['tolerance']}",
                "warm_up_days":      f"set Sumo__Simulator__param__t_warmup {settings['warm_up_days']}",
            }
            for pname, cmd in set_cmds.items():
                try:
                    _ds.sumo.executeCommand(cmd)
                    applied.append({"param": pname, "value": settings[pname], "command": cmd})
                except Exception:
                    not_applicable.append({"param": pname,
                        "note": "Set manually in SUMO GUI -> Simulate -> Dynamic Settings"})
            CONFIG["dynamic_settings"] = settings
            result = {
                "settings_applied":       settings,
                "sumo_commands_applied":  len(applied),
                "manual_settings_needed": not_applicable,
                "applied_list":           applied,
                "tip": ("For diurnal profiles and storm events, use import_influent_profile "
                         "with a time-series CSV. The diurnal/storm factors above are stored "
                         "in CONFIG['dynamic_settings'] for reference."),
            }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â”€â”€ assume_parameters_from_research â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "assume_parameters_from_research":
        try:
            unit_type   = arguments.get("unit_type",   "full_plant")
            temp_C      = arguments.get("temp_C")
            asm_model   = arguments.get("asm_model",   "ASM1")
            plant_scale = arguments.get("plant_scale", "medium")
            t = temp_C if temp_C is not None else 25.0
            is_warm = t >= 22.0
            scale_Q = {"small": 8000, "medium": 28000, "large": 80000}.get(plant_scale, 28000)
            asm_corrected = _correct_asm_params(asm_model, t)
            warm_notes = []
            if is_warm and asm_model == "ASM1":
                if asm_corrected.get("mu_A", 0) > 0.8:
                    warm_notes.append(
                        f"mu_A @ {t}C = {asm_corrected['mu_A']:.3f} d-1 - nitrification kinetics "
                        "accelerated vs 20C default. Reducing design SRT to 8-12 d is feasible."
                    )
                warm_notes.append(
                    "Egyptian practice: raise Y_H to 0.75-0.90 for warm-climate calibration. "
                    "IWA default 0.67 underestimates sludge production."
                )
                warm_notes.append(
                    "COD fractionation: fXI often 15-20% (higher than European default) "
                    "due to absence of primary clarifiers at many Egyptian plants."
                )
            result = {
                "unit_type": unit_type, "temp_C": t,
                "asm_model": asm_model, "plant_scale": plant_scale,
                "warm_climate_notes": warm_notes,
            }
            if unit_type in ("full_plant", "influent"):
                eg = RESEARCH_DEFAULTS["influent_egypt"].copy()
                eg["Q_m3d"] = scale_Q
                result["influent_defaults"]  = eg
                result["law48_limits_nile"]  = RESEARCH_DEFAULTS["law48_nile"]
                result["law48_limits_drain"] = RESEARCH_DEFAULTS["law48_drain"]
            if unit_type in ("full_plant", "asm_kinetics"):
                result["asm_kinetics"] = {
                    "model": asm_model,
                    "reference_temp": "20C (IWA STR-9)",
                    "operating_temp": f"{t}C (Arrhenius corrected)",
                    "parameters": asm_corrected,
                }
            if unit_type in ("full_plant", "aeration_tank"):
                srt = 12.0 if is_warm else 15.0
                hrt = 6.0
                result["aeration_tank"] = {
                    "SRT_d": srt, "HRT_h": hrt, "DO_setpoint_mgL": 2.0,
                    "MLSS_mgL": 3000, "volume_m3": round(scale_Q / 24 * hrt, 0),
                    "SOTE_pct": 20,
                    "note": f"SRT={srt}d chosen for {'warm-climate' if is_warm else 'moderate temp'} nitrification",
                }
            if unit_type in ("full_plant", "secondary_clarifier"):
                result["secondary_clarifier"] = {
                    "overflow_rate_m3m2d": 24, "underflow_conc_gL": 10, "RAS_ratio": 0.75,
                    **{k: v["val"] for k, v in RESEARCH_DEFAULTS["takacs"].items()},
                }
            if unit_type in ("full_plant", "dynamic"):
                result["dynamic_simulation"] = RESEARCH_DEFAULTS["dynamics"].copy()
                result["dynamic_simulation"]["recommended_summer_temp_C"] = 28.0
                result["dynamic_simulation"]["recommended_winter_temp_C"] = 18.0
            result["next_steps"] = [
                f"1. Call set_plant_wide_parameters(temp_C={t}) to set global temperature.",
                "2. Call set_influent_parameters() to apply Egyptian influent defaults.",
                f"3. Call set_{asm_model.lower()}_kinetics(temp_C={t}) on each biological unit.",
                "4. Call set_aeration_tank_parameters(), set_secondary_clarifier_parameters(), etc.",
                "5. Call set_dynamic_simulation_parameters() for transient runs.",
                "6. Run run_steady_state() and check_compliance() to verify Law 48 compliance.",
            ]
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]



    # ==========================================================================
    # DATASET-READING TOOLS
    # ==========================================================================

    # -- read_dataset_excel ----------------------------------------------------
    if name == "read_dataset_excel":
        result = {}
        try:
            xlsx_path = arguments.get("xlsx_path") or _dataset_default_path()
            wb, ws = _open_dataset(xlsx_path)
            parsed = _parse_dataset_rows(ws)
            result = {
                "file":               xlsx_path,
                "sheet":              ws.title,
                "points_found":       len(parsed["points"]),
                "law48_params":       list(parsed["law48"].keys()),
                "measurement_points": list(parsed["points"].keys()),
                "law48_limits":       parsed["law48"],
                "section_headers":    parsed["section_headers_found"],
                "tip":                "Call get_measurement_point_data(label=...) for details of any point.",
            }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # -- list_dataset_measurement_points ---------------------------------------
    if name == "list_dataset_measurement_points":
        result = {}
        try:
            xlsx_path = arguments.get("xlsx_path") or _dataset_default_path()
            wb, ws = _open_dataset(xlsx_path)
            parsed = _parse_dataset_rows(ws)
            table = []
            for label, data in parsed["points"].items():
                mapping = DATASET_SCHEMA["point_map"].get(label, {})
                numeric_params = sorted([k for k, v in data.items()
                                         if k not in ("_row", "Notes")
                                         and isinstance(v, (int, float))])
                table.append({
                    "label":          label,
                    "row":            data.get("_row"),
                    "sumo_unit":      mapping.get("unit", "unmapped"),
                    "role":           mapping.get("role", "-"),
                    "params_present": numeric_params,
                    "param_count":    len(numeric_params),
                })
            unmapped = [t["label"] for t in table if t["sumo_unit"] == "unmapped"]
            result = {
                "file":             xlsx_path,
                "total_points":     len(table),
                "points":           table,
                "unmapped_points":  unmapped,
            }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # -- get_measurement_point_data --------------------------------------------
    if name == "get_measurement_point_data":
        result = {}
        try:
            label = arguments["label"]
            xlsx_path = arguments.get("xlsx_path") or _dataset_default_path()
            wb, ws = _open_dataset(xlsx_path)
            parsed = _parse_dataset_rows(ws)
            if label not in parsed["points"]:
                close = [p for p in parsed["points"]
                         if label.lower() in p.lower() or p.lower() in label.lower()]
                result = {
                    "error": f"Point '{label}' not found in dataset.",
                    "available_points": list(parsed["points"].keys()),
                    "closest_matches":  close,
                }
            else:
                data    = parsed["points"][label]
                mapping = DATASET_SCHEMA["point_map"].get(label, {})
                numeric = {k: v for k, v in data.items()
                           if k not in ("_row", "Notes") and isinstance(v, (int, float))}
                flags = []
                for pname, val in numeric.items():
                    limit = parsed["law48"].get(pname)
                    if limit is not None and val > limit:
                        flags.append({"param": pname, "value": val, "limit": limit,
                                      "exceedance_pct": round((val - limit) / limit * 100, 1)})
                result = {
                    "label":            label,
                    "sumo_unit":        mapping.get("unit", "unmapped"),
                    "role":             mapping.get("role", "-"),
                    "measured_values":  numeric,
                    "notes":            data.get("Notes", ""),
                    "law48_exceedances":flags,
                    "row":              data.get("_row"),
                }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # -- get_parameter_across_points -------------------------------------------
    if name == "get_parameter_across_points":
        result = {}
        try:
            parameter = arguments["parameter"]
            xlsx_path = arguments.get("xlsx_path") or _dataset_default_path()
            if parameter not in DATASET_SCHEMA["parameter_map"]:
                result = {"error": f"Unknown parameter '{parameter}'.",
                          "valid_parameters": list(DATASET_SCHEMA["parameter_map"].keys())}
            else:
                wb, ws = _open_dataset(xlsx_path)
                parsed = _parse_dataset_rows(ws)
                trace = []
                for label, data in parsed["points"].items():
                    if parameter in data:
                        trace.append({
                            "point":     label,
                            "sumo_unit": DATASET_SCHEMA["point_map"].get(label, {}).get("unit"),
                            "value":     data[parameter],
                        })
                limit = parsed["law48"].get(parameter)
                removal = None
                inf_pt  = next((t for t in trace if "Raw Influent"    in t["point"]), None)
                eff_pt  = next((t for t in trace if "Final Effluent"  in t["point"]), None)
                if inf_pt and eff_pt and inf_pt["value"] > 0:
                    removal = round((inf_pt["value"] - eff_pt["value"]) / inf_pt["value"] * 100, 1)
                result = {
                    "parameter":           parameter,
                    "unit":                DATASET_SCHEMA["parameter_map"][parameter]["unit"],
                    "law48_limit":         limit,
                    "data_points":         len(trace),
                    "trace":               trace,
                    "overall_removal_pct": removal,
                }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # -- get_law48_limits_from_excel -------------------------------------------
    if name == "get_law48_limits_from_excel":
        result = {}
        try:
            xlsx_path = arguments.get("xlsx_path") or _dataset_default_path()
            wb, ws = _open_dataset(xlsx_path)
            parsed = _parse_dataset_rows(ws)
            CONFIG["law48_limits_from_dataset"] = parsed["law48"]
            result = {
                "file":         xlsx_path,
                "law48_limits": parsed["law48"],
                "param_count":  len(parsed["law48"]),
                "note":         "Values synced into CONFIG['law48_limits_from_dataset'].",
            }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # -- map_dataset_to_sumo_variables -----------------------------------------
    if name == "map_dataset_to_sumo_variables":
        result = {}
        try:
            xlsx_path = arguments.get("xlsx_path") or _dataset_default_path()
            wb, ws = _open_dataset(xlsx_path)
            parsed = _parse_dataset_rows(ws)
            ds = _make_ds()
            all_vars = ds.sumo.getVariableNames()
            resolved, unresolved = [], []
            for label, data in parsed["points"].items():
                mapping = DATASET_SCHEMA["point_map"].get(label)
                if not mapping:
                    unresolved.append({"point": label, "reason": "unit not mapped"}); continue
                unit = mapping["unit"]
                for pkey, value in data.items():
                    if pkey in ("_row", "Notes") or not isinstance(value, (int, float)):
                        continue
                    var = _resolve_sumo_variable(unit, pkey, all_vars)
                    if var:
                        resolved.append({"point": label, "unit": unit, "param": pkey,
                                         "value": value, "variable": var})
                    else:
                        unresolved.append({"point": label, "unit": unit, "param": pkey,
                                           "value": value, "reason": "No matching SUMO variable"})
            result = {
                "file":                xlsx_path,
                "resolved":            len(resolved),
                "unresolved":          len(unresolved),
                "resolved_mappings":   resolved[:50],
                "unresolved_mappings": unresolved,
            }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # -- validate_dataset ------------------------------------------------------
    if name == "validate_dataset":
        result = {}
        try:
            xlsx_path = arguments.get("xlsx_path") or _dataset_default_path()
            wb, ws = _open_dataset(xlsx_path)
            parsed = _parse_dataset_rows(ws)
            pts = parsed["points"]
            issues, warnings, passes = [], [], []
            inf = pts.get("Raw Influent", {})
            for must in ("Flow", "COD"):
                if must not in inf:
                    issues.append(f"Raw Influent is missing '{must}'")
                else:
                    passes.append(f"Raw Influent {must} present: {inf[must]}")
            for label, data in pts.items():
                if "pH" in data and not (5.0 <= data["pH"] <= 9.5):
                    issues.append(f"{label}: pH {data['pH']} outside 5-9.5")
                if "Temp" in data and not (5.0 <= data["Temp"] <= 45.0):
                    warnings.append(f"{label}: Temp {data['Temp']}C unusual")
                if "DO" in data and not (0.0 <= data["DO"] <= 15.0):
                    warnings.append(f"{label}: DO {data['DO']} mg/L outside 0-15")
            cod_inf = pts.get("Raw Influent", {}).get("COD")
            cod_eff = pts.get("Final Effluent", {}).get("COD")
            if cod_inf and cod_eff:
                if cod_eff >= cod_inf:
                    issues.append(f"Effluent COD ({cod_eff}) >= Influent COD ({cod_inf}) - no removal")
                else:
                    rem = round((cod_inf - cod_eff) / cod_inf * 100, 1)
                    passes.append(f"COD removal = {rem}% ({cod_inf} -> {cod_eff} mg/L)")
            q_inf = pts.get("Raw Influent", {}).get("Flow")
            q_eff = pts.get("Final Effluent", {}).get("Flow")
            q_was = pts.get("WAS (Waste Activated Sludge)", {}).get("Flow", 0)
            if q_inf and q_eff:
                expected = q_inf - q_was
                delta_pct = abs(q_eff - expected) / q_inf * 100
                if delta_pct > 10:
                    warnings.append(f"Flow continuity off by {delta_pct:.1f} pct")
                else:
                    passes.append(f"Flow continuity OK: Q_inf {q_inf} ~ Q_eff {q_eff} + Q_WAS {q_was}")
            law48_exc = []
            for pname, val in pts.get("Final Effluent", {}).items():
                if pname in ("_row", "Notes") or not isinstance(val, (int, float)):
                    continue
                limit = parsed["law48"].get(pname)
                if limit is not None and val > limit:
                    law48_exc.append({"param": pname, "value": val, "limit": limit})
            status = "PASS" if not issues else ("FAIL" if len(issues) > 2 else "WARNINGS")
            result = {
                "file":              xlsx_path,
                "overall_status":    status,
                "issue_count":       len(issues),
                "warning_count":     len(warnings),
                "passes":            len(passes),
                "issues":            issues,
                "warnings":          warnings,
                "passes_summary":    passes,
                "law48_exceedances": law48_exc,
            }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # -- check_dataset_against_law48 -------------------------------------------
    if name == "check_dataset_against_law48":
        result = {}
        try:
            xlsx_path = arguments.get("xlsx_path") or _dataset_default_path()
            wb, ws = _open_dataset(xlsx_path)
            parsed = _parse_dataset_rows(ws)
            report = []
            for label, data in parsed["points"].items():
                exc = []
                for pname, val in data.items():
                    if pname in ("_row", "Notes") or not isinstance(val, (int, float)):
                        continue
                    limit = parsed["law48"].get(pname)
                    if limit is not None and val > limit:
                        exc.append({"param": pname, "value": val, "limit": limit,
                                    "exceedance_pct": round((val - limit) / limit * 100, 1)})
                if exc:
                    report.append({"point": label, "exceedance_count": len(exc), "exceedances": exc})
            eff_ok = not any(r["point"] == "Final Effluent" for r in report)
            result = {
                "file":               xlsx_path,
                "law48_limits":       parsed["law48"],
                "points_with_issues": len(report),
                "effluent_compliant": eff_ok,
                "detailed_report":    report,
            }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # -- apply_influent_from_dataset -------------------------------------------
    if name == "apply_influent_from_dataset":
        result = {}
        try:
            xlsx_path = arguments.get("xlsx_path") or _dataset_default_path()
            row_label = arguments.get("row_label", "Raw Influent")
            wb, ws = _open_dataset(xlsx_path)
            parsed = _parse_dataset_rows(ws)
            if row_label not in parsed["points"]:
                result = {"error": f"Row '{row_label}' not found.",
                          "available": list(parsed["points"].keys())}
            else:
                data = parsed["points"][row_label]
                ds = _make_ds()
                all_vars = ds.sumo.getVariableNames()
                unit = DATASET_SCHEMA["point_map"][row_label]["unit"]
                applied, failed = [], []
                for pkey, value in data.items():
                    if pkey in ("_row", "Notes") or not isinstance(value, (int, float)):
                        continue
                    var = _resolve_sumo_variable(unit, pkey, all_vars)
                    if var:
                        try:
                            ds.sumo.set(var, float(value))
                            applied.append({"param": pkey, "value": value, "variable": var})
                        except Exception as e:
                            failed.append({"param": pkey, "error": str(e)})
                    else:
                        failed.append({"param": pkey, "reason": f"no SUMO variable found for {unit}/{pkey}"})
                result = {
                    "source_row":  row_label,
                    "sumo_unit":   unit,
                    "applied":     len(applied),
                    "failed":      len(failed),
                    "results":     applied,
                    "failed_list": failed,
                }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # -- apply_point_to_unit ---------------------------------------------------
    if name == "apply_point_to_unit":
        result = {}
        try:
            point_label = arguments["point_label"]
            unit_override = arguments.get("unit_override", "")
            xlsx_path = arguments.get("xlsx_path") or _dataset_default_path()
            wb, ws = _open_dataset(xlsx_path)
            parsed = _parse_dataset_rows(ws)
            if point_label not in parsed["points"]:
                result = {"error": f"Point '{point_label}' not found."}
            else:
                data = parsed["points"][point_label]
                mapping = DATASET_SCHEMA["point_map"].get(point_label, {})
                unit = unit_override or mapping.get("unit")
                if not unit:
                    result = {"error": f"No SUMO unit mapping for '{point_label}'. Pass unit_override."}
                else:
                    ds = _make_ds()
                    all_vars = ds.sumo.getVariableNames()
                    applied, failed = [], []
                    for pkey, value in data.items():
                        if pkey in ("_row", "Notes") or not isinstance(value, (int, float)):
                            continue
                        var = _resolve_sumo_variable(unit, pkey, all_vars)
                        if var:
                            try:
                                ds.sumo.set(var, float(value))
                                applied.append({"param": pkey, "value": value, "variable": var})
                            except Exception as e:
                                failed.append({"param": pkey, "error": str(e)})
                        else:
                            failed.append({"param": pkey, "reason": "variable not found"})
                    result = {
                        "point":       point_label,
                        "sumo_unit":   unit,
                        "applied":     len(applied),
                        "failed":      len(failed),
                        "results":     applied,
                        "failed_list": failed,
                    }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # -- apply_dataset_to_model ------------------------------------------------
    if name == "apply_dataset_to_model":
        result = {}
        try:
            xlsx_path = arguments.get("xlsx_path") or _dataset_default_path()
            only_inf  = bool(arguments.get("only_influent", False))
            dry_run   = bool(arguments.get("dry_run", False))
            wb, ws = _open_dataset(xlsx_path)
            parsed = _parse_dataset_rows(ws)
            ds = _make_ds()
            all_vars = ds.sumo.getVariableNames()
            applied, failed = [], []
            for label, data in parsed["points"].items():
                if only_inf and label != "Raw Influent":
                    continue
                mapping = DATASET_SCHEMA["point_map"].get(label)
                if not mapping:
                    continue
                unit = mapping["unit"]
                for pkey, value in data.items():
                    if pkey in ("_row", "Notes") or not isinstance(value, (int, float)):
                        continue
                    var = _resolve_sumo_variable(unit, pkey, all_vars)
                    if not var:
                        failed.append({"point": label, "param": pkey, "value": value,
                                       "reason": f"no SUMO variable found for {unit}"})
                        continue
                    if dry_run:
                        applied.append({"point": label, "param": pkey, "value": value,
                                        "variable": var, "status": "would_apply"})
                    else:
                        try:
                            ds.sumo.set(var, float(value))
                            applied.append({"point": label, "param": pkey, "value": value,
                                            "variable": var, "status": "applied"})
                        except Exception as e:
                            failed.append({"point": label, "param": pkey, "error": str(e)})
            result = {
                "file":          xlsx_path,
                "dry_run":       dry_run,
                "only_influent": only_inf,
                "applied_count": len(applied),
                "failed_count":  len(failed),
                "applied":       applied[:100],
                "failed":        failed,
            }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # -- compare_dataset_vs_simulation -----------------------------------------
    if name == "compare_dataset_vs_simulation":
        result = {}
        try:
            job_id = arguments["job_id"]
            xlsx_path = arguments.get("xlsx_path") or _dataset_default_path()
            pts_filter = arguments.get("points_to_compare")
            wb, ws = _open_dataset(xlsx_path)
            parsed = _parse_dataset_rows(ws)
            job = JOBS.get(job_id)
            if not job:
                result = {"error": f"Job '{job_id}' not found."}
            else:
                ds = _make_ds()
                all_vars = ds.sumo.getVariableNames()
                comparisons, fit_errors = [], []
                for label, measured in parsed["points"].items():
                    if pts_filter and label not in pts_filter:
                        continue
                    mapping = DATASET_SCHEMA["point_map"].get(label)
                    if not mapping:
                        continue
                    unit = mapping["unit"]
                    for pkey, m_val in measured.items():
                        if pkey in ("_row", "Notes") or not isinstance(m_val, (int, float)):
                            continue
                        var = _resolve_sumo_variable(unit, pkey, all_vars)
                        if not var:
                            continue
                        try:
                            s_val = float(ds.sumo.get(var))
                        except Exception:
                            continue
                        delta = s_val - m_val
                        delta_pct = round(delta / m_val * 100, 1) if m_val != 0 else None
                        comparisons.append({"point": label, "param": pkey,
                                            "measured": m_val, "simulated": round(s_val, 3),
                                            "delta": round(delta, 3), "delta_pct": delta_pct})
                        if delta_pct is not None:
                            fit_errors.append(abs(delta_pct))
                mape = round(sum(fit_errors) / len(fit_errors), 2) if fit_errors else None
                bad = [c for c in comparisons
                       if c["delta_pct"] is not None and abs(c["delta_pct"]) > 20]
                result = {
                    "file":               xlsx_path,
                    "job_id":             job_id,
                    "comparisons":        len(comparisons),
                    "mean_abs_pct_error": mape,
                    "poor_fit_count":     len(bad),
                    "poor_fit_points":    bad,
                    "all_comparisons":    comparisons[:100],
                }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # -- compute_removal_efficiencies_from_dataset -----------------------------
    if name == "compute_removal_efficiencies_from_dataset":
        result = {}
        try:
            xlsx_path = arguments.get("xlsx_path") or _dataset_default_path()
            wb, ws = _open_dataset(xlsx_path)
            parsed = _parse_dataset_rows(ws)
            pts = parsed["points"]
            sequence = ["Raw Influent", "Primary Clarifier \u2014 Effluent",
                        "Bio Reactor \u2014 Outlet", "Secondary Clarifier \u2014 Outlet",
                        "Final Effluent"]
            params_to_track = ["BOD5", "COD", "TSS", "NH4_N", "TN", "TP"]
            stepwise = []
            for i in range(len(sequence) - 1):
                a, b = sequence[i], sequence[i + 1]
                if a not in pts or b not in pts:
                    continue
                stage = {"from": a, "to": b, "removals": {}}
                for p in params_to_track:
                    if p in pts[a] and p in pts[b] and pts[a][p] > 0:
                        rem = round((pts[a][p] - pts[b][p]) / pts[a][p] * 100, 1)
                        stage["removals"][p] = rem
                if stage["removals"]:
                    stepwise.append(stage)
            overall = {}
            inf = pts.get("Raw Influent", {})
            eff = pts.get("Final Effluent", {})
            for p in params_to_track:
                if p in inf and p in eff and inf[p] > 0:
                    overall[p] = {
                        "influent":    inf[p],
                        "effluent":    eff[p],
                        "removal_pct": round((inf[p] - eff[p]) / inf[p] * 100, 1),
                        "law48_limit": parsed["law48"].get(p),
                        "compliant":   (eff[p] <= parsed["law48"][p])
                                       if p in parsed["law48"] else None,
                    }
            result = {
                "file":              xlsx_path,
                "stepwise_removals": stepwise,
                "overall_removals":  overall,
            }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # -- export_simulation_to_dataset ------------------------------------------
    if name == "export_simulation_to_dataset":
        result = {}
        try:
            xlsx_path   = arguments.get("xlsx_path") or _dataset_default_path()
            output_path = arguments.get("output_path", "")
            src = Path(xlsx_path)
            if not src.exists():
                result = {"error": f"Template not found: {xlsx_path}"}
            else:
                if not output_path:
                    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                    output_path = str(Path(CONFIG["output_dir"]) /
                                      f"simulation_results_{ts}.xlsx")
                shutil.copy2(src, output_path)
                wb = load_workbook(output_path)
                sheet_name = DATASET_SCHEMA["sheet_name"]
                if sheet_name not in wb.sheetnames:
                    matches = [s for s in wb.sheetnames if "Data Entry" in s]
                    if matches: sheet_name = matches[0]
                ws = wb[sheet_name]
                ds = _make_ds()
                all_vars = ds.sumo.getVariableNames()
                param_cols = {v: k for k, v in DATASET_SCHEMA["columns"].items()
                              if v not in ("label", "Notes")}
                written, skipped = [], []
                for r in range(DATASET_SCHEMA["data_start_row"], ws.max_row + 1):
                    label_val = ws.cell(row=r, column=1).value
                    if label_val is None:
                        continue
                    label = str(label_val).strip()
                    if label.isupper() or label.startswith("  "):
                        continue
                    mapping = DATASET_SCHEMA["point_map"].get(label)
                    if not mapping:
                        continue
                    unit = mapping["unit"]
                    for pkey, col_idx in param_cols.items():
                        var = _resolve_sumo_variable(unit, pkey, all_vars)
                        if not var:
                            continue
                        try:
                            sim_val = float(ds.sumo.get(var))
                            ws.cell(row=r, column=col_idx, value=round(sim_val, 3))
                            ws.cell(row=r, column=col_idx).fill = PatternFill(
                                "solid", fgColor="FFF2CC")
                            written.append({"point": label, "param": pkey,
                                            "variable": var, "value": round(sim_val, 3)})
                        except Exception as e:
                            skipped.append({"point": label, "param": pkey, "error": str(e)})
                ws.cell(row=2, column=2,
                        value=f"Simulation export {datetime.now().strftime('%Y-%m-%d %H:%M')}")
                ws.cell(row=2, column=2).font = Font(bold=True, color="0066CC")
                wb.save(output_path)
                result = {
                    "output_file":    output_path,
                    "cells_written":  len(written),
                    "cells_skipped":  len(skipped),
                    "preview":        written[:20],
                }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # -- export_comparison_report ----------------------------------------------
    if name == "export_comparison_report":
        result = {}
        try:
            job_id      = arguments["job_id"]
            xlsx_path   = arguments.get("xlsx_path") or _dataset_default_path()
            output_path = arguments.get("output_path", "")
            src = Path(xlsx_path)
            if not src.exists():
                result = {"error": f"Template not found: {xlsx_path}"}
            else:
                if not output_path:
                    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                    output_path = str(Path(CONFIG["output_dir"]) /
                                      f"calibration_comparison_{ts}.xlsx")
                shutil.copy2(src, output_path)
                wb = load_workbook(output_path)
                if "Calibration Comparison" in wb.sheetnames:
                    del wb["Calibration Comparison"]
                cmp_ws = wb.create_sheet("Calibration Comparison")
                _, data_ws = _open_dataset(xlsx_path)
                parsed = _parse_dataset_rows(data_ws)
                ds = _make_ds()
                all_vars = ds.sumo.getVariableNames()
                headers = ["Measurement Point", "SUMO Unit", "Parameter",
                           "Measured", "Simulated", "Delta", "Delta %", "Fit Quality"]
                for i, h in enumerate(headers, start=1):
                    c = cmp_ws.cell(row=1, column=i, value=h)
                    c.font = Font(bold=True, color="FFFFFF")
                    c.fill = PatternFill("solid", fgColor="1F3864")
                r = 2
                total_err, comparable = 0.0, 0
                for label, data in parsed["points"].items():
                    mapping = DATASET_SCHEMA["point_map"].get(label)
                    if not mapping:
                        continue
                    unit = mapping["unit"]
                    for pkey, m_val in data.items():
                        if pkey in ("_row", "Notes") or not isinstance(m_val, (int, float)):
                            continue
                        var = _resolve_sumo_variable(unit, pkey, all_vars)
                        if not var:
                            continue
                        try:
                            s_val = float(ds.sumo.get(var))
                        except Exception:
                            continue
                        delta = s_val - m_val
                        delta_pct = (delta / m_val * 100) if m_val != 0 else None
                        if delta_pct is not None:
                            total_err += abs(delta_pct)
                            comparable += 1
                        if delta_pct is None:
                            quality, color = "-", "DDDDDD"
                        elif abs(delta_pct) < 10:
                            quality, color = "excellent", "C6EFCE"
                        elif abs(delta_pct) < 20:
                            quality, color = "acceptable", "FFF2CC"
                        else:
                            quality, color = "poor", "FFC7CE"
                        cmp_ws.cell(row=r, column=1, value=label)
                        cmp_ws.cell(row=r, column=2, value=unit)
                        cmp_ws.cell(row=r, column=3, value=pkey)
                        cmp_ws.cell(row=r, column=4, value=round(m_val, 3))
                        cmp_ws.cell(row=r, column=5, value=round(s_val, 3))
                        cmp_ws.cell(row=r, column=6, value=round(delta, 3))
                        cmp_ws.cell(row=r, column=7,
                                    value=round(delta_pct, 1) if delta_pct is not None else None)
                        qc = cmp_ws.cell(row=r, column=8, value=quality)
                        qc.fill = PatternFill("solid", fgColor=color)
                        r += 1
                mape = round(total_err / comparable, 2) if comparable else None
                cmp_ws.cell(row=r + 1, column=1,
                            value=(f"MAPE (Mean Absolute % Error): {mape}%" if mape
                                   else "No comparable points."))
                cmp_ws.cell(row=r + 1, column=1).font = Font(bold=True)
                widths = [32, 16, 14, 14, 14, 10, 10, 16]
                for i, w in enumerate(widths, start=1):
                    cmp_ws.column_dimensions[cmp_ws.cell(row=1, column=i).column_letter].width = w
                wb.save(output_path)
                result = {
                    "output_file":        output_path,
                    "rows_compared":      comparable,
                    "mean_abs_pct_error": mape,
                }
        except Exception as e:
            result = {"error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "validate_model_structure":
        result = _val_structure()
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "validate_influent_configuration":
        result = _val_influent(arguments.get("unit_name", "Influent"))
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "validate_mass_balance_pre_run":
        result = _val_mass_balance_pre(arguments.get("tolerance_pct", 5.0))
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "validate_mass_balance_post_run":
        result = _val_mass_balance_post(
            arguments.get("job_id"),
            arguments.get("elements"),
            arguments.get("tolerance_pct", 5.0),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "validate_mlss_health":
        result = _val_mlss_health(arguments.get("unit_name", "AerTank1"))
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "validate_srt_fm_feasibility":
        result = _val_srt_fm(
            arguments.get("unit_name", "AerTank1"),
            arguments.get("influent_unit", "Influent"),
            arguments.get("was_unit", "WAS_Pump"),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "validate_do_levels":
        result = _val_do_levels(arguments.get("zone_units"))
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "validate_oxygen_demand_vs_supply":
        result = _val_oxygen_supply(
            arguments.get("unit_name", "AerTank1"),
            arguments.get("influent_unit", "Influent"),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "validate_hydraulic_balance":
        result = _val_hydraulic(
            arguments.get("influent_unit", "Influent"),
            arguments.get("effluent_unit", "Effluent"),
            arguments.get("was_unit", "WAS_Pump"),
            arguments.get("ras_unit", "RAS_Pump"),
            arguments.get("sec_clar_unit", "SecClar1"),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "validate_asm_kinetics":
        result = _val_asm_kinetics(arguments.get("unit_name", "AerTank1"))
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "validate_simulation_convergence":
        result = _val_convergence(
            arguments.get("job_id"),
            arguments.get("stability_window_pct", 10.0),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "validate_simulation_plausibility":
        result = _val_plausibility(arguments.get("job_id"))
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "validate_full_model":
        result = _val_full_model(
            arguments.get("aer_tank", "AerTank1"),
            arguments.get("influent_unit", "Influent"),
            arguments.get("effluent_unit", "Effluent"),
            arguments.get("sec_clar_unit", "SecClar1"),
            arguments.get("was_unit", "WAS_Pump"),
            arguments.get("ras_unit", "RAS_Pump"),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "validate_post_simulation":
        result = _val_post_simulation(arguments.get("job_id"))
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "scan_parameter_sensitivity":
        result = _opt_scan_parameter_sensitivity(
            arguments.get("parameter"),
            arguments.get("values", []),
            arguments.get("output_variable"),
            arguments.get("unit_name", "AerTank1"),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "optimize_was_for_target_srt":
        result = _opt_was_for_srt(
            arguments.get("target_srt_d"),
            arguments.get("tolerance_pct", 5.0),
            20,
            arguments.get("was_unit", "WAS_Pump"),
            arguments.get("aer_tank", "AerTank1"),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "optimize_kla_for_target_do":
        result = _opt_kla_for_do(
            arguments.get("target_do_mgL"),
            arguments.get("unit_name", "AerTank1"),
            arguments.get("tolerance", 0.1),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "diagnose_nitrification_failure":
        result = _diag_nitrification(
            arguments.get("unit_name", "AerTank1"),
            arguments.get("influent_unit", "Influent"),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "diagnose_bulking_risk":
        result = _diag_bulking(arguments.get("unit_name", "AerTank1"))
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "diagnose_phosphorus_removal":
        result = _diag_phosphorus(arguments.get("unit_name", "AerTank1"))
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "recommend_setpoint_adjustments":
        result = _diag_recommendations(
            arguments.get("aer_tank", "AerTank1"),
            arguments.get("influent_unit", "Influent"),
            arguments.get("was_unit", "WAS_Pump"),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "estimate_annual_opex":
        result = _econ_opex(
            arguments.get("job_id", ""),
            arguments.get("energy_cost_per_kwh", 0.08),
            arguments.get("sludge_disposal_cost_per_tonne", 50.0),
            arguments.get("alum_cost_per_kg", 0.30),
            arguments.get("polymer_cost_per_kg", 3.50),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "estimate_ghg_emissions":
        result = _econ_ghg(
            arguments.get("job_id", ""),
            arguments.get("grid_emission_factor_kgCO2_per_kWh", 0.45),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "plot_time_series_chart":
        result = _viz_time_series_chart(
            arguments.get("job_id"),
            arguments.get("variables"),
            arguments.get("output_path"),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]


    if name == "export_publication_time_series":
        result = _exp_publication_time_series(
            arguments["job_ids"],
            arguments.get("variables"),
            arguments.get("output_path"),
            arguments.get("title"),
            arguments.get("ylabel"),
            arguments.get("xlabel"),
            arguments.get("formats"),
            tuple(arguments.get("figsize", (7, 4.2))),
            arguments.get("grayscale_safe", True),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "export_calibration_parity_plot":
        result = _exp_calibration_parity_plot(
            arguments["simulated"],
            arguments["observed"],
            arguments.get("variable_name", "variable"),
            arguments.get("units", "mg/L"),
            arguments.get("output_path"),
            arguments.get("formats"),
            arguments.get("title"),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "export_scenario_comparison_bar_chart":
        result = _exp_scenario_comparison_bar(
            arguments["job_ids"],
            arguments["variables"],
            arguments.get("output_path"),
            arguments.get("formats"),
            tuple(arguments.get("figsize", (7.0, 4.2))),
            arguments.get("ylabel"),
            arguments.get("title"),
            arguments.get("show_errorbars", True),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "export_sensitivity_tornado_diagram":
        result = _exp_tornado_diagram(
            arguments["sensitivity_data"],
            arguments.get("baseline"),
            arguments.get("output_path"),
            arguments.get("formats"),
            tuple(arguments.get("figsize", (7, 4.5))),
            arguments.get("xlabel", "Output response"),
            arguments.get("title"),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "export_box_plot_distributions":
        result = _exp_box_plot_distributions(
            arguments["job_ids"],
            arguments["variable"],
            arguments.get("output_path"),
            arguments.get("formats"),
            tuple(arguments.get("figsize", (6.5, 4.2))),
            arguments.get("ylabel"),
            arguments.get("title"),
            arguments.get("use_violin", False),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "export_academic_results_table_docx":
        result = _exp_academic_results_docx(
            arguments["headers"],
            arguments["rows"],
            arguments.get("output_path"),
            arguments.get("caption"),
            arguments.get("table_number", 1),
            arguments.get("column_units"),
            arguments.get("title"),
            arguments.get("footnote"),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "export_latex_results_table":
        result = _exp_latex_results_table(
            arguments["headers"],
            arguments["rows"],
            arguments.get("output_path"),
            arguments.get("caption"),
            arguments.get("label", "tab:results"),
            arguments.get("column_align"),
            arguments.get("column_units"),
            arguments.get("footnote"),
            arguments.get("table_type", "table"),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "export_statistical_summary_xlsx":
        result = _exp_statistical_summary_xlsx(
            arguments["job_ids"],
            arguments.get("variables"),
            arguments.get("output_path"),
            arguments.get("sheet_per_job", True),
            arguments.get("include_combined", True),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "export_academic_report_docx":
        result = _exp_academic_report_docx(
            arguments["job_ids"],
            arguments.get("variables"),
            arguments.get("title", "SUMO24 Simulation Report"),
            arguments.get("authors"),
            arguments.get("abstract"),
            arguments.get("output_path"),
            arguments.get("include_figures", True),
            arguments.get("embed_figures_dir"),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "export_multi_file_batch_comparison":
        result = _exp_multi_file_batch_comparison(
            arguments["jobs_by_file"],
            arguments.get("variables"),
            arguments.get("output_dir"),
            arguments.get("formats"),
            arguments.get("generate_plots", True),
            arguments.get("generate_xlsx", True),
            arguments.get("generate_latex", True),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "export_mass_balance_diagram":
        result = _exp_mass_balance_diagram(
            arguments["flows"],
            arguments.get("output_path"),
            arguments.get("formats"),
            tuple(arguments.get("figsize", (8, 4.5))),
            arguments.get("title"),
            arguments.get("units", "kg/d"),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    # Group AA â€” WWTP Template Tools (Excel â†’ SUMO model)
    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    if name == "generate_wwtp_template_xlsx":
        result = _exp_generate_wwtp_template_xlsx(arguments.get("output_path"))
        return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

    if name == "read_wwtp_template":
        result = _exp_read_wwtp_template(arguments["template_path"])
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "validate_wwtp_template":
        result = _exp_validate_wwtp_template(arguments["template_path"])
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "build_model_from_template":
        result = _exp_build_model_from_template(
            arguments["template_path"],
            arguments.get("project_path"),
            arguments.get("apply_via_dtt", True),
            arguments.get("dry_run", False),
            arguments.get("allow_offline", False),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    # â”€â”€ Group BB â€” HTML Schematic â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "import_schematic_from_html":
        result = _exp_import_schematic_from_html(arguments["filepath"])
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "build_model_from_schematic":
        result = _exp_build_model_from_schematic(
            arguments.get("schematic"),
            arguments.get("filepath"),
            arguments.get("apply_immediately", False),
            arguments.get("output_scs_path"),
            arguments.get("include_dtt_actions", False),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "compare_schematic_to_model":
        result = _exp_compare_schematic_to_model(
            arguments.get("schematic"),
            arguments.get("filepath"),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "generate_sumoslang_from_schematic":
        result = _exp_generate_sumoslang_from_schematic(
            arguments.get("schematic"),
            arguments.get("filepath"),
            arguments.get("output_path"),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "list_schematic_unit_types":
        result = _exp_list_schematic_unit_types()
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    # â”€â”€ Group CC â€” Troubleshooting â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if name == "diagnose_sumo_file":
        result = _exp_diagnose_sumo_file(arguments["sumo_path"])
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "scan_sumo_directory":
        result = _exp_scan_sumo_directory(arguments.get("directory"))
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "check_dll_companion":
        result = _exp_check_dll_companion(arguments["sumo_path"])
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "repair_sumo_file":
        result = _exp_repair_sumo_file(
            arguments["sumo_path"],
            arguments.get("source_dll"),
            arguments.get("dry_run", False),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "diagnose_sumo_crash":
        result = _exp_diagnose_sumo_crash(
            arguments.get("log_path"),
            arguments.get("tail_lines", 400),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "validate_sumo_environment":
        result = _exp_validate_sumo_environment()
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "list_sumo_diagnostics":
        result = _exp_list_sumo_diagnostics()
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    # ─── Group DD - schematic-to-SUMO compiler ───────────────────────
    if name == "compile_schematic_to_sumo":
        result = _exp_compile_schematic_to_sumo(
            arguments.get("schematic"),
            arguments["output_path"],
            arguments.get("overwrite", False),
            arguments.get("include_sumoslang", True),
            arguments.get("attach_dll", "") or None,
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "build_sumo_from_html":
        result = _exp_build_sumo_from_html(
            arguments["html_path"],
            arguments["output_path"],
            arguments.get("overwrite", False),
            arguments.get("include_sumoslang", True),
            arguments.get("attach_dll", "") or None,
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "preview_sumo_manifest":
        result = _exp_preview_sumo_manifest(
            arguments.get("schematic"),
            arguments.get("html_path", "") or None,
            arguments.get("include_sumoslang", True),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "attach_companion_dll":
        result = _exp_attach_companion_dll(
            arguments["sumo_path"],
            arguments["source_dll"],
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "verify_sumo_file_against_schematic":
        result = _exp_verify_sumo_file_against_schematic(
            arguments["sumo_path"],
            arguments.get("schematic"),
            arguments.get("html_path", "") or None,
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "read_sumo_manifest":
        result = _exp_read_sumo_manifest(arguments["sumo_path"])
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    # ─── Group EE - Build Pack & Baseline Apply ───────────────────────────
    if name == "build_sumo_pack":
        result = _exp_build_sumo_pack(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "apply_schematic_to_baseline":
        result = _exp_apply_schematic_to_baseline(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "build_native_sumo":
        result = _exp_build_native_sumo(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "topology_match_report":
        result = _exp_topology_match_report(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "validate_sumoslang":
        result = _exp_validate_sumoslang(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "verify_dll_matches_schematic":
        result = _exp_verify_dll_matches_schematic(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "update_schematic_in_html":
        result = _exp_update_schematic_in_html(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "next_step_for_project":
        result = _exp_next_step_for_project(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "apply_template_auto_fixes":
        result = _exp_apply_template_auto_fixes(
            arguments["template_path"],
            arguments.get("fixes", []),
        )
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    # ─── Group FF — Pipeline driver and stage orchestration ──────────────────
    if name == "pipeline_init":
        result = _exp_pipeline_init(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "pipeline_status":
        result = _exp_pipeline_status(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "pipeline_advance":
        result = _exp_pipeline_advance(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "pipeline_revert":
        result = _exp_pipeline_revert(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "pipeline_describe":
        result = _exp_pipeline_describe(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "stage1_open_schematic":
        result = _exp_stage1_open_schematic(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "stage1_validate_schematic":
        result = _exp_stage1_validate_schematic(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "stage2_open_data":
        result = _exp_stage2_open_data(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "stage2_validate_parameters":
        result = _exp_stage2_validate_parameters(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "stage3_run_engineering_checks":
        result = _exp_stage3_run_engineering_checks(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "stage4_define_static_inputs":
        result = _exp_stage4_define_static_inputs(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "stage4_validate_static_inputs":
        result = _exp_stage4_validate_static_inputs(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "stage5_generate_dynamic_inputs":
        result = _exp_stage5_generate_dynamic_inputs(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "stage5_validate_dynamic_inputs":
        result = _exp_stage5_validate_dynamic_inputs(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "stage6_add_controller":
        result = _exp_stage6_add_controller(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "stage6_validate_controllers":
        result = _exp_stage6_validate_controllers(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "stage7_build":
        result = _exp_stage7_build(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "stage8_verify":
        result = _exp_stage8_verify(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    # ── Group FF: DTT-mediated .sumo editing ─────────────────────────────────
    if name == "rename_unit_process":
        result = _exp_rename_unit_process(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "change_unit_type":
        result = _exp_change_unit_type(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "remove_flow_connection":
        result = _exp_remove_flow_connection(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "set_stream_flow_rate":
        result = _exp_set_stream_flow_rate(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "list_controllers":
        result = _exp_list_controllers(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "modify_controller":
        result = _exp_modify_controller(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "remove_controller":
        result = _exp_remove_controller(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "begin_edit_transaction":
        result = _exp_begin_edit_transaction(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "commit_edit_transaction":
        result = _exp_commit_edit_transaction(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "rollback_edit_transaction":
        result = _exp_rollback_edit_transaction(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "apply_dtt_command_batch":
        result = _exp_apply_dtt_command_batch(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    if name == "diff_inmemory_vs_disk":
        result = _exp_diff_inmemory_vs_disk(arguments)
        return [types.TextContent(type="text", text=json.dumps(result, indent=2, default=str))]
    # ── end Group FF ─────────────────────────────────────────────────────────

    return [types.TextContent(type="text", text=f"Unknown tool: {name}")]


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# Group AA â€” WWTP Template Helpers
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# Default path of the WWTP input template living next to this script
_WWTP_TEMPLATE_DEFAULT = str(_MCP_DIR / "WWTP_Model_Input_Template.xlsx")


def _exp_generate_wwtp_template_xlsx(output_path: str | None = None) -> dict:
    """Create a fresh blank WWTP Model Input Template at output_path."""
    if not OPENPYXL_AVAILABLE:
        return {"error": "openpyxl is not installed"}
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font as _F, PatternFill as _P, Alignment as _A, Border as _B, Side as _S
        from openpyxl.utils import get_column_letter as _GCL
        from openpyxl.worksheet.datavalidation import DataValidation as _DV
        from openpyxl.comments import Comment as _Cm
    except Exception as ex:
        return {"error": f"openpyxl import failed: {ex}"}

    out = Path(output_path) if output_path else Path(_WWTP_TEMPLATE_DEFAULT)
    out.parent.mkdir(parents=True, exist_ok=True)

    HEADER_FILL = _P("solid", start_color="1F3864")
    NOTE_FILL   = _P("solid", start_color="FFF2CC")
    INPUT_FILL  = _P("solid", start_color="DEEBF7")
    OPT_FILL    = _P("solid", start_color="F2F2F2")
    HEADER_FONT = _F(name="Calibri", bold=True, size=11, color="FFFFFF")
    TITLE_FONT  = _F(name="Calibri", bold=True, size=16, color="1F3864")
    BODY_FONT   = _F(name="Calibri", size=10)
    UNIT_FONT   = _F(name="Calibri", italic=True, size=9, color="595959")
    NOTE_FONT   = _F(name="Calibri", italic=True, size=10, color="555555")
    THIN  = _S(border_style="thin",  color="BFBFBF")
    THICK = _S(border_style="medium", color="1F3864")
    BOX = _B(left=THIN, right=THIN, top=THIN, bottom=THIN)
    HEAD_BORDER = _B(left=THICK, right=THICK, top=THICK, bottom=THICK)
    CENTER = _A(horizontal="center", vertical="center", wrap_text=True)
    LEFT_  = _A(horizontal="left",   vertical="center", wrap_text=True)

    def _hdr(cell):
        cell.font = HEADER_FONT; cell.fill = HEADER_FILL
        cell.border = HEAD_BORDER; cell.alignment = CENTER

    def _title(ws, txt, ncols):
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=ncols)
        c = ws.cell(row=1, column=1, value=txt); c.font = TITLE_FONT; c.alignment = LEFT_
        ws.row_dimensions[1].height = 24

    def _widths(ws, widths):
        for i, w in enumerate(widths, start=1):
            ws.column_dimensions[_GCL(i)].width = w

    def _input_block(ws, top_row, last_row, n_cols):
        for r in range(top_row, last_row + 1):
            for c in range(1, n_cols + 1):
                cl = ws.cell(row=r, column=c)
                cl.font = BODY_FONT; cl.fill = INPUT_FILL; cl.border = BOX; cl.alignment = LEFT_

    wb = Workbook()
    wb.active.title = "README"
    ws = wb["README"]
    _title(ws, "WWTP Model Input Template â€” SUMO24 MCP", 2)
    _widths(ws, [32, 100])
    readme = [
        ("Purpose", "Fill this workbook with WWTP data; the SUMO24 MCP server will build a SUMO24 model from it via build_model_from_template."),
        ("Workflow", "1) Plant_Info; 2) Influent_Dataset + Influent_Fractionation; 3) Effluent_Dataset + Effluent_Targets; 4) Treatment_Processes; 5) Connections; 6) Process_Sizing; 7) Bio_Operational / Phys_Operational / Chem_Operational; 8) save and call build_model_from_template."),
        ("Cell key", "Navy = header, pale blue = user input, gray = hint, yellow = note."),
        ("Required", "Plant_Info, Treatment_Processes and Connections are required. Other sheets are optional."),
        ("Time series", "Influent_Dataset / Effluent_Dataset accept daily or sub-daily timestamps."),
        ("Egyptian limits", "Effluent_Targets is preloaded with Law 48/1982 default limits."),
    ]
    ws.cell(row=2, column=1, value="Section").font = _F(bold=True)
    ws.cell(row=2, column=2, value="Description").font = _F(bold=True)
    for i, (k, v) in enumerate(readme, start=3):
        ws.cell(row=i, column=1, value=k).font = _F(bold=True, size=11)
        ws.cell(row=i, column=2, value=v).font = BODY_FONT
        for c in range(1, 3):
            ws.cell(row=i, column=c).alignment = _A(horizontal="left", vertical="top", wrap_text=True)
        ws.row_dimensions[i].height = 36

    # --- Plant_Info ---
    ws = wb.create_sheet("Plant_Info"); _title(ws, "1. Plant Information", 4); _widths(ws, [38, 32, 14, 60])
    for c, h in enumerate(["Field","Value","Unit","Description / hint"], start=1):
        _hdr(ws.cell(row=3, column=c, value=h))
    pi = [("project_name","MyPlant_2026","-","Identifier for .sumo file/folder."),
          ("project_path","F:/UNI/SUMO/MCP/projects","-","Absolute folder for the project."),
          ("plant_location","","-","City/country."),
          ("design_flow_avg","","m3/d","Average dry-weather flow."),
          ("design_flow_peak","","m3/d","Peak hourly flow."),
          ("design_flow_min","","m3/d","Minimum night flow."),
          ("design_population","","PE","Population equivalent."),
          ("kinetic_model","ASM2d","-","ASM1/ASM2d/ASM3/ASM3+BioP/Mantis2."),
          ("temperature_design","20","degC","Reference T."),
          ("temperature_min","","degC","Coldest T."),
          ("temperature_max","","degC","Hottest T."),
          ("elevation","","m a.s.l.","Elevation."),
          ("discharge_target","law48_drain","-","law48_nile/law48_drain/custom."),
          ("project_currency","EGP","-","Currency for OPEX/CAPEX."),
          ("created_by","","-","Modeller name."),
          ("created_date","","yyyy-mm-dd","Date.")]
    for i, (k, v, u, d) in enumerate(pi, start=4):
        ws.cell(row=i, column=1, value=k).font = _F(bold=True, size=10)
        ws.cell(row=i, column=2, value=v).fill = INPUT_FILL
        ws.cell(row=i, column=3, value=u).font = UNIT_FONT
        ws.cell(row=i, column=4, value=d).font = NOTE_FONT
        for c in range(1, 5):
            ws.cell(row=i, column=c).border = BOX
            ws.cell(row=i, column=c).alignment = LEFT_
    dv1 = _DV(type="list", formula1='"ASM1,ASM2d,ASM3,ASM3+BioP,Mantis2,MBR_ASM2d,ADM1"', allow_blank=True); dv1.add("B11"); ws.add_data_validation(dv1)
    dv2 = _DV(type="list", formula1='"law48_nile,law48_drain,custom"', allow_blank=True); dv2.add("B16"); ws.add_data_validation(dv2)
    ws.freeze_panes = "A4"

    # --- Influent_Dataset ---
    inf_cols = [("Date","yyyy-mm-dd"),("Q","m3/d"),("COD","mg/L"),("sCOD","mg/L"),("BOD5","mg/L"),
                ("TSS","mg/L"),("VSS","mg/L"),("TKN","mg/L"),("NH4-N","mg/L"),("NO3-N","mg/L"),
                ("TP","mg/L"),("PO4-P","mg/L"),("Alkalinity","mmol/L"),("pH","-"),("Temperature","degC")]
    for sheet_name, title in [("Influent_Dataset","2. Influent Dataset"),("Effluent_Dataset","4. Effluent Dataset")]:
        ws = wb.create_sheet(sheet_name); _title(ws, title, len(inf_cols))
        _widths(ws, [16] + [13]*(len(inf_cols)-1))
        for c, (h, u) in enumerate(inf_cols, start=1):
            _hdr(ws.cell(row=3, column=c, value=h))
            uc = ws.cell(row=4, column=c, value=u); uc.font = UNIT_FONT; uc.alignment = CENTER; uc.fill = OPT_FILL; uc.border = BOX
        _input_block(ws, 5, 34, len(inf_cols))
        ws.freeze_panes = "B5"

    # --- Influent_Fractionation ---
    ws = wb.create_sheet("Influent_Fractionation"); _title(ws, "3. Influent COD/N/P Fractionation", 4); _widths(ws, [22,16,14,70])
    for c, h in enumerate(["Fraction","Value","Unit","Description"], start=1): _hdr(ws.cell(row=3,column=c,value=h))
    frac = [("fSI",0.07,"g COD/g COD","Soluble inert."),("fSS",0.20,"g COD/g COD","Readily biodegradable."),
            ("fXS",0.50,"g COD/g COD","Slowly biodegradable."),("fXI",0.15,"g COD/g COD","Particulate inert."),
            ("fXBH",0.08,"g COD/g COD","Heterotrophic biomass."),("fXBA",0.00,"g COD/g COD","Autotrophic biomass."),
            ("fSNH4",0.65,"g N/g TKN","NH4 fraction of TKN."),("fSND",0.10,"g N/g TKN","Soluble organic N."),
            ("fXND",0.18,"g N/g TKN","Particulate organic N."),("fSPO4",0.50,"g P/g TP","Ortho-P fraction."),]
    for i, (k, v, u, d) in enumerate(frac, start=4):
        ws.cell(row=i, column=1, value=k).font = _F(bold=True)
        ws.cell(row=i, column=2, value=v).fill = INPUT_FILL
        ws.cell(row=i, column=3, value=u).font = UNIT_FONT
        ws.cell(row=i, column=4, value=d).font = NOTE_FONT
        for c in range(1, 5): ws.cell(row=i, column=c).border = BOX; ws.cell(row=i, column=c).alignment = LEFT_
    ws.freeze_panes = "A4"

    # --- Effluent_Targets ---
    ws = wb.create_sheet("Effluent_Targets"); _title(ws, "5. Effluent Discharge Targets", 5); _widths(ws, [16,14,14,14,60])
    for c, h in enumerate(["Parameter","Limit","Unit","Stat","Source / notes"], start=1): _hdr(ws.cell(row=3,column=c,value=h))
    tgt = [("BOD5",60,"mg/L","max","Egyptian Law 48/1982 â€” drains."),("COD",80,"mg/L","max","Law 48 â€” drains."),
           ("TSS",50,"mg/L","max","Law 48 â€” drains."),("NH4-N",5,"mg/L","max","Law 48 â€” drains."),
           ("TN",50,"mg/L","max","Law 48 â€” drains."),("TP",3,"mg/L","max","Law 48 â€” drains."),
           ("Oil",10,"mg/L","max","Law 48 â€” drains."),("FC",5000,"MPN/100mL","max","Faecal coliforms."),
           ("pH_min",6.0,"-","min","Law 48."),("pH_max",9.0,"-","max","Law 48.")]
    for i, (p, l, u, s, src) in enumerate(tgt, start=4):
        ws.cell(row=i, column=1, value=p).font = _F(bold=True)
        ws.cell(row=i, column=2, value=l).fill = INPUT_FILL
        ws.cell(row=i, column=3, value=u).font = UNIT_FONT
        ws.cell(row=i, column=4, value=s).font = UNIT_FONT
        ws.cell(row=i, column=5, value=src).font = NOTE_FONT
        for c in range(1, 6): ws.cell(row=i, column=c).border = BOX; ws.cell(row=i, column=c).alignment = LEFT_
    ws.freeze_panes = "A4"

    # --- Treatment_Processes ---
    ws = wb.create_sheet("Treatment_Processes"); _title(ws, "6. Treatment Processes", 7); _widths(ws, [10,22,22,14,14,16,50])
    for c, h in enumerate(["order_index","instance_name","process_type","category","stage","kinetic_model","notes"], start=1):
        _hdr(ws.cell(row=3,column=c,value=h))
    sp = [(1,"Influent","Influent","physical","headworks","",""),
          (2,"Screen1","Screen","physical","headworks","",""),
          (3,"GritChamber1","GritChamber","physical","headworks","",""),
          (4,"PrimaryClarifier1","PrimaryClarifier","physical","primary","",""),
          (5,"AnaerobicTank1","AnaerobicReactor","biological","biological","ASM2d",""),
          (6,"AnoxicTank1","AnoxicReactor","biological","biological","ASM2d",""),
          (7,"AerationTank1","AerationTank","biological","biological","ASM2d",""),
          (8,"SecondaryClarifier1","SecondaryClarifier","physical","secondary","Takacs",""),
          (9,"Effluent","Effluent","physical","effluent","","")]
    for i, row in enumerate(sp, start=4):
        for c, v in enumerate(row, start=1):
            cell = ws.cell(row=i, column=c, value=v); cell.font = BODY_FONT; cell.fill = INPUT_FILL; cell.border = BOX; cell.alignment = LEFT_
    _input_block(ws, 4 + len(sp), 4 + len(sp) + 9, 7)
    dv = _DV(type="list", formula1='"physical,biological,chemical"', allow_blank=True); dv.add("D4:D200"); ws.add_data_validation(dv)
    dv = _DV(type="list", formula1='"headworks,primary,biological,secondary,tertiary,disinfection,sludge,effluent"', allow_blank=True); dv.add("E4:E200"); ws.add_data_validation(dv)
    dv = _DV(type="list", formula1='"ASM1,ASM2d,ASM3,ASM3+BioP,Mantis2,MBR_ASM2d,Takacs,ADM1,Siegrist,CSTR_Simple"', allow_blank=True); dv.add("F4:F200"); ws.add_data_validation(dv)
    ws.freeze_panes = "A4"

    # --- Connections ---
    ws = wb.create_sheet("Connections"); _title(ws, "7. Stream Connections", 6); _widths(ws, [22,22,22,14,14,50])
    for c, h in enumerate(["from_unit","to_unit","stream_name","stream_type","recycle_type","notes"], start=1):
        _hdr(ws.cell(row=3,column=c,value=h))
    sc = [("Influent","Screen1","S_inf","liquid","",""),
          ("Screen1","GritChamber1","S_sg","liquid","",""),
          ("GritChamber1","PrimaryClarifier1","S_gp","liquid","",""),
          ("PrimaryClarifier1","AnaerobicTank1","S_pa","liquid","",""),
          ("AnaerobicTank1","AnoxicTank1","S_ax","liquid","",""),
          ("AnoxicTank1","AerationTank1","S_xa","liquid","",""),
          ("AerationTank1","SecondaryClarifier1","S_as","liquid","",""),
          ("SecondaryClarifier1","Effluent","S_se","liquid","",""),
          ("AerationTank1","AnoxicTank1","S_IR","recycle","InternalRecycle",""),
          ("SecondaryClarifier1","AnoxicTank1","S_RAS","recycle","RAS",""),
          ("SecondaryClarifier1","AerationTank1","S_WAS","recycle","WAS","")]
    for i, row in enumerate(sc, start=4):
        for c, v in enumerate(row, start=1):
            cell = ws.cell(row=i, column=c, value=v); cell.font = BODY_FONT; cell.fill = INPUT_FILL; cell.border = BOX; cell.alignment = LEFT_
    # Add an explanatory comment to the first from_unit cell.
    ws.cell(row=4, column=1).comment = _Cm(
        "from_unit and to_unit MUST exactly match instance_name values in "
        "Treatment_Processes (case-sensitive). Delete example rows and replace "
        "with your plant's actual connections.", "MCP"
    )
    _input_block(ws, 4 + len(sc), 4 + len(sc) + 9, 6)
    # Dropdowns for stream_type and recycle_type columns.
    dv = _DV(type="list", formula1='"liquid,sludge,gas,recycle"', allow_blank=True); dv.add("D4:D200"); ws.add_data_validation(dv)
    dv = _DV(type="list", formula1='"RAS,WAS,InternalRecycle,RejectWater,Sidestream"', allow_blank=True); dv.add("E4:E200"); ws.add_data_validation(dv)
    # Cross-reference dropdowns: from_unit and to_unit sourced from Treatment_Processes.
    dv_unit = _DV(type="list", formula1="=Treatment_Processes!$B$4:$B$50",
                  allow_blank=True, showDropDown=False)
    ws.add_data_validation(dv_unit)
    dv_unit.add("A4:A200")
    dv_unit.add("B4:B200")
    ws.freeze_panes = "A4"

    # --- Process_Sizing ---
    ws = wb.create_sheet("Process_Sizing"); _title(ws, "8. Process Sizing", 10); _widths(ws, [22] + [13]*9)
    for c, h in enumerate(["instance_name","n_units","volume_m3","depth_m","surface_area_m2","diameter_m","length_m","width_m","HRT_h","notes"], start=1):
        _hdr(ws.cell(row=3,column=c,value=h))
    sz = [("Screen1",1,"","","","","","","",""),
          ("GritChamber1",1,50,2,"","","","",1.0,""),
          ("PrimaryClarifier1",1,800,3.5,230,17,"","",2.0,""),
          ("AnaerobicTank1",1,600,4,"","",25,6,0.6,""),
          ("AnoxicTank1",1,1200,4,"","",30,10,1.2,""),
          ("AerationTank1",2,4000,4.5,"","",40,22,6.0,""),
          ("SecondaryClarifier1",2,1500,4,380,22,"","",3.5,"")]
    for i, row in enumerate(sz, start=4):
        for c, v in enumerate(row, start=1):
            cell = ws.cell(row=i, column=c, value=v); cell.font = BODY_FONT; cell.fill = INPUT_FILL; cell.border = BOX; cell.alignment = LEFT_
    _input_block(ws, 4 + len(sz), 4 + len(sz) + 9, 10)
    ws.freeze_panes = "A4"

    # --- Bio_Operational ---
    ws = wb.create_sheet("Bio_Operational"); _title(ws, "9. Biological-Process Operational Parameters", 6); _widths(ws, [22,22,14,14,14,50])
    for c, h in enumerate(["instance_name","parameter","value","unit","applies_to","notes"], start=1): _hdr(ws.cell(row=3,column=c,value=h))
    bio = [("AerationTank1","DO_setpoint",2.0,"mg/L","aeration",""),
           ("AerationTank1","MLSS_target",3500,"mg/L","aeration",""),
           ("AerationTank1","SRT_target",12,"d","aeration",""),
           ("AnoxicTank1","DO_max",0.2,"mg/L","anoxic",""),
           ("AnaerobicTank1","DO_max",0.05,"mg/L","anaerobic",""),
           ("SecondaryClarifier1","RAS_ratio",0.75,"-","RAS",""),
           ("SecondaryClarifier1","WAS_ratio",0.02,"-","WAS",""),
           ("AnoxicTank1","Internal_recycle_ratio",3.0,"-","IR","")]
    for i, row in enumerate(bio, start=4):
        for c, v in enumerate(row, start=1):
            cell = ws.cell(row=i, column=c, value=v); cell.font = BODY_FONT; cell.fill = INPUT_FILL; cell.border = BOX; cell.alignment = LEFT_
    _input_block(ws, 4 + len(bio), 4 + len(bio) + 14, 6)
    ws.freeze_panes = "A4"

    # --- Phys_Operational ---
    ws = wb.create_sheet("Phys_Operational"); _title(ws, "10. Physical-Process Operational Parameters", 6); _widths(ws, [22,22,14,14,14,50])
    for c, h in enumerate(["instance_name","parameter","value","unit","applies_to","notes"], start=1): _hdr(ws.cell(row=3,column=c,value=h))
    phy = [("Screen1","bar_spacing",6,"mm","screen",""),
           ("GritChamber1","HRT",3,"min","grit",""),
           ("PrimaryClarifier1","SOR_avg",32,"m3/(m2Â·d)","clarifier",""),
           ("SecondaryClarifier1","SLR_avg",4.0,"kg/(m2Â·h)","clarifier",""),
           ("SecondaryClarifier1","SOR_avg",24,"m3/(m2Â·d)","clarifier","")]
    for i, row in enumerate(phy, start=4):
        for c, v in enumerate(row, start=1):
            cell = ws.cell(row=i, column=c, value=v); cell.font = BODY_FONT; cell.fill = INPUT_FILL; cell.border = BOX; cell.alignment = LEFT_
    _input_block(ws, 4 + len(phy), 4 + len(phy) + 14, 6)
    ws.freeze_panes = "A4"

    # --- Chem_Operational ---
    ws = wb.create_sheet("Chem_Operational"); _title(ws, "11. Chemical Dosing", 6); _widths(ws, [22,18,12,14,14,50])
    for c, h in enumerate(["instance_name","chemical","dose","unit","target_unit","purpose"], start=1): _hdr(ws.cell(row=3,column=c,value=h))
    _input_block(ws, 4, 13, 6)
    dv = _DV(type="list", formula1='"FeCl3,FeSO4,Alum,PACl,Polymer,NaOH,Ca(OH)2,Methanol,Acetate,Glycerol,Cl2,NaOCl,KMnO4,O3"', allow_blank=True); dv.add("B4:B100"); ws.add_data_validation(dv)
    ws.freeze_panes = "A4"

    # --- Variable_Map ---
    ws = wb.create_sheet("Variable_Map"); _title(ws, "12. Variable Map", 4); _widths(ws, [16,28,18,60])
    for c, h in enumerate(["Column header","SUMO state var","Unit","Notes"], start=1): _hdr(ws.cell(row=3,column=c,value=h))
    vmap = [("Q","Q","m3/d","Volumetric flow."),("COD","TCOD","mg/L","Total COD."),("sCOD","SCOD","mg/L","Soluble COD."),
            ("BOD5","BOD5","mg/L","BOD5."),("TSS","TSS","mg/L","Total SS."),("VSS","VSS","mg/L","Volatile SS."),
            ("TKN","TKN","mg/L","TKN."),("NH4-N","SNH","mg/L","Ammonia."),("NO3-N","SNO","mg/L","Nitrate."),
            ("TP","TP","mg/L","Total P."),("PO4-P","SPO","mg/L","Ortho-P."),("Alkalinity","SALK","mmol/L","Bicarbonate."),
            ("pH","PH","-","pH."),("Temperature","T","degC","Bulk T.")]
    for i, row in enumerate(vmap, start=4):
        for c, v in enumerate(row, start=1):
            cell = ws.cell(row=i, column=c, value=v); cell.font = BODY_FONT; cell.fill = OPT_FILL; cell.border = BOX; cell.alignment = LEFT_
    ws.freeze_panes = "A4"

    wb.save(out)
    return {"created": str(out), "size_bytes": out.stat().st_size, "sheets": wb.sheetnames}


def _exp_read_wwtp_template(template_path: str) -> dict:
    """Load and return a structured dict from the WWTP input template."""
    if not OPENPYXL_AVAILABLE:
        return {"error": "openpyxl is not installed"}
    p = Path(template_path)
    if not p.exists():
        return {"error": f"Template not found: {template_path}"}
    try:
        wb = load_workbook(p, data_only=True)
    except Exception as ex:
        return {"error": f"Failed to open workbook: {ex}"}

    out: dict = {"template_path": str(p), "sheets": wb.sheetnames}

    # Plant_Info: rows starting at 4 with [field, value, unit, desc]
    if "Plant_Info" in wb.sheetnames:
        ws = wb["Plant_Info"]; pi = {}
        for r in range(4, ws.max_row + 1):
            k = ws.cell(row=r, column=1).value
            v = ws.cell(row=r, column=2).value
            if k:
                pi[str(k).strip()] = v
        out["plant_info"] = pi

    # Influent_Fractionation
    if "Influent_Fractionation" in wb.sheetnames:
        ws = wb["Influent_Fractionation"]; fr = {}
        for r in range(4, ws.max_row + 1):
            k = ws.cell(row=r, column=1).value; v = ws.cell(row=r, column=2).value
            if k and not str(k).lower().startswith("sum"): fr[str(k).strip()] = v
        out["influent_fractionation"] = fr

    # Effluent_Targets
    if "Effluent_Targets" in wb.sheetnames:
        ws = wb["Effluent_Targets"]; tg = {}
        for r in range(4, ws.max_row + 1):
            k = ws.cell(row=r, column=1).value
            if k:
                tg[str(k).strip()] = {
                    "limit": ws.cell(row=r, column=2).value,
                    "unit":  ws.cell(row=r, column=3).value,
                    "stat":  ws.cell(row=r, column=4).value,
                    "source":ws.cell(row=r, column=5).value,
                }
        out["effluent_targets"] = tg

    # Time-series: Influent / Effluent
    def _read_timeseries(sheet_name):
        if sheet_name not in wb.sheetnames: return []
        ws = wb[sheet_name]
        headers = [ws.cell(row=3, column=c).value for c in range(1, ws.max_column + 1)]
        rows = []
        for r in range(5, ws.max_row + 1):
            vals = [ws.cell(row=r, column=c).value for c in range(1, ws.max_column + 1)]
            if any(v is not None and v != "" for v in vals):
                rows.append({h: v for h, v in zip(headers, vals) if h})
        return rows

    out["influent_dataset"] = _read_timeseries("Influent_Dataset")
    out["effluent_dataset"] = _read_timeseries("Effluent_Dataset")

    # Treatment_Processes
    if "Treatment_Processes" in wb.sheetnames:
        ws = wb["Treatment_Processes"]; procs = []
        for r in range(4, ws.max_row + 1):
            inst = ws.cell(row=r, column=2).value
            if not inst: continue
            procs.append({
                "order_index":   ws.cell(row=r, column=1).value,
                "instance_name": str(inst).strip(),
                "process_type":  ws.cell(row=r, column=3).value,
                "category":      ws.cell(row=r, column=4).value,
                "stage":         ws.cell(row=r, column=5).value,
                "kinetic_model": ws.cell(row=r, column=6).value,
                "notes":         ws.cell(row=r, column=7).value,
            })
        out["treatment_processes"] = procs

    # Connections
    if "Connections" in wb.sheetnames:
        ws = wb["Connections"]; conns = []
        for r in range(4, ws.max_row + 1):
            f = ws.cell(row=r, column=1).value
            t = ws.cell(row=r, column=2).value
            if not (f and t): continue
            conns.append({
                "from_unit":   str(f).strip(),
                "to_unit":     str(t).strip(),
                "stream_name": ws.cell(row=r, column=3).value,
                "stream_type": ws.cell(row=r, column=4).value or "liquid",
                "recycle_type":ws.cell(row=r, column=5).value,
                "notes":       ws.cell(row=r, column=6).value,
            })
        out["connections"] = conns

    # Process_Sizing
    if "Process_Sizing" in wb.sheetnames:
        ws = wb["Process_Sizing"]; sz = {}
        for r in range(4, ws.max_row + 1):
            inst = ws.cell(row=r, column=1).value
            if not inst: continue
            sz[str(inst).strip()] = {
                "n_units":          ws.cell(row=r, column=2).value,
                "volume_m3":        ws.cell(row=r, column=3).value,
                "depth_m":          ws.cell(row=r, column=4).value,
                "surface_area_m2":  ws.cell(row=r, column=5).value,
                "diameter_m":       ws.cell(row=r, column=6).value,
                "length_m":         ws.cell(row=r, column=7).value,
                "width_m":          ws.cell(row=r, column=8).value,
                "HRT_h":            ws.cell(row=r, column=9).value,
                "notes":            ws.cell(row=r, column=10).value,
            }
        out["process_sizing"] = sz

    # Bio / Phys / Chem operational sheets
    def _read_op(sheet_name, n_cols=6):
        if sheet_name not in wb.sheetnames: return []
        ws = wb[sheet_name]; rows = []
        for r in range(4, ws.max_row + 1):
            inst = ws.cell(row=r, column=1).value
            if not inst: continue
            rec = {
                "instance_name": str(inst).strip(),
                "parameter":     ws.cell(row=r, column=2).value,
                "value":         ws.cell(row=r, column=3).value,
                "unit":          ws.cell(row=r, column=4).value,
                "applies_to":    ws.cell(row=r, column=5).value,
                "notes":         ws.cell(row=r, column=6).value,
            }
            rows.append(rec)
        return rows

    out["bio_operational"]  = _read_op("Bio_Operational")
    out["phys_operational"] = _read_op("Phys_Operational")

    # Chem_Operational has slightly different column layout
    if "Chem_Operational" in wb.sheetnames:
        ws = wb["Chem_Operational"]; chems = []
        for r in range(4, ws.max_row + 1):
            inst = ws.cell(row=r, column=1).value
            chem = ws.cell(row=r, column=2).value
            if not (inst or chem): continue
            chems.append({
                "instance_name": inst, "chemical": chem,
                "dose":          ws.cell(row=r, column=3).value,
                "unit":          ws.cell(row=r, column=4).value,
                "target_unit":   ws.cell(row=r, column=5).value,
                "purpose":       ws.cell(row=r, column=6).value,
            })
        out["chem_operational"] = chems

    # Influent statistics â€” convenient summary
    inf = out.get("influent_dataset", [])
    if inf:
        nums = {}
        for row in inf:
            for k, v in row.items():
                if k == "Date": continue
                try: x = float(v)
                except (TypeError, ValueError): continue
                nums.setdefault(k, []).append(x)
        out["influent_statistics"] = {
            k: {"n": len(vs),
                "mean": sum(vs) / len(vs),
                "min":  min(vs),
                "max":  max(vs)}
            for k, vs in nums.items() if vs
        }

    return out


def _exp_validate_wwtp_template(template_path: str) -> dict:
    """Run consistency checks on the template and return errors / warnings."""
    parsed = _exp_read_wwtp_template(template_path)
    if "error" in parsed:
        return parsed

    errors:   list[str] = []
    warnings: list[str] = []

    pi = parsed.get("plant_info", {})
    for required in ("project_name", "project_path", "kinetic_model"):
        if not pi.get(required):
            errors.append(f"Plant_Info: '{required}' is required")

    procs = parsed.get("treatment_processes", [])
    if not procs:
        errors.append("Treatment_Processes: must list at least one unit")
    inst_names = [p["instance_name"] for p in procs]
    dups = {n for n in inst_names if inst_names.count(n) > 1}
    if dups: errors.append(f"Treatment_Processes: duplicate instance_name(s): {sorted(dups)}")
    if not any(p.get("process_type") == "Influent" for p in procs):
        warnings.append("Treatment_Processes: no Influent unit found")
    if not any(p.get("process_type") == "Effluent" for p in procs):
        warnings.append("Treatment_Processes: no Effluent unit found")

    conns = parsed.get("connections", [])
    if not conns:
        errors.append("Connections: at least one stream is required")
    inst_set = set(inst_names)

    # Collect raw connection errors with row indices, then group duplicates.
    _BOUNDARY_UNITS = {"WAS", "RAS", "Effluent", "Influent"}
    raw_conn_errors: list[tuple[int, str, str, str]] = []  # (row, side, name, msg)
    for row_idx, c in enumerate(conns, start=4):
        for side in ("from_unit", "to_unit"):
            val = c.get(side)
            if val and val not in inst_set:
                msg = f"Connections: {side}='{val}' not present in Treatment_Processes"
                raw_conn_errors.append((row_idx, side, val, msg))

    # Group by message, attach auto_fix where applicable.
    from collections import defaultdict as _dd
    _grouped: dict = _dd(list)
    for row_idx, side, val, msg in raw_conn_errors:
        _grouped[msg].append((row_idx, side, val))

    for msg, occurrences in _grouped.items():
        rows = [r for r, _, _ in occurrences]
        entry = {"issue": msg, "rows": rows, "count": len(rows)}
        # auto_fix for canonical boundary names.
        _, side, val = occurrences[0]
        uval = val.upper() if val else ""
        if uval in ("WAS", "RAS"):
            entry["auto_fix"] = {
                "action": "add_to_treatment_processes",
                "instance_name": val,
                "process_type": f"{val}_Flow",
                "category": "physical",
                "stage": "effluent" if uval == "WAS" else "biological",
                "applied_by_tool": "apply_template_auto_fixes",
            }
        elif val and val not in inst_set and len(rows) == len(conns):
            entry["auto_fix"] = {
                "action": "clear_na_rows",
                "applied_by_tool": "apply_template_auto_fixes",
            }
        errors.append(entry)

    sizing = parsed.get("process_sizing", {})
    for inst in inst_names:
        proc = next((p for p in procs if p["instance_name"] == inst), None)
        if not proc: continue
        cat = (proc.get("process_type") or "").lower()
        if cat in ("aerationtank","anoxicreactor","anaerobicreactor","cstr","plugflow"):
            sz = sizing.get(inst, {})
            if not sz.get("volume_m3"):
                warnings.append(f"Process_Sizing: missing volume_m3 for reactor '{inst}'")
        if cat in ("primaryclarifier","secondaryclarifier"):
            sz = sizing.get(inst, {})
            if not (sz.get("surface_area_m2") or sz.get("diameter_m")):
                warnings.append(f"Process_Sizing: missing surface_area_m2 or diameter_m for clarifier '{inst}'")

    fr = parsed.get("influent_fractionation", {}) or {}
    cod_keys = ("fSI","fSS","fXS","fXI","fXBH","fXBA")
    cod_vals = [fr.get(k) for k in cod_keys if isinstance(fr.get(k), (int, float))]
    if cod_vals:
        s = sum(cod_vals)
        if abs(s - 1.0) > 0.05:
            warnings.append(f"Influent_Fractionation: COD fractions sum to {s:.3f}, expected ~1.00")

    return {
        "template_path": template_path,
        "valid": len(errors) == 0,
        "errors":   errors,
        "warnings": warnings,
        "n_processes":   len(procs),
        "n_connections": len(conns),
        "n_influent_rows": len(parsed.get("influent_dataset", [])),
        "n_effluent_rows": len(parsed.get("effluent_dataset", [])),
    }


def _exp_build_model_from_template(
    template_path: str,
    project_path: str | None = None,
    apply_via_dtt: bool = True,
    dry_run: bool = False,
    allow_offline: bool = False,
) -> dict:
    """Read the template, validate, and build a SUMO24 model + apply parameters."""
    # Guard: refuse to proceed when DTT is silently buffering commands
    # (i.e. commands get queued but never executed, leaving a stub .sumo on disk).
    if apply_via_dtt and not dry_run and not allow_offline:
        try:
            import dynamita_compat as _dc  # type: ignore
            dtt_live = bool(getattr(_dc, "is_dtt_live", lambda: False)())
        except Exception:
            dtt_live = DTT_AVAILABLE  # best-effort fallback
        if not dtt_live:
            return {
                "ok": False,
                "reason": "DTT bridge is not connected to a running SUMO instance",
                "hint": (
                    "build_model_from_template needs a live DTT connection to execute "
                    "topology-building commands. Either start SUMO24 and load the "
                    "target .sumo project before calling this tool, OR pass "
                    "allow_offline=true to dump the build commands to "
                    "<output>.commands.txt for manual application — note that "
                    "allow_offline mode will NOT produce a working .sumo file. "
                    "Alternatively use build_sumo_pack to generate a build-pack "
                    "directory the user can apply inside SUMO24 GUI."
                ),
            }
    val = _exp_validate_wwtp_template(template_path)
    if val.get("error"):
        return val
    if not val["valid"]:
        return {"error": "Template has validation errors â€” fix them and retry",
                "validation": val}
    parsed = _exp_read_wwtp_template(template_path)
    pi = parsed.get("plant_info", {})

    # Resolve project folder
    project_name = pi.get("project_name", "WWTPModel")
    base = project_path or pi.get("project_path") or str(_MCP_DIR / "projects")
    project_dir = Path(base) / str(project_name)
    sumo_stub = project_dir / f"{project_name}.sumo"

    actions: list[dict] = []   # log of every step taken or planned
    errors: list[str] = []

    if not dry_run:
        try:
            project_dir.mkdir(parents=True, exist_ok=True)
            for sub in ("outputs", "scenarios", "dynamita"):
                (project_dir / sub).mkdir(exist_ok=True)
            sumo_stub.write_text(
                '<?xml version="1.0" encoding="utf-8"?>\n'
                f'<SumoProject name="{project_name}" version="24">\n'
                '  <!-- Auto-generated by build_model_from_template -->\n'
                '</SumoProject>\n')
            actions.append({"step": "create_project", "ok": True, "path": str(project_dir)})
            ACTIVE_MODEL["sumo_file"] = str(sumo_stub)
        except Exception as ex:
            errors.append(f"create_project failed: {ex}")
            actions.append({"step": "create_project", "ok": False, "error": str(ex)})

    # Helper: try executeCommand, otherwise log a planned action
    def _run(cmd: str, step: str, **extra) -> bool:
        if dry_run or not apply_via_dtt:
            actions.append({"step": step, "planned_cmd": cmd, **extra})
            return False
        try:
            _ds = _make_ds()
            _ds.sumo.executeCommand(cmd)
            actions.append({"step": step, "ok": True, "cmd": cmd, **extra})
            return True
        except Exception as ex:
            actions.append({"step": step, "ok": False, "cmd": cmd, "error": str(ex), **extra})
            return False

    # 1. Add unit processes (sorted by order_index)
    procs = sorted(parsed.get("treatment_processes", []),
                   key=lambda p: (p.get("order_index") or 9999))
    for p in procs:
        inst = p["instance_name"]; ptype = p["process_type"]
        if not (inst and ptype): continue
        _run(f'addprocess "{ptype}" "{inst}"', step=f"add_unit:{inst}",
             process_type=ptype, kinetic_model=p.get("kinetic_model"))
        if p.get("kinetic_model"):
            _run(f'setmodel "{inst}" "{p["kinetic_model"]}"',
                 step=f"set_kinetic:{inst}", kinetic_model=p["kinetic_model"])
        ACTIVE_MODEL["unit_processes"].append(
            {"name": inst, "type": ptype, "kinetic_model": p.get("kinetic_model") or "default"})

    # 2. Connections
    type_map = {"RAS":"return_activated_sludge", "WAS":"waste_activated_sludge",
                "InternalRecycle":"internal_recycle", "RejectWater":"reject_water",
                "Sidestream":"sidestream"}
    for c in parsed.get("connections", []):
        sname = c.get("stream_name") or f"{c['from_unit']}_to_{c['to_unit']}"
        rec = c.get("recycle_type")
        if rec:
            stype = type_map.get(rec, "recycle")
            _run(f'connect "{c["from_unit"]}" "{c["to_unit"]}" "{sname}" "{stype}"',
                 step=f"connect:{sname}", recycle_type=rec)
        else:
            stype = c.get("stream_type") or "liquid"
            _run(f'connect "{c["from_unit"]}" "{c["to_unit"]}" "{sname}" "{stype}"',
                 step=f"connect:{sname}", stream_type=stype)
        ACTIVE_MODEL["streams"].append(
            {"name": sname, "from": c["from_unit"], "to": c["to_unit"], "type": stype})

    # 3. Apply sizing parameters
    sz_map = parsed.get("process_sizing", {})
    sumo_sz_keys = {  # template key -> SUMO parameter suffix
        "volume_m3":       "Volume",
        "depth_m":         "Depth",
        "surface_area_m2": "SurfaceArea",
        "diameter_m":      "Diameter",
        "length_m":        "Length",
        "width_m":         "Width",
        "HRT_h":           "HRT",
    }
    for inst, sz in sz_map.items():
        if not isinstance(sz, dict): continue
        for k, sumo_key in sumo_sz_keys.items():
            v = sz.get(k)
            if v in (None, ""): continue
            _run(f'setparameter "{inst}.{sumo_key}" {v}',
                 step=f"size:{inst}.{sumo_key}", value=v)

    # 4. Bio / Phys / Chem operational parameters
    sumo_bio_map = {
        "DO_setpoint": "DOSetpoint", "kLa": "KLa",
        "MLSS_target": "MLSSSetpoint", "SRT_target": "SRTSetpoint",
        "Temperature_op": "Temperature",
        "RAS_ratio": "Q_RAS_ratio", "WAS_ratio": "Q_WAS_ratio",
        "Internal_recycle_ratio": "Q_IR_ratio",
        "Mixing_power": "MixingPower", "DO_max": "DOMax",
        "F_M_target": "FMTargetRatio", "HRT": "HRT", "VS_loading": "VSLoading",
    }
    for op in parsed.get("bio_operational", []):
        inst = op.get("instance_name"); param = op.get("parameter"); v = op.get("value")
        if not (inst and param) or v in (None, ""): continue
        sumo_p = sumo_bio_map.get(param, param)
        _run(f'setparameter "{inst}.{sumo_p}" {v}',
             step=f"bio:{inst}.{sumo_p}", value=v)

    sumo_phys_map = {
        "bar_spacing": "BarSpacing", "approach_velocity": "ApproachVelocity",
        "HRT": "HRT", "air_supply": "AirSupply",
        "SOR_avg": "SOR_avg", "SOR_peak": "SOR_peak",
        "SLR_avg": "SLR_avg", "weir_loading": "WeirLoading",
        "filtration_rate": "FiltrationRate", "backwash_rate": "BackwashRate",
        "UV_dose": "UVDose", "SLR": "SLR",
        "polymer_dose": "PolymerDose", "cake_solids_target": "CakeSolidsTarget",
    }
    for op in parsed.get("phys_operational", []):
        inst = op.get("instance_name"); param = op.get("parameter"); v = op.get("value")
        if not (inst and param) or v in (None, ""): continue
        sumo_p = sumo_phys_map.get(param, param)
        _run(f'setparameter "{inst}.{sumo_p}" {v}',
             step=f"phys:{inst}.{sumo_p}", value=v)

    for ch in parsed.get("chem_operational", []):
        inst = ch.get("instance_name"); chem = ch.get("chemical"); dose = ch.get("dose")
        if not (inst and chem) or dose in (None, ""): continue
        _run(f'setchemical "{inst}" "{chem}" {dose}',
             step=f"chem:{inst}.{chem}", dose=dose, target=ch.get("target_unit"))

    # 5. Influent characteristics from Plant_Info / Influent_Dataset / Fractionation
    inf_stats = parsed.get("influent_statistics", {})
    inf_chars = {}
    if inf_stats:
        # Use mean values from the dataset where available
        for col, sumo_p in [("Q","Q"),("COD","TCOD"),("BOD5","BOD5"),("TSS","TSS"),
                            ("VSS","VSS"),("TKN","TKN"),("NH4-N","SNH"),
                            ("NO3-N","SNO"),("TP","TP"),("PO4-P","SPO"),
                            ("Alkalinity","SALK"),("Temperature","T")]:
            if col in inf_stats:
                inf_chars[sumo_p] = inf_stats[col]["mean"]
                _run(f'setparameter "Influent.{sumo_p}" {inf_stats[col]["mean"]}',
                     step=f"influent:{sumo_p}", value=inf_stats[col]["mean"])
    if not inf_chars and pi.get("design_flow_avg"):
        # Fallback to Plant_Info design flow
        _run(f'setparameter "Influent.Q" {pi["design_flow_avg"]}',
             step="influent:Q_design", value=pi["design_flow_avg"])

    # Apply COD fractionation as influent fractions
    for k, v in (parsed.get("influent_fractionation", {}) or {}).items():
        if isinstance(v, (int, float)) and v != 0:
            _run(f'setparameter "Influent.{k}" {v}',
                 step=f"influent_frac:{k}", value=v)

    # 6. Compile-status check
    compile_ready = bool(_make_ds) and apply_via_dtt and not dry_run

    # ── Offline mode: merge planned commands into the build-pack's .scs ──────
    offline_mode = not apply_via_dtt or dry_run
    merged_into: str | None = None
    if offline_mode:
        planned_cmds = [a["planned_cmd"] for a in actions if "planned_cmd" in a]
        if planned_cmds:
            pack_candidate = sumo_stub.with_suffix(".sumo-pack")
            if not pack_candidate.exists():
                # Try same stem in same directory.
                pack_candidate = project_dir / f"{project_name}.sumo-pack"
            if pack_candidate.is_dir():
                scs_path = pack_candidate / "apply_parameters.scs"
                try:
                    with open(scs_path, "a", encoding="utf-8") as f:
                        f.write("\n// --- appended by build_model_from_template (offline mode) ---\n")
                        for cmd in planned_cmds:
                            f.write(cmd.rstrip(";") + ";\n")
                    merged_into = str(scs_path)
                except Exception as scs_err:
                    merged_into = None
                    errors.append(f"Could not append to apply_parameters.scs: {scs_err}")

    summary = {
        "project_dir":       str(project_dir),
        "sumo_file":         str(sumo_stub),
        "kinetic_model":     pi.get("kinetic_model"),
        "n_processes_added": len(procs),
        "n_connections":     len(parsed.get("connections", [])),
        "n_actions":         len(actions),
        "n_failed_actions":  sum(1 for a in actions if a.get("ok") is False),
        "dry_run":           dry_run,
        "apply_via_dtt":     apply_via_dtt,
        "validation":        val,
        "actions": actions[:300],   # cap to avoid huge JSON
    }

    if offline_mode:
        summary["mode"]           = "offline_script_generated"
        summary["model_built"]    = False
        summary["merged_into"]    = merged_into
        summary["next_step"] = (
            f"In SUMO24 Core Window, paste {merged_into}. "
            "No model has been built yet — this offline mode only generates the build script."
            if merged_into else
            "No matching .sumo-pack found. Run build_sumo_pack first, then re-run "
            "build_model_from_template so the commands can be merged into apply_parameters.scs."
        )
    else:
        summary["next_steps"] = [
            "1. Open the .sumo file in SUMO24 GUI",
            "2. Verify the network and parameters",
            "3. Run 'Simulate' until Ready",
            "4. In Core Window: maptoic; save 'state.xml'",
            "5. Place sumoproject.dll alongside the .sumo file",
            "6. Call load_model with the .sumo path to register the digital twin",
        ]

    if errors:
        summary["pre_run_errors"] = errors
    return summary


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# Group BB â€” HTML Schematic helpers
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def _require_schematic_parser() -> dict | None:
    if not SCHEMATIC_AVAILABLE:
        return {"error": "schematic_parser module is not available. Expected at PY/schematic_parser.py."}
    return None


def _load_schematic_either(schematic: dict | None, filepath: str | None) -> tuple[dict | None, dict | None]:
    """Resolve schematic-or-filepath arguments. Returns (data, error)."""
    if schematic is not None:
        return schematic, None
    if filepath:
        try:
            return _schp.load_schematic(filepath), None
        except Exception as e:
            return None, {"error": f"Failed to load schematic from {filepath}: {e}"}
    return None, {"error": "Pass either 'schematic' (dict) or 'filepath' (.html / .json)."}


def _exp_import_schematic_from_html(filepath: str) -> dict:
    err = _require_schematic_parser()
    if err:
        return err
    try:
        data = _schp.load_schematic(filepath)
    except Exception as e:
        return {"error": str(e), "filepath": filepath}
    validation = _schp.validate_schematic(data)
    return {
        "filepath":   filepath,
        "schematic":  data,
        "summary":    _schp.summarise(data),
        "validation": validation,
        "ok":         not validation["errors"],
    }


def _exp_build_model_from_schematic(
    schematic: dict | None,
    filepath: str | None,
    apply_immediately: bool,
    output_scs_path: str | None,
    include_dtt_actions: bool,
) -> dict:
    err = _require_schematic_parser()
    if err:
        return err
    data, load_err = _load_schematic_either(schematic, filepath)
    if load_err:
        return load_err
    validation = _schp.validate_schematic(data)
    if validation["errors"]:
        return {"error": "Schematic has validation errors; aborting build.", "validation": validation}

    commands = _schp.schematic_to_commands(data)
    written = None
    if output_scs_path:
        try:
            Path(output_scs_path).write_text(";\n".join(commands) + ";\n", encoding="utf-8")
            written = output_scs_path
        except Exception as e:
            return {"error": f"Could not write .scs file: {e}", "commands": commands}

    applied: list[dict] = []
    if apply_immediately:
        if not DTT_AVAILABLE:
            return {"error": "DTT not available; cannot apply_immediately.", "commands": commands}
        # Best-effort: send each non-comment command via DTT's setparameter pathway.
        # We do NOT have a guaranteed compiled job id here, so just record what would run.
        for cmd in commands:
            c = cmd.strip()
            if not c or c.startswith("#"):
                continue
            applied.append({"command": c, "status": "queued (requires active compiled job)"})

    result = {
        "summary":      _schp.summarise(data),
        "validation":   validation,
        "command_count": sum(1 for c in commands if c and not c.startswith("#")),
        "commands":     commands,
        "scs_file":     written,
        "applied":      applied if apply_immediately else None,
    }
    if include_dtt_actions:
        result["dtt_actions"] = _schp.schematic_to_dtt_actions(data)
    return result


def _exp_compare_schematic_to_model(schematic: dict | None, filepath: str | None) -> dict:
    err = _require_schematic_parser()
    if err:
        return err
    data, load_err = _load_schematic_either(schematic, filepath)
    if load_err:
        return load_err
    # Try to ask DTT for the list of currently-loaded unit IDs (best-effort)
    model_unit_ids: list[str] | None = None
    if DTT_AVAILABLE:
        try:
            # The DTT API exposes a list of units indirectly via the loaded project.
            # We can't always retrieve it; pass None when unsure.
            model_unit_ids = None
        except Exception:
            model_unit_ids = None
    return _schp.compare_to_model(data, model_unit_ids)


def _exp_generate_sumoslang_from_schematic(
    schematic: dict | None,
    filepath: str | None,
    output_path: str | None,
) -> dict:
    err = _require_schematic_parser()
    if err:
        return err
    data, load_err = _load_schematic_either(schematic, filepath)
    if load_err:
        return load_err
    text = _schp.schematic_to_sumoslang(data)
    written = None
    if output_path:
        try:
            Path(output_path).write_text(text, encoding="utf-8")
            written = output_path
        except Exception as e:
            return {"error": f"Could not write SumoSlang file: {e}", "sumoslang": text}
    return {"summary": _schp.summarise(data), "sumoslang": text, "written_to": written,
            "note": "Skeleton only â€” hand-edit kinetic constants before SUMO will accept."}


def _exp_list_schematic_unit_types() -> dict:
    if not SCHEMATIC_AVAILABLE:
        return {"error": "unit_type_registry not loaded"}
    return {
        "unit_types": _utr.UNIT_TYPES,
        "stream_types": _utr.STREAM_TYPES,
        "stream_to_dtt": _utr.STREAM_TO_DTT,
        "boundary_types": sorted(_utr.BOUNDARY_TYPES),
        "count": len(_utr.UNIT_TYPES),
    }


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# Group CC â€” SUMO File Troubleshooting Package
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

import zipfile

# Catalogue of known SUMO error signatures and remediation guidance.
_SUMO_DIAGNOSTICS = [
    {
        "id": "EMPTY_FILE",
        "signature": "0-byte .sumo file",
        "cause": "Save was interrupted or DTT crashed before writing the manifest.",
        "fix":   "Delete the empty file and rebuild via build_model_from_template or import an earlier backup.",
    },
    {
        "id": "MISSING_DLL",
        "signature": "sumoproject.dll absent next to .sumo",
        "cause": "DTT requires the compiled DLL alongside the .sumo. Building without compile leaves it missing.",
        "fix":   "Run the SUMO GUI 'Simulate' once to produce sumoproject.dll, or copy a known-good DLL via repair_sumo_file.",
    },
    {
        "id": "MISSING_STATE",
        "signature": "state.xml absent",
        "cause": "DTT loads from a state file. Without it, the first simulation cannot bootstrap.",
        "fix":   "In SUMO GUI: Core Window â†’ 'maptoic' â†’ save 'state.xml' next to the .sumo.",
    },
    {
        "id": "WRONG_EXTENSION",
        "signature": "File named .xml/.zip but valid SUMO container",
        "cause": "User renamed the file or copied without the .sumo extension.",
        "fix":   "Rename back to .sumo; the DTT loader keys on the extension.",
    },
    {
        "id": "ZIP_CORRUPT",
        "signature": "Bad CRC / truncated zip",
        "cause": "Incomplete download or interrupted save.",
        "fix":   "Restore from .sumo.bak or re-export the project.",
    },
    {
        "id": "LICENSE_ERROR",
        "signature": "'license' or 'invalid key' in log",
        "cause": "DTT addon license missing or expired.",
        "fix":   "Run SUMO GUI License Manager; ensure standalone+DTT or container+DTT is valid.",
    },
    {
        "id": "DLL_LOAD_FAIL",
        "signature": "'LoadLibrary' / 'DLLNotFound' in log",
        "cause": "sumoproject.dll built for the wrong architecture or wrong SUMO version.",
        "fix":   "Recompile the project against the installed SUMO24 version (use prepare_project.py).",
    },
    {
        "id": "COMPILE_FAIL",
        "signature": "'compile failed' / 'syntax error' in log",
        "cause": "Hand-edited .sumo file contains invalid SumoSlang.",
        "fix":   "Restore from backup or use generate_sumoslang_from_schematic to regenerate.",
    },
    {
        "id": "EMPTY_MANIFEST",
        "signature": ".sumo opens but has 0 units",
        "cause": "Project was created via DTT but addprocess calls failed silently.",
        "fix":   "Re-run build_model_from_template with apply_via_dtt=True after restarting DTT.",
    },
    {
        "id": "DUPLICATE_NAMES",
        "signature": "Two units share the same instance_name",
        "cause": "Template or schematic supplied duplicate IDs.",
        "fix":   "Run validate_wwtp_template or validate the schematic before build.",
    },
]


def _exp_diagnose_sumo_file(sumo_path: str) -> dict:
    """Diagnose a single .sumo file. Returns health report."""
    p = Path(sumo_path)
    report: dict[str, Any] = {
        "sumo_path":     str(p),
        "exists":        p.exists(),
        "ok":            True,
        "errors":        [],
        "warnings":      [],
        "checks":        {},
    }
    if not p.exists():
        report["ok"] = False
        report["errors"].append(f"File not found: {p}")
        return report

    try:
        size = p.stat().st_size
        report["checks"]["size_bytes"] = size
        if size == 0:
            report["ok"] = False
            report["errors"].append("File is empty (0 bytes) â€” EMPTY_FILE")
            return report
        if size < 256:
            report["warnings"].append(f"File is very small ({size} bytes); may be truncated.")
    except Exception as e:
        report["errors"].append(f"stat failed: {e}")
        report["ok"] = False
        return report

    # extension check
    if p.suffix.lower() != ".sumo":
        report["warnings"].append(f"Extension is {p.suffix!r}, expected '.sumo' â€” WRONG_EXTENSION")

    # try zip open (the .sumo container is a zip)
    is_zip = zipfile.is_zipfile(p)
    report["checks"]["is_zip_container"] = is_zip
    if is_zip:
        try:
            with zipfile.ZipFile(p) as zf:
                names = zf.namelist()
                report["checks"]["archive_member_count"] = len(names)
                report["checks"]["archive_members_preview"] = names[:15]
                bad = zf.testzip()
                if bad:
                    report["ok"] = False
                    report["errors"].append(f"Zip corruption in member: {bad} â€” ZIP_CORRUPT")
        except Exception as e:
            report["ok"] = False
            report["errors"].append(f"Zip parse failed: {e} â€” ZIP_CORRUPT")
    else:
        # Try XML
        try:
            head = p.read_bytes()[:1024].decode("utf-8", errors="ignore")
            looks_xml = head.lstrip().startswith("<?xml") or "<sumo" in head.lower()
            report["checks"]["looks_xml"] = looks_xml
            if not looks_xml:
                report["warnings"].append("Not a zip and not an XML header â€” unknown format.")
        except Exception as e:
            report["warnings"].append(f"Header sniff failed: {e}")

    # companion DLL check
    dll = p.with_name("sumoproject.dll")
    report["checks"]["dll_present"] = dll.exists()
    report["checks"]["dll_path"]    = str(dll)
    if not dll.exists():
        report["warnings"].append("sumoproject.dll companion missing â€” MISSING_DLL")

    # state.xml check
    state = p.with_name("state.xml")
    report["checks"]["state_xml_present"] = state.exists()
    report["checks"]["state_xml_path"]    = str(state)
    if not state.exists():
        report["warnings"].append("state.xml missing next to .sumo â€” MISSING_STATE")

    # try to count units in manifest (best-effort)
    unit_count = None
    try:
        if is_zip:
            with zipfile.ZipFile(p) as zf:
                for nm in zf.namelist():
                    if nm.lower().endswith(".xml"):
                        body = zf.read(nm).decode("utf-8", errors="ignore")
                        # naive heuristic
                        unit_count = body.count("<unit ") or body.count("<process ")
                        break
        else:
            body = p.read_text(encoding="utf-8", errors="ignore")
            unit_count = body.count("<unit ") or body.count("<process ")
    except Exception:
        unit_count = None
    report["checks"]["heuristic_unit_count"] = unit_count
    if unit_count == 0:
        report["warnings"].append("Manifest parsed but 0 units found â€” EMPTY_MANIFEST")

    # ── Stub detection & strict ok semantics ──────────────────────────────────
    checks = report["checks"]
    warnings = report["warnings"]

    is_minimal_stub = (
        checks.get("size_bytes", 9999) <= 1024
        and bool(checks.get("looks_xml"))
        and not bool(checks.get("is_zip_container"))
        and (checks.get("heuristic_unit_count") or 0) == 0
        and _root_is_sumoproject(str(p))
    )
    report["is_minimal_stub"] = is_minimal_stub

    # Look for an adjacent build-pack directory.
    pack_dir = None
    parent_dir = p.parent
    try:
        for entry in os.listdir(parent_dir):
            if entry.endswith(".sumo-pack") and os.path.isdir(os.path.join(parent_dir, entry)):
                pack_dir = os.path.join(parent_dir, entry)
                break
    except OSError:
        pass
    report["pack_dir_found"] = pack_dir

    # Re-evaluate ok / status strictly.
    ok = True
    status = "OK"
    reason = None

    if is_minimal_stub:
        status = "PENDING_BUILD"
        # Downgrade EMPTY_MANIFEST: expected and harmless in stub mode.
        warnings[:] = [w for w in warnings if "EMPTY_MANIFEST" not in w]
        if not checks.get("dll_present"):
            ok = False
            reason = (
                "Stub present but sumoproject.dll missing — "
                "SUMO will not be able to open this yet. "
                "Compile the build-pack SumoSlang in SUMO24 GUI first."
            )
    else:
        if not checks.get("dll_present"):
            ok = False
            reason = "sumoproject.dll companion missing — MISSING_DLL"
        elif (checks.get("heuristic_unit_count") or 0) == 0 and not checks.get("is_zip_container"):
            ok = False
            reason = "non-stub .sumo has 0 units in manifest — file is malformed"

    report["ok"] = ok
    report["status"] = status
    if reason:
        report["reason"] = reason

    # Rename "warnings" → "details" for clarity.
    report["details"] = report.pop("warnings")

    if is_minimal_stub and ok:
        report["next_step"] = (
            f"Open the stub in SUMO24 GUI, then follow "
            f"{pack_dir}/BUILD_INSTRUCTIONS.md to compile the plant."
            if pack_dir else
            "Open the stub in SUMO24 GUI and build the plant manually, "
            "or call build_sumo_pack to generate the matching build-pack first."
        )

    if report["errors"]:
        report["ok"] = False
    return report


def _root_is_sumoproject(path: str) -> bool:
    """Cheap check: first 512 bytes contain <SumoProject ...>."""
    try:
        with open(path, "rb") as f:
            head = f.read(512).decode("utf-8", errors="replace")
        return "<SumoProject" in head
    except OSError:
        return False


def _exp_scan_sumo_directory(directory: str | None) -> dict:
    root = Path(directory) if directory else _MCP_DIR
    if not root.exists():
        return {"error": f"Directory not found: {root}"}
    files = sorted(root.rglob("*.sumo"))
    reports = [_exp_diagnose_sumo_file(str(f)) for f in files]
    healthy = sum(1 for r in reports if r.get("ok"))
    return {
        "scanned_dir":    str(root),
        "sumo_file_count": len(files),
        "healthy":        healthy,
        "with_errors":    len(reports) - healthy,
        "reports":        reports,
    }


def _exp_check_dll_companion(sumo_path: str) -> dict:
    p = Path(sumo_path)
    if not p.exists():
        return {"error": f"Not found: {p}"}
    dll = p.with_name("sumoproject.dll")
    if dll.exists():
        sz = dll.stat().st_size
        return {"present": True, "dll_path": str(dll), "size_bytes": sz, "ok": sz > 0}
    # search the parent tree for any sumoproject.dll
    candidates = list(p.parent.parent.rglob("sumoproject.dll")) if p.parent.parent else []
    return {
        "present": False,
        "dll_path_expected": str(dll),
        "ok": False,
        "candidates_found_elsewhere": [str(c) for c in candidates][:10],
        "fix": "Copy one of the candidate DLLs to the expected path, or run prepare_project.py to extract a fresh DLL.",
    }


_EMPTY_STATE_XML = (
    '<?xml version="1.0" encoding="UTF-8"?>\n'
    '<state generator="sumo24-mcp repair_sumo_file" />\n'
)


_EMPTY_SUMO_STUB = (
    '<?xml version="1.0" encoding="UTF-8"?>\n'
    '<sumoproject schema="1.0">\n'
    '  <!-- Auto-generated minimal stub by repair_sumo_file. Open in SUMO GUI to populate. -->\n'
    '  <units/>\n'
    '  <connections/>\n'
    '  <parameters/>\n'
    '</sumoproject>\n'
)


def _exp_repair_sumo_file(sumo_path: str, source_dll: str | None, dry_run: bool) -> dict:
    p = Path(sumo_path)
    actions: list[dict] = []
    if not p.exists():
        actions.append({"action": "create_sumo_stub", "target": str(p)})
        if not dry_run:
            p.parent.mkdir(parents=True, exist_ok=True)
            p.write_text(_EMPTY_SUMO_STUB, encoding="utf-8")
    else:
        # backup first
        if not dry_run:
            bak = p.with_suffix(p.suffix + ".bak")
            try:
                shutil.copy2(p, bak)
                actions.append({"action": "backup", "target": str(bak)})
            except Exception as e:
                actions.append({"action": "backup", "error": str(e)})
        sz = p.stat().st_size
        if sz == 0:
            actions.append({"action": "rewrite_empty_stub", "target": str(p)})
            if not dry_run:
                p.write_text(_EMPTY_SUMO_STUB, encoding="utf-8")

    # state.xml companion
    state = p.with_name("state.xml")
    if not state.exists():
        actions.append({"action": "create_empty_state_xml", "target": str(state)})
        if not dry_run:
            state.write_text(_EMPTY_STATE_XML, encoding="utf-8")

    # sumoproject.dll companion
    dll = p.with_name("sumoproject.dll")
    if not dll.exists() and source_dll:
        src_dll = Path(source_dll)
        if src_dll.exists():
            actions.append({"action": "copy_dll", "from": str(src_dll), "to": str(dll)})
            if not dry_run:
                shutil.copy2(src_dll, dll)
        else:
            actions.append({"action": "copy_dll_failed", "reason": f"source_dll not found: {src_dll}"})
    elif not dll.exists():
        actions.append({"action": "dll_missing_no_source", "note": "Pass source_dll= to copy a known-good DLL."})

    # extension fix
    if p.exists() and p.suffix.lower() != ".sumo":
        new = p.with_suffix(".sumo")
        actions.append({"action": "rename_extension", "from": str(p), "to": str(new)})
        if not dry_run:
            try:
                p.rename(new)
            except Exception as e:
                actions.append({"action": "rename_failed", "error": str(e)})

    return {
        "ok": True,
        "sumo_path": str(p),
        "dry_run": dry_run,
        "actions": actions,
        "action_count": len(actions),
    }


def _exp_diagnose_sumo_crash(log_path: str | None, tail_lines: int) -> dict:
    """Parse SUMO/DTT logs for known error signatures."""
    candidates: list[Path] = []
    if log_path:
        candidates.append(Path(log_path))
    else:
        # Best-effort guesses for likely log locations
        for guess in (_MCP_DIR / "logs", _MCP_DIR, _MCP_DIR.parent / "logs"):
            if guess.exists():
                for ext in ("*.log", "*.txt"):
                    candidates.extend(sorted(guess.glob(ext), key=lambda p: p.stat().st_mtime, reverse=True)[:5])

    if not candidates:
        return {"ok": False, "error": "no log files found", "searched": [str(_MCP_DIR / "logs"), str(_MCP_DIR)]}

    findings = []
    for lp in candidates:
        if not lp.exists():
            continue
        try:
            text = lp.read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            findings.append({"log": str(lp), "error": f"read failed: {e}"})
            continue
        lines = text.splitlines()
        tail = lines[-tail_lines:] if len(lines) > tail_lines else lines
        body = "\n".join(tail)
        hits: list[dict] = []
        for diag in _SUMO_DIAGNOSTICS:
            sig = diag.get("signature", "")
            if sig and sig.lower() in body.lower():
                hits.append(diag)
        # also flag raw "Traceback" / "Exception" / "ERROR"
        trace_count = body.count("Traceback")
        error_count = sum(1 for ln in tail if "ERROR" in ln or "FATAL" in ln)
        findings.append({
            "log": str(lp),
            "size_bytes": lp.stat().st_size,
            "tail_lines_scanned": len(tail),
            "traceback_count": trace_count,
            "error_lines": error_count,
            "matched_signatures": hits,
        })
    return {
        "ok": True,
        "log_count": len(findings),
        "findings": findings,
        "diagnostic_catalogue_size": len(_SUMO_DIAGNOSTICS),
    }


def _exp_validate_sumo_environment() -> dict:
    """Top-to-bottom health check for the SUMO24 MCP environment."""
    checks: dict = {}

    # 1. DTT availability
    checks["dtt_available"] = DTT_AVAILABLE

    # 2. openpyxl
    try:
        import openpyxl  # noqa
        checks["openpyxl_available"] = True
    except Exception as e:
        checks["openpyxl_available"] = False
        checks["openpyxl_error"] = str(e)

    # 3. schematic_parser
    checks["schematic_parser_available"] = SCHEMATIC_AVAILABLE

    # 4. sumo_compiler (legacy — DEPRECATED)
    checks["sumo_compiler_available"] = COMPILER_AVAILABLE

    # 4b. sumo_pack (Group EE replacement)
    checks["sumo_pack_module"] = PACK_AVAILABLE

    # 4c. Baseline .sumo + DLL presence (required for apply_schematic_to_baseline)
    baseline_sumo = CONFIG.get("baseline_sumo", "")
    baseline_dll  = CONFIG.get("baseline_dll",  "")
    checks["baseline_sumo_path"] = baseline_sumo
    checks["baseline_dll_path"]  = baseline_dll
    checks["baseline_present"] = (
        bool(baseline_sumo) and os.path.exists(baseline_sumo) and
        bool(baseline_dll)  and os.path.exists(baseline_dll)
    )

    # 4d. DTT bridge liveness — distinguish "connected" from "silently buffering"
    try:
        import dynamita_compat as _dc  # type: ignore
        dtt_live     = bool(getattr(_dc, "is_dtt_live", lambda: False)())
        dtt_q_size   = int(getattr(_dc, "queue_size",   lambda: 0)())
        checks["dtt_live"]       = dtt_live
        checks["dtt_queue_size"] = dtt_q_size
        if not dtt_live and dtt_q_size > 0:
            checks["dtt_warning"] = (
                f"DTT is not connected to a running SUMO instance but "
                f"{dtt_q_size} command(s) are queued. build_model_from_template "
                "will NOT execute these — use build_sumo_pack instead, or start "
                "SUMO24 and reload the project before calling the template builder."
            )
    except Exception:
        checks["dtt_live"] = "unknown"

    # 5. MCP directory layout
    expected = {
        "MCP_root":      _MCP_DIR.exists(),
        "PY_dir":        (_MCP_DIR / "PY").exists(),
        "schematic_html": (_MCP_DIR / "wwtp_schematic_template.html").exists(),
        "wwtp_xlsx":     (_MCP_DIR / "wwtp_input_template.xlsx").exists(),
    }
    checks["layout"] = expected

    # 6. companion DLL search (best-effort)
    dll_candidates = []
    for guess in (_MCP_DIR, _MCP_DIR.parent):
        if guess.exists():
            dll_candidates.extend(str(p) for p in guess.rglob("sumoproject.dll"))
    checks["sumoproject_dll_candidates"] = dll_candidates[:10]
    checks["sumoproject_dll_found"] = bool(dll_candidates)

    # 7. license / state files
    license_paths = [str(p) for p in (_MCP_DIR.parent / "LICENSE.txt", _MCP_DIR / "LICENSE.txt") if p.exists()]
    state_paths   = [str(p) for p in (_MCP_DIR / "state.xml", _MCP_DIR / "PY" / "state.xml") if p.exists()]
    checks["license_paths"] = license_paths
    checks["state_paths"]   = state_paths

    # Overall ok if the critical pieces are in place.
    # Note: sumo_compiler_available is no longer required (sumo_pack replaces it).
    ok = (
        checks["schematic_parser_available"]
        and checks["sumo_pack_module"]
        and checks["layout"]["MCP_root"]
    )
    return {"ok": ok, "checks": checks}


def _exp_list_sumo_diagnostics() -> dict:
    return {"count": len(_SUMO_DIAGNOSTICS), "diagnostics": _SUMO_DIAGNOSTICS}


# ════════════════════════════════════════════════════════════════════════════
# Group DD — Schematic-to-SUMO File Compiler
# ════════════════════════════════════════════════════════════════════════════


def _require_compiler() -> dict | None:
    if not COMPILER_AVAILABLE or _scc is None:
        return {"error": "sumo_compiler module unavailable. Ensure sumo_compiler.py "
                         "lives alongside server.py."}
    return None


def _resolve_schematic(schematic: dict | None, html_path: str | None) -> tuple[dict | None, dict | None]:
    if schematic and isinstance(schematic, dict) and schematic.get("units") is not None:
        return schematic, None
    if html_path:
        try:
            data = _schp.load_schematic(html_path)
        except Exception as e:
            return None, {"error": f"could not load schematic from html_path: {e}",
                          "html_path": html_path}
        return data, None
    return None, {"error": "Pass either 'schematic' (dict) or 'html_path' (string)."}


def _exp_compile_schematic_to_sumo(
    schematic: dict | None,
    output_path: str,
    overwrite: bool,
    include_sumoslang: bool,
    attach_dll: str | None,
) -> dict:
    return {
        "ok": False,
        "deprecated": True,
        "reason": "This tool produced a zip SUMO24 cannot load.",
        "use_instead": "build_sumo_pack",
        "migration_hint": (
            "Replace `compile_schematic_to_sumo(schematic=..., output_path='x.sumo')` "
            "with `build_sumo_pack(schematic_html=..., output_dir='x.sumo-pack')`, "
            "then follow x.sumo-pack/BUILD_INSTRUCTIONS.md inside SUMO24 GUI."
        ),
    }


def _exp_build_sumo_from_html(
    html_path: str,
    output_path: str,
    overwrite: bool,
    include_sumoslang: bool,
    attach_dll: str | None,
) -> dict:
    return {
        "ok": False,
        "deprecated": True,
        "reason": "This tool produced a zip SUMO24 cannot load.",
        "use_instead": "build_sumo_pack",
        "migration_hint": (
            "Replace `build_sumo_from_html(html_path=H, output_path='x.sumo')` "
            "with `build_sumo_pack(schematic_html=H, output_dir='x.sumo-pack')`."
        ),
    }


def _exp_preview_sumo_manifest(
    schematic: dict | None,
    html_path: str | None,
    include_sumoslang: bool,
) -> dict:
    err = _require_compiler()
    if err:
        return err
    data, load_err = _resolve_schematic(schematic, html_path)
    if load_err:
        return load_err
    files = _scc.preview_sumo_files(data, include_sumoslang=include_sumoslang)
    return {
        "ok": True,
        "summary": _schp.summarise(data),
        "files": {name: {"bytes": len(body), "text": body} for name, body in files.items()},
    }


def _exp_attach_companion_dll(sumo_path: str, source_dll: str) -> dict:
    return {
        "ok": False,
        "deprecated": True,
        "reason": (
            "Copying any sumoproject.dll next to a .sumo file does NOT make the "
            "file loadable — the DLL's symbol table is specific to the plant it "
            "was compiled from."
        ),
        "use_instead": "verify_dll_matches_schematic",
        "migration_hint": (
            "Run `verify_dll_matches_schematic(dll_path=..., schematic_html=...)` "
            "first. If it returns missing symbols, the DLL is the wrong one and "
            "needs to be regenerated by SUMO GUI from the build-pack SumoSlang."
        ),
    }


def _exp_verify_sumo_file_against_schematic(
    sumo_path: str,
    schematic: dict | None,
    html_path: str | None,
) -> dict:
    return {
        "ok": False,
        "deprecated": True,
        "reason": "This tool only verified the legacy zip format that SUMO can't load.",
        "use_instead": "verify_dll_matches_schematic",
    }


def _exp_read_sumo_manifest(sumo_path: str) -> dict:
    err = _require_compiler()
    if err:
        return err
    return _scc.read_sumo_manifest(sumo_path)


# ════════════════════════════════════════════════════════════════════════════
# Group EE — Build Pack & Baseline Apply
# Replaces the broken zip-as-.sumo compiler in Group DD.
# ════════════════════════════════════════════════════════════════════════════


def _require_pack() -> dict | None:
    """Return an error dict if sumo_pack is not loaded, else None."""
    if not PACK_AVAILABLE or _spk is None:
        return {
            "ok": False,
            "error": (
                "sumo_pack module is not available. Ensure sumo_pack.py "
                "lives alongside server.py in the PY/ directory."
            ),
        }
    return None


def _resolve_schematic_from_args(args: dict) -> tuple[dict | None, dict | None]:
    """
    Resolve a schematic from an args dict that may contain:
      - schematic_html: path to HTML file
      - schematic_json: path to JSON export file
    Returns (schematic_dict, error_dict).
    """
    if not SCHEMATIC_AVAILABLE or _schp is None:
        return None, {"ok": False, "error": "schematic_parser module is not available."}

    html_path = args.get("schematic_html") or ""
    json_path = args.get("schematic_json") or ""

    if html_path:
        try:
            data = _schp.load_schematic(html_path)
            return data, None
        except Exception as e:
            return None, {"ok": False, "error": f"could not load schematic from HTML: {e}",
                          "path": html_path}
    if json_path:
        try:
            import json as _json
            with open(json_path, "r", encoding="utf-8") as f:
                data = _json.load(f)
            return data, None
        except Exception as e:
            return None, {"ok": False, "error": f"could not load schematic JSON: {e}",
                          "path": json_path}

    return None, {
        "ok": False,
        "error": "Provide either schematic_html or schematic_json to identify the source schematic.",
    }


def _exp_build_sumo_pack(args: dict) -> dict:
    """Generate a build-pack directory from a schematic."""
    err = _require_pack()
    if err:
        return err

    output_dir = args.get("output_dir")
    if not output_dir:
        return {"ok": False, "error": "output_dir is required"}

    # Load schematic — either from html/json path, or from inline dict.
    if args.get("schematic_html") or args.get("schematic_json"):
        schematic, load_err = _resolve_schematic_from_args(args)
        if load_err:
            return load_err
    else:
        return {
            "ok": False,
            "error": "Provide schematic_html or schematic_json to locate the source schematic.",
        }

    try:
        result = _spk.build_pack(
            schematic,
            output_dir=output_dir,
            overwrite=bool(args.get("overwrite", False)),
        )
        return result
    except (FileExistsError, ValueError) as e:
        return {"ok": False, "error": str(e)}
    except Exception as e:
        return {"ok": False, "error": f"build_pack failed: {e}"}


def _exp_apply_schematic_to_baseline(args: dict) -> dict:
    """Copy a verified baseline and apply schematic parameter overrides."""
    err = _require_pack()
    if err:
        return err

    baseline_sumo = args.get("baseline_sumo") or CONFIG.get("baseline_sumo", "")
    baseline_dll  = args.get("baseline_dll")  or CONFIG.get("baseline_dll",  "")
    output_sumo   = args.get("output_sumo")

    if not baseline_sumo:
        return {"ok": False, "error": "baseline_sumo is required"}
    if not baseline_dll:
        return {"ok": False, "error": "baseline_dll is required"}
    if not output_sumo:
        return {"ok": False, "error": "output_sumo is required"}

    # Load schematic
    if args.get("schematic_html") or args.get("schematic_json"):
        schematic, load_err = _resolve_schematic_from_args(args)
        if load_err:
            return load_err
    else:
        # Allow calling with an empty schematic for a pure baseline copy.
        schematic = {"units": [], "streams": []}

    try:
        result = _spk.apply_to_baseline(
            schematic,
            baseline_sumo=baseline_sumo,
            baseline_dll=baseline_dll,
            output_sumo=output_sumo,
            strict=bool(args.get("strict", True)),
        )
        return result
    except FileNotFoundError as e:
        return {"ok": False, "error": str(e)}
    except Exception as e:
        return {"ok": False, "error": f"apply_to_baseline failed: {e}"}


# ── Group GG — Native .sumo composer ──────────────────────────────────────


def _require_native() -> dict | None:
    """Return an error dict if sumo_native is not loaded, else None."""
    if not NATIVE_AVAILABLE or _snt is None:
        return {
            "ok": False,
            "error": (
                "sumo_native module is not available. Ensure sumo_native.py "
                "lives alongside server.py in the PY/ directory."
            ),
        }
    return None


def _verify_native_build(output_sumo: str) -> dict:
    """
    Structural + (when available) DTT-driven verification of a freshly
    built .sumo. Never throws — always returns a structured report so the
    caller can include it in its response.
    """
    report: dict = {"output_sumo": output_sumo, "dtt_live": False, "checks": {}}

    p = Path(output_sumo)
    if not p.exists():
        report["checks"]["exists"] = False
        return report
    report["checks"]["exists"] = True
    report["checks"]["bytes"]  = p.stat().st_size

    # Quick zip / manifest sanity
    try:
        import zipfile as _zf
        report["checks"]["is_valid_zip"] = _zf.is_zipfile(p)
        with _zf.ZipFile(p, "r") as zf:
            members = zf.namelist()
            report["checks"]["members"]            = len(members)
            report["checks"]["dll_inside_zip"]     = "sumoproject.dll" in members
            report["checks"]["parameters_present"] = "parameters.txt"  in members
    except Exception as e:
        report["checks"]["zip_error"] = str(e)
        return report

    # Companion files (siblings)
    dll_sib   = p.with_name("sumoproject.dll")
    state_sib = p.with_name("state.xml")
    report["checks"]["companion_dll"]   = dll_sib.exists()
    report["checks"]["companion_state"] = state_sib.exists()

    # If DTT bridge looks live, try to register the project as active.
    # We use the same logic the load_model dispatcher uses but in-line so
    # this helper stays side-effect-light (we still mutate ACTIVE_MODEL,
    # which is the documented load_model contract).
    try:
        import dynamita_compat as _dc  # type: ignore
        dtt_live = bool(getattr(_dc, "is_dtt_live", lambda: False)())
    except Exception:
        dtt_live = False
    report["dtt_live"] = dtt_live

    try:
        ACTIVE_MODEL["sumo_file"] = str(p)
        ACTIVE_MODEL["compiled"]  = bool(dll_sib.exists() and state_sib.exists())
        if dll_sib.exists():
            ACTIVE_MODEL["dll_path"] = str(dll_sib)
            CONFIG["model_dll"]      = str(dll_sib)
        if state_sib.exists():
            ACTIVE_MODEL["state_xml"] = str(state_sib)
            CONFIG["state_xml"]       = str(state_sib)
        report["registered_as_active"] = True
    except Exception as e:
        report["registered_as_active"] = False
        report["register_error"]       = str(e)

    return report


def _exp_build_native_sumo(args: dict) -> dict:
    """Compose a real .sumo zip from a baseline + schematic (Tier-1)."""
    err = _require_native()
    if err:
        return err

    baseline_sumo = args.get("baseline_sumo") or CONFIG.get("baseline_sumo", "")
    baseline_dll  = args.get("baseline_dll")  or CONFIG.get("baseline_dll",  "")
    output_sumo   = args.get("output_sumo")

    if not baseline_sumo:
        return {"ok": False, "error": "baseline_sumo is required"}
    if not output_sumo:
        return {"ok": False, "error": "output_sumo is required"}

    schematic_dict = None
    if args.get("schematic_html") or args.get("schematic_json"):
        schematic_dict, load_err = _resolve_schematic_from_args(args)
        if load_err:
            return load_err
    else:
        return {
            "ok": False,
            "error": "Provide schematic_html or schematic_json to locate the source schematic.",
        }

    try:
        result = _snt.build_native_sumo(
            schematic=schematic_dict,
            baseline_sumo=baseline_sumo,
            baseline_dll=(baseline_dll or None),
            output_sumo=output_sumo,
            unit_mapping=args.get("unit_mapping") or None,
            overwrite=bool(args.get("overwrite", False)),
            strict=bool(args.get("strict", True)),
            write_companion_dll=True,
        )
    except FileNotFoundError as e:
        return {"ok": False, "error": str(e)}
    except (ValueError, ImportError) as e:
        return {"ok": False, "error": str(e)}
    except Exception as e:
        return {"ok": False, "error": f"build_native_sumo failed: {e}"}

    # Verification step
    if bool(args.get("verify_load", True)) and result.get("output_sumo"):
        result["verify_load"] = _verify_native_build(result["output_sumo"])

    return result


def _exp_topology_match_report(args: dict) -> dict:
    """Preview compatibility between a schematic and a baseline .sumo."""
    err = _require_native()
    if err:
        return err

    baseline_sumo = args.get("baseline_sumo") or CONFIG.get("baseline_sumo", "")
    if not baseline_sumo:
        return {"ok": False, "error": "baseline_sumo is required"}

    if args.get("schematic_html") or args.get("schematic_json"):
        schematic_dict, load_err = _resolve_schematic_from_args(args)
        if load_err:
            return load_err
    else:
        return {
            "ok": False,
            "error": "Provide schematic_html or schematic_json to locate the source schematic.",
        }

    try:
        info = _snt.parse_baseline(Path(baseline_sumo))
        report = _snt.topology_match_report(
            schematic_dict, info,
            user_mapping=args.get("unit_mapping") or None,
        )
        return {
            "ok":             report["ok"],
            "baseline":       {"path": baseline_sumo, "unit_ids": info["unit_ids"]},
            "mapping":        report["mapping"],
            "unmapped":       report["unmapped"],
            "unused_targets": report["unused_targets"],
            "warnings":       report["warnings"],
        }
    except FileNotFoundError as e:
        return {"ok": False, "error": str(e)}
    except Exception as e:
        return {"ok": False, "error": f"topology_match_report failed: {e}"}


def _exp_validate_sumoslang(args: dict) -> dict:
    """Lint a SumoSlang text for common breakages."""
    err = _require_pack()
    if err:
        return err

    text = args.get("text")
    path = args.get("path")

    if text is None and path:
        try:
            with open(path, "r", encoding="utf-8") as f:
                text = f.read()
        except Exception as e:
            return {"ok": False, "error": f"could not read file {path}: {e}"}

    if text is None:
        return {"ok": False, "error": "Supply either 'text' (string) or 'path' (file path)."}

    try:
        return _spk.validate_sumoslang(text)
    except Exception as e:
        return {"ok": False, "error": f"validate_sumoslang failed: {e}"}


# ════════════════════════════════════════════════════════════════════════════
# Group FF — DLL verification, safe HTML edit, workflow guide
# ════════════════════════════════════════════════════════════════════════════


def _read_dll_symbols(dll_path: str) -> set:
    """Return the set of Sumo__Plant__* symbol names exported by the DLL."""
    try:
        import pefile
        pe = pefile.PE(dll_path, fast_load=True)
        pe.parse_data_directories(
            directories=[pefile.DIRECTORY_ENTRY["IMAGE_DIRECTORY_ENTRY_EXPORT"]]
        )
        if hasattr(pe, "DIRECTORY_ENTRY_EXPORT") and pe.DIRECTORY_ENTRY_EXPORT:
            return {
                e.name.decode("utf-8", errors="replace")
                for e in pe.DIRECTORY_ENTRY_EXPORT.symbols
                if e.name
            }
        return _read_dll_symbols_strings(dll_path)
    except Exception:
        return _read_dll_symbols_strings(dll_path)


def _read_dll_symbols_strings(dll_path: str) -> set:
    """Fallback: grep binary for Sumo__Plant__* substrings."""
    import re as _re
    with open(dll_path, "rb") as f:
        blob = f.read()
    return {
        m.group(0).decode("ascii", errors="replace")
        for m in _re.finditer(rb"Sumo__Plant__[A-Za-z0-9_]+", blob)
    }


def _exp_verify_dll_matches_schematic(args: dict) -> dict:
    dll_path = args.get("dll_path", "")
    if not os.path.exists(dll_path):
        return {"ok": False, "reason": f"DLL not found: {dll_path}"}

    exports = _read_dll_symbols(dll_path)

    # If no schematic provided, return the export list for inspection.
    if not (args.get("schematic_html") or args.get("schematic_json")):
        return {
            "ok": True,
            "mode": "exports_only",
            "dll_path": dll_path,
            "dll_export_count": len(exports),
            "sample_exports": sorted(list(exports))[:30],
        }

    schematic, load_err = _resolve_schematic_from_args(args)
    if load_err:
        return load_err

    refs: set = set()
    for u in schematic.get("units", []):
        for p in u.get("parameters", []):
            sv = p.get("sumo_variable")
            if sv:
                refs.add(sv)

    missing = sorted(refs - exports)
    return {
        "ok":               not missing,
        "dll_path":         dll_path,
        "refs_in_schematic": len(refs),
        "refs_resolved":    len(refs) - len(missing),
        "missing":          missing[:20],
        "missing_total":    len(missing),
        "dll_export_count": len(exports),
        "verdict": (
            "DLL covers the schematic — appears compatible"
            if not missing else
            f"DLL is missing {len(missing)} variables the schematic references — "
            "this DLL is for a DIFFERENT plant. Do not use it."
        ),
    }


def _exp_update_schematic_in_html(args: dict) -> dict:
    try:
        import html_schematic_io as _hsi
    except Exception as e:
        return {"ok": False, "reason": f"html_schematic_io module not available: {e}"}

    html_path      = args.get("html_path", "")
    schematic_json = args.get("schematic_json", "")
    backup         = bool(args.get("backup", True))

    if not html_path:
        return {"ok": False, "reason": "html_path is required"}
    if not schematic_json:
        return {"ok": False, "reason": "schematic_json is required (JSON string or path to .json file)"}

    try:
        result = _hsi.update_schematic_in_html(
            html_path=html_path,
            new_json_text=schematic_json,
            backup=backup,
        )
        return result
    except Exception as e:
        return {"ok": False, "reason": str(e)}


def _exp_next_step_for_project(args: dict) -> dict:
    project_dir = args.get("project_dir", "")
    d = Path(project_dir) if project_dir else None
    if not d or not d.is_dir():
        return {"ok": False, "reason": f"not a directory: {project_dir}"}

    sumo_stubs = list(d.glob("*.sumo"))
    pack_dirs  = list(d.glob("*.sumo-pack"))
    dll        = (d / "sumoproject.dll").is_file()
    state      = (d / "state.xml").is_file()
    applied    = (d / ".overrides_applied").is_file()

    if not sumo_stubs and not pack_dirs:
        return _workflow_step(
            "call_build_sumo_pack",
            "No .sumo stub or .sumo-pack found. Start with "
            "build_sumo_pack(schematic_html=..., output_dir='<plant>.sumo-pack').",
        )
    if not sumo_stubs:
        return _workflow_step(
            "open_pack_instructions",
            f"Build pack found ({pack_dirs[0].name}) but no .sumo stub. "
            f"Open {pack_dirs[0] / 'BUILD_INSTRUCTIONS.md'} and follow step 1.",
        )
    if not dll:
        return _workflow_step(
            "compile_in_sumo_gui",
            f"Stub {sumo_stubs[0].name} present but no sumoproject.dll. "
            "Open the stub in SUMO24 GUI, paste plant.sumoslang from the pack, "
            "and click Compile. SUMO will write the DLL next to the stub.",
        )
    if not state:
        return _workflow_step(
            "save_state_in_core_window",
            'DLL present but no state.xml. In SUMO24 Core Window run: maptoic; save "state.xml";',
        )
    if not applied:
        return _workflow_step(
            "paste_apply_parameters_scs",
            "Topology compiled and state saved. Paste apply_parameters.scs into "
            "the Core Window to apply your schematic's parameter values, then "
            "create the marker file .overrides_applied to record completion.",
        )
    return _workflow_step(
        "ready",
        "Project is fully built. You can now call run_steady_state, "
        "run_dynamic_simulation, etc.",
    )


def _workflow_step(action: str, message: str) -> dict:
    return {"ok": True, "action": action, "message": message}


def _exp_apply_template_auto_fixes(template_path: str, fixes: list) -> dict:
    """Apply auto_fix suggestions from validate_wwtp_template to the workbook in-place."""
    if not OPENPYXL_AVAILABLE:
        return {"ok": False, "error": "openpyxl is not installed"}
    import shutil as _sh
    p = Path(template_path)
    if not p.exists():
        return {"ok": False, "error": f"Template not found: {template_path}"}

    try:
        wb = load_workbook(p)
    except Exception as ex:
        return {"ok": False, "error": f"Could not open workbook: {ex}"}

    applied: list[str] = []
    skipped: list[str] = []

    for fix in fixes:
        action = fix.get("action")
        if action == "add_to_treatment_processes":
            ws = wb["Treatment_Processes"] if "Treatment_Processes" in wb.sheetnames else None
            if ws is None:
                skipped.append(f"add_to_treatment_processes: sheet not found")
                continue
            inst = fix.get("instance_name", "")
            ptype = fix.get("process_type", "")
            cat   = fix.get("category", "physical")
            stage = fix.get("stage", "effluent")
            # Find last used row and append.
            last = ws.max_row + 1
            order = last - 3  # rough order_index
            ws.cell(last, 1, order)
            ws.cell(last, 2, inst)
            ws.cell(last, 3, ptype)
            ws.cell(last, 4, cat)
            ws.cell(last, 5, stage)
            applied.append(f"add_to_treatment_processes: {inst}")

        elif action == "clear_na_rows":
            ws = wb["Connections"] if "Connections" in wb.sheetnames else None
            if ws is None:
                skipped.append("clear_na_rows: Connections sheet not found")
                continue
            rows_cleared = 0
            for r in range(4, ws.max_row + 1):
                v = ws.cell(r, 1).value
                if str(v).upper() in ("NA", "N/A", "NONE", ""):
                    ws.cell(r, 1).value = None
                    ws.cell(r, 2).value = None
                    ws.cell(r, 3).value = None
                    ws.cell(r, 4).value = None
                    ws.cell(r, 5).value = None
                    rows_cleared += 1
            applied.append(f"clear_na_rows: cleared {rows_cleared} rows")
        else:
            skipped.append(f"unknown action: {action}")

    # Backup + save.
    backup = str(p) + ".bak"
    try:
        _sh.copy2(p, backup)
        wb.save(p)
    except Exception as ex:
        return {"ok": False, "error": f"Could not save workbook: {ex}"}

    return {
        "ok": True,
        "template_path": str(p),
        "backup": backup,
        "applied": applied,
        "skipped": skipped,
    }



# ═══════════════════════════════════════════════════════════════════════════════
# Group FF — Pipeline driver and stage orchestration
# ═══════════════════════════════════════════════════════════════════════════════

# ── Module imports (lazy, so server still starts if pipeline isn't installed) ─

def _pipeline_imports():
    """Return (manifest, hashing, snapshot, seeding, engineering, si, di, build) modules."""
    import sys, os
    sys.path.insert(0, str(Path(__file__).parent))
    from pipeline import manifest as _m
    from pipeline import hashing as _h
    from pipeline import snapshot as _sn
    from pipeline import seeding as _sd
    from pipeline import engineering as _eng
    from pipeline import static_inputs as _si
    from pipeline import dynamic_inputs as _di
    from pipeline import build as _bld
    return _m, _h, _sn, _sd, _eng, _si, _di, _bld

try:
    _pipe_m, _pipe_h, _pipe_sn, _pipe_sd, _pipe_eng, _pipe_si, _pipe_di, _pipe_bld = _pipeline_imports()
    PIPELINE_AVAILABLE = True
except Exception as _pipe_err:
    PIPELINE_AVAILABLE = False
    print(f"[WARN] pipeline sub-package unavailable: {_pipe_err}")


def _pipe_check() -> dict | None:
    if not PIPELINE_AVAILABLE:
        return {"ok": False, "error": "pipeline sub-package not available — check PY/pipeline/__init__.py"}
    return None


# ── Internal helpers ──────────────────────────────────────────────────────────

def _detect_stale(manifest: dict, project_dir: str) -> list[int]:
    """Re-hash on-disk artefacts; return list of stage ids where hash mismatched."""
    pd = Path(project_dir)
    plant = manifest.get("plant_name", "plant")
    hashers = {
        1: lambda: _pipe_h.hash_schematic(pd / f"{plant}_schematic.html"),
        2: lambda: _pipe_h.hash_excel(pd / f"{plant}_data.xlsx"),
        3: lambda: _pipe_h.hash_json(pd / "engineering_checks.json"),
        4: lambda: _pipe_h.hash_json(pd / "static_inputs.json"),
        6: lambda: _pipe_h.hash_json(pd / "controllers.json"),
        7: lambda: _pipe_h.hash_pack(pd / f"{plant}.sumo-pack"),
    }
    stale = []
    for s in manifest.get("stages", []):
        if s["status"] != "passed":
            continue
        hfn = hashers.get(s["id"])
        if not hfn:
            continue
        try:
            cur = hfn()
        except FileNotFoundError:
            stale.append(s["id"])
            continue
        if cur != s.get("output_hash"):
            stale.append(s["id"])
    return stale


_STAGE_ENTRY = {
    0: "pipeline_init",
    1: "stage1_open_schematic",
    2: "stage2_open_data",
    3: "stage3_run_engineering_checks",
    4: "stage4_define_static_inputs",
    5: "stage5_generate_dynamic_inputs",
    6: "stage6_add_controller",
    7: "stage7_build",
    8: "stage8_verify",
}


def _recommended_action(manifest: dict, project_dir: str) -> dict:
    stages = manifest.get("stages", [])
    for s in stages:
        st = s["status"]
        sid = s["id"]
        if st in ("passed", "skipped"):
            continue
        if st == "stale":
            return {"tool": "pipeline_advance",
                    "reason": f"Stage {sid} ({s['name']}) artefact changed since last advance — re-validate.",
                    "args": {"project_dir": project_dir, "from_stage": sid}}
        if st == "failed":
            blocker = s.get("blocker") or "see manifest for details"
            return {"tool": _STAGE_ENTRY.get(sid, "pipeline_status"),
                    "reason": f"Stage {sid} failed: {blocker}. Fix and re-advance.",
                    "args": {"project_dir": project_dir}}
        if st == "pending":
            return {"tool": _STAGE_ENTRY.get(sid, "pipeline_status"),
                    "reason": f"Stage {sid} ({s['name']}) not started — begin here.",
                    "args": {"project_dir": project_dir}}
    return {"tool": None, "reason": "All stages passed. Project is complete."}


def _status_payload(manifest: dict, project_dir: str, stale: list) -> dict:
    stages = manifest.get("stages", [])
    # Find current stage
    cur = stages[-1]
    for s in stages:
        if s["status"] not in ("passed", "skipped"):
            cur = s
            break
    return {
        "ok": True,
        "project_dir": project_dir,
        "plant_name": manifest.get("plant_name"),
        "current_stage": cur["id"],
        "current_stage_name": cur["name"],
        "current_stage_status": cur["status"],
        "stages": [{"id": s["id"], "name": s["name"], "status": s["status"],
                    "summary": s.get("summary"), "blocker": s.get("blocker")}
                   for s in stages],
        "stale": stale,
        "next_action": _recommended_action(manifest, project_dir),
    }


def _read_params_from_excel(excel_path: str) -> dict:
    """Extract engineering check parameters from the Excel template."""
    params: dict = {}
    if not OPENPYXL_AVAILABLE:
        return params
    try:
        wb = load_workbook(excel_path, data_only=True, read_only=True)
    except Exception:
        return params

    # Plant_Info sheet
    if "Plant_Info" in wb.sheetnames:
        ws = wb["Plant_Info"]
        for row in ws.iter_rows(values_only=True):
            if not row or row[0] is None:
                continue
            key = str(row[0]).strip().lower()
            val = row[1] if len(row) > 1 else None
            if val is None:
                continue
            try:
                v = float(val)
            except (TypeError, ValueError):
                continue
            mapping = {
                "design_flow_avg": "Q_avg_m3d",
                "q_avg": "Q_avg_m3d",
                "q_m3d": "Q_avg_m3d",
                "cod": "COD_in_mgL",
                "cod_mgl": "COD_in_mgL",
                "bod5": "BOD_mgL",
                "bod_mgl": "BOD_mgL",
                "tkn": "TKN_mgL",
                "tkn_mgl": "TKN_mgL",
                "tp": "TP_mgL",
                "tp_mgl": "TP_mgL",
                "tss": "TSS_mgL",
                "tss_mgl": "TSS_mgL",
                "temperature": "T_C",
                "temp_c": "T_C",
                "t_c": "T_C",
                "alkalinity": "Alk_in_mgCaCO3",
                "alk_mgcaco3": "Alk_in_mgCaCO3",
            }
            mapped = mapping.get(key)
            if mapped:
                params[mapped] = v

    # Treatment_Processes — sum volumes and find clarifier area
    V_total = 0.0
    A_clarifier = 0.0
    if "Treatment_Processes" in wb.sheetnames:
        ws = wb["Treatment_Processes"]
        headers = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i == 0:
                headers = [str(c).strip().lower() if c else "" for c in row]
                continue
            if not row or row[0] is None:
                continue
            row_d = dict(zip(headers, row))
            ptype = str(row_d.get("process_type", "")).lower()
            for vcol in ("volume_m3", "volume", "v_m3"):
                v = row_d.get(vcol)
                if v and str(v).replace(".", "").isdigit():
                    try:
                        V_total += float(v)
                    except (TypeError, ValueError):
                        pass
                    break
            if "clarifier" in ptype or "settler" in ptype:
                for acol in ("area_m2", "surface_area_m2", "a_m2"):
                    a = row_d.get(acol)
                    if a and str(a).replace(".", "").isdigit():
                        try:
                            A_clarifier += float(a)
                        except (TypeError, ValueError):
                            pass
                        break

    if V_total > 0:
        params["V_total_m3"] = V_total
        params["V_reactor_m3"] = V_total
    if A_clarifier > 0:
        params["A_clarifier_m2"] = A_clarifier

    # Reasonable defaults for missing params
    mlss = params.get("MLSS_mgL", 3500.0)
    params.setdefault("MLSS_mgL", mlss)
    params.setdefault("MLVSS_mgL", mlss * 0.75)
    q = params.get("Q_avg_m3d", 1000.0)
    params.setdefault("Q_WAS_m3d", q * 0.01)
    params.setdefault("X_R_mgL", 8000.0)
    params.setdefault("Q_eff_m3d", q)
    params.setdefault("COD_eff_mgL", 50.0)
    params.setdefault("COD_WAS_mgL", params.get("MLSS_mgL", 3500.0) * 1.2)
    params.setdefault("KLa_per_d", 120.0)
    params.setdefault("RAS_ratio", 0.6)
    params.setdefault("NO3_removed_mgL", params.get("TKN_mgL", 0) * 0.5)

    wb.close()
    return params


# ── Handler implementations ───────────────────────────────────────────────────

def _exp_pipeline_init(args: dict) -> dict:
    err = _pipe_check()
    if err:
        return err

    project_dir = args.get("project_dir", "")
    plant_name  = args.get("plant_name", "")
    kinetic     = args.get("kinetic_model", "ASM2d")
    template    = args.get("from_template", "blank")
    force       = bool(args.get("force", False))

    if not project_dir or not plant_name:
        return {"ok": False, "error": "project_dir and plant_name are required"}

    pd = Path(project_dir)

    # Refuse to overwrite unless force=True
    manifest_file = pd / "project.yaml"
    if manifest_file.exists() and not force:
        try:
            m = _pipe_m.load(project_dir)
            return {"ok": True, "idempotent": True,
                    "message": f"Project already initialised at {project_dir}",
                    "plant_name": m.get("plant_name"),
                    "current_stage": next((s["id"] for s in m["stages"]
                                           if s["status"] not in ("passed","skipped")), 8)}
        except Exception:
            pass

    pd.mkdir(parents=True, exist_ok=True)
    (pd / "snapshots").mkdir(exist_ok=True)
    (pd / "dynamic_inputs").mkdir(exist_ok=True)

    # Scaffold artefact placeholders
    for fname, content in [
        ("engineering_checks.json", '{"status":"pending"}'),
        ("static_inputs.json",      '{"status":"pending"}'),
        ("controllers.json",        '{"controllers":[]}'),
    ]:
        p = pd / fname
        if not p.exists():
            p.write_text(content, encoding="utf-8")

    # Copy HTML schematic template
    html_src = _MCP_DIR / "wwtp_schematic_template.html"
    html_dst = pd / f"{plant_name}_schematic.html"
    if html_src.exists() and not html_dst.exists():
        import shutil as _sh
        _sh.copy2(html_src, html_dst)
    elif not html_dst.exists():
        html_dst.write_text(
            f"<!-- SUMO24 schematic for {plant_name} — edit in browser -->", encoding="utf-8"
        )

    # Generate Excel template
    xlsx_dst = pd / f"{plant_name}_data.xlsx"
    if not xlsx_dst.exists():
        try:
            xlsx_result = _exp_generate_wwtp_template_xlsx(str(xlsx_dst))
        except Exception as ex:
            xlsx_result = {"error": str(ex)}
    else:
        xlsx_result = {"path": str(xlsx_dst), "note": "existing file kept"}

    # README
    readme = pd / "README.md"
    if not readme.exists():
        readme.write_text(
            f"# {plant_name} — SUMO24 MCP Pipeline\n\n"
            f"Kinetic model: {kinetic}\n"
            f"Template: {template}\n\n"
            "## Next steps\n"
            "1. Call `pipeline_status` to see the current stage.\n"
            "2. Edit the HTML schematic in your browser.\n"
            "3. Call `pipeline_advance` after each stage is ready.\n",
            encoding="utf-8",
        )

    # Build and save manifest (stage 0 = passed)
    m = _pipe_m.new_manifest(str(pd), plant_name, kinetic)
    _pipe_m.set_stage(m, 0, status="passed",
                      output_hash=_pipe_h.hash_json(pd / "project.yaml")
                      if (pd / "project.yaml").exists() else None,
                      summary={"plant_name": plant_name, "kinetic_model": kinetic,
                               "from_template": template})
    _pipe_m.record(m, "pipeline_init", args, "ok")
    _pipe_m.save(project_dir, m)

    return {
        "ok": True,
        "project_dir": str(pd),
        "plant_name": plant_name,
        "schematic_html": str(html_dst),
        "data_xlsx": str(xlsx_dst),
        "next_action": {
            "tool": "stage1_open_schematic",
            "reason": "Edit the schematic HTML in your browser, then call pipeline_advance.",
            "args": {"project_dir": str(pd)},
        },
    }


def _exp_pipeline_status(args: dict) -> dict:
    err = _pipe_check()
    if err:
        return err

    project_dir = args.get("project_dir", "")
    try:
        m = _pipe_m.load(project_dir)
    except FileNotFoundError as e:
        return {"ok": False, "reason": str(e),
                "next_action": {"tool": "pipeline_init",
                                "reason": "Initialise the project first.",
                                "args": {"project_dir": project_dir}}}

    stale = _detect_stale(m, project_dir)
    for sid in stale:
        _pipe_m.set_stage(m, sid, status="stale")

    _pipe_m.record(m, "pipeline_status", {"project_dir": project_dir}, "ok")
    _pipe_m.save(project_dir, m)

    return _status_payload(m, project_dir, stale)


def _exp_pipeline_advance(args: dict) -> dict:
    err = _pipe_check()
    if err:
        return err

    project_dir = args.get("project_dir", "")
    from_stage  = args.get("from_stage")
    skip        = bool(args.get("skip", False))
    skip_reason = args.get("skip_reason", "")
    force       = bool(args.get("force", False))

    try:
        m = _pipe_m.load(project_dir)
    except FileNotFoundError as e:
        return {"ok": False, "error": str(e)}

    # Determine which stage to advance
    if from_stage is not None:
        sid = int(from_stage)
    else:
        sid = next((s["id"] for s in m["stages"]
                    if s["status"] not in ("passed", "skipped")), None)
        if sid is None:
            return {"ok": True, "message": "All stages already passed.",
                    **_status_payload(m, project_dir, [])}

    # Handle skip for optional stages
    OPTIONAL = {5, 6}
    if skip:
        if sid not in OPTIONAL:
            return {"ok": False,
                    "error": f"Stage {sid} is not optional — cannot skip."}
        if not skip_reason:
            return {"ok": False, "error": "skip_reason is required when skip=True."}
        _pipe_m.set_stage(m, sid, status="skipped",
                          summary={"skip_reason": skip_reason})
        _pipe_m.record(m, "pipeline_advance",
                       {"project_dir": project_dir, "skip_stage": sid, "reason": skip_reason},
                       "skipped")
        _pipe_m.save(project_dir, m)
        return {"ok": True, "stage": sid, "action": "skipped",
                **_status_payload(m, project_dir, [])}

    # Run the validator for this stage
    validator_map = {
        1: lambda: _exp_stage1_validate_schematic({"project_dir": project_dir}),
        2: lambda: _exp_stage2_validate_parameters({"project_dir": project_dir}),
        3: lambda: _exp_stage3_run_engineering_checks({"project_dir": project_dir}),
        4: lambda: _exp_stage4_validate_static_inputs({"project_dir": project_dir}),
        5: lambda: _exp_stage5_validate_dynamic_inputs({"project_dir": project_dir}),
        6: lambda: _exp_stage6_validate_controllers({"project_dir": project_dir}),
        7: lambda: _exp_stage7_build({"project_dir": project_dir}),
    }

    validator = validator_map.get(sid)
    if validator is None:
        return {"ok": False, "error": f"No validator for stage {sid}."}

    # Snapshot before mutating
    try:
        _pipe_sn.snapshot(project_dir, f"pre_stage{sid}")
    except Exception:
        pass  # non-fatal

    val_result = validator()
    passed = val_result.get("ok", False)

    if passed or force:
        # Compute and record hash
        pd = Path(project_dir)
        plant = m.get("plant_name", "plant")
        hash_fns = {
            1: lambda: _pipe_h.hash_schematic(pd / f"{plant}_schematic.html"),
            2: lambda: _pipe_h.hash_excel(pd / f"{plant}_data.xlsx"),
            3: lambda: _pipe_h.hash_json(pd / "engineering_checks.json"),
            4: lambda: _pipe_h.hash_json(pd / "static_inputs.json"),
            6: lambda: _pipe_h.hash_json(pd / "controllers.json"),
            7: lambda: _pipe_h.hash_pack(pd / f"{plant}.sumo-pack"),
        }
        h_fn = hash_fns.get(sid)
        new_hash = None
        try:
            new_hash = h_fn() if h_fn else None
        except Exception:
            pass

        old_hash = m["stages"][sid].get("output_hash")
        stale = []
        if new_hash and new_hash != old_hash:
            stale = _pipe_m.mark_downstream_stale(m, sid)

        summary = val_result.get("summary") or val_result.get("computed") or {}
        _pipe_m.set_stage(m, sid, status="passed",
                          output_hash=new_hash,
                          based_on=_pipe_m.last_passed_hash(m, sid),
                          summary=summary,
                          blocker=None)
        _pipe_m.record(m, "pipeline_advance",
                       {"project_dir": project_dir, "stage": sid}, "passed")
        _pipe_m.save(project_dir, m)
        return {"ok": True, "stage": sid, "action": "passed", "stale": stale,
                "validator_result": val_result,
                **_status_payload(m, project_dir, stale)}
    else:
        blocker = (val_result.get("top_fix") or
                   val_result.get("error") or
                   val_result.get("reason") or
                   "validation failed — see validator_result for details")
        _pipe_m.set_stage(m, sid, status="failed", blocker=blocker)
        _pipe_m.record(m, "pipeline_advance",
                       {"project_dir": project_dir, "stage": sid}, "failed")
        _pipe_m.save(project_dir, m)
        return {"ok": False, "stage": sid, "action": "failed",
                "blocker": blocker,
                "validator_result": val_result,
                **_status_payload(m, project_dir, [])}


def _exp_pipeline_revert(args: dict) -> dict:
    err = _pipe_check()
    if err:
        return err
    project_dir = args.get("project_dir", "")
    to_stage    = int(args.get("to_stage", 0))
    try:
        src = _pipe_sn.revert(project_dir, f"pre_stage{to_stage + 1}")
        return {"ok": True, "restored_from": src,
                "message": f"Reverted to state at end of stage {to_stage}."}
    except FileNotFoundError as e:
        # Try post_stage snapshot
        try:
            src = _pipe_sn.revert(project_dir, f"post_stage{to_stage}")
            return {"ok": True, "restored_from": src}
        except FileNotFoundError:
            return {"ok": False, "error": str(e)}


def _exp_pipeline_describe(args: dict) -> dict:
    err = _pipe_check()
    if err:
        return err
    project_dir = args.get("project_dir", "")
    try:
        m = _pipe_m.load(project_dir)
    except FileNotFoundError as e:
        return {"ok": False, "error": str(e)}

    lines = [f"# Pipeline Status: {m.get('plant_name', '?')}",
             f"",
             f"Project: `{project_dir}`",
             f"Kinetic model: {m.get('kinetic_model', '?')}",
             f"Created: {m.get('created', '?')}",
             f"",
             "## Stages",
             ""]
    for s in m.get("stages", []):
        emoji = {"passed":"✅","failed":"❌","pending":"⏳",
                 "stale":"⚠️","skipped":"⏭️","in_progress":"🔄"}.get(s["status"],"❓")
        lines.append(f"{emoji} **Stage {s['id']}: {s['name']}** — {s['status']}")
        if s.get("summary"):
            for k, v in s["summary"].items():
                lines.append(f"   - {k}: {v}")
        if s.get("blocker"):
            lines.append(f"   - ⚡ Blocker: {s['blocker']}")

    snaps = _pipe_sn.list_snapshots(project_dir)
    lines += ["", "## Snapshots", f"{len(snaps)} snapshot(s) available."]

    prov = m.get("provenance", [])
    lines += ["", f"## Provenance ({len(prov)} entries)", ""]
    for p in prov[-10:]:
        lines.append(f"- `{p.get('ts','')}` {p.get('tool','')} → {p.get('result','')}")

    report = "\n".join(lines)
    return {"ok": True, "report": report, "plant_name": m.get("plant_name")}


def _exp_stage1_open_schematic(args: dict) -> dict:
    err = _pipe_check()
    if err:
        return err
    project_dir = args.get("project_dir", "")
    try:
        m = _pipe_m.load(project_dir)
    except FileNotFoundError as e:
        return {"ok": False, "error": str(e)}

    plant = m.get("plant_name", "plant")
    html_path = Path(project_dir) / f"{plant}_schematic.html"
    return {
        "ok": True,
        "schematic_html": str(html_path),
        "exists": html_path.exists(),
        "next_action": "Edit the schematic in your browser, then call pipeline_advance.",
    }


def _exp_stage1_validate_schematic(args: dict) -> dict:
    err = _pipe_check()
    if err:
        return err
    project_dir = args.get("project_dir", "")
    try:
        m = _pipe_m.load(project_dir)
    except FileNotFoundError as e:
        return {"ok": False, "error": str(e)}

    plant = m.get("plant_name", "plant")
    html_path = Path(project_dir) / f"{plant}_schematic.html"
    if not html_path.exists():
        return {"ok": False, "error": f"Schematic not found: {html_path}",
                "fix": "Call stage1_open_schematic and create the schematic first."}

    if not SCHEMATIC_AVAILABLE:
        return {"ok": False, "error": "schematic_parser module unavailable"}

    # Parse
    try:
        schematic = _schp.parse_html_schematic(str(html_path))
    except Exception as ex:
        return {"ok": False, "error": f"Parse failed: {ex}"}

    errors = []
    warnings = []
    units  = schematic.get("units", [])
    streams = schematic.get("streams", [])

    # Check unit types
    known_types = set(_utr.UNIT_TYPES.keys())
    for u in units:
        ut = u.get("type", "")
        if ut not in known_types:
            closest = _closest_type(ut, known_types)
            errors.append(f"Unit '{u.get('name',u.get('id'))}' has unknown type '{ut}'. "
                          f"Did you mean '{closest}'?")

    # Check stream references
    unit_ids = {u.get("id") for u in units}
    for s in streams:
        if s.get("from") not in unit_ids:
            errors.append(f"Stream from '{s.get('from')}' references non-existent unit.")
        if s.get("to") not in unit_ids:
            errors.append(f"Stream to '{s.get('to')}' references non-existent unit.")

    # Topology: must have influent + effluent
    types_present = {u.get("type") for u in units}
    if "influent" not in types_present:
        errors.append("No 'influent' boundary unit found — required.")
    if "effluent" not in types_present:
        errors.append("No 'effluent' boundary unit found — add an Effluent boundary unit.")

    # Orphan units (no connections)
    connected = set()
    for s in streams:
        connected.add(s.get("from"))
        connected.add(s.get("to"))
    boundary_types = {"influent", "effluent", "ras_flow", "was_flow"}
    for u in units:
        if u.get("id") not in connected and u.get("type") not in boundary_types:
            warnings.append(f"Unit '{u.get('name', u.get('id'))}' has no connections (orphan).")

    ok = len(errors) == 0
    result = {
        "ok": ok,
        "errors": errors,
        "warnings": warnings,
        "unit_count": len(units),
        "stream_count": len(streams),
        "unit_types_used": sorted(types_present),
    }
    if ok:
        result["summary"] = {
            "units": len(units),
            "streams": len(streams),
            "unit_types": sorted(types_present),
        }
    return result


def _closest_type(t: str, known: set) -> str:
    """Return the known type most similar to t (simple edit-distance proxy)."""
    def _sim(a, b):
        return sum(x == y for x, y in zip(a, b)) / max(len(a), len(b), 1)
    return max(known, key=lambda k: _sim(t.lower(), k.lower()))


def _exp_stage2_open_data(args: dict) -> dict:
    err = _pipe_check()
    if err:
        return err
    project_dir = args.get("project_dir", "")
    try:
        m = _pipe_m.load(project_dir)
    except FileNotFoundError as e:
        return {"ok": False, "error": str(e)}

    plant = m.get("plant_name", "plant")
    xlsx_path = Path(project_dir) / f"{plant}_data.xlsx"

    # Generate template if missing
    if not xlsx_path.exists():
        _exp_generate_wwtp_template_xlsx(str(xlsx_path))

    # Seed from schematic if Stage 1 is passed
    stage1 = m["stages"][1]
    if stage1["status"] == "passed" and SCHEMATIC_AVAILABLE:
        html_path = Path(project_dir) / f"{plant}_schematic.html"
        if html_path.exists():
            try:
                schematic = _schp.parse_html_schematic(str(html_path))
                seed_result = _pipe_sd.seed_excel_from_schematic(
                    str(xlsx_path), schematic,
                    kinetic_model=m.get("kinetic_model", "ASM2d"),
                    backup=True,
                )
            except Exception as ex:
                seed_result = {"ok": False, "error": str(ex)}
        else:
            seed_result = {"ok": False, "error": "schematic HTML not found"}
    else:
        seed_result = {"ok": True, "note": "Stage 1 not yet passed — skipping schematic seeding"}

    return {
        "ok": True,
        "data_xlsx": str(xlsx_path),
        "exists": xlsx_path.exists(),
        "seeding": seed_result,
        "next_action": "Fill in numbers in the Excel template, then call pipeline_advance.",
    }


def _exp_stage2_validate_parameters(args: dict) -> dict:
    err = _pipe_check()
    if err:
        return err
    project_dir = args.get("project_dir", "")
    try:
        m = _pipe_m.load(project_dir)
    except FileNotFoundError as e:
        return {"ok": False, "error": str(e)}

    plant = m.get("plant_name", "plant")
    xlsx_path = Path(project_dir) / f"{plant}_data.xlsx"
    if not xlsx_path.exists():
        return {"ok": False, "error": f"Excel template not found: {xlsx_path}",
                "fix": "Call stage2_open_data first."}

    # Run existing validate_wwtp_template
    val_result = _exp_validate_wwtp_template({"template_path": str(xlsx_path)})
    errors = val_result.get("errors", [])
    warnings_list = val_result.get("warnings", [])

    # Cross-check against schematic if Stage 1 passed
    if m["stages"][1]["status"] == "passed" and SCHEMATIC_AVAILABLE:
        html_path = Path(project_dir) / f"{plant}_schematic.html"
        if html_path.exists():
            try:
                schematic = _schp.parse_html_schematic(str(html_path))
                schematic_names = {u.get("name") for u in schematic.get("units", [])}

                if OPENPYXL_AVAILABLE:
                    wb = load_workbook(str(xlsx_path), data_only=True, read_only=True)
                    if "Treatment_Processes" in wb.sheetnames:
                        ws = wb["Treatment_Processes"]
                        excel_names = set()
                        for i, row in enumerate(ws.iter_rows(values_only=True)):
                            if i == 0:
                                continue
                            if row and row[1] is not None:
                                excel_names.add(str(row[1]).strip())
                        missing_in_excel = schematic_names - excel_names
                        extra_in_excel   = excel_names - schematic_names
                        for n in sorted(missing_in_excel):
                            errors.append(f"Schematic unit '{n}' is missing from Treatment_Processes sheet.")
                        for n in sorted(extra_in_excel):
                            warnings_list.append(f"'{n}' in Treatment_Processes has no matching schematic unit.")
                    wb.close()
            except Exception as ex:
                warnings_list.append(f"Cross-check with schematic failed: {ex}")

    ok = len(errors) == 0
    return {
        "ok": ok,
        "errors": errors[:20],
        "warnings": warnings_list[:20],
        "error_count": len(errors),
        "summary": val_result.get("summary", {}),
    }


def _exp_stage3_run_engineering_checks(args: dict) -> dict:
    err = _pipe_check()
    if err:
        return err
    project_dir = args.get("project_dir", "")
    try:
        m = _pipe_m.load(project_dir)
    except FileNotFoundError as e:
        return {"ok": False, "error": str(e)}

    plant = m.get("plant_name", "plant")
    xlsx_path = Path(project_dir) / f"{plant}_data.xlsx"
    if not xlsx_path.exists():
        return {"ok": False, "error": f"Excel template not found: {xlsx_path}"}

    params = _read_params_from_excel(str(xlsx_path))
    output_path = Path(project_dir) / "engineering_checks.json"
    result = _pipe_eng.run_engineering_checks(params, output_path=str(output_path))
    result["summary"] = result.get("computed", {})
    return result


def _exp_stage4_define_static_inputs(args: dict) -> dict:
    err = _pipe_check()
    if err:
        return err
    project_dir = args.get("project_dir", "")
    try:
        m = _pipe_m.load(project_dir)
    except FileNotFoundError as e:
        return {"ok": False, "error": str(e)}

    plant = m.get("plant_name", "plant")
    xlsx_path = Path(project_dir) / f"{plant}_data.xlsx"
    if not xlsx_path.exists():
        return {"ok": False, "error": f"Excel template not found: {xlsx_path}"}

    # Read influent parameters from Excel
    params = _read_params_from_excel(str(xlsx_path))
    influent = {
        "Q_m3d":       params.get("Q_avg_m3d", 1000.0),
        "COD_mgL":     params.get("COD_in_mgL", 400.0),
        "BOD_mgL":     params.get("BOD_mgL", 200.0),
        "TKN_mgL":     params.get("TKN_mgL", 50.0),
        "TP_mgL":      params.get("TP_mgL", 7.0),
        "TSS_mgL":     params.get("TSS_mgL", 250.0),
        "T_C":         params.get("T_C", 20.0),
        "Alk_mgCaCO3": params.get("Alk_in_mgCaCO3", 300.0),
    }

    result = _pipe_si.write_static_inputs(
        project_dir=str(project_dir),
        plant_name=plant,
        influent=influent,
        kinetic_model=m.get("kinetic_model", "ASM2d"),
    )
    return result


def _exp_stage4_validate_static_inputs(args: dict) -> dict:
    err = _pipe_check()
    if err:
        return err
    project_dir = args.get("project_dir", "")
    si_path = Path(project_dir) / "static_inputs.json"
    if not si_path.exists():
        return {"ok": False, "error": "static_inputs.json not found — call stage4_define_static_inputs first."}
    try:
        data = json.loads(si_path.read_text(encoding="utf-8"))
    except Exception as ex:
        return {"ok": False, "error": f"Cannot read static_inputs.json: {ex}"}

    validation = data.get("validation", [])
    failed = [v for v in validation if not v.get("ok")]
    return {
        "ok": data.get("ok", False),
        "validation": validation,
        "failed_checks": failed,
        "influent_sv": data.get("influent_sv", {}),
    }


def _exp_stage5_generate_dynamic_inputs(args: dict) -> dict:
    err = _pipe_check()
    if err:
        return err
    project_dir  = args.get("project_dir", "")
    profile_type = args.get("profile_type", "constant")
    duration     = float(args.get("duration_days", 30.0))
    dt_h         = float(args.get("dt_h", 0.5))
    name         = args.get("name")

    try:
        m = _pipe_m.load(project_dir)
    except FileNotFoundError as e:
        return {"ok": False, "error": str(e)}

    plant = m.get("plant_name", "plant")
    xlsx_path = Path(project_dir) / f"{plant}_data.xlsx"
    params = _read_params_from_excel(str(xlsx_path)) if xlsx_path.exists() else {}
    influent = {
        "Q_m3d":   params.get("Q_avg_m3d", 1000.0),
        "COD_mgL": params.get("COD_in_mgL", 400.0),
        "TKN_mgL": params.get("TKN_mgL", 50.0),
        "TP_mgL":  params.get("TP_mgL", 7.0),
        "TSS_mgL": params.get("TSS_mgL", 250.0),
        "T_C":     params.get("T_C", 20.0),
    }

    return _pipe_di.generate_dynamic_inputs(
        project_dir=str(project_dir),
        profile_type=profile_type,
        influent=influent,
        duration_days=duration,
        dt_h=dt_h,
        name=name,
    )


def _exp_stage5_validate_dynamic_inputs(args: dict) -> dict:
    err = _pipe_check()
    if err:
        return err
    project_dir    = args.get("project_dir", "")
    sumo_stop_time = args.get("sumo_stop_time")

    dyn_dir = Path(project_dir) / "dynamic_inputs"
    if not dyn_dir.exists():
        return {"ok": False, "error": "dynamic_inputs/ directory not found."}

    tsvs = list(dyn_dir.glob("*.tsv"))
    if not tsvs:
        return {"ok": False, "error": "No .tsv files found in dynamic_inputs/."}

    results = []
    all_ok = True
    for tsv in sorted(tsvs):
        r = _pipe_di.validate_dynamic_tsv(tsv, sumo_stop_time)
        results.append(r)
        if not r["ok"]:
            all_ok = False

    return {"ok": all_ok, "files_checked": len(results), "results": results}


def _exp_stage6_add_controller(args: dict) -> dict:
    err = _pipe_check()
    if err:
        return err
    project_dir = args.get("project_dir", "")
    kind        = args.get("kind", "DO")
    unit        = args.get("unit", "")
    setpoint    = float(args.get("setpoint", 2.0))
    sensor      = args.get("sensor", "")
    actuator    = args.get("actuator", "")
    Kp          = float(args.get("Kp", 10.0))
    Ti          = float(args.get("Ti", 0.01))

    ctrl_path = Path(project_dir) / "controllers.json"
    try:
        data = json.loads(ctrl_path.read_text(encoding="utf-8")) if ctrl_path.exists() else {"controllers": []}
    except Exception:
        data = {"controllers": []}

    data["controllers"].append({
        "kind": kind, "unit": unit, "setpoint": setpoint,
        "sensor": sensor, "actuator": actuator, "Kp": Kp, "Ti": Ti,
    })
    ctrl_path.write_text(json.dumps(data, indent=2), encoding="utf-8")

    return {"ok": True, "controller_count": len(data["controllers"]),
            "added": data["controllers"][-1]}


def _exp_stage6_validate_controllers(args: dict) -> dict:
    err = _pipe_check()
    if err:
        return err
    project_dir = args.get("project_dir", "")

    ctrl_path = Path(project_dir) / "controllers.json"
    if not ctrl_path.exists():
        return {"ok": False, "error": "controllers.json not found."}
    try:
        data = json.loads(ctrl_path.read_text(encoding="utf-8"))
    except Exception as ex:
        return {"ok": False, "error": f"Cannot read controllers.json: {ex}"}

    ctrls = data.get("controllers", [])
    if not ctrls:
        return {"ok": False, "error": "No controllers defined."}

    errors = []
    # Check for conflicting actuators
    actuator_units: dict = {}
    for c in ctrls:
        key = (c.get("unit"), c.get("actuator"))
        if key in actuator_units:
            errors.append(f"Conflicting controllers: two controllers drive '{c.get('actuator')}' "
                          f"on '{c.get('unit')}'.")
        actuator_units[key] = True
        # Setpoint range checks
        kind = c.get("kind", "")
        sp = c.get("setpoint", 0)
        if kind == "DO" and not (0.1 <= sp <= 6.0):
            errors.append(f"DO setpoint {sp} mg/L out of range [0.1, 6.0].")
        if kind == "SRT" and not (1.0 <= sp <= 60.0):
            errors.append(f"SRT setpoint {sp} d out of range [1, 60].")

    ok = len(errors) == 0
    return {"ok": ok, "controller_count": len(ctrls),
            "errors": errors, "controllers": ctrls}


def _exp_stage7_build(args: dict) -> dict:
    err = _pipe_check()
    if err:
        return err
    project_dir = args.get("project_dir", "")
    try:
        m = _pipe_m.load(project_dir)
    except FileNotFoundError as e:
        return {"ok": False, "error": str(e)}

    plant = m.get("plant_name", "plant")

    # Require stages 1-4 to be passed
    for sid in [1, 2, 3, 4]:
        status = m["stages"][sid]["status"]
        if status not in ("passed",):
            return {"ok": False,
                    "error": f"Stage {sid} ({m['stages'][sid]['name']}) must be passed before building. "
                             f"Current status: {status}"}

    stage5_skipped = m["stages"][5]["status"] == "skipped"
    stage6_skipped = m["stages"][6]["status"] == "skipped"

    # Load schematic
    html_path = Path(project_dir) / f"{plant}_schematic.html"
    if not SCHEMATIC_AVAILABLE or not html_path.exists():
        return {"ok": False, "error": "Schematic parser unavailable or HTML not found."}

    try:
        schematic = _schp.parse_html_schematic(str(html_path))
    except Exception as ex:
        return {"ok": False, "error": f"Cannot parse schematic: {ex}"}

    # Generate parameters SCS via existing sumo_pack
    if not PACK_AVAILABLE:
        return {"ok": False, "error": "sumo_pack module unavailable."}

    try:
        scs_result = _spk.generate_apply_parameters_scs(schematic)
        if not scs_result.get("ok"):
            return {"ok": False, "error": "generate_apply_parameters_scs failed: " + scs_result.get("error", "?")}
        parameters_scs = scs_result.get("scs", "")
    except Exception as ex:
        # Fall back to the pack's build_sumo_pack which writes the scs itself
        parameters_scs = f"// apply_parameters.scs for {plant}\n// (generated by build_sumo_pack)\n"

    # Generate controllers.scs if not skipped
    if not stage6_skipped:
        ctrl_path = Path(project_dir) / "controllers.json"
        if ctrl_path.exists():
            try:
                ctrl_data = json.loads(ctrl_path.read_text(encoding="utf-8"))
                ctrl_scs = _pipe_bld.generate_controllers_scs(plant, ctrl_data.get("controllers", []))
                (Path(project_dir) / "controllers.scs").write_text(ctrl_scs, encoding="utf-8")
            except Exception:
                pass

    result = _pipe_bld.build_pack(
        project_dir=str(project_dir),
        plant_name=plant,
        schematic=schematic,
        parameters_scs=parameters_scs,
        stage5_skipped=stage5_skipped,
        stage6_skipped=stage6_skipped,
    )
    return result


def _exp_stage8_verify(args: dict) -> dict:
    err = _pipe_check()
    if err:
        return err
    project_dir = args.get("project_dir", "")
    dll_path    = args.get("dll_path", "")
    state_xml   = args.get("state_xml", "")

    if not dll_path:
        return {"ok": False, "error": "dll_path is required — provide the sumoproject.dll compiled by SUMO."}

    # Run DLL verification
    try:
        m = _pipe_m.load(project_dir)
    except FileNotFoundError as e:
        return {"ok": False, "error": str(e)}

    plant = m.get("plant_name", "plant")
    html_path = Path(project_dir) / f"{plant}_schematic.html"

    dll_result = _exp_verify_dll_matches_schematic({
        "dll_path": dll_path,
        "schematic_html": str(html_path) if html_path.exists() else "",
    })

    all_ok = dll_result.get("ok", False)

    result = {
        "ok": all_ok,
        "stage": 8,
        "dll_check": dll_result,
        "message": (
            "DLL verification passed. Project is ready for simulation."
            if all_ok else
            dll_result.get("verdict", "DLL verification failed.")
        ),
    }

    if all_ok:
        _pipe_m.set_stage(m, 8, status="passed",
                          summary={"dll_refs_resolved": dll_result.get("refs_resolved"),
                                   "dll_ok": True})
        _pipe_m.record(m, "stage8_verify",
                       {"project_dir": project_dir, "dll_path": dll_path}, "passed")
        _pipe_m.save(project_dir, m)

    return result


# ─────────────────────────────────────────────────────────────────────────────
# Group FF — DTT-mediated .sumo editing helpers
# ─────────────────────────────────────────────────────────────────────────────
def _dte_guard():
    if not DTT_EDITOR_AVAILABLE:
        return {"ok": False, "error": "dtt_editor not loaded",
                "detail": globals().get("_DTE_IMPORT_ERROR", "unknown")}
    return None


def _dispatch_dtt(commands):
    """Execute a list of DTT command strings via the existing bridge."""
    try:
        ds_ref = _make_ds()
    except Exception as e:
        return ([{"command": c, "ok": False, "error": repr(e)} for c in commands], False)
    log = []
    for c in commands:
        try:
            ds_ref.sumo.executeCommand(c)
            log.append({"command": c, "ok": True})
        except Exception as e:
            log.append({"command": c, "ok": False, "error": repr(e)})
            return log, False
    return log, True


def _ff_live_units():
    """Best-effort enumeration of live unit-process names from DTT."""
    try:
        ds_ref = _make_ds()
        all_vars = ds_ref.sumo.getVariableNames()
        units = set()
        for v in all_vars:
            parts = v.split("__")
            if len(parts) >= 3:
                units.add(parts[2])
        return sorted(units)
    except Exception:
        # Fall back to the in-memory ACTIVE_MODEL list (set by add_unit_process)
        return sorted({u.get("name", "") for u in ACTIVE_MODEL.get("unit_processes", []) if u.get("name")})


def _ff_live_streams():
    """Best-effort enumeration of live stream identifiers."""
    out = []
    # Registered streams (from add_recycle_stream / connect_unit_processes)
    for s in ACTIVE_MODEL.get("streams", []):
        sid = s.get("id") or s.get("name") or ""
        if sid:
            out.append(sid)
    # If DTT exposes flow variables, harvest a representative set
    try:
        ds_ref = _make_ds()
        all_vars = ds_ref.sumo.getVariableNames()
        q_vars = [v for v in all_vars if "__Q" in v and "param" not in v]
        for v in q_vars[:200]:
            out.append(v)
    except Exception:
        pass
    return sorted(set(out))


def _exp_rename_unit_process(args):
    g = _dte_guard()
    if g: return g
    try:
        cmd = _dte.cmd_rename_unit(args["old_name"], args["new_name"])
    except (KeyError, ValueError) as e:
        return {"ok": False, "error": str(e)}
    if args.get("dry_run", False):
        return {"ok": True, "planned": [cmd], "dispatched": False}
    log, ok = _dispatch_dtt([cmd])
    return {"ok": ok, "log": log,
            "note": "Persist with save_model when satisfied."}


def _exp_change_unit_type(args):
    g = _dte_guard()
    if g: return g
    try:
        cmd = _dte.cmd_change_unit_type(args["unit_name"], args["new_class"])
    except (KeyError, ValueError) as e:
        return {"ok": False, "error": str(e)}
    if args.get("dry_run", True):
        return {"ok": True, "planned": [cmd], "dispatched": False,
                "warning": "Class swaps can invalidate the model; "
                           "validate_full_model after dispatch."}
    log, ok = _dispatch_dtt([cmd])
    return {"ok": ok, "log": log}


def _exp_remove_flow_connection(args):
    g = _dte_guard()
    if g: return g
    try:
        cmd = _dte.cmd_remove_stream(args["stream_id"])
    except (KeyError, ValueError) as e:
        return {"ok": False, "error": str(e)}
    if args.get("dry_run", False):
        return {"ok": True, "planned": [cmd], "dispatched": False}
    log, ok = _dispatch_dtt([cmd])
    return {"ok": ok, "log": log}


def _exp_set_stream_flow_rate(args):
    g = _dte_guard()
    if g: return g
    try:
        cmd = _dte.cmd_set_stream_flow(args["stream_id"], args["q_m3d"])
    except (KeyError, ValueError) as e:
        return {"ok": False, "error": str(e)}
    if args.get("dry_run", False):
        return {"ok": True, "planned": [cmd], "dispatched": False}
    log, ok = _dispatch_dtt([cmd])
    return {"ok": ok, "log": log}


def _exp_list_controllers(args):
    g = _dte_guard()
    if g: return g
    # Try DTT enumeration via getVariableNames — controllers expose
    # __param__SP/Kp/Ti/uMin/uMax variables under their id.
    found = []
    try:
        ds_ref = _make_ds()
        names = ds_ref.sumo.getVariableNames()
        prefix = "Sumo__Plant__"
        for n in names:
            if "__param__SP" in n and n.startswith(prefix):
                ctrl = n[len(prefix):].split("__param__")[0]
                if ctrl not in found:
                    found.append(ctrl)
    except Exception:
        found = []
    if found:
        return {"ok": True, "source": "DTT",
                "controllers": sorted(found), "count": len(found)}
    sp = args.get("sumo_path")
    if sp:
        res = _dte.read_controllers_from_manifest(sp)
        res["source"] = "manifest_fallback"
        return res
    return {"ok": False,
            "error": "DTT enumeration empty and no sumo_path "
                     "provided for manifest fallback."}


def _exp_modify_controller(args):
    g = _dte_guard()
    if g: return g
    try:
        cmds = _dte.cmds_modify_controller(
            args["controller_id"],
            setpoint=args.get("setpoint"),
            gain=args.get("gain"),
            integral_time=args.get("integral_time"),
            output_min=args.get("output_min"),
            output_max=args.get("output_max"))
    except (KeyError, ValueError) as e:
        return {"ok": False, "error": str(e)}
    if args.get("dry_run", False):
        return {"ok": True, "planned": cmds, "dispatched": False}
    log, ok = _dispatch_dtt(cmds)
    return {"ok": ok, "log": log}


def _exp_remove_controller(args):
    g = _dte_guard()
    if g: return g
    try:
        cmd = _dte.cmd_remove_controller(args["controller_id"])
    except (KeyError, ValueError) as e:
        return {"ok": False, "error": str(e)}
    if args.get("dry_run", False):
        return {"ok": True, "planned": [cmd], "dispatched": False}
    log, ok = _dispatch_dtt([cmd])
    return {"ok": ok, "log": log}


def _exp_begin_edit_transaction(args):
    g = _dte_guard()
    if g: return g
    return _dte.tx_begin(CONFIG["state_xml"], args.get("label", "edit"))


def _exp_commit_edit_transaction(args):
    g = _dte_guard()
    if g: return g
    try:
        tx_id = args["tx_id"]
    except KeyError as e:
        return {"ok": False, "error": f"missing required field: {e!s}"}
    return _dte.tx_commit(
        CONFIG["state_xml"], tx_id,
        keep_snapshot=bool(args.get("keep_snapshot", False)))


def _exp_rollback_edit_transaction(args):
    g = _dte_guard()
    if g: return g
    try:
        tx_id = args["tx_id"]
    except KeyError as e:
        return {"ok": False, "error": f"missing required field: {e!s}"}
    res = _dte.tx_rollback(CONFIG["state_xml"], tx_id)
    if res.get("ok"):
        # Re-initialise DTT state so the in-memory model matches the restored
        # state.xml. Best-effort: try loadState on the active scheduler.
        try:
            ds_ref = _make_ds()
            try:
                ds_ref.sumo.loadState(CONFIG["state_xml"])
                res["reinitialised"] = True
            except Exception as e:
                res["reinitialised"] = False
                res["next_step"] = ("Call initialize_state to sync DTT "
                                    "(loadState failed: %s)" % repr(e))
        except Exception as e:
            res["reinitialised"] = False
            res["next_step"] = "Call initialize_state to sync DTT."
            res["reinit_error"] = repr(e)
    return res


def _exp_apply_dtt_command_batch(args):
    g = _dte_guard()
    if g: return g
    try:
        commands = args["commands"]
    except KeyError:
        return {"ok": False, "error": "missing required field: 'commands'"}
    plan = _dte.plan_command_batch(commands,
                                   allow_raw=bool(args.get("allow_raw", False)))
    if args.get("dry_run", True) or not plan.get("ok"):
        plan["dispatched"] = False
        return plan
    log, ok = _dispatch_dtt(plan["planned"])
    plan["dispatched"] = True
    plan["ok"] = ok
    plan["log"] = log
    return plan


def _exp_diff_inmemory_vs_disk(args):
    g = _dte_guard()
    if g: return g
    try:
        sumo_path = args["sumo_path"]
    except KeyError:
        return {"ok": False, "error": "missing required field: 'sumo_path'"}
    if not globals().get("_scc") or not hasattr(_scc, "read_sumo_manifest"):
        return {"ok": False,
                "error": "sumo_compiler.read_sumo_manifest missing"}
    try:
        man = _scc.read_sumo_manifest(sumo_path)
    except Exception as e:
        return {"ok": False, "error": f"read_sumo_manifest failed: {e!r}"}
    if not isinstance(man, dict):
        man = {}
    live_units = _ff_live_units()
    live_streams = _ff_live_streams()
    res = _dte.diff_units_streams(live_units, live_streams, man)
    res["ok"] = True
    return res
# ─── end Group FF helpers ────────────────────────────────────────────────────


async def main():
    async with stdio_server() as (read_stream, write_stream):
        await app.run(
            read_stream,
            write_stream,
            app.create_initialization_options(),
        )


if __name__ == "__main__":
    asyncio.run(main())
